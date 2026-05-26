"""Late DOCX/TXT/ODT layout bridge for OCR export.

This module intentionally does not touch OCR, LM revision, or worker scope.  It
only routes the normal Word export through the same column/table block detection
that the late TXT/ODT exporter already uses, so all three office/text exports
share one reading-order model.
"""

from .shared import *
from .ui_components import *
from .workers import *
from .dialogs import *
from .image_edit import *
from .main_window import MainWindow


def _bk_fix64_docx_set_cell_border(cell, **kwargs):
    try:
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
    except Exception:
        return
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        edge_data = kwargs.get(edge)
        if edge_data is None:
            continue
        tag = "w:%s" % edge
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        for key, value in edge_data.items():
            element.set(qn("w:%s" % key), str(value))


def _bk_fix64_docx_add_run(paragraph, text: str, font_name: str, font_size_pt: float):
    run = paragraph.add_run(str(text or ""))
    try:
        from docx.shared import Pt
        run.font.name = font_name
        run.font.size = Pt(font_size_pt)
    except Exception:
        pass
    return run


def _bk_fix64_docx_clean_text(value) -> str:
    try:
        return _bk_fix62_clean(value)
    except Exception:
        try:
            return _bk_fix36_clean_text(value)
        except Exception:
            return " ".join(str(value or "").replace("\r", "\n").split())


def _bk_fix64_docx_export_blocks(record_views, image_size):
    try:
        blocks = _bk_fix62_export_blocks(record_views, image_size)
        if blocks:
            return blocks
    except Exception:
        pass
    try:
        page_w = int((image_size or (0, 0))[0]) or 1600
        return _bk_fix51_split_blocks(list(record_views or []), page_w)
    except Exception:
        return [
            {"type": "paragraph", "text": _bk_fix64_docx_clean_text(getattr(rv, "text", ""))}
            for rv in (record_views or [])
            if _bk_fix64_docx_clean_text(getattr(rv, "text", ""))
        ]


def _bk_fix64_is_wide_register_columns(block) -> bool:
    cols = [list(c or []) for c in block.get("columns") or []]
    if len(cols) < 2:
        return False
    non_empty = [sum(1 for x in col if _bk_fix64_docx_clean_text(x)) for col in cols]
    return sum(1 for n in non_empty if n >= 3) >= 2


def _bk_fix64_docx_write_columns(doc, block, usable_width_in: float):
    from docx.shared import Inches, Pt
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT

    cols = [list(c or []) for c in block.get("columns") or []]
    if not cols:
        return
    table = doc.add_table(rows=1, cols=len(cols))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    col_width = Inches(max(1.0, usable_width_in / max(1, len(cols))))
    for c_idx, col in enumerate(cols):
        cell = table.rows[0].cells[c_idx]
        try:
            cell.width = col_width
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            cell.margin_top = cell.margin_bottom = cell.margin_left = cell.margin_right = 0
        except Exception:
            pass
        _bk_fix64_docx_set_cell_border(
            cell,
            top={"val": "nil"}, left={"val": "nil"}, bottom={"val": "nil"}, right={"val": "nil"},
            insideH={"val": "nil"}, insideV={"val": "nil"},
        )
        first = True
        for line in col:
            text = _bk_fix64_docx_clean_text(line)
            if not text:
                continue
            p = cell.paragraphs[0] if first else cell.add_paragraph()
            first = False
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.space_before = Pt(0)
            _bk_fix64_docx_add_run(p, text, "Arial", 7.2)
    doc.add_paragraph()


def _bk_fix64_docx_write_table(doc, block):
    from docx.shared import Pt
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT

    grid = block.get("rows") or []
    cols = max((len(r) for r in grid), default=0)
    if not grid or cols < 2:
        return
    table = doc.add_table(rows=len(grid), cols=cols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = True
    font_size = 6.2 if cols >= 8 else 7.4 if cols >= 5 else 8.0
    for r_idx, row in enumerate(grid):
        for c_idx in range(cols):
            cell = table.rows[r_idx].cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.space_before = Pt(0)
            text = row[c_idx] if c_idx < len(row) else ""
            _bk_fix64_docx_add_run(p, text, "Arial", font_size)
            try:
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            except Exception:
                pass
    doc.add_paragraph()


def _bk_fix64_write_docx(path: str, item: TaskItem, export_image: Image.Image, record_views: List[RecordView]):
    try:
        from docx import Document
        from docx.shared import Inches, Pt
        from docx.enum.section import WD_ORIENT
    except Exception as e:
        raise RuntimeError(translation.translate(translation.DEFAULT_LANGUAGE, "err_no_docx_package_short")) from e

    image_size = getattr(export_image, "size", None) or (0, 0)
    page_w_px, page_h_px = image_size
    blocks = _bk_fix64_docx_export_blocks(list(record_views or []), image_size)
    doc = Document()
    section = doc.sections[0]
    landscape = page_w_px > page_h_px * 1.15
    if landscape:
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width = Inches(11.69)
        section.page_height = Inches(8.27)
    else:
        section.page_width = Inches(8.27)
        section.page_height = Inches(11.69)
    margin = 0.45
    section.left_margin = section.right_margin = section.top_margin = section.bottom_margin = Inches(margin)
    usable_width_in = float(section.page_width.inches) - (2 * margin)
    try:
        normal = doc.styles["Normal"]
        normal.font.name = "Arial"
        normal.font.size = Pt(8.5)
    except Exception:
        pass

    for block in blocks:
        typ = block.get("type")
        if typ == "columns" and _bk_fix64_is_wide_register_columns(block):
            _bk_fix64_docx_write_columns(doc, block, usable_width_in)
        elif typ == "table":
            _bk_fix64_docx_write_table(doc, block)
        else:
            text = _bk_fix64_docx_clean_text(block.get("text", ""))
            if text:
                p = doc.add_paragraph(text)
                p.paragraph_format.space_after = Pt(1)
                p.paragraph_format.space_before = Pt(0)
    doc.save(path)


try:
    _BK_FIX64_PREV_RENDER_FILE = MainWindow._render_file
except Exception:
    _BK_FIX64_PREV_RENDER_FILE = None


def _bk_fix64_render_file(self, path: str, fmt: str, item: TaskItem):
    fmt_l = str(fmt or "").lower()
    if fmt_l in {"docx", ".docx"}:
        if not item or not getattr(item, "results", None):
            return
        _text, _kr, pil_image, record_views = item.results
        try:
            export_image = _load_image_color(item.path)
        except Exception:
            export_image = pil_image
        return _bk_fix64_write_docx(path, item, export_image, record_views)
    if callable(_BK_FIX64_PREV_RENDER_FILE):
        return _BK_FIX64_PREV_RENDER_FILE(self, path, fmt, item)
    return None


try:
    MainWindow._render_file = _bk_fix64_render_file
except Exception:
    pass
