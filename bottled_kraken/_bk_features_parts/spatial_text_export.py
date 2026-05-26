"""Zusatzmodul: Ditto-Zeichen, Tabellen-/Datenexporte, Overlay-Box-Hilfen und UI-Fixes.

Der Patch ist bewusst defensiv: Er hängt sich nur an vorhandene Methoden, wenn sie in der
jeweiligen Bottled-Kraken-Version existieren. Dadurch bleibt die Datei auch mit älteren
Zwischenständen lauffähig.
"""

from .shared import *
from .ui_components import *
from .workers import *
from .dialogs import *
from .image_edit import *
from .main_window import MainWindow

_BK_FIX38_DITTO_MARK_RE = re.compile(r'(?<![A-Za-zÀ-ÿÄÖÜäöüß0-9])(?:[-–—]\s*)?["„“”]{1,4}(?:\s*[-–—])?(?![A-Za-zÀ-ÿÄÖÜäöüß0-9])')

_BK_FIX38_ATTACHED_DITTO_RE = re.compile(r'(?<![A-Za-zÀ-ÿÄÖÜäöüß0-9])(?:[-–—]\s*)?["„“”]{1,4}(?:\s*[-–—])?(?=[A-Za-zÀ-ÿÄÖÜäöüß0-9])')

def _bk_fix38_token_fields(line: str) -> List[Tuple[int, int, str]]:
    fields = []
    for m in re.finditer(r"[A-Za-zÀ-ÿÄÖÜäöüß0-9][A-Za-zÀ-ÿÄÖÜäöüß0-9./,'()\-]*", str(line or "")):
        token = m.group(0).strip(" ,;:")
        if token:
            fields.append((m.start(), m.end(), token))
    return fields

def _bk_fix38_prev_value(prev_line: str, pos: int, field_index: int = -1, next_token: str = "") -> str:
    fields = _bk_fix38_token_fields(prev_line)
    if not fields:
        return ""
    if 0 <= field_index < len(fields):
        return fields[field_index][2]
    # Prefer the token covering the same visual/character column.
    covering = [f for f in fields if f[0] - 3 <= pos <= f[1] + 3]
    if covering:
        return covering[-1][2]
    nearest = min(fields, key=lambda f: min(abs(pos - f[0]), abs(pos - f[1])))
    if min(abs(pos - nearest[0]), abs(pos - nearest[1])) <= 24:
        return nearest[2]
    # If the next token also occurred in the previous line, repeat it. This handles
    # OCR like ""Beltzkey where the mark is glued to the following repeated value.
    nt = _bk_fix36_clean_text(next_token).strip(".,;:")
    if nt:
        for _s, _e, tok in fields:
            if tok.strip(".,;:").lower() == nt.lower():
                return tok
    return ""

def _bk_fix38_resolve_ditto_marks_in_lines(lines: List[str]) -> List[str]:
    """Resolve ditto marks more aggressively.

    A literal ", "", -"- or - " - is never kept as text when it appears as a
    table/register abbreviation. It is replaced by the value in the previous line
    at the same field/column whenever possible. When OCR glued the mark to the
    following repeated value, the mark is removed instead of being output literally.
    """
    out: List[str] = []
    prev_line = ""
    for raw in lines or []:
        line = str(raw or "")
        if not line.strip():
            out.append(line)
            continue

        # Determine field index of each mark by token count before the mark.
        def standalone_repl(match):
            before = line[:match.start()]
            field_index = len(_bk_fix38_token_fields(before))
            # look ahead one token for glued/nearby known repeated values
            after = line[match.end():]
            m_next = re.search(r"[A-Za-zÀ-ÿÄÖÜäöüß0-9][A-Za-zÀ-ÿÄÖÜäöüß0-9./,'()\-]*", after)
            next_token = m_next.group(0) if m_next else ""
            copied = _bk_fix38_prev_value(prev_line, match.start(), field_index, next_token)
            return copied if copied else ""

        line2 = _BK_FIX38_DITTO_MARK_RE.sub(standalone_repl, line)

        def attached_repl(match):
            # For cases like ""Beltzkey: first try a previous value; if none is
            # reliable, remove the mark so the following word remains clean.
            before = line[:match.start()]
            field_index = len(_bk_fix38_token_fields(before))
            after = line[match.end():]
            m_next = re.search(r"[A-Za-zÀ-ÿÄÖÜäöüß0-9][A-Za-zÀ-ÿÄÖÜäöüß0-9./,'()\-]*", after)
            next_token = m_next.group(0) if m_next else ""
            copied = _bk_fix38_prev_value(prev_line, match.start(), field_index, next_token)
            # Avoid duplicate 'BeltzkeyBeltzkey' when the next token is already that value.
            if copied and next_token and copied.strip(".,;:").lower() == next_token.strip(".,;:").lower():
                return ""
            return (copied + " ") if copied else ""

        line2 = _BK_FIX38_ATTACHED_DITTO_RE.sub(attached_repl, line2)
        line2 = re.sub(r"\s{2,}", " ", line2).strip()
        line2 = re.sub(r"\s+([.,;:])", r"\1", line2)
        out.append(line2)
        if line2:
            prev_line = line2
    return out

_bk_fix36_resolve_ditto_marks_in_lines = _bk_fix38_resolve_ditto_marks_in_lines

_bk_fix37_resolve_ditto_marks_in_lines = _bk_fix38_resolve_ditto_marks_in_lines

def _bk_fix38_expand_ditto_text(text: str) -> str:
    return "\n".join(_bk_fix38_resolve_ditto_marks_in_lines(str(text or "").splitlines()))

_bk_fix37_expand_ditto_text = _bk_fix38_expand_ditto_text

def _bk_fix38_spatial_text_from_recs(recs) -> str:
    recs = list(recs or [])
    recs = _bk_fix36_resolve_ditto_marks_in_recs(recs)
    with_boxes = [rv for rv in recs if getattr(rv, "bbox", None) and _bk_fix36_clean_text(getattr(rv, "text", ""))]
    if not with_boxes:
        return "\n".join(_bk_fix36_clean_text(x) for x in _bk_fix38_resolve_ditto_marks_in_lines([getattr(rv, "text", "") for rv in recs]) if _bk_fix36_clean_text(x)) + "\n"

    min_x = min(float(rv.bbox[0]) for rv in with_boxes)
    max_x = max(float(rv.bbox[2]) for rv in with_boxes)
    min_y = min(float(rv.bbox[1]) for rv in with_boxes)
    max_y = max(float(rv.bbox[3]) for rv in with_boxes)
    width_px = max(1.0, max_x - min_x)
    target_cols = max(90, min(220, int(width_px / 5.5)))
    px_per_col = max(1.0, width_px / float(target_cols))

    heights = sorted(max(1, int(rv.bbox[3] - rv.bbox[1])) for rv in with_boxes)
    median_h = heights[len(heights)//2] if heights else 18
    row_threshold = max(7, median_h * 0.62)
    rows: List[Tuple[float, List[Any]]] = []
    for rv in sorted(with_boxes, key=lambda r: ((r.bbox[1] + r.bbox[3]) / 2.0, r.bbox[0])):
        cy = (rv.bbox[1] + rv.bbox[3]) / 2.0
        for idx, (rcy, row) in enumerate(rows):
            if abs(cy - rcy) <= row_threshold:
                row.append(rv)
                rows[idx] = ((rcy * (len(row)-1) + cy) / max(1, len(row)), row)
                break
        else:
            rows.append((cy, [rv]))

    output_lines: List[str] = []
    for _cy, row in rows:
        row = sorted(row, key=lambda r: r.bbox[0])
        canvas = [" "] * (target_cols + 80)
        for rv in row:
            bb = rv.bbox
            txt = _bk_fix36_clean_text(getattr(rv, "text", ""))
            if not txt:
                continue
            col = max(0, int((float(bb[0]) - min_x) / px_per_col))
            # If the target area is occupied, shift right to keep both cells.
            while col < len(canvas) and any(ch != " " for ch in canvas[col:col+min(len(txt), 12)]):
                col += 1
            end = min(len(canvas), col + len(txt))
            for i, ch in enumerate(txt[:max(0, end-col)]):
                canvas[col+i] = ch
        output_lines.append("".join(canvas).rstrip())
    return "\n".join(line for line in output_lines if line.strip()).rstrip() + "\n"

_bk_fix36_table_text_from_recs = _bk_fix38_spatial_text_from_recs

def _bk_fix38_default_export_name(task, suffix: str) -> str:
    return os.path.splitext(os.path.basename(getattr(task, "path", "ocr") or "ocr"))[0] + suffix

def _bk_fix38_export_task_as_spatial_txt(self, task):
    if not task or not getattr(task, "results", None):
        QMessageBox.information(self, _bk_fix36_tr(self, "info_title"), _bk_fix36_tr(self, "warn_no_ocr_results"))
        return
    _text, _kr, _im, recs = task.results
    out_text = _bk_fix38_spatial_text_from_recs(recs)
    default_name = _bk_fix38_default_export_name(task, ".txt")
    start_dir = getattr(self, "current_export_dir", "") or os.path.dirname(getattr(task, "path", "") or "") or os.getcwd()
    path, _ = QFileDialog.getSaveFileName(self, _bk_fix36_tr(self, "export_format_txt_plain", "Text (.txt)"), os.path.join(start_dir, default_name), _bk_fix36_tr(self, "filter_text_files", "Text Files (*.txt)"))
    if not path:
        return
    if not path.lower().endswith(".txt"):
        path += ".txt"
    with open(path, "w", encoding="utf-8") as fh:
        fh.write(out_text)
    try:
        self.current_export_dir = os.path.dirname(path)
        self.status_bar.showMessage(_bk_fix36_tr(self, "msg_origami_table_export_done", "Table-layout export finished."), 5000)
    except Exception:
        pass

def _bk_fix38_export_task_as_spatial_docx(self, task):
    if not task or not getattr(task, "results", None):
        QMessageBox.information(self, _bk_fix36_tr(self, "info_title"), _bk_fix36_tr(self, "warn_no_ocr_results"))
        return
    _text, _kr, _im, recs = task.results
    out_text = _bk_fix38_spatial_text_from_recs(recs)
    default_name = _bk_fix38_default_export_name(task, ".docx")
    start_dir = getattr(self, "current_export_dir", "") or os.path.dirname(getattr(task, "path", "") or "") or os.getcwd()
    path, _ = QFileDialog.getSaveFileName(self, _bk_fix36_tr(self, "export_format_docx", "Word (.docx)"), os.path.join(start_dir, default_name), _bk_fix36_tr(self, "filter_word_files", "Word Files (*.docx)"))
    if not path:
        return
    if not path.lower().endswith(".docx"):
        path += ".docx"
    try:
        from docx import Document
        from docx.shared import Pt
        doc = Document()
        style = doc.styles["Normal"]
        style.font.name = "Courier New"
        style.font.size = Pt(8)
        for line in out_text.splitlines() or [""]:
            p = doc.add_paragraph()
            r = p.add_run(line)
            r.font.name = "Courier New"
            r.font.size = Pt(8)
        doc.save(path)
    except Exception as exc:
        alt = path.rsplit(".", 1)[0] + ".txt"
        with open(alt, "w", encoding="utf-8") as fh:
            fh.write(out_text)
        QMessageBox.information(self, _bk_fix36_tr(self, "info_title"), _bk_fix36_tr(self, "msg_docx_missing_txt_saved_with_error", exc, alt))
        return
    try:
        self.current_export_dir = os.path.dirname(path)
        self.status_bar.showMessage(_bk_fix36_tr(self, "msg_origami_table_export_done", "Table-layout export finished."), 5000)
    except Exception:
        pass

MainWindow.bk_export_origami_table_txt = lambda self: _bk_fix38_export_task_as_spatial_txt(self, _bk_fix36_current_task(self))

MainWindow.bk_export_origami_table_docx = lambda self: _bk_fix38_export_task_as_spatial_docx(self, _bk_fix36_current_task(self))

def _bk_fix38_is_standard_txt_docx(fmt: str) -> bool:
    f = str(fmt or "").lower()
    if f in {"pdf", "csv", "json", "alto", "hocr", "xml", "html"}:
        return False
    return f in {"txt", "text", "txt_plain", "txt_boxes", "docx", "word"} or f.endswith(".txt") or f.endswith(".docx") or "docx" in f or (f.startswith("txt") and "origami" not in f)

def _bk_fix38_export_format_items(self):
    items = []
    if callable(_BK_FIX36_PREV_EXPORT_FORMAT_ITEMS):
        try:
            items = list(_BK_FIX36_PREV_EXPORT_FORMAT_ITEMS(self) or [])
        except Exception:
            items = []
    filtered = []
    for name, fmt in items:
        if str(fmt) in {"txt_origami_table", "docx_origami_table"}:
            continue
        if str(fmt).lower() in {"txt_plain", "txt", "text"}:
            name = _bk_fix36_tr(self, "export_format_txt_plain", str(name))
        if "docx" in str(fmt).lower() or str(fmt).lower() == "word":
            name = _bk_fix36_tr(self, "export_format_docx", str(name))
        filtered.append((name, fmt))
    return filtered

try:
    MainWindow._export_format_items = _bk_fix38_export_format_items
except Exception:
    pass

try:
    _BK_FIX38_PREV_EXPORT_SINGLE_INTERACTIVE = MainWindow._export_single_interactive
except Exception:
    _BK_FIX38_PREV_EXPORT_SINGLE_INTERACTIVE = None

def _bk_fix38_export_single_interactive(self, task, fmt, *args, **kwargs):
    if _bk_fix38_is_standard_txt_docx(fmt):
        if "docx" in str(fmt).lower() or str(fmt).lower() == "word":
            return _bk_fix38_export_task_as_spatial_docx(self, task)
        return _bk_fix38_export_task_as_spatial_txt(self, task)
    if callable(_BK_FIX38_PREV_EXPORT_SINGLE_INTERACTIVE):
        return _BK_FIX38_PREV_EXPORT_SINGLE_INTERACTIVE(self, task, fmt, *args, **kwargs)
    return None

if callable(_BK_FIX38_PREV_EXPORT_SINGLE_INTERACTIVE):
    MainWindow._export_single_interactive = _bk_fix38_export_single_interactive

_BK_FIX38_PREV_INIT = MainWindow.__init__

def _bk_fix38_mainwindow_init(self, *args, **kwargs):
    _BK_FIX38_PREV_INIT(self, *args, **kwargs)
    try:
        _bk_fix37_ensure_sqlite_menu_action(self)
    except Exception:
        pass

_BK_FIX38_PREV_RETRANSLATE = MainWindow.retranslate_ui

def _bk_fix38_retranslate_ui(self, *args, **kwargs):
    _BK_FIX38_PREV_RETRANSLATE(self, *args, **kwargs)
    try:
        _bk_fix37_ensure_sqlite_menu_action(self)
        if hasattr(self, "export_format_actions"):
            for fmt in ("txt_origami_table", "docx_origami_table"):
                act = self.export_format_actions.get(fmt)
                if act is not None:
                    try:
                        self.export_menu.removeAction(act)
                    except Exception:
                        pass
    except Exception:
        pass

MainWindow.__init__ = _bk_fix38_mainwindow_init

MainWindow.retranslate_ui = _bk_fix38_retranslate_ui

def _bk_fix38_cleanup_task_ditto(self, task=None):
    try:
        if task is None:
            task = _bk_fix36_current_task(self)
        if task is not None and getattr(task, "results", None):
            text, kr_records, im, recs = task.results
            before = [str(getattr(r, "text", "") or "") for r in (recs or [])]
            _bk_fix36_resolve_ditto_marks_in_recs(recs)
            after = [str(getattr(r, "text", "") or "") for r in (recs or [])]
            if before != after:
                task.results = ("\n".join(after), kr_records, im, recs)
                task.edited = True
    except Exception:
        pass

try:
    _BK_FIX38_PREV_SYNC_AFTER_RECS = MainWindow._sync_ui_after_recs_change
except Exception:
    _BK_FIX38_PREV_SYNC_AFTER_RECS = None

def _bk_fix38_sync_ui_after_recs_change(self, task, *args, **kwargs):
    _bk_fix38_cleanup_task_ditto(self, task)
    if callable(_BK_FIX38_PREV_SYNC_AFTER_RECS):
        return _BK_FIX38_PREV_SYNC_AFTER_RECS(self, task, *args, **kwargs)
    return None

if callable(_BK_FIX38_PREV_SYNC_AFTER_RECS):
    MainWindow._sync_ui_after_recs_change = _bk_fix38_sync_ui_after_recs_change

def _bk_fix38_remove_old_origami_actions(self):
    try:
        if hasattr(self, "export_menu") and hasattr(self, "export_format_actions"):
            for fmt in ("txt_origami_table", "docx_origami_table"):
                act = self.export_format_actions.pop(fmt, None)
                if act is not None:
                    self.export_menu.removeAction(act)
    except Exception:
        pass

_BK_FIX38B_PREV_INIT = MainWindow.__init__

def _bk_fix38b_mainwindow_init(self, *args, **kwargs):
    _BK_FIX38B_PREV_INIT(self, *args, **kwargs)
    try:
        _bk_fix38_remove_old_origami_actions(self)
        _bk_fix38_ensure_lm_page_boxes_action(self)
        _bk_fix37_ensure_sqlite_menu_action(self)
        _bk_fix38_cleanup_task_ditto(self)
    except Exception:
        pass

_BK_FIX38B_PREV_RETRANSLATE = MainWindow.retranslate_ui

def _bk_fix38b_retranslate_ui(self, *args, **kwargs):
    _BK_FIX38B_PREV_RETRANSLATE(self, *args, **kwargs)
    try:
        _bk_fix38_remove_old_origami_actions(self)
        _bk_fix38_ensure_lm_page_boxes_action(self)
        _bk_fix37_ensure_sqlite_menu_action(self)
    except Exception:
        pass

MainWindow.__init__ = _bk_fix38b_mainwindow_init

MainWindow.retranslate_ui = _bk_fix38b_retranslate_ui
