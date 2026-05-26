"""Zusatzmodul: Ditto-Zeichen, Tabellen-/Datenexporte, Overlay-Box-Hilfen und UI-Fixes.

Der Patch ist bewusst defensiv: Er hängt sich nur an vorhandene Methoden, wenn sie in der
jeweiligen Bottled-Kraken-Version existieren. Dadurch bleibt die Datei auch mit älteren
Zwischenständen lauffähig.
"""

from .shared import *
from .ui_components import *
from .workers import *
from .dialogs import *
from .image_edit import *
from .main_window import MainWindow

def _bk_fix40_group_rows(recs, page_width: int = 0):
    recs = [rv for rv in (recs or []) if _bk_fix36_clean_text(getattr(rv, "text", ""))]
    with_boxes = [rv for rv in recs if getattr(rv, "bbox", None)]
    if not with_boxes:
        return [[rv] for rv in recs]
    try:
        return group_rows_by_y(with_boxes, max(1, int(page_width or 1)))
    except Exception:
        return _bk_fix36_group_recs_into_table_rows(with_boxes)

def _bk_fix40_column_anchors(rows, page_width: int) -> List[float]:
    xs = []
    for row in rows:
        for rv in row:
            if not getattr(rv, "bbox", None):
                continue
            txt = _bk_fix36_clean_text(getattr(rv, "text", ""))
            if not txt:
                continue
            x0, _y0, x1, _y1 = rv.bbox
            if (x1 - x0) > max(1, page_width) * 0.72 and len(row) <= 1:
                continue
            xs.append(float(x0))
    if not xs:
        return []
    xs.sort()
    clusters = []
    threshold = max(18.0, min(48.0, page_width * 0.035))
    for x in xs:
        if not clusters or abs(x - clusters[-1][-1]) > threshold:
            clusters.append([x])
        else:
            clusters[-1].append(x)
    anchors = [sum(c) / len(c) for c in clusters if len(c) >= 2 or len(xs) <= 12]
    if len(anchors) > 10:
        step = len(anchors) / 10.0
        anchors = [anchors[int(i * step)] for i in range(10)]
    return anchors

def _bk_fix40_table_block_score(rows, page_width: int) -> float:
    if len(rows) < 3:
        return 0.0
    multi = 0
    wide = 0
    numeric = 0
    total_cells = 0
    for row in rows:
        cells = [rv for rv in row if _bk_fix36_clean_text(getattr(rv, "text", ""))]
        total_cells += len(cells)
        if len(cells) >= 3:
            multi += 1
        boxes = [rv.bbox for rv in cells if getattr(rv, "bbox", None)]
        if boxes and (max(bb[2] for bb in boxes) - min(bb[0] for bb in boxes)) > page_width * 0.45:
            wide += 1
        for rv in cells:
            if re.search(r"\d", str(getattr(rv, "text", ""))):
                numeric += 1
    if total_cells <= 0:
        return 0.0
    return (multi / len(rows)) + (wide / len(rows)) + min(0.6, numeric / total_cells)

def _bk_fix40_split_docx_blocks(recs, page_width: int) -> List[Dict[str, Any]]:
    _bk_fix40_resolve_ditto_marks_in_recs(recs)
    rows = _bk_fix40_group_rows(recs, page_width)
    blocks = []
    i = 0
    while i < len(rows):
        row = rows[i]
        j = i
        run = []
        while j < len(rows):
            cells = [rv for rv in rows[j] if _bk_fix36_clean_text(getattr(rv, "text", ""))]
            if len(cells) >= 2:
                run.append(rows[j]); j += 1
            else:
                break
        score = _bk_fix40_table_block_score(run, page_width) if run else 0.0
        anchors = _bk_fix40_column_anchors(run, page_width) if run else []
        if len(run) >= 3 and len(anchors) >= 3 and score >= 1.35:
            grid = []
            for rr in run:
                line = [""] * len(anchors)
                for rv in sorted(rr, key=lambda r: r.bbox[0] if getattr(r, "bbox", None) else 0):
                    txt = _bk_fix36_clean_text(getattr(rv, "text", ""))
                    if not txt:
                        continue
                    c = min(range(len(anchors)), key=lambda k: abs((rv.bbox[0] if getattr(rv, "bbox", None) else 0) - anchors[k]))
                    line[c] = (line[c] + " " + txt).strip() if line[c] else txt
                grid.append(line)
            boxes = [rv.bbox for rr in run for rv in rr if getattr(rv, "bbox", None)]
            bb = (min(b[0] for b in boxes), min(b[1] for b in boxes), max(b[2] for b in boxes), max(b[3] for b in boxes)) if boxes else None
            blocks.append({"type": "table", "rows": grid, "anchors": anchors, "bbox": bb})
            i = j
            continue

        txt = " ".join(_bk_fix36_clean_text(getattr(rv, "text", "")) for rv in sorted(row, key=lambda r: r.bbox[0] if getattr(r, "bbox", None) else 0)).strip()
        if txt:
            boxes = [rv.bbox for rv in row if getattr(rv, "bbox", None)]
            bb = (min(b[0] for b in boxes), min(b[1] for b in boxes), max(b[2] for b in boxes), max(b[3] for b in boxes)) if boxes else None
            blocks.append({"type": "paragraph", "text": txt, "bbox": bb})
        i += 1
    return blocks

def _bk_fix40_write_structured_txt_export(path: str, record_views: List[RecordView], image_size=None):
    _bk_fix40_resolve_ditto_marks_in_recs(record_views)
    img_w, img_h = image_size or (0, 0)
    with open(path, "w", encoding="utf-8") as f:
        f.write("# Bottled Kraken line export\n")
        f.write("# idx\tx\ty\twidth\theight\tx0\ty0\tx1\ty1\tx0_norm\ty0_norm\tx1_norm\ty1_norm\ttext\n")
        for i, rv in enumerate(record_views or []):
            txt = str(getattr(rv, "text", "") or "")
            if getattr(rv, "bbox", None):
                x0, y0, x1, y1 = [int(v) for v in rv.bbox]
                w, h = max(0, x1 - x0), max(0, y1 - y0)
                norms = [
                    (x0 / img_w) if img_w else "",
                    (y0 / img_h) if img_h else "",
                    (x1 / img_w) if img_w else "",
                    (y1 / img_h) if img_h else "",
                ]
                norm_cols = [("" if n == "" else f"{float(n):.6f}") for n in norms]
                cols = [i, x0, y0, w, h, x0, y0, x1, y1, *norm_cols, json.dumps(txt, ensure_ascii=False)]
            else:
                cols = [i, "", "", "", "", "", "", "", "", "", "", "", "", json.dumps(txt, ensure_ascii=False)]
            f.write("\t".join(str(c) for c in cols) + "\n")

def _bk_fix40_write_docx_export(path: str, item: TaskItem, export_image: Image.Image, record_views: List[RecordView]):
    try:
        from docx import Document
        from docx.shared import Inches, Pt
        from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
        from docx.enum.section import WD_ORIENT
    except Exception as e:
        raise RuntimeError(translation.translate(translation.DEFAULT_LANGUAGE, "err_no_docx_package_short")) from e

    _bk_fix40_resolve_ditto_marks_in_recs(record_views)
    page_w_px, page_h_px = export_image.size
    blocks = _bk_fix40_split_docx_blocks(list(record_views or []), page_w_px)
    doc = Document()
    section = doc.sections[0]
    landscape = page_w_px > page_h_px * 1.15
    if landscape:
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width = Inches(11.69)
        section.page_height = Inches(8.27)
    else:
        section.page_width = Inches(8.27)
        section.page_height = Inches(11.69)
    margin = 0.35
    section.left_margin = section.right_margin = section.top_margin = section.bottom_margin = Inches(margin)
    try:
        normal = doc.styles["Normal"]
        normal.font.name = "Arial"
        normal.font.size = Pt(9)
    except Exception:
        pass

    for block in blocks:
        if block.get("type") == "table":
            rows = block.get("rows") or []
            if not rows:
                continue
            cols = max(len(r) for r in rows)
            table = doc.add_table(rows=len(rows), cols=cols)
            table.style = "Table Grid"
            table.alignment = WD_TABLE_ALIGNMENT.LEFT
            table.autofit = True
            for r_idx, row in enumerate(rows):
                for c_idx in range(cols):
                    cell = table.rows[r_idx].cells[c_idx]
                    cell.text = ""
                    p = cell.paragraphs[0]
                    p.paragraph_format.space_after = Pt(0)
                    run = p.add_run(row[c_idx] if c_idx < len(row) else "")
                    run.font.name = "Arial"
                    run.font.size = Pt(7.5 if cols >= 7 else 8.5)
                    try:
                        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                    except Exception:
                        pass
            doc.add_paragraph()
        else:
            text = str(block.get("text") or "").strip()
            if text:
                p = doc.add_paragraph(text)
                p.paragraph_format.space_after = Pt(2)

    doc.save(path)

try:
    _BK_FIX40_PREV_RENDER_FILE = MainWindow._render_file
except Exception:
    _BK_FIX40_PREV_RENDER_FILE = None

def _bk_fix40_render_file(self, path: str, fmt: str, item: TaskItem):
    if not item.results:
        return
    text, kr_records, pil_image, record_views = item.results
    export_image = _load_image_color(item.path)
    if pil_image is None:
        pil_image = export_image
    fmt_l = str(fmt or "").lower()
    if fmt_l == "txt_boxes":
        return _bk_fix40_write_structured_txt_export(path, record_views, export_image.size)
    if fmt_l == "docx":
        return _bk_fix40_write_docx_export(path, item, export_image, record_views)
    if fmt_l == "txt":
        _bk_fix40_resolve_ditto_marks_in_recs(record_views)
        with open(path, "w", encoding="utf-8") as f:
            for rv in record_views or []:
                line = str(getattr(rv, "text", "") or "").replace("\r\n", " ").replace("\r", " ").replace("\n", " ").strip()
                if line:
                    f.write(line + "\n")
        return
    if callable(_BK_FIX40_PREV_RENDER_FILE):
        return _BK_FIX40_PREV_RENDER_FILE(self, path, fmt, item)
    return None

try:
    MainWindow._render_file = _bk_fix40_render_file
except Exception:
    pass

def _bk_fix40_export_single_interactive(self, item: TaskItem, fmt: str):
    base_name = self._export_default_stem(item, fmt) if hasattr(self, "_export_default_stem") else os.path.splitext(getattr(item, "display_name", "ocr"))[0]
    base_dir = getattr(self, "current_export_dir", "") or os.path.dirname(getattr(item, "path", "") or "") or os.getcwd()
    ext = self._export_ext(fmt) if hasattr(self, "_export_ext") else str(fmt).lower()
    dest_path, _ = QFileDialog.getSaveFileName(
        self,
        self._tr("dlg_save"),
        os.path.join(base_dir, base_name),
        self._export_filter(fmt) if hasattr(self, "_export_filter") else "All (*.*)",
    )
    if not dest_path:
        return
    if ext and not dest_path.lower().endswith(f".{ext}"):
        dest_path += f".{ext}"
    try:
        self._render_file(dest_path, fmt, item)
    except Exception as e:
        QMessageBox.critical(self, self._tr("err_title"), str(e))
        return
    try:
        self.current_export_dir = os.path.dirname(dest_path)
        self._log(self._tr_log("log_export_single", item.display_name, dest_path))
        self.status_bar.showMessage(self._tr("msg_exported", os.path.basename(dest_path)))
    except Exception:
        pass

try:
    MainWindow._export_single_interactive = _bk_fix40_export_single_interactive
except Exception:
    pass
