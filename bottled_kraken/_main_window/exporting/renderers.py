from bottled_kraken.common import _load_image_color
from bottled_kraken.common import (
    Any,
    BytesIO,
    Dict,
    Image,
    List,
    RecordView,
    TaskItem,
    base64,
    containers,
    csv,
    docx_layout_blocks,
    html,
    json,
    os,
    pdf_canvas,
    serialization,
    table_to_rows,
)
class MainWindowExportRenderersMixin:
        def _build_kraken_segmentation_for_export(
                self,
                image_path: str,
                record_views: List[RecordView]
        ):
            export_lines = []
            for i, rv in enumerate(record_views):
                if not rv.bbox:
                    continue
                x0, y0, x1, y1 = rv.bbox
                export_lines.append(
                    containers.BBoxLine(
                        id=f"line_{i + 1:04d}",
                        bbox=(int(x0), int(y0), int(x1), int(y1)),
                        text=str(rv.text or ""),
                        base_dir=None,
                        imagename=image_path,
                        regions=None,
                        tags=None,
                        split=None,
                        text_direction="horizontal-lr",
                    )
                )
            if not export_lines:
                return None
            return containers.Segmentation(
                type="bbox",
                imagename=image_path,
                text_direction="horizontal-lr",
                script_detection=False,
                lines=export_lines,
                regions=None,
                line_orders=None,
            )
        def _render_hocr_html(self, path: str, item: TaskItem, export_image: Image.Image, record_views: List[RecordView]):
            buf = BytesIO()
            export_image.save(buf, format="PNG")
            img_b64 = base64.b64encode(buf.getvalue()).decode("ascii")
            width, height = export_image.size
            page_name = html.escape(os.path.basename(item.path))
            line_blocks = []
            for i, rv in enumerate(record_views):
                if not rv.bbox:
                    continue
                x0, y0, x1, y1 = rv.bbox
                w = max(1, x1 - x0)
                h = max(1, y1 - y0)
                txt = html.escape(rv.text or "")
                line_blocks.append(f"""
                <span class="ocr_line"
                      id="line_{i + 1:04d}"
                      title="bbox {x0} {y0} {x1} {y1}"
                      style="left:{x0}px; top:{y0}px; width:{w}px; height:{h}px;">
                    <span class="ocrx_word"
                          id="word_{i + 1:04d}"
                          title="bbox {x0} {y0} {x1} {y1}">{txt}</span>
                </span>
                """)
            html_doc = f"""<!DOCTYPE html>
        <html>
        <head>
        <meta charset="utf-8">
        <meta name="ocr-system" content="Bottled Kraken">
        <meta name="ocr-capabilities" content="ocr_page ocr_line ocrx_word">
        <title>{page_name}</title>
        <style>
            body {{
                margin: 0;
                padding: 20px;
                background: #f3f3f3;
                font-family: Arial, sans-serif;
            }}
            .page-wrap {{
                display: inline-block;
                position: relative;
                box-shadow: 0 2px 16px rgba(0,0,0,0.18);
                background: white;
            }}
            .ocr_page {{
                position: relative;
                width: {width}px;
                height: {height}px;
                overflow: hidden;
                background: white;
            }}
            .ocr_page img {{
                position: absolute;
                left: 0;
                top: 0;
                width: {width}px;
                height: {height}px;
                display: block;
            }}
            .ocr_line {{
                position: absolute;
                box-sizing: border-box;
                border: 1px solid rgba(220, 38, 38, 0.45);
                background: rgba(255, 255, 255, 0.10);
                overflow: hidden;
                white-space: nowrap;
            }}
            .ocrx_word {{
                position: absolute;
                left: 0;
                top: 0;
                font-size: 12px;
                line-height: 1.1;
                color: rgba(180, 0, 0, 0.92);
                background: rgba(255, 255, 255, 0.55);
                padding: 0 2px;
            }}
        </style>
        </head>
        <body>
        <div class="page-wrap">
            <div class="ocr_page" title="image {page_name}; bbox 0 0 {width} {height}">
                <img src="data:image/png;base64,{img_b64}" alt="{page_name}">
                {''.join(line_blocks)}
            </div>
        </div>
        </body>
        </html>
        """
            with open(path, "w", encoding="utf-8") as f:
                f.write(html_doc)
        def _line_export_entry(self, rv: RecordView, fallback_idx: int) -> Dict[str, Any]:
            idx = int(getattr(rv, "idx", fallback_idx))
            entry: Dict[str, Any] = {
                "idx": idx,
                "text": str(getattr(rv, "text", "") or ""),
            }
            if rv.bbox:
                x0, y0, x1, y1 = [int(v) for v in rv.bbox]
                entry.update({
                    "x": x0,
                    "y": y0,
                    "width": max(0, x1 - x0),
                    "height": max(0, y1 - y0),
                    "bbox": [x0, y0, x1, y1],
                })
            else:
                entry.update({"x": None, "y": None, "width": None, "height": None, "bbox": None})
            return entry
        def _write_plain_txt_export(self, path: str, record_views: List[RecordView]):
            with open(path, "w", encoding="utf-8") as f:
                for rv in record_views:
                    line = str(getattr(rv, "text", "") or "")
                    line = line.replace("\r\n", " ").replace("\r", " ").replace("\n", " ")
                    f.write(line.strip() + "\n")
        def _write_structured_txt_export(self, path: str, record_views: List[RecordView]):
            with open(path, "w", encoding="utf-8") as f:
                f.write("# Bottled Kraken line export\n")
                f.write("# idx\tx\ty\twidth\theight\ttext\n")
                for i, rv in enumerate(record_views):
                    entry = self._line_export_entry(rv, i)
                    cols = [
                        str(entry["idx"]),
                        "" if entry["x"] is None else str(entry["x"]),
                        "" if entry["y"] is None else str(entry["y"]),
                        "" if entry["width"] is None else str(entry["width"]),
                        "" if entry["height"] is None else str(entry["height"]),
                        json.dumps(entry["text"], ensure_ascii=False),
                    ]
                    f.write("\t".join(cols) + "\n")
        def _docx_set_cell_text(self, cell, text: str):
            from docx.shared import Pt
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(text or ""))
            try:
                run.font.size = Pt(8.5)
            except Exception:
                pass
        def _docx_set_cell_width(self, cell, width_twips: int):
            try:
                from docx.oxml import OxmlElement
                from docx.oxml.ns import qn
                tc = cell._tc
                tc_pr = tc.get_or_add_tcPr()
                tc_w = tc_pr.first_child_found_in("w:tcW")
                if tc_w is None:
                    tc_w = OxmlElement("w:tcW")
                    tc_pr.append(tc_w)
                tc_w.set(qn("w:w"), str(max(120, int(width_twips))))
                tc_w.set(qn("w:type"), "dxa")
            except Exception:
                pass
        def _docx_set_table_layout_fixed(self, table):
            try:
                from docx.oxml import OxmlElement
                from docx.oxml.ns import qn
                tbl = table._tbl
                tbl_pr = tbl.tblPr
                layout = tbl_pr.first_child_found_in("w:tblLayout")
                if layout is None:
                    layout = OxmlElement("w:tblLayout")
                    tbl_pr.append(layout)
                layout.set(qn("w:type"), "fixed")
            except Exception:
                pass
        def _write_docx_export(self, path: str, item: TaskItem, export_image: Image.Image, record_views: List[RecordView]):
            try:
                from docx import Document
                from docx.shared import Inches, Pt
                from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
            except Exception as e:
                raise RuntimeError(self._tr("err_docx_missing")) from e
            page_w_px, page_h_px = export_image.size
            blocks = docx_layout_blocks(record_views, page_w_px, page_h_px)
            document = Document()
            section = document.sections[0]
            page_width_in = 8.27
            page_height_in = max(9.0, min(14.0, page_width_in * (page_h_px / max(1, page_w_px))))
            section.page_width = Inches(page_width_in)
            section.page_height = Inches(page_height_in)
            margin_in = 0.35
            section.left_margin = Inches(margin_in)
            section.right_margin = Inches(margin_in)
            section.top_margin = Inches(margin_in)
            section.bottom_margin = Inches(margin_in)
            usable_width_in = max(1.0, page_width_in - (2 * margin_in))
            usable_height_in = max(1.0, page_height_in - (2 * margin_in))
            scale_x = usable_width_in / max(1, page_w_px)
            scale_y = usable_height_in / max(1, page_h_px)
            last_bottom = 0
            try:
                normal_style = document.styles["Normal"]
            except Exception:
                normal_style = None
            if normal_style is not None:
                normal_style.font.name = "Arial"
                normal_style.font.size = Pt(9)
            for block in blocks:
                bb = block.get("bbox")
                top = int(block.get("top") or (bb[1] if bb else last_bottom))
                gap_px = max(0, top - int(last_bottom))
                if block.get("type") == "table":
                    rows = block.get("rows") or []
                    if not rows:
                        continue
                    if gap_px > 4:
                        spacer = document.add_paragraph()
                        spacer.paragraph_format.space_before = Pt(min(36, gap_px * scale_y * 72.0 * 0.35))
                        spacer.paragraph_format.space_after = Pt(0)
                    col_count = max(len(row) for row in rows)
                    table = document.add_table(rows=len(rows), cols=col_count)
                    table.style = "Table Grid"
                    table.alignment = WD_TABLE_ALIGNMENT.LEFT
                    table.autofit = False
                    self._docx_set_table_layout_fixed(table)
                    anchors = block.get("anchors") or []
                    if len(anchors) == col_count and page_w_px > 0:
                        bounds = []
                        for c, a in enumerate(anchors):
                            if c == 0:
                                left = float(bb[0] if bb else 0)
                            else:
                                left = (float(anchors[c - 1]) + float(a)) / 2.0
                            if c == col_count - 1:
                                right = float(bb[2] if bb else page_w_px)
                            else:
                                right = (float(a) + float(anchors[c + 1])) / 2.0
                            bounds.append(max(1.0, right - left))
                        total = max(1.0, sum(bounds))
                        col_width_twips = [int((w / total) * usable_width_in * 1440) for w in bounds]
                    else:
                        col_width_twips = [int((usable_width_in / max(1, col_count)) * 1440)] * col_count
                    for r_idx, row in enumerate(rows):
                        cells = table.rows[r_idx].cells
                        for c_idx in range(col_count):
                            txt = row[c_idx] if c_idx < len(row) else ""
                            self._docx_set_cell_text(cells[c_idx], txt)
                            self._docx_set_cell_width(cells[c_idx], col_width_twips[min(c_idx, len(col_width_twips) - 1)])
                            try:
                                cells[c_idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                            except Exception:
                                pass
                    if bb:
                        last_bottom = int(bb[3])
                    continue
                text = str(block.get("text") or "").strip()
                if not text:
                    continue
                p = document.add_paragraph()
                pf = p.paragraph_format
                if bb:
                    pf.left_indent = Inches(max(0.0, float(bb[0]) * scale_x))
                    line_height_pt = max(7.0, min(18.0, (float(bb[3] - bb[1]) * scale_y * 72.0) * 0.78))
                    pf.line_spacing = Pt(line_height_pt)
                    last_bottom = int(bb[3])
                if gap_px > 4:
                    pf.space_before = Pt(min(24, gap_px * scale_y * 72.0 * 0.25))
                pf.space_after = Pt(0)
                run = p.add_run(text)
                try:
                    run.font.size = Pt(9)
                except Exception:
                    pass
            try:
                document.save(path)
            except PermissionError as e:
                raise PermissionError(self._tr("export_permission_error", path)) from e
        def _render_file(self, path: str, fmt: str, item: TaskItem):
            if not item.results:
                return
            text, kr_records, pil_image, record_views = item.results
            # BK-OPT: txt/txt_boxes never need the rasterized page image, so we
            # avoid loading it from disk at all for these two formats (previously
            # it was loaded unconditionally even for a plain-text export).
            if fmt == "txt":
                self._write_plain_txt_export(path, record_views)
                return
            if fmt == "txt_boxes":
                self._write_structured_txt_export(path, record_views)
                return
            export_image = _load_image_color(item.path)
            # BK-OPT: everything below used to leave `export_image` (a PIL Image /
            # open file handle) to be garbage-collected implicitly. It is now
            # closed deterministically via try/finally, regardless of which
            # branch/return is taken or whether an exception is raised. Outputs
            # are unchanged - only resource cleanup timing is now explicit.
            try:
                if pil_image is None:
                    pil_image = export_image
                if fmt == "csv":
                    grid = table_to_rows(record_views, pil_image.size[0]) if any(rv.bbox for rv in record_views) else [
                        [rv.text] for rv in record_views
                    ]
                    with open(path, "w", newline="", encoding="utf-8") as f:
                        w = csv.writer(f)
                        w.writerows(grid)
                    return
                if fmt == "json":
                    grid = table_to_rows(record_views, pil_image.size[0]) if any(rv.bbox for rv in record_views) else [
                        [rv.text] for rv in record_views
                    ]
                    payload = {
                        "format": "bottled_kraken_lines",
                        "version": 2,
                        "image": {
                            "file": os.path.basename(item.path),
                            "width": int(pil_image.size[0]),
                            "height": int(pil_image.size[1]),
                        },
                        "lines": [self._line_export_entry(rv, i) for i, rv in enumerate(record_views)],
                        "rows": grid,
                    }
                    with open(path, "w", encoding="utf-8") as f:
                        json.dump(payload, f, indent=2, ensure_ascii=False)
                    return
                if fmt == "docx":
                    self._write_docx_export(path, item, export_image, record_views)
                    return
                if fmt == "alto":
                    seg_result = self._build_kraken_segmentation_for_export(
                        image_path=item.path,
                        record_views=record_views
                    )
                    if seg_result is None:
                        raise ValueError(self._tr("err_alto_requires_boxes"))
                    xml = serialization.serialize(
                        seg_result,
                        image_size=export_image.size,
                        template="alto"
                    )
                    with open(path, "w", encoding="utf-8") as f:
                        f.write(xml)
                    return
                if fmt == "hocr":
                    self._render_hocr_html(path, item, export_image, record_views)
                    return
                if fmt == "pdf":
                    c = pdf_canvas.Canvas(path, pagesize=(1, 1))
                    self._render_pdf_page_to_canvas(c, item)
                    try:
                        c.save()
                    except PermissionError as e:
                        raise PermissionError(
                            self._tr("err_pdf_save_locked").format(path)
                        ) from e
                    return
            finally:
                try:
                    export_image.close()
                except Exception:
                    pass
