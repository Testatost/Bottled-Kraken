"""Zusatzmodul: Ditto-Zeichen, Tabellen-/Datenexporte, Overlay-Box-Hilfen und UI-Fixes.

Der Patch ist bewusst defensiv: Er hängt sich nur an vorhandene Methoden, wenn sie in der
jeweiligen Bottled-Kraken-Version existieren. Dadurch bleibt die Datei auch mit älteren
Zwischenständen lauffähig.
"""

from .shared import *
from .ui_components import *
from .workers import *
from .dialogs import *
from .image_edit import *
from .main_window import MainWindow

def _bk_fix51_focus_is_text_input() -> bool:
    try:
        fw = QApplication.focusWidget()
        text_classes = (QLineEdit, QTextEdit, QPlainTextEdit, QSpinBox, QDoubleSpinBox, QComboBox)
        return isinstance(fw, text_classes)
    except Exception:
        return False

def _bk_fix51_begin_draw_box_for_current_line(self):
    """Startet Box-Zeichnen für die aktuell ausgewählte erkannte Zeile."""
    if _bk_fix51_focus_is_text_input():
        return
    try:
        task = self._ensure_overlay_possible() if hasattr(self, '_ensure_overlay_possible') else self._current_task()
        if not task or not getattr(task, 'results', None):
            return
        row = -1
        try:
            row = int(self.list_lines.currentRow())
        except Exception:
            row = -1
        try:
            rows = self._selected_line_rows()
            if rows and (row < 0 or row not in rows):
                row = int(rows[0])
        except Exception:
            pass
        _text, _kr, _im, recs = task.results
        if not (0 <= row < len(recs)):
            return
        self._pending_new_line_box = False
        self._pending_box_for_row = row
        try:
            self.show_overlay = True
            if hasattr(self.canvas, 'set_overlay_enabled'):
                self.canvas.set_overlay_enabled(True)
            self._refresh_overlay_display(recs)
            self.canvas.select_idx(row, center=True)
            self.list_lines.setCurrentRow(row)
        except Exception:
            pass
        try:
            self.canvas.start_draw_box_mode()
            self.canvas.setFocus(Qt.ShortcutFocusReason)
            self.canvas.viewport().setFocus(Qt.ShortcutFocusReason)
        except Exception:
            pass
        try:
            self.status_bar.showMessage(self._tr('line_menu_draw_box'), 3500)
        except Exception:
            pass
    except Exception as exc:
        try:
            print(f'FIX8.51 B shortcut failed: {exc}')
        except Exception:
            pass

def _bk_fix51_install_b_shortcut(self):
    try:
        if getattr(self, '_bk_fix51_b_action_installed', False):
            return
        act = QAction(self)
        act.setShortcut(QKeySequence('B'))
        act.setShortcutContext(Qt.ApplicationShortcut)
        act.triggered.connect(lambda _checked=False, w=self: _bk_fix51_begin_draw_box_for_current_line(w))
        self.addAction(act)
        self._bk_fix51_b_action = act
        self._bk_fix51_b_action_installed = True
    except Exception as exc:
        try:
            print(f'FIX8.51 install B shortcut failed: {exc}')
        except Exception:
            pass

try:
    _BK_FIX51_PREV_MAINWINDOW_INIT = MainWindow.__init__
except Exception:
    _BK_FIX51_PREV_MAINWINDOW_INIT = None

if callable(_BK_FIX51_PREV_MAINWINDOW_INIT) and not getattr(MainWindow.__init__, '_bk_fix51_init_wrapped', False):
    def _bk_fix51_mainwindow_init(self, *args, **kwargs):
        _BK_FIX51_PREV_MAINWINDOW_INIT(self, *args, **kwargs)
        try:
            QTimer.singleShot(0, lambda w=self: _bk_fix51_install_b_shortcut(w))
        except Exception:
            _bk_fix51_install_b_shortcut(self)
    _bk_fix51_mainwindow_init._bk_fix51_init_wrapped = True
    MainWindow.__init__ = _bk_fix51_mainwindow_init

try:
    MainWindow.bk_draw_box_for_current_line_shortcut = _bk_fix51_begin_draw_box_for_current_line
except Exception:
    pass

def _bk_fix51_group_rows(recs, page_width: int = 0):
    with_boxes = [rv for rv in (recs or []) if getattr(rv, 'bbox', None) and _bk_fix36_clean_text(getattr(rv, 'text', ''))]
    if not with_boxes:
        return [[rv] for rv in (recs or []) if _bk_fix36_clean_text(getattr(rv, 'text', ''))]
    heights = sorted(max(2, int(rv.bbox[3] - rv.bbox[1])) for rv in with_boxes)
    med_h = heights[len(heights)//2] if heights else 12
    gap = max(4, min(22, int(med_h * 0.72)))
    rows: List[List[Any]] = []
    for rv in sorted(with_boxes, key=lambda r: (((r.bbox[1]+r.bbox[3])/2.0), r.bbox[0])):
        cy = (rv.bbox[1] + rv.bbox[3]) / 2.0
        if rows:
            last = rows[-1]
            lcy = sum((x.bbox[1]+x.bbox[3])/2.0 for x in last) / max(1, len(last))
            if abs(cy - lcy) <= gap:
                last.append(rv)
                continue
        rows.append([rv])
    for row in rows:
        row.sort(key=lambda r: r.bbox[0])
    return rows

def _bk_fix51_text_columns(rows, page_width: int, max_cols: int = 160) -> str:
    rows = _bk_fix51_group_rows([rv for row in rows for rv in row], page_width)
    boxes = [rv.bbox for row in rows for rv in row if getattr(rv, 'bbox', None)]
    if not boxes:
        return '\n'.join(' '.join(_bk_fix36_clean_text(getattr(rv, 'text', '')) for rv in row).strip() for row in rows)
    min_x = min(b[0] for b in boxes)
    max_x = max(b[2] for b in boxes)
    width = max(1, max_x - min_x)
    max_cols = max(80, min(220, int(max_cols)))
    lines = []
    for row in rows:
        cells = [' '] * (max_cols + 8)
        for rv in sorted(row, key=lambda r: r.bbox[0]):
            txt = _bk_fix36_clean_text(getattr(rv, 'text', ''))
            if not txt:
                continue
            col = int(round(((rv.bbox[0] - min_x) / width) * (max_cols - 1)))
            col = max(0, min(max_cols - 1, col))
            # Nicht brutal überschreiben: bei Kollision rechts weitersuchen.
            while col < len(cells) and cells[col] != ' ':
                col += 1
            for ch in txt:
                if col >= len(cells):
                    cells.append(' ')
                cells[col] = ch
                col += 1
        lines.append(''.join(cells).rstrip())
    return '\n'.join(lines).rstrip()

def _bk_fix51_column_anchors(rows, page_width: int) -> List[float]:
    xs = []
    for row in rows:
        for rv in row:
            if not getattr(rv, 'bbox', None):
                continue
            txt = _bk_fix36_clean_text(getattr(rv, 'text', ''))
            if not txt:
                continue
            x0, y0, x1, y1 = rv.bbox
            # Sehr breite Textzeilen/Header nicht als Spaltenanker nutzen.
            if (x1 - x0) > max(1, page_width) * 0.42 and len(row) <= 2:
                continue
            xs.append(float(x0))
    if not xs:
        return []
    xs.sort()
    clusters = []
    threshold = max(18.0, min(58.0, page_width * 0.028))
    for x in xs:
        if not clusters or abs(x - (sum(clusters[-1]) / len(clusters[-1]))) > threshold:
            clusters.append([x])
        else:
            clusters[-1].append(x)
    anchors = [sum(c)/len(c) for c in clusters]
    # Zu viele Wortanker auf maximal sinnvolle Tabellenspalten verdichten.
    if len(anchors) > 12:
        # Behalte große Lücken als Spaltentrennung.
        gaps = [(anchors[i+1]-anchors[i], i) for i in range(len(anchors)-1)]
        keep_breaks = sorted(i for _gap, i in sorted(gaps, reverse=True)[:11])
        groups = []
        start = 0
        for br in keep_breaks:
            groups.append(anchors[start:br+1])
            start = br + 1
        groups.append(anchors[start:])
        anchors = [sum(g)/len(g) for g in groups if g]
    return anchors

def _bk_fix51_rows_are_table(rows, page_width: int) -> bool:
    if len(rows) < 3:
        return False
    rows = [row for row in rows if row]
    multi = sum(1 for row in rows if len([rv for rv in row if _bk_fix36_clean_text(getattr(rv, 'text', ''))]) >= 3)
    numeric = sum(1 for row in rows for rv in row if re.search(r'\d', str(getattr(rv, 'text', ''))))
    spread = 0
    for row in rows:
        bbs = [rv.bbox for rv in row if getattr(rv, 'bbox', None)]
        if bbs and (max(b[2] for b in bbs) - min(b[0] for b in bbs)) > page_width * 0.38:
            spread += 1
    return (multi >= max(2, int(len(rows) * 0.35)) and spread >= max(2, int(len(rows) * 0.35))) or (numeric >= 6 and multi >= 2 and spread >= 2)

def _bk_fix51_split_blocks(recs, page_width: int):
    try:
        _bk_fix43_resolve_ditto_marks_in_recs(recs)
    except Exception:
        pass
    rows = _bk_fix51_group_rows(recs, page_width)
    blocks = []
    i = 0
    while i < len(rows):
        # Tabellenlauf: mehrere Zeilen mit mehreren Zellen und horizontaler Streuung.
        j = i
        run = []
        while j < len(rows):
            row = rows[j]
            cells = [rv for rv in row if _bk_fix36_clean_text(getattr(rv, 'text', ''))]
            if len(cells) >= 2:
                run.append(row)
                j += 1
            else:
                break
        if _bk_fix51_rows_are_table(run, page_width):
            anchors = _bk_fix51_column_anchors(run, page_width)
            if len(anchors) >= 2:
                grid = []
                for rr in run:
                    line = [''] * len(anchors)
                    for rv in sorted(rr, key=lambda r: r.bbox[0] if getattr(r, 'bbox', None) else 0):
                        txt = _bk_fix36_clean_text(getattr(rv, 'text', ''))
                        if not txt or not getattr(rv, 'bbox', None):
                            continue
                        cx = (rv.bbox[0] + rv.bbox[2]) / 2.0
                        c = min(range(len(anchors)), key=lambda k: abs(cx - anchors[k]))
                        line[c] = (line[c] + ' ' + txt).strip() if line[c] else txt
                    grid.append(line)
                blocks.append({'type': 'table', 'rows': grid, 'anchors': anchors})
                i = j
                continue
        # Normaler Absatz / Kopfzeile.
        row = rows[i]
        txt = ' '.join(_bk_fix36_clean_text(getattr(rv, 'text', '')) for rv in sorted(row, key=lambda r: r.bbox[0] if getattr(r, 'bbox', None) else 0)).strip()
        if txt:
            blocks.append({'type': 'paragraph', 'text': txt})
        i += 1
    return blocks

def _bk_fix51_write_txt(path: str, record_views: List[RecordView], image_size=None):
    img_w, _img_h = image_size or (0, 0)
    rows = _bk_fix51_group_rows(record_views, img_w)
    with open(path, 'w', encoding='utf-8') as f:
        txt = _bk_fix51_text_columns(rows, img_w or 1600)
        if not _bk_fix36_clean_text(txt):
            txt = '\n'.join(_bk_fix36_clean_text(getattr(rv, 'text', '')) for rv in (record_views or []) if _bk_fix36_clean_text(getattr(rv, 'text', '')))
        f.write(txt.rstrip() + '\n')

def _bk_fix51_write_docx(path: str, item: TaskItem, export_image: Image.Image, record_views: List[RecordView]):
    try:
        from docx import Document
        from docx.shared import Inches, Pt
        from docx.enum.section import WD_ORIENT
        from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
    except Exception as e:
        raise RuntimeError(translation.translate(translation.DEFAULT_LANGUAGE, "err_no_docx_package_short")) from e
    page_w_px, page_h_px = export_image.size
    blocks = _bk_fix51_split_blocks(list(record_views or []), page_w_px)
    doc = Document()
    section = doc.sections[0]
    if page_w_px > page_h_px * 1.15:
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width = Inches(11.69)
        section.page_height = Inches(8.27)
    else:
        section.page_width = Inches(8.27)
        section.page_height = Inches(11.69)
    section.left_margin = section.right_margin = Inches(0.30)
    section.top_margin = section.bottom_margin = Inches(0.35)
    try:
        normal = doc.styles['Normal']
        normal.font.name = 'Arial'
        normal.font.size = Pt(8.5)
    except Exception:
        pass
    for block in blocks:
        if block.get('type') == 'table':
            grid = block.get('rows') or []
            if not grid:
                continue
            cols = max(len(r) for r in grid)
            if cols < 2:
                continue
            table = doc.add_table(rows=len(grid), cols=cols)
            table.style = 'Table Grid'
            table.alignment = WD_TABLE_ALIGNMENT.LEFT
            table.autofit = True
            for r_idx, row in enumerate(grid):
                for c_idx in range(cols):
                    cell = table.rows[r_idx].cells[c_idx]
                    cell.text = ''
                    p = cell.paragraphs[0]
                    p.paragraph_format.space_after = Pt(0)
                    p.paragraph_format.space_before = Pt(0)
                    txt = row[c_idx] if c_idx < len(row) else ''
                    run = p.add_run(txt)
                    run.font.name = 'Arial'
                    run.font.size = Pt(6.8 if cols >= 8 else 7.8)
                    try:
                        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                    except Exception:
                        pass
            doc.add_paragraph()
        else:
            txt = str(block.get('text') or '').strip()
            if txt:
                p = doc.add_paragraph(txt)
                p.paragraph_format.space_after = Pt(1)
    doc.save(path)

try:
    _BK_FIX51_PREV_RENDER_FILE = MainWindow._render_file
except Exception:
    _BK_FIX51_PREV_RENDER_FILE = None

def _bk_fix51_render_file(self, path: str, fmt: str, item: TaskItem):
    if not item or not getattr(item, 'results', None):
        return
    text, kr_records, pil_image, record_views = item.results
    export_image = _load_image_color(item.path)
    fmt_l = str(fmt or '').lower()
    if fmt_l == 'txt':
        return _bk_fix51_write_txt(path, record_views, export_image.size)
    if fmt_l == 'docx':
        return _bk_fix51_write_docx(path, item, export_image, record_views)
    if fmt_l == 'txt_boxes':
        return _bk_fix40_write_structured_txt_export(path, record_views, export_image.size)
    if callable(_BK_FIX51_PREV_RENDER_FILE):
        return _BK_FIX51_PREV_RENDER_FILE(self, path, fmt, item)
    return None

try:
    MainWindow._render_file = _bk_fix51_render_file
except Exception:
    pass
