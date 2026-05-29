import csv
import html
import math
import zipfile
from typing import Any
from bottled_kraken.common import _load_image_color, translation
def _clean_text(value: Any) -> str:
    text = str(value or "").replace("\r", "\n")
    return " ".join(text.split())
def _records_from_views(record_views):
    records = []
    for index, view in enumerate(record_views or []):
        text = _clean_text(getattr(view, "text", ""))
        if not text:
            continue
        bbox = getattr(view, "bbox", None)
        if bbox and len(bbox) >= 4:
            try:
                x0, y0, x1, y1 = [float(v) for v in bbox[:4]]
                if x1 > x0 and y1 > y0:
                    records.append({
                        "text": text,
                        "bbox": (x0, y0, x1, y1),
                        "x0": x0,
                        "y0": y0,
                        "x1": x1,
                        "y1": y1,
                        "cx": (x0 + x1) / 2.0,
                        "cy": (y0 + y1) / 2.0,
                        "w": x1 - x0,
                        "h": y1 - y0,
                        "index": index,
                    })
                    continue
            except Exception:
                pass
        records.append({
            "text": text,
            "bbox": None,
            "x0": 0.0,
            "y0": float(index),
            "x1": 0.0,
            "y1": float(index + 1),
            "cx": 0.0,
            "cy": float(index),
            "w": 0.0,
            "h": 1.0,
            "index": index,
        })
    records.sort(key=lambda item: (float(item.get("y0", 0.0)), float(item.get("x0", 0.0)), int(item.get("index", 0))))
    return records
def _page_size(image_size, records):
    try:
        page_width = float((image_size or (0, 0))[0] or 0)
        page_height = float((image_size or (0, 0))[1] or 0)
    except Exception:
        page_width = page_height = 0.0
    boxed = [record for record in records if record.get("bbox")]
    if boxed:
        page_width = max(page_width, max(float(record["x1"]) for record in boxed))
        page_height = max(page_height, max(float(record["y1"]) for record in boxed))
    return max(1.0, page_width or 1600.0), max(1.0, page_height or 2200.0)
def _content_bounds(records, page_width, page_height):
    boxed = [record for record in records if record.get("bbox")]
    if not boxed:
        return 0.0, 0.0, float(page_width), float(page_height)
    return (
        max(0.0, min(float(record["x0"]) for record in boxed)),
        max(0.0, min(float(record["y0"]) for record in boxed)),
        min(float(page_width), max(float(record["x1"]) for record in boxed)),
        min(float(page_height), max(float(record["y1"]) for record in boxed)),
    )
def _median_height(records):
    heights = sorted(
        float(record.get("h", 0.0) or 0.0)
        for record in records
        if record.get("bbox") and float(record.get("h", 0.0) or 0.0) > 0
    )
    if not heights:
        return 12.0
    return max(2.0, heights[len(heights) // 2])
def _group_rows(records, tolerance):
    rows = []
    ordered = sorted(
        records or [],
        key=lambda record: (
            float(record.get("cy", record.get("y0", 0.0)) or 0.0),
            float(record.get("x0", 0.0) or 0.0),
            int(record.get("index", 0) or 0),
        ),
    )
    for record in ordered:
        cy = float(record.get("cy", record.get("y0", 0.0)) or 0.0)
        for row in rows:
            if abs(cy - row["cy"]) <= tolerance:
                row["items"].append(record)
                row["cy"] = sum(float(item.get("cy", item.get("y0", 0.0)) or 0.0) for item in row["items"]) / len(row["items"])
                break
        else:
            rows.append({"cy": cy, "items": [record]})
    rows.sort(key=lambda row: (row["cy"], min(float(item.get("x0", 0.0) or 0.0) for item in row["items"])))
    for row in rows:
        row["items"].sort(key=lambda record: (float(record.get("x0", 0.0) or 0.0), int(record.get("index", 0) or 0)))
    return rows
def _place_text(canvas, position, text):
    if not text:
        return
    position = max(0, min(int(position), max(0, len(canvas) - 1)))
    if 0 < position < len(canvas) and canvas[position] != " ":
        while position < len(canvas) - 1 and canvas[position] != " ":
            position += 1
    for offset, character in enumerate(str(text)):
        index = position + offset
        if index >= len(canvas):
            break
        canvas[index] = character
def build_spatial_text(record_views, image_size=None, width_chars=156):
    records = _records_from_views(record_views)
    if not records:
        return ""
    page_width, page_height = _page_size(image_size, records)
    min_x, min_y, max_x, max_y = _content_bounds(records, page_width, page_height)
    span_x = max(1.0, max_x - min_x)
    median_height = _median_height(records)
    rows = _group_rows(records, max(2.0, median_height * 0.55))
    output = []
    previous_y = None
    for row in rows:
        if previous_y is not None:
            gap = float(row["cy"]) - float(previous_y)
            if gap > median_height * 1.9:
                blanks = min(3, max(1, int(round(gap / max(1.0, median_height * 1.35))) - 1))
                output.extend([""] * blanks)
        canvas = [" "] * max(80, int(width_chars))
        for record in row["items"]:
            text = _clean_text(record.get("text", ""))
            if not text:
                continue
            if record.get("bbox"):
                position = round(((float(record.get("x0", min_x)) - min_x) / span_x) * (width_chars - 1))
            else:
                position = 0
            _place_text(canvas, position, text)
        line = "".join(canvas).rstrip()
        if line.strip():
            output.append(line)
        previous_y = row["cy"]
    return "\n".join(output).rstrip() + "\n"
def write_spatial_txt(path, record_views, image_size=None):
    text = build_spatial_text(record_views, image_size, width_chars=156)
    with open(path, "w", encoding="utf-8", newline="") as handle:
        handle.write(text.rstrip() + "\n")
def _run_font(run, font_size_pt=6.0, font_name="Consolas"):
    try:
        from docx.shared import Pt
        run.font.name = font_name
        run.font.size = Pt(font_size_pt)
    except Exception:
        pass
def _textbox_xml(text, left_pt, top_pt, width_pt, height_pt, font_size_pt=6.0, box_id=1):
    safe_text = html.escape(str(text or ""), quote=False)
    left_pt = max(0.0, float(left_pt))
    top_pt = max(0.0, float(top_pt))
    width_pt = max(8.0, float(width_pt))
    height_pt = max(6.0, float(height_pt))
    font_twips = max(2, int(round(float(font_size_pt) * 2)))
    line_twips = max(120, int(round(float(font_size_pt) * 20.0 * 1.10)))
    return f"""
<w:pict xmlns:w=\"http://schemas.openxmlformats.org/wordprocessingml/2006/main\" xmlns:v=\"urn:schemas-microsoft-com:vml\" xmlns:o=\"urn:schemas-microsoft-com:office:office\">
  <v:shape id=\"bk_textbox_{box_id}\" type=\"#_x0000_t202\" style=\"position:absolute;margin-left:{left_pt:.2f}pt;margin-top:{top_pt:.2f}pt;width:{width_pt:.2f}pt;height:{height_pt:.2f}pt;z-index:{box_id};mso-position-horizontal:absolute;mso-position-vertical:absolute\" stroked=\"f\" filled=\"f\" o:allowincell=\"f\">
    <v:textbox inset=\"0,0,0,0\" style=\"mso-fit-shape-to-text:false\">
      <w:txbxContent>
        <w:p>
          <w:pPr><w:spacing w:before=\"0\" w:after=\"0\" w:line=\"{line_twips}\" w:lineRule=\"exact\"/></w:pPr>
          <w:r>
            <w:rPr><w:rFonts w:ascii=\"Consolas\" w:hAnsi=\"Consolas\" w:cs=\"Consolas\"/><w:sz w:val=\"{font_twips}\"/><w:szCs w:val=\"{font_twips}\"/></w:rPr>
            <w:t xml:space=\"preserve\">{safe_text}</w:t>
          </w:r>
        </w:p>
      </w:txbxContent>
    </v:textbox>
  </v:shape>
</w:pict>"""
def write_positioned_docx(path, item, export_image, record_views):
    try:
        from docx import Document
        from docx.shared import Inches, Pt
        from docx.oxml import parse_xml
    except Exception as exc:
        raise RuntimeError(translation.translate(translation.DEFAULT_LANGUAGE, "err_no_docx_package_short")) from exc
    image_size = getattr(export_image, "size", None) or getattr(item, "image_size", None) or (0, 0)
    records = _records_from_views(record_views)
    if not records:
        document = Document()
        document.save(path)
        return
    page_width, page_height = _page_size(image_size, records)
    min_x, min_y, max_x, max_y = _content_bounds(records, page_width, page_height)
    content_width = max(1.0, max_x - min_x)
    content_height = max(1.0, max_y - min_y)
    document = Document()
    section = document.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    margin_in = 0.28
    section.left_margin = section.right_margin = Inches(margin_in)
    section.top_margin = section.bottom_margin = Inches(margin_in)
    usable_width_pt = (8.27 - margin_in * 2.0) * 72.0
    usable_height_pt = (11.69 - margin_in * 2.0) * 72.0
    scale = min(usable_width_pt / content_width, usable_height_pt / content_height)
    offset_x_pt = (usable_width_pt - content_width * scale) / 2.0
    offset_y_pt = (usable_height_pt - content_height * scale) / 2.0
    font_size = 6.0
    median_height = _median_height(records)
    try:
        normal_style = document.styles["Normal"]
        normal_style.font.name = "Consolas"
        normal_style.font.size = Pt(font_size)
        normal_style.paragraph_format.space_before = Pt(0)
        normal_style.paragraph_format.space_after = Pt(0)
    except Exception:
        pass
    anchor = document.add_paragraph()
    try:
        anchor.paragraph_format.space_before = Pt(0)
        anchor.paragraph_format.space_after = Pt(0)
        anchor.paragraph_format.line_spacing = Pt(1)
    except Exception:
        pass
    run = anchor.add_run()
    for index, record in enumerate(records, start=1):
        if not record.get("bbox"):
            continue
        text = _clean_text(record.get("text", ""))
        if not text:
            continue
        left = offset_x_pt + (float(record["x0"]) - min_x) * scale
        top = offset_y_pt + (float(record["y0"]) - min_y) * scale
        width = max(float(record["w"]) * scale + 6.0, min(usable_width_pt, max(18.0, len(text) * font_size * 0.43)))
        height = max(float(record["h"]) * scale * 1.25, font_size * 1.35, median_height * scale * 0.95)
        try:
            run._r.append(parse_xml(_textbox_xml(text, left, top, width, height, font_size, index)))
        except Exception:
            continue
    filler = document.add_paragraph()
    _run_font(filler.add_run(" "), 1.0)
    document.save(path)
def _odt_xml_text(value: Any) -> str:
    return html.escape(str(value or ""), quote=False)
def _odt_file_xml(name: str, data: str) -> bytes:
    return data.encode("utf-8")
def _odt_content_xml(records, page_width, page_height, min_x, min_y, max_x, max_y):
    content_width = max(1.0, max_x - min_x)
    content_height = max(1.0, max_y - min_y)
    page_w_cm = 21.0
    page_h_cm = 29.7
    margin_cm = 0.7
    usable_w_cm = page_w_cm - margin_cm * 2.0
    usable_h_cm = page_h_cm - margin_cm * 2.0
    scale = min(usable_w_cm / content_width, usable_h_cm / content_height)
    offset_x = (usable_w_cm - content_width * scale) / 2.0
    offset_y = (usable_h_cm - content_height * scale) / 2.0
    frames = []
    median_height = _median_height(records)
    for index, record in enumerate(records, start=1):
        if not record.get("bbox"):
            continue
        text = _clean_text(record.get("text", ""))
        if not text:
            continue
        left = margin_cm + offset_x + (float(record["x0"]) - min_x) * scale
        top = margin_cm + offset_y + (float(record["y0"]) - min_y) * scale
        width = max(float(record.get("w", 0.0)) * scale + 0.08, min(usable_w_cm, max(0.35, len(text) * 0.055)))
        height = max(float(record.get("h", 0.0)) * scale * 1.25, median_height * scale * 0.95, 0.18)
        frames.append(
            '<draw:frame draw:style-name="BkTextFrame" draw:name="bk_text_%04d" text:anchor-type="page" '
            'svg:x="%.4fcm" svg:y="%.4fcm" svg:width="%.4fcm" svg:height="%.4fcm" draw:z-index="%d">'
            '<draw:text-box fo:min-height="%.4fcm"><text:p text:style-name="BkOcrText">%s</text:p></draw:text-box></draw:frame>' % (
                index, max(0.0, left), max(0.0, top), max(0.1, width), max(0.1, height), index, max(0.1, height), _odt_xml_text(text)
            )
        )
    return ''.join([
        '<?xml version="1.0" encoding="UTF-8"?>',
        '<office:document-content ',
        'xmlns:office="urn:oasis:names:tc:opendocument:xmlns:office:1.0" ',
        'xmlns:style="urn:oasis:names:tc:opendocument:xmlns:style:1.0" ',
        'xmlns:text="urn:oasis:names:tc:opendocument:xmlns:text:1.0" ',
        'xmlns:draw="urn:oasis:names:tc:opendocument:xmlns:drawing:1.0" ',
        'xmlns:fo="urn:oasis:names:tc:opendocument:xmlns:xsl-fo-compatible:1.0" ',
        'xmlns:svg="urn:oasis:names:tc:opendocument:xmlns:svg-compatible:1.0" ',
        'office:version="1.2">',
        '<office:scripts/>',
        '<office:font-face-decls><style:font-face style:name="Calibri" svg:font-family="Calibri"/></office:font-face-decls>',
        '<office:automatic-styles>',
        '<style:style style:name="BkPageP" style:family="paragraph"><style:paragraph-properties fo:margin-top="0cm" fo:margin-bottom="0cm" fo:line-height="100%"/></style:style>',
        '<style:style style:name="BkOcrText" style:family="paragraph"><style:paragraph-properties fo:margin-top="0cm" fo:margin-bottom="0cm" fo:line-height="100%"/><style:text-properties fo:font-size="6pt" style:font-name="Consolas"/></style:style>',
        '<style:style style:name="BkTextFrame" style:family="graphic"><style:graphic-properties draw:fill="none" draw:stroke="none" style:wrap="run-through" style:vertical-pos="from-top" style:vertical-rel="page" style:horizontal-pos="from-left" style:horizontal-rel="page" fo:padding="0cm" fo:margin="0cm"/></style:style>',
        '</office:automatic-styles>',
        '<office:body><office:text><text:p text:style-name="BkPageP">',
        ''.join(frames),
        '</text:p></office:text></office:body></office:document-content>'
    ])
def _odt_styles_xml() -> str:
    return ''.join([
        '<?xml version="1.0" encoding="UTF-8"?>',
        '<office:document-styles ',
        'xmlns:office="urn:oasis:names:tc:opendocument:xmlns:office:1.0" ',
        'xmlns:style="urn:oasis:names:tc:opendocument:xmlns:style:1.0" ',
        'xmlns:text="urn:oasis:names:tc:opendocument:xmlns:text:1.0" ',
        'xmlns:fo="urn:oasis:names:tc:opendocument:xmlns:xsl-fo-compatible:1.0" ',
        'xmlns:svg="urn:oasis:names:tc:opendocument:xmlns:svg-compatible:1.0" ',
        'office:version="1.2">',
        '<office:font-face-decls><style:font-face style:name="Calibri" svg:font-family="Calibri"/></office:font-face-decls>',
        '<office:styles><style:default-style style:family="paragraph"><style:text-properties fo:font-size="6pt" style:font-name="Consolas"/></style:default-style></office:styles>',
        '<office:automatic-styles><style:page-layout style:name="BkPageLayout"><style:page-layout-properties fo:page-width="21cm" fo:page-height="29.7cm" fo:margin-top="0.7cm" fo:margin-bottom="0.7cm" fo:margin-left="0.7cm" fo:margin-right="0.7cm"/></style:page-layout></office:automatic-styles>',
        '<office:master-styles><style:master-page style:name="Standard" style:page-layout-name="BkPageLayout"/></office:master-styles>',
        '</office:document-styles>'
    ])
def _odt_manifest_xml() -> str:
    return '<?xml version="1.0" encoding="UTF-8"?><manifest:manifest xmlns:manifest="urn:oasis:names:tc:opendocument:xmlns:manifest:1.0" manifest:version="1.2"><manifest:file-entry manifest:full-path="/" manifest:media-type="application/vnd.oasis.opendocument.text"/><manifest:file-entry manifest:full-path="content.xml" manifest:media-type="text/xml"/><manifest:file-entry manifest:full-path="styles.xml" manifest:media-type="text/xml"/><manifest:file-entry manifest:full-path="meta.xml" manifest:media-type="text/xml"/><manifest:file-entry manifest:full-path="settings.xml" manifest:media-type="text/xml"/></manifest:manifest>'
def write_positioned_odt(path, item, export_image, record_views):
    image_size = getattr(export_image, "size", None) or getattr(item, "image_size", None) or (0, 0)
    records = _records_from_views(record_views)
    page_width, page_height = _page_size(image_size, records)
    min_x, min_y, max_x, max_y = _content_bounds(records, page_width, page_height)
    content = _odt_content_xml(records, page_width, page_height, min_x, min_y, max_x, max_y)
    styles = _odt_styles_xml()
    meta = '<?xml version="1.0" encoding="UTF-8"?><office:document-meta xmlns:office="urn:oasis:names:tc:opendocument:xmlns:office:1.0" xmlns:meta="urn:oasis:names:tc:opendocument:xmlns:meta:1.0" office:version="1.2"><office:meta><meta:generator>Bottled Kraken</meta:generator></office:meta></office:document-meta>'
    settings = '<?xml version="1.0" encoding="UTF-8"?><office:document-settings xmlns:office="urn:oasis:names:tc:opendocument:xmlns:office:1.0" office:version="1.2"><office:settings/></office:document-settings>'
    with zipfile.ZipFile(path, "w") as archive:
        info = zipfile.ZipInfo("mimetype")
        info.date_time = (2020, 1, 1, 0, 0, 0)
        info.compress_type = zipfile.ZIP_STORED
        archive.writestr(info, "application/vnd.oasis.opendocument.text")
        for name, data in (("content.xml", content), ("styles.xml", styles), ("meta.xml", meta), ("settings.xml", settings), ("META-INF/manifest.xml", _odt_manifest_xml())):
            zi = zipfile.ZipInfo(name)
            zi.date_time = (2020, 1, 1, 0, 0, 0)
            zi.compress_type = zipfile.ZIP_DEFLATED
            archive.writestr(zi, data.encode("utf-8"))
def _column_anchors(records, max_columns=16):
    xs = sorted(float(record.get("x0", 0.0) or 0.0) for record in records if record.get("bbox"))
    if not xs:
        return []
    anchors = []
    for x in xs:
        if not anchors or abs(x - anchors[-1]) > 18:
            anchors.append(x)
        else:
            anchors[-1] = (anchors[-1] + x) / 2.0
    if len(anchors) > max_columns:
        step = max(1, math.ceil(len(anchors) / max_columns))
        anchors = anchors[::step]
    return anchors
def _row_to_csv_cells(row_items, anchors):
    if not anchors:
        return [_clean_text(item.get("text", "")) for item in row_items]
    cells = [""] * len(anchors)
    for item in row_items:
        text = _clean_text(item.get("text", ""))
        if not text:
            continue
        x = float(item.get("x0", 0.0) or 0.0)
        column = min(range(len(anchors)), key=lambda index: abs(anchors[index] - x))
        cells[column] = (cells[column] + " " + text).strip() if cells[column] else text
    while cells and not cells[-1]:
        cells.pop()
    return cells
def write_plain_csv(path, record_views, image_size=None):
    records = _records_from_views(record_views)
    if not records:
        with open(path, "w", encoding="utf-8-sig", newline="") as handle:
            csv.writer(handle).writerow(["text"])
        return
    rows = _group_rows(records, max(2.0, _median_height(records) * 0.55))
    anchors = _column_anchors(records)
    with open(path, "w", encoding="utf-8-sig", newline="") as handle:
        writer = csv.writer(handle)
        for row in rows:
            cells = _row_to_csv_cells(row["items"], anchors)
            if any(cell.strip() for cell in cells):
                writer.writerow(cells)

def _xlsx_col_name(index):
    name = ""
    index = int(index)
    while index > 0:
        index, rem = divmod(index - 1, 26)
        name = chr(65 + rem) + name
    return name

def _xml(value):
    return html.escape(str(value or ""), quote=True)

def _merge_numeric_edges(values, minimum, maximum, tolerance):
    raw = [float(minimum), float(maximum)]
    raw.extend(float(value) for value in values if value is not None)
    raw = [max(float(minimum), min(float(maximum), value)) for value in raw]
    raw = sorted(raw)
    merged = []
    for value in raw:
        if not merged or abs(value - merged[-1]) > tolerance:
            merged.append(value)
        else:
            merged[-1] = (merged[-1] + value) / 2.0
    if len(merged) < 2:
        merged = [float(minimum), float(maximum)]
    if merged[0] > float(minimum) + tolerance:
        merged.insert(0, float(minimum))
    else:
        merged[0] = float(minimum)
    if merged[-1] < float(maximum) - tolerance:
        merged.append(float(maximum))
    else:
        merged[-1] = float(maximum)
    cleaned = [merged[0]]
    for value in merged[1:]:
        if value > cleaned[-1] + 0.5:
            cleaned.append(value)
    if len(cleaned) < 2:
        cleaned = [float(minimum), float(maximum)]
    return cleaned

def _nearest_edge_index(edges, value):
    value = float(value)
    return min(range(len(edges)), key=lambda index: abs(edges[index] - value))

def _column_width_from_pixels(pixels, scale):
    width = max(0.35, float(pixels) * float(scale) / 5.25)
    return min(28.0, width)

def _row_height_from_pixels(pixels, scale):
    height = max(3.0, float(pixels) * float(scale))
    return min(80.0, height)

def _spreadsheet_layout(records, image_size=None):
    records = [record for record in records or [] if _clean_text(record.get("text", ""))]
    if not records:
        return 1, 1, {}, {}, [18.0], [18.0]
    page_width, page_height = _page_size(image_size, records)
    min_x, min_y, max_x, max_y = _content_bounds(records, page_width, page_height)
    content_width = max(1.0, max_x - min_x)
    median_height = _median_height(records)
    boxed_widths = sorted(float(record.get("w", 0.0) or 0.0) for record in records if record.get("bbox") and float(record.get("w", 0.0) or 0.0) > 0)
    median_width = boxed_widths[len(boxed_widths) // 2] if boxed_widths else max(40.0, content_width / 18.0)
    col_count = int(round(content_width / max(38.0, median_width * 0.55)))
    col_count = max(22, min(42, col_count))
    col_widths = [4.6] * col_count
    rows = _group_rows(records, max(3.0, median_height * 0.62))
    cells = {}
    merges = {}
    row_heights = []
    out_row = 1
    previous_y = None
    previous_height = median_height
    def to_col(x_value):
        position = (float(x_value) - min_x) / content_width
        return max(1, min(col_count, int(math.floor(position * col_count)) + 1))
    def to_col_end(x_value):
        position = (float(x_value) - min_x) / content_width
        return max(1, min(col_count, int(math.ceil(position * col_count))))
    for row in rows:
        row_y = float(row.get("cy", 0.0) or 0.0)
        if previous_y is not None:
            gap = row_y - previous_y
            if gap > previous_height * 2.4:
                blank_count = min(3, max(1, int(round(gap / max(1.0, median_height * 2.8))) - 1))
                for _ in range(blank_count):
                    row_heights.append(10.0)
                    out_row += 1
        items = sorted(row.get("items", []), key=lambda item: (float(item.get("x0", 0.0) or 0.0), int(item.get("index", 0) or 0)))
        occupied = [False] * (col_count + 1)
        tallest = max((float(item.get("h", median_height) or median_height) for item in items), default=median_height)
        estimated_lines = 1
        for item in items:
            text_value = _clean_text(item.get("text", ""))
            if not text_value:
                continue
            if item.get("bbox"):
                start_col = to_col(item.get("x0", min_x))
                end_col = to_col_end(item.get("x1", min_x + content_width))
            else:
                start_col = 1
                end_col = col_count
            if end_col < start_col:
                end_col = start_col
            span = max(1, end_col - start_col + 1)
            wanted = max(span, min(col_count - start_col + 1, int(math.ceil(len(text_value) / 18.0))))
            wanted = max(1, min(wanted, col_count - start_col + 1))
            best_col = start_col
            best_span = wanted
            found = False
            for shrink in range(0, max(0, wanted - 1) + 1):
                trial_span = wanted - shrink
                left = max(1, start_col - 2)
                right = min(col_count - trial_span + 1, start_col + 2)
                for candidate in range(left, right + 1):
                    if all(not occupied[col] for col in range(candidate, candidate + trial_span)):
                        best_col = candidate
                        best_span = trial_span
                        found = True
                        break
                if found:
                    break
            if not found:
                free_cols = [col for col in range(1, col_count + 1) if not occupied[col]]
                if free_cols:
                    best_col = min(free_cols, key=lambda col: abs(col - start_col))
                    best_span = 1
                else:
                    best_col = max(1, min(col_count, start_col))
                    best_span = 1
                    existing = cells.get((out_row, best_col), "")
                    cells[(out_row, best_col)] = (existing + " " + text_value).strip() if existing else text_value
                    estimated_lines = max(estimated_lines, int(math.ceil(len(cells[(out_row, best_col)]) / 22.0)))
                    continue
            for col in range(best_col, min(col_count + 1, best_col + best_span)):
                occupied[col] = True
            cells[(out_row, best_col)] = text_value
            if best_span > 1:
                merges[(out_row, best_col)] = (1, best_span)
            estimated_lines = max(estimated_lines, int(math.ceil(len(text_value) / max(10.0, best_span * 12.0))))
        row_height = max(12.0, min(42.0, tallest * 0.34, 18.0 * estimated_lines))
        if items and any(_clean_text(item.get("text", "")) for item in items):
            row_heights.append(row_height)
            out_row += 1
        previous_y = row_y
        previous_height = tallest
    row_count = max(1, out_row - 1)
    if not row_heights:
        row_heights = [18.0]
    return row_count, col_count, cells, merges, col_widths, row_heights

def _spreadsheet_grid(records, image_size=None, columns=120):
    row_count, col_count, cells, _merges, _col_widths, _row_heights = _spreadsheet_layout(records, image_size)
    return row_count, col_count, cells

def _xlsx_content_types():
    return ''.join([
        '<?xml version="1.0" encoding="UTF-8"?>',
        '<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">',
        '<Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>',
        '<Default Extension="xml" ContentType="application/xml"/>',
        '<Override PartName="/xl/workbook.xml" ContentType="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet.main+xml"/>',
        '<Override PartName="/xl/worksheets/sheet1.xml" ContentType="application/vnd.openxmlformats-officedocument.spreadsheetml.worksheet+xml"/>',
        '<Override PartName="/xl/styles.xml" ContentType="application/vnd.openxmlformats-officedocument.spreadsheetml.styles+xml"/>',
        '<Override PartName="/docProps/core.xml" ContentType="application/vnd.openxmlformats-package.core-properties+xml"/>',
        '<Override PartName="/docProps/app.xml" ContentType="application/vnd.openxmlformats-officedocument.extended-properties+xml"/>',
        '</Types>'
    ])

def _xlsx_root_rels():
    return ''.join([
        '<?xml version="1.0" encoding="UTF-8"?>',
        '<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">',
        '<Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument" Target="xl/workbook.xml"/>',
        '<Relationship Id="rId2" Type="http://schemas.openxmlformats.org/package/2006/relationships/metadata/core-properties" Target="docProps/core.xml"/>',
        '<Relationship Id="rId3" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/extended-properties" Target="docProps/app.xml"/>',
        '</Relationships>'
    ])

def _xlsx_workbook_xml():
    return ''.join([
        '<?xml version="1.0" encoding="UTF-8"?>',
        '<workbook xmlns="http://schemas.openxmlformats.org/spreadsheetml/2006/main" xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships">',
        '<sheets><sheet name="OCR Layout" sheetId="1" r:id="rId1"/></sheets>',
        '</workbook>'
    ])

def _xlsx_workbook_rels():
    return ''.join([
        '<?xml version="1.0" encoding="UTF-8"?>',
        '<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">',
        '<Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/worksheet" Target="worksheets/sheet1.xml"/>',
        '<Relationship Id="rId2" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/styles" Target="styles.xml"/>',
        '</Relationships>'
    ])

def _xlsx_styles_xml():
    return ''.join([
        '<?xml version="1.0" encoding="UTF-8"?>',
        '<styleSheet xmlns="http://schemas.openxmlformats.org/spreadsheetml/2006/main">',
        '<fonts count="2"><font><sz val="11"/><name val="Calibri"/></font><font><sz val="8"/><name val="Calibri"/></font></fonts>',
        '<fills count="2"><fill><patternFill patternType="none"/></fill><fill><patternFill patternType="gray125"/></fill></fills>',
        '<borders count="1"><border><left/><right/><top/><bottom/><diagonal/></border></borders>',
        '<cellStyleXfs count="1"><xf numFmtId="0" fontId="0" fillId="0" borderId="0"/></cellStyleXfs>',
        '<cellXfs count="2"><xf numFmtId="0" fontId="0" fillId="0" borderId="0" xfId="0"/><xf numFmtId="0" fontId="1" fillId="0" borderId="0" xfId="0" applyFont="1" applyAlignment="1"><alignment horizontal="left" vertical="top" wrapText="1"/></xf></cellXfs>',
        '<cellStyles count="1"><cellStyle name="Normal" xfId="0" builtinId="0"/></cellStyles>',
        '</styleSheet>'
    ])

def _xlsx_merge_ref(row, col, row_span, col_span):
    start = f'{_xlsx_col_name(col)}{row}'
    end = f'{_xlsx_col_name(col + col_span - 1)}{row + row_span - 1}'
    return start if start == end else f'{start}:{end}'

def _xlsx_sheet_xml(row_count, col_count, cells, merges=None, col_widths=None, row_heights=None):
    merges = merges or {}
    col_widths = col_widths or [2.5] * col_count
    row_heights = row_heights or [12.0] * row_count
    col_xml = ''.join(f'<col min="{i}" max="{i}" width="{max(0.35, min(28.0, float(col_widths[i - 1]))):.3f}" customWidth="1"/>' for i in range(1, col_count + 1))
    rows_xml = []
    for row in range(1, row_count + 1):
        row_cells = []
        for (r, c), text in sorted(cells.items()):
            if r != row:
                continue
            ref = f'{_xlsx_col_name(c)}{r}'
            row_cells.append(f'<c r="{ref}" s="1" t="inlineStr"><is><t xml:space="preserve">{_xml(text)}</t></is></c>')
        height = max(3.0, min(80.0, float(row_heights[row - 1])))
        rows_xml.append(f'<row r="{row}" ht="{height:.2f}" customHeight="1">{"".join(row_cells)}</row>')
    merge_refs = [_xlsx_merge_ref(row, col, span[0], span[1]) for (row, col), span in sorted(merges.items()) if span[0] > 1 or span[1] > 1]
    merge_xml = ''
    if merge_refs:
        merge_xml = '<mergeCells count="%d">%s</mergeCells>' % (len(merge_refs), ''.join(f'<mergeCell ref="{ref}"/>' for ref in merge_refs))
    return ''.join([
        '<?xml version="1.0" encoding="UTF-8"?>',
        '<worksheet xmlns="http://schemas.openxmlformats.org/spreadsheetml/2006/main" xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships">',
        '<sheetViews><sheetView workbookViewId="0" zoomScale="100" zoomScaleNormal="100"/></sheetViews>',
        '<sheetFormatPr defaultRowHeight="15"/>',
        '<cols>', col_xml, '</cols>',
        '<sheetData>', ''.join(rows_xml), '</sheetData>',
        merge_xml,
        '<pageMargins left="0.25" right="0.25" top="0.25" bottom="0.25" header="0" footer="0"/>',
        '<pageSetup paperSize="9" orientation="portrait" fitToWidth="1" fitToHeight="0"/>',
        '</worksheet>'
    ])

def write_positioned_xlsx(path, item, export_image, record_views):
    image_size = getattr(export_image, "size", None) or getattr(item, "image_size", None) or (0, 0)
    records = _records_from_views(record_views)
    row_count, col_count, cells, merges, col_widths, row_heights = _spreadsheet_layout(records, image_size)
    core = '<?xml version="1.0" encoding="UTF-8"?><cp:coreProperties xmlns:cp="http://schemas.openxmlformats.org/package/2006/metadata/core-properties" xmlns:dc="http://purl.org/dc/elements/1.1/"><dc:creator>Bottled Kraken</dc:creator></cp:coreProperties>'
    app = '<?xml version="1.0" encoding="UTF-8"?><Properties xmlns="http://schemas.openxmlformats.org/officeDocument/2006/extended-properties"><Application>Bottled Kraken</Application></Properties>'
    with zipfile.ZipFile(path, "w", zipfile.ZIP_DEFLATED) as archive:
        archive.writestr("[Content_Types].xml", _xlsx_content_types())
        archive.writestr("_rels/.rels", _xlsx_root_rels())
        archive.writestr("xl/workbook.xml", _xlsx_workbook_xml())
        archive.writestr("xl/_rels/workbook.xml.rels", _xlsx_workbook_rels())
        archive.writestr("xl/styles.xml", _xlsx_styles_xml())
        archive.writestr("xl/worksheets/sheet1.xml", _xlsx_sheet_xml(row_count, col_count, cells, merges, col_widths, row_heights))
        archive.writestr("docProps/core.xml", core)
        archive.writestr("docProps/app.xml", app)

def _ods_cell_xml(text, row_span=1, col_span=1):
    if not text and row_span <= 1 and col_span <= 1:
        return '<table:table-cell/>'
    span = ''
    if col_span > 1:
        span += ' table:number-columns-spanned="%d"' % int(col_span)
    if row_span > 1:
        span += ' table:number-rows-spanned="%d"' % int(row_span)
    if not text:
        return '<table:table-cell table:style-name="ce1"%s/>' % span
    return '<table:table-cell table:style-name="ce1" office:value-type="string"%s><text:p>%s</text:p></table:table-cell>' % (span, _odt_xml_text(text))

def _ods_sheet_content(row_count, col_count, cells, merges=None):
    merges = merges or {}
    covered = set()
    for (row, col), (row_span, col_span) in merges.items():
        for r in range(row, row + row_span):
            for c in range(col, col + col_span):
                if (r, c) != (row, col):
                    covered.add((r, c))
    rows = []
    for row in range(1, row_count + 1):
        parts = []
        col = 1
        while col <= col_count:
            key = (row, col)
            if key in covered:
                parts.append('<table:covered-table-cell/>')
                col += 1
                continue
            row_span, col_span = merges.get(key, (1, 1))
            parts.append(_ods_cell_xml(cells.get(key, ""), row_span, col_span))
            col += 1
        rows.append('<table:table-row table:style-name="ro%d">%s</table:table-row>' % (row, ''.join(parts)))
    return ''.join(rows)

def _ods_content_xml(row_count, col_count, cells, merges=None, col_widths=None, row_heights=None):
    col_widths = col_widths or [2.5] * col_count
    row_heights = row_heights or [12.0] * row_count
    column_styles = []
    columns = []
    row_styles = []
    for index, width in enumerate(col_widths, start=1):
        cm = max(0.04, min(3.0, float(width) * 0.18))
        column_styles.append('<style:style style:name="co%d" style:family="table-column"><style:table-column-properties style:column-width="%.3fcm"/></style:style>' % (index, cm))
        columns.append('<table:table-column table:style-name="co%d"/>' % index)
    for index, height in enumerate(row_heights, start=1):
        cm = max(0.04, min(2.0, float(height) * 0.0353))
        row_styles.append('<style:style style:name="ro%d" style:family="table-row"><style:table-row-properties style:row-height="%.3fcm" fo:break-before="auto"/></style:style>' % (index, cm))
    return ''.join([
        '<?xml version="1.0" encoding="UTF-8"?>',
        '<office:document-content ',
        'xmlns:office="urn:oasis:names:tc:opendocument:xmlns:office:1.0" ',
        'xmlns:style="urn:oasis:names:tc:opendocument:xmlns:style:1.0" ',
        'xmlns:text="urn:oasis:names:tc:opendocument:xmlns:text:1.0" ',
        'xmlns:table="urn:oasis:names:tc:opendocument:xmlns:table:1.0" ',
        'xmlns:fo="urn:oasis:names:tc:opendocument:xmlns:xsl-fo-compatible:1.0" ',
        'xmlns:svg="urn:oasis:names:tc:opendocument:xmlns:svg-compatible:1.0" ',
        'office:version="1.2">',
        '<office:font-face-decls><style:font-face style:name="Calibri" svg:font-family="Calibri"/></office:font-face-decls>',
        '<office:automatic-styles>',
        ''.join(column_styles),
        ''.join(row_styles),
        '<style:style style:name="ce1" style:family="table-cell"><style:table-cell-properties fo:padding="0.06cm"/><style:text-properties fo:font-size="8pt" style:font-name="Calibri"/><style:paragraph-properties fo:margin-top="0cm" fo:margin-bottom="0cm" fo:line-height="100%"/></style:style>',
        '</office:automatic-styles>',
        '<office:body><office:spreadsheet><table:table table:name="OCR Layout">',
        ''.join(columns),
        _ods_sheet_content(row_count, col_count, cells, merges),
        '</table:table></office:spreadsheet></office:body></office:document-content>'
    ])

def _ods_spreadsheet_styles_xml():
    return ''.join([
        '<?xml version="1.0" encoding="UTF-8"?>',
        '<office:document-styles xmlns:office="urn:oasis:names:tc:opendocument:xmlns:office:1.0" xmlns:style="urn:oasis:names:tc:opendocument:xmlns:style:1.0" xmlns:fo="urn:oasis:names:tc:opendocument:xmlns:xsl-fo-compatible:1.0" office:version="1.2">',
        '<office:styles><style:default-style style:family="table-cell"><style:text-properties fo:font-size="8pt" style:font-name="Calibri"/></style:default-style></office:styles>',
        '</office:document-styles>'
    ])

def _ods_spreadsheet_manifest_xml():
    return '<?xml version="1.0" encoding="UTF-8"?><manifest:manifest xmlns:manifest="urn:oasis:names:tc:opendocument:xmlns:manifest:1.0" manifest:version="1.2"><manifest:file-entry manifest:full-path="/" manifest:media-type="application/vnd.oasis.opendocument.spreadsheet"/><manifest:file-entry manifest:full-path="content.xml" manifest:media-type="text/xml"/><manifest:file-entry manifest:full-path="styles.xml" manifest:media-type="text/xml"/><manifest:file-entry manifest:full-path="meta.xml" manifest:media-type="text/xml"/></manifest:manifest>'

def write_positioned_ods(path, item, export_image, record_views):
    image_size = getattr(export_image, "size", None) or getattr(item, "image_size", None) or (0, 0)
    records = _records_from_views(record_views)
    row_count, col_count, cells, merges, col_widths, row_heights = _spreadsheet_layout(records, image_size)
    meta = '<?xml version="1.0" encoding="UTF-8"?><office:document-meta xmlns:office="urn:oasis:names:tc:opendocument:xmlns:office:1.0" xmlns:meta="urn:oasis:names:tc:opendocument:xmlns:meta:1.0" office:version="1.2"><office:meta><meta:generator>Bottled Kraken</meta:generator></office:meta></office:document-meta>'
    with zipfile.ZipFile(path, "w") as archive:
        info = zipfile.ZipInfo("mimetype")
        info.date_time = (2020, 1, 1, 0, 0, 0)
        info.compress_type = zipfile.ZIP_STORED
        archive.writestr(info, "application/vnd.oasis.opendocument.spreadsheet")
        for name, data in (("content.xml", _ods_content_xml(row_count, col_count, cells, merges, col_widths, row_heights)), ("styles.xml", _ods_spreadsheet_styles_xml()), ("meta.xml", meta), ("META-INF/manifest.xml", _ods_spreadsheet_manifest_xml())):
            zi = zipfile.ZipInfo(name)
            zi.date_time = (2020, 1, 1, 0, 0, 0)
            zi.compress_type = zipfile.ZIP_DEFLATED
            archive.writestr(zi, data.encode("utf-8"))

def install_export_layout(main_window_cls):
    previous_render_file = getattr(main_window_cls, "_render_file", None)
    def render_file(self, path, fmt, item):
        fmt_name = str(fmt or "").lower().lstrip(".")
        if fmt_name in {"docx", "word", "odt", "xlsx", "excel", "ods", "calc", "txt", "text", "txt_plain", "csv"}:
            if not item or not getattr(item, "results", None):
                return None
            _text, _kraken_records, pil_image, record_views = item.results
            try:
                export_image = _load_image_color(item.path)
            except Exception:
                export_image = pil_image
            image_size = getattr(export_image, "size", None) or getattr(pil_image, "size", None)
            if fmt_name in {"txt", "text", "txt_plain"}:
                return write_spatial_txt(path, record_views, image_size)
            if fmt_name == "csv":
                return write_plain_csv(path, record_views, image_size)
            if fmt_name == "odt":
                return write_positioned_odt(path, item, export_image, record_views)
            if fmt_name in {"xlsx", "excel"}:
                return write_positioned_xlsx(path, item, export_image, record_views)
            if fmt_name in {"ods", "calc"}:
                return write_positioned_ods(path, item, export_image, record_views)
            return write_positioned_docx(path, item, export_image, record_views)
        if callable(previous_render_file):
            return previous_render_file(self, path, fmt, item)
        return None
    main_window_cls._render_file = render_file
    return render_file
