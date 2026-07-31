from bottled_kraken.module_registry import register_globals, seed_globals
seed_globals('bk', globals())
from bottled_kraken.common import (
    Any,
    List,
    QFileDialog,
    QMessageBox,
    os,
    re,
)
from bottled_kraken.main_window import MainWindow
def _bk_fix36_tr(window, key: str, *args) -> str:
    lang = getattr(window, "current_lang", translation.DEFAULT_LANGUAGE)
    return translation.translate(lang, key, *args)
def _bk_fix36_clean_text(value) -> str:
    return re.sub(r"\s+", " ", str(value or "")).strip()
def _bk_fix36_is_ditto(value: str) -> bool:
    txt = _bk_fix36_clean_text(value)
    compact = re.sub(r"\s+", "", txt)
    return compact in {'"', "''", "„", "“", "”", "-\"-", "–\"–", "—\"—", "-''-", "--\"--"}
def _bk_fix36_resolve_ditto_marks_in_lines(lines: List[str]) -> List[str]:
    out: List[str] = []
    prev_tokens: List[str] = []
    for raw in lines or []:
        text = str(raw or "")
        if not text.strip():
            out.append(text)
            continue
        tokens = re.split(r"(\s+)", text)
        word_index = 0
        changed = False
        for i, token in enumerate(tokens):
            if token.isspace() or token == "":
                continue
            if _bk_fix36_is_ditto(token) and word_index < len(prev_tokens):
                tokens[i] = prev_tokens[word_index]
                changed = True
            word_index += 1
        new_text = "".join(tokens) if changed else text
        out.append(new_text)
        words = [t for t in re.split(r"\s+", new_text.strip()) if t]
        if words:
            prev_tokens = words
    return out
def _bk_fix36_resolve_ditto_marks_in_recs(recs):
    try:
        texts = [getattr(rv, "text", "") for rv in recs]
        resolved = _bk_fix36_resolve_ditto_marks_in_lines(texts)
        for rv, txt in zip(recs, resolved):
            try:
                rv.text = txt
            except Exception:
                pass
    except Exception:
        pass
    return recs
def _bk_fix36_group_recs_into_table_rows(recs) -> List[List[Any]]:
    with_boxes = [rv for rv in (recs or []) if getattr(rv, "bbox", None)]
    if not with_boxes:
        return [[rv] for rv in (recs or [])]
    heights = [max(1, rv.bbox[3] - rv.bbox[1]) for rv in with_boxes]
    median_h = sorted(heights)[len(heights)//2] if heights else 18
    threshold = max(8, median_h * 0.65)
    rows: List[List[Any]] = []
    for rv in sorted(with_boxes, key=lambda r: ((r.bbox[1] + r.bbox[3]) / 2.0, r.bbox[0])):
        cy = (rv.bbox[1] + rv.bbox[3]) / 2.0
        placed = False
        for row in rows:
            rcy = sum((x.bbox[1] + x.bbox[3]) / 2.0 for x in row) / max(1, len(row))
            if abs(cy - rcy) <= threshold:
                row.append(rv)
                placed = True
                break
        if not placed:
            rows.append([rv])
    for row in rows:
        row.sort(key=lambda r: r.bbox[0] if getattr(r, "bbox", None) else 0)
    return rows
def _bk_fix36_table_text_from_recs(recs) -> str:
    recs = list(recs or [])
    recs = _bk_fix36_resolve_ditto_marks_in_recs(recs)
    if not recs:
        return ""
    if not any(getattr(rv, "bbox", None) for rv in recs):
        return "\n".join(_bk_fix36_clean_text(getattr(rv, "text", "")) for rv in recs)
    rows = _bk_fix36_group_recs_into_table_rows(recs)
    anchors = []
    for rv in recs:
        bb = getattr(rv, "bbox", None)
        if not bb:
            continue
        x = int(bb[0])
        if not any(abs(x - a) < 28 for a in anchors):
            anchors.append(x)
    anchors = sorted(anchors)
    if len(anchors) < 2:
        return "\n".join(_bk_fix36_clean_text(getattr(rv, "text", "")) for row in rows for rv in row)
    matrix = []
    for row in rows:
        cells = [""] * len(anchors)
        for rv in row:
            bb = getattr(rv, "bbox", None)
            text = _bk_fix36_clean_text(getattr(rv, "text", ""))
            if not bb or not text:
                continue
            idx = min(range(len(anchors)), key=lambda i: abs(bb[0] - anchors[i]))
            cells[idx] = (cells[idx] + " " + text).strip() if cells[idx] else text
        matrix.append(cells)
    widths = []
    for col in range(len(anchors)):
        widths.append(min(44, max(3, max((len(row[col]) for row in matrix), default=3))))
    out = []
    for row in matrix:
        out.append(" | ".join((row[i] or "").ljust(widths[i]) for i in range(len(widths))).rstrip())
    return "\n".join(out).rstrip() + "\n"
def _bk_fix36_current_task(self):
    try:
        if hasattr(self, "_current_task"):
            return self._current_task()
    except Exception:
        return None
    return None
def _bk_export_origami_table_txt(self):
    task = _bk_fix36_current_task(self)
    if not task or not getattr(task, "results", None):
        QMessageBox.information(self, _bk_fix36_tr(self, "info_title"), _bk_fix36_tr(self, "warn_no_ocr_results"))
        return
    text, kr_records, im, recs = task.results
    out_text = _bk_fix36_table_text_from_recs(recs)
    default_name = os.path.splitext(os.path.basename(getattr(task, "path", "ocr")))[0] + "_origami_table.txt"
    path, _ = QFileDialog.getSaveFileName(self, _bk_fix36_tr(self, "export_format_txt_origami_table"), default_name, _bk_fix36_tr(self, "filter_text_files"))
    if not path:
        return
    if not path.lower().endswith(".txt"):
        path += ".txt"
    with open(path, "w", encoding="utf-8") as fh:
        fh.write(out_text)
    try:
        self.status_bar.showMessage(_bk_fix36_tr(self, "msg_origami_table_export_done"), 5000)
    except Exception:
        pass
def _bk_export_origami_table_docx(self):
    task = _bk_fix36_current_task(self)
    if not task or not getattr(task, "results", None):
        QMessageBox.information(self, _bk_fix36_tr(self, "info_title"), _bk_fix36_tr(self, "warn_no_ocr_results"))
        return
    text, kr_records, im, recs = task.results
    out_text = _bk_fix36_table_text_from_recs(recs)
    default_name = os.path.splitext(os.path.basename(getattr(task, "path", "ocr")))[0] + "_origami_table.docx"
    path, _ = QFileDialog.getSaveFileName(self, _bk_fix36_tr(self, "export_format_docx_origami_table"), default_name, _bk_fix36_tr(self, "filter_word_files"))
    if not path:
        return
    if not path.lower().endswith(".docx"):
        path += ".docx"
    try:
        from docx import Document
        from docx.shared import Pt
        doc = Document()
        para = doc.add_paragraph()
        run = para.add_run(out_text)
        run.font.name = "Courier New"
        run.font.size = Pt(9)
        doc.save(path)
    except Exception:
        alt = path.rsplit(".", 1)[0] + ".txt"
        with open(alt, "w", encoding="utf-8") as fh:
            fh.write(out_text)
        QMessageBox.information(self, _bk_fix36_tr(self, "info_title"), _bk_fix36_tr(self, "msg_docx_missing_txt_saved", alt))
        return
    try:
        self.status_bar.showMessage(_bk_fix36_tr(self, "msg_origami_table_export_done"), 5000)
    except Exception:
        pass
MainWindow.bk_export_origami_table_txt = _bk_export_origami_table_txt
MainWindow.bk_export_origami_table_docx = _bk_export_origami_table_docx
MainWindow.bk_resolve_ditto_marks_in_recs = lambda self, recs: _bk_fix36_resolve_ditto_marks_in_recs(recs)
def _bk_fix36_wrap_text_export_method(method_name: str):
    old = getattr(MainWindow, method_name, None)
    if not callable(old) or getattr(old, "_bk_fix36_wrapped", False):
        return
    def wrapper(self, *args, **kwargs):
        try:
            task = _bk_fix36_current_task(self)
            if task is not None and getattr(task, "results", None):
                text, kr_records, im, recs = task.results
                _bk_fix36_resolve_ditto_marks_in_recs(recs)
                task.results = (text, kr_records, im, recs)
        except Exception:
            pass
        return old(self, *args, **kwargs)
    wrapper._bk_fix36_wrapped = True
    setattr(MainWindow, method_name, wrapper)
for _name in (
    "export_txt", "export_text", "export_results_txt", "export_current_txt",
    "export_docx", "export_word", "export_results_docx", "export_current_docx",
):
    _bk_fix36_wrap_text_export_method(_name)
try:
    _BK_FIX36_PREV_EXPORT_FORMAT_ITEMS = MainWindow._export_format_items
except Exception:
    _BK_FIX36_PREV_EXPORT_FORMAT_ITEMS = None
def _bk_fix36_export_format_items(self):
    items = []
    if callable(_BK_FIX36_PREV_EXPORT_FORMAT_ITEMS):
        try:
            items = list(_BK_FIX36_PREV_EXPORT_FORMAT_ITEMS(self) or [])
        except Exception:
            items = []
    existing = {fmt for _name, fmt in items}
    if "txt_origami_table" not in existing:
        items.append((_bk_fix36_tr(self, "export_format_txt_origami_table"), "txt_origami_table"))
    if "docx_origami_table" not in existing:
        items.append((_bk_fix36_tr(self, "export_format_docx_origami_table"), "docx_origami_table"))
    return items
try:
    if callable(_BK_FIX36_PREV_EXPORT_FORMAT_ITEMS):
        MainWindow._export_format_items = _bk_fix36_export_format_items
except Exception:
    pass
try:
    _BK_FIX36_PREV_EXPORT_FLOW = MainWindow.export_flow
except Exception:
    _BK_FIX36_PREV_EXPORT_FLOW = None
def _bk_fix36_export_flow(self, fmt, *args, **kwargs):
    if fmt == "txt_origami_table":
        return self.bk_export_origami_table_txt()
    if fmt == "docx_origami_table":
        return self.bk_export_origami_table_docx()
    if callable(_BK_FIX36_PREV_EXPORT_FLOW):
        return _BK_FIX36_PREV_EXPORT_FLOW(self, fmt, *args, **kwargs)
    return None
try:
    if callable(_BK_FIX36_PREV_EXPORT_FLOW):
        MainWindow.export_flow = _bk_fix36_export_flow
except Exception:
    pass
__all__ = [
    '_bk_export_origami_table_docx',
    '_bk_export_origami_table_txt',
    '_bk_fix36_clean_text',
    '_bk_fix36_current_task',
    '_bk_fix36_export_flow',
    '_bk_fix36_export_format_items',
    '_bk_fix36_group_recs_into_table_rows',
    '_bk_fix36_is_ditto',
    '_bk_fix36_resolve_ditto_marks_in_lines',
    '_bk_fix36_resolve_ditto_marks_in_recs',
    '_bk_fix36_table_text_from_recs',
    '_bk_fix36_tr',
    '_bk_fix36_wrap_text_export_method',
    '_name',
]
register_globals('bk', globals(), __all__)
