"""KI-Export über die lokale KI: DOCX / XLSX / ODS / ODT mit ECHTEN Tabellen.

Ablauf:
1. Die erkannten Zeilen (Overlay-Boxen) des aktuellen Eintrags werden in
   Lesereihenfolge an die lokale KI geschickt.
2. Die KI strukturiert den Inhalt als striktes JSON-Blockschema:
       {"blocks": [
           {"type": "heading",   "text": "..."},
           {"type": "paragraph", "text": "..."},
           {"type": "table",     "rows": [["Zelle", ...], ...]}
       ]}
   Der Zeilentext muss dabei wortgetreu erhalten bleiben; die KI entscheidet
   nur über die Struktur (welche Zeilen bilden zusammen eine Tabelle, wo
   verlaufen die Spalten).
3. Das Programm rendert daraus echte Dokumente: python-docx-Tabellen (DOCX),
   odfpy-Tabellen (ODT), Tabellenkalkulations-Zellen (XLSX/ODS) - keine
   Textfelder, keine Pseudo-Tabellen aus Leerzeichen.

Die LM-Anbindung nutzt denselben Endpoint wie die Zeilen-Revision
(self.ai_endpoint) und dieselben Fehlertexte (ai_err_*-Schlüssel).
"""

from __future__ import annotations

import http.client
import json
import os
import re
import select
import socket
import time
import urllib.parse
import urllib.request
import urllib.error
from typing import Any, Dict, List, Optional

from PySide6.QtCore import Qt, QThread, Signal
from PySide6.QtGui import QAction, QGuiApplication, QCursor
from PySide6.QtWidgets import QFileDialog, QMenu, QMessageBox

from bottled_kraken.dialogs import ProgressStatusDialog

from bottled_kraken.module_registry import register_globals, seed_globals
from bottled_kraken.common.chain_consolidation import (
    register_init_delta,
    register_retranslate_delta,
)

seed_globals("bk", globals())

try:
    from bottled_kraken.main_window import MainWindow
except Exception:  # pragma: no cover
    MainWindow = None

_BK_LMX_FORMATS = ("docx", "xlsx", "odt", "ods")
_BK_LMX_EMPTY_CELL = "[[EMPTY]]"


def _bk_lmx_tr(self, key: str, default: str) -> str:
    try:
        value = self._tr(key)
        if value and value != key:
            return str(value)
    except Exception:
        pass
    return default


def _bk_lmx_prompt(self, fmt: str, role: str) -> str:
    """Return the saved prompt for one LM Office export format.

    The prompt editor stores overrides under lm_prompts/<language>/<key>.
    This helper uses exactly those settings and falls back to the translated
    format-specific defaults. Older installations still fall back to the
    former generic structured-export prompts.
    """
    fmt = str(fmt or "docx").strip().lower()
    if fmt not in _BK_LMX_FORMATS:
        fmt = "docx"
    role = "user" if str(role).strip().lower() == "user" else "system"
    key = f"ai_prompt_lmx_{fmt}_{role}"
    try:
        override = _bk_lm_prompt_override(self, key)
    except Exception:
        override = ""
    if str(override or ""):
        return str(override)
    try:
        lang = getattr(self, "current_lang", translation.DEFAULT_LANGUAGE)
        default = _bk_lm_default_prompt(lang, key)
    except Exception:
        default = ""
    if default and default != key:
        return str(default)
    fallback_key = "ai_prompt_structured_user" if role == "user" else "ai_prompt_structured_system"
    fallback_default = "{}" if role == "user" else ""
    return _bk_lmx_tr(self, fallback_key, fallback_default)


# ---------------------------------------------------------------------------
# LM-Anfrage
# ---------------------------------------------------------------------------

def _bk_lmx_endpoint(self) -> str:
    endpoint = str(getattr(self, "ai_endpoint", "") or "").strip()
    if not endpoint:
        endpoint = "http://127.0.0.1:1234/v1/chat/completions"
    return endpoint


def _bk_lmx_model(self) -> str:
    for attr in ("ai_model", "ai_model_name", "lm_model", "selected_ai_model"):
        value = str(getattr(self, attr, "") or "").strip()
        if value:
            return value
    return "local"


_BK_LMX_BLOCKS_SCHEMA = {
    "type": "json_schema",
    "json_schema": {
        "name": "structured_blocks",
        "schema": {
            "type": "object",
            "properties": {
                "blocks": {
                    "type": "array",
                    "items": {
                        "type": "object",
                        "properties": {
                            "type": {"type": "string", "enum": ["heading", "paragraph", "table"]},
                            "text": {"type": "string"},
                            "header": {"type": "array", "items": {"type": "string"}},
                            "rows": {"type": "array", "items": {"type": "array", "items": {"type": "string"}}},
                        },
                        "required": ["type"],
                    },
                }
            },
            "required": ["blocks"],
        },
    },
}


def _bk_lmx_chat(self, system: str, user: str, timeout: int = 300, cancel_check=None) -> str:
    """Abbruchfaehige LM-Anfrage fuer den KI-Export.

    Drei Ursachen des bisherigen "Ungueltige JSON-Antwort"-Fehlers werden hier
    beseitigt: (1) response_format=json_schema zwingt den Server per Grammatik
    zu gueltigem JSON, (2) max_tokens wird an das kleine Kontextfenster (z. B.
    8000) angepasst statt fix 16000 zu fordern (was die Antwort abschnitt),
    (3) waehrend der Generierung wird per select() gepollt, sodass Abbrechen
    sofort greift und die GUI nie blockiert."""
    body_obj = {
        "model": _bk_lmx_model(self),
        "messages": [
            {"role": "system", "content": "/no_think\n" + system},
            {"role": "user", "content": "/no_think\n" + user},
        ],
        "temperature": 0,
        "response_format": _BK_LMX_BLOCKS_SCHEMA,
        "reasoning": {"effort": "none"},
        "reasoning_effort": "none",
        "enable_thinking": False,
        "chat_template_kwargs": {"enable_thinking": False},
    }
    # Token-Budget an das Kontextfenster anpassen: grobe Eingabeschaetzung
    # (Zeichen/3 wegen JSON-Overhead) von einem konservativen 8000er-Fenster
    # abziehen; die Ausgabe ist etwa so lang wie die Eingabezeilen.
    est_in = (len(system) + len(user)) // 3 + 64
    body_obj["max_tokens"] = max(1200, min(7000, 7800 - est_in))
    body = json.dumps(body_obj).encode("utf-8")

    parsed = urllib.parse.urlparse(_bk_lmx_endpoint(self))
    host = parsed.hostname or "127.0.0.1"
    port = parsed.port
    path = parsed.path or "/v1/chat/completions"
    conn = None
    try:
        if parsed.scheme == "https":
            conn = http.client.HTTPSConnection(host, port or 443, timeout=30)
        else:
            conn = http.client.HTTPConnection(host, port or 80, timeout=30)
        conn.request("POST", path, body=body, headers={
            "Content-Type": "application/json",
            "Authorization": "Bearer lm-studio",
        })
        deadline = time.monotonic() + max(60, int(timeout))
        sock = getattr(conn, "sock", None)
        if sock is not None:
            while True:
                if callable(cancel_check) and cancel_check():
                    raise RuntimeError(_bk_lmx_tr(self, "msg_ai_cancelled", "KI-Vorgang abgebrochen."))
                if time.monotonic() > deadline:
                    raise RuntimeError(_bk_lmx_tr(self, "ai_err_timeout", "Zeitüberschreitung beim LM-Server."))
                try:
                    readable, _, _ = select.select([sock], [], [], 0.4)
                except Exception:
                    break
                if readable:
                    break
        resp = conn.getresponse()
        raw = resp.read().decode("utf-8", errors="replace")
        if callable(cancel_check) and cancel_check():
            raise RuntimeError(_bk_lmx_tr(self, "msg_ai_cancelled", "KI-Vorgang abgebrochen."))
        if resp.status >= 400:
            raise RuntimeError(_bk_lmx_tr(self, "ai_err_http", "HTTP {}: {}").format(resp.status, raw[:400]))
        data = json.loads(raw)
    except (OSError, socket.timeout) as exc:
        if callable(cancel_check) and cancel_check():
            raise RuntimeError(_bk_lmx_tr(self, "msg_ai_cancelled", "KI-Vorgang abgebrochen.")) from exc
        raise RuntimeError(_bk_lmx_tr(self, "ai_err_server_unreachable",
                                      "LM-Server nicht erreichbar: {}").format(exc)) from exc
    finally:
        try:
            if conn is not None:
                conn.close()
        except Exception:
            pass
    choices = data.get("choices") or []
    if not choices:
        raise RuntimeError(_bk_lmx_tr(self, "ai_err_no_choices",
                                      "LM-Server lieferte keine choices. Antwort:\n{}").format(str(data)[:400]))
    message = choices[0].get("message") or {}
    content = str(message.get("content") or "").strip()
    if not content and message.get("reasoning_content"):
        # Notanker fuer Denk-Modelle: Steckt das JSON versehentlich im
        # reasoning_content, wird es dort geborgen statt hart zu scheitern.
        reasoning = str(message.get("reasoning_content") or "")
        start, end = reasoning.find("{"), reasoning.rfind("}")
        if start >= 0 and end > start:
            candidate = reasoning[start:end + 1]
            try:
                json.loads(candidate)
                return candidate
            except Exception:
                pass
        raise RuntimeError(_bk_lmx_tr(self, "ai_err_reasoning_only",
                                      "Das Modell lieferte nur reasoning_content."))
    return content


def _bk_lmx_repair_json(text: str) -> Optional[Dict[str, Any]]:
    """Rettet abgeschnittenes Block-JSON: schrittweise bis zum letzten
    vollstaendigen Element zurueckschneiden und offene Klammern schliessen.
    Lieber die bereits gelieferten Bloecke verwenden als hart scheitern."""
    for cut in range(len(text), max(0, len(text) - 6000), -1):
        candidate = text[:cut].rstrip().rstrip(",")
        opens = candidate.count("{") - candidate.count("}")
        opens_sq = candidate.count("[") - candidate.count("]")
        if opens < 0 or opens_sq < 0:
            continue
        closed = candidate + ("]" * opens_sq) + ("}" * opens)
        try:
            data = json.loads(closed)
        except Exception:
            continue
        if isinstance(data, dict) and isinstance(data.get("blocks"), list) and data["blocks"]:
            return data
    return None


def _bk_lmx_extract_json(self, content: str) -> Dict[str, Any]:
    text = str(content or "").strip()
    text = re.sub(r"^```(?:json)?\s*", "", text)
    text = re.sub(r"\s*```$", "", text)
    start, end = text.find("{"), text.rfind("}")
    if start >= 0 and end > start:
        text = text[start:end + 1]
    try:
        data = json.loads(text)
    except Exception as exc:
        repaired = _bk_lmx_repair_json(str(content or ""))
        if repaired is not None:
            return repaired
        raise RuntimeError(_bk_lmx_tr(self, "ai_err_invalid_json",
                                      "Ungültige JSON-Antwort vom LM-Server: {}").format(exc)) from exc
    if not isinstance(data, dict) or not isinstance(data.get("blocks"), list):
        raise RuntimeError(_bk_lmx_tr(self, "lmx_err_no_blocks",
                                      "Die KI-Antwort enthält kein gültiges 'blocks'-Feld."))
    return data


def _bk_lmx_normalize_blocks(data: Dict[str, Any]) -> List[Dict[str, Any]]:
    blocks: List[Dict[str, Any]] = []
    for raw in data.get("blocks") or []:
        if not isinstance(raw, dict):
            continue
        kind = str(raw.get("type") or "").strip().lower()
        if kind == "table":
            rows_in = raw.get("rows") or []
            rows: List[List[str]] = []
            width = 0
            header: List[str] = []
            def _cell(value):
                text = str(value if value is not None else "").strip()
                if text.upper() in {"[[EMPTY]]", "[EMPTY]", "<EMPTY>", "EMPTY", "LEER", "[LEER]"}:
                    return ""
                return text
            raw_header = raw.get("header")
            if isinstance(raw_header, (list, tuple)):
                header = [_cell(c) for c in raw_header]
                width = max(width, len(header))
            for row in rows_in:
                if isinstance(row, (list, tuple)):
                    cells = [_cell(c) for c in row]
                else:
                    cells = [_cell(row)]
                rows.append(cells)
                width = max(width, len(cells))
            if rows and width:
                for cells in rows:
                    cells.extend([""] * (width - len(cells)))
                # Leere Endspalten abschneiden (einzelne Ausreisserzeilen erzeugen
                # sonst eine ueberzaehlige leere Spalte fuer die ganze Tabelle).
                while width > 1 and all(not str(r[width - 1]).strip() for r in rows) \
                        and (len(header) < width or not str(header[width - 1]).strip()):
                    for r in rows:
                        del r[width - 1]
                    if len(header) >= width:
                        del header[width - 1]
                    width -= 1
                block = {"type": "table", "rows": rows}
                if header and any(h.strip() for h in header):
                    header.extend([""] * (width - len(header)))
                    block["header"] = header[:width]
                blocks.append(block)
        elif kind in ("heading", "paragraph"):
            text = str(raw.get("text") or "").strip()
            if text:
                blocks.append({"type": kind, "text": text})
    return blocks


# ---------------------------------------------------------------------------
# Renderer mit echten Tabellen
# ---------------------------------------------------------------------------

def _bk_lmx_cell_is_numeric(value: str) -> bool:
    v = str(value or "").strip()
    if not v:
        return False
    return bool(re.fullmatch(r"[\d IVXLCDMivxlcdm.,;:/\-]+\.?;?", v)) and bool(re.search(r"\d", v))


def _bk_lmx_column_alignments(rows: List[List[str]]) -> List[str]:
    """Spalten, die überwiegend Zahlen/Daten enthalten, werden rechtsbündig -
    so wirken Register und Rechnungen wie von Hand gesetzt."""
    if not rows:
        return []
    width = max(len(r) for r in rows)
    aligns = []
    for c in range(width):
        vals = [str(r[c]).strip() for r in rows if c < len(r) and str(r[c]).strip()]
        numeric = sum(1 for v in vals if _bk_lmx_cell_is_numeric(v))
        aligns.append("right" if vals and numeric / len(vals) >= 0.6 else "left")
    return aligns


def _bk_lmx_write_docx(path: str, blocks: List[Dict[str, Any]]) -> None:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    for block in blocks:
        if block["type"] == "heading":
            doc.add_heading(block["text"], level=2)
        elif block["type"] == "paragraph":
            doc.add_paragraph(block["text"])
        else:
            rows = block["rows"]
            header = block.get("header") or []
            aligns = _bk_lmx_column_alignments(rows)
            total = len(rows) + (1 if header else 0)
            table = doc.add_table(rows=total, cols=len(rows[0]))
            table.style = "Table Grid"
            offset = 0
            if header:
                offset = 1
                for c, value in enumerate(header[:len(rows[0])]):
                    cell = table.cell(0, c)
                    cell.text = value
                    for p in cell.paragraphs:
                        for run in p.runs:
                            run.bold = True
            for r, cells in enumerate(rows):
                for c, value in enumerate(cells):
                    cell = table.cell(r + offset, c)
                    cell.text = value
                    if c < len(aligns) and aligns[c] == "right":
                        for p in cell.paragraphs:
                            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    doc.save(path)




def _bk_lmx_blocks_to_layout(blocks: List[Dict[str, Any]]) -> Dict[str, Any]:
    """Blöcke in ein Layout für die Tabellen-Writer übersetzen: Überschriften
    und Absätze als verbundene fette Zeilen (spans), Tabellenköpfe fett
    (bold_rows), Zahlen-/Datumsspalten rechtsbündig (aligns)."""
    matrix: List[List[str]] = []
    spans: set = set()
    bold_rows: set = set()
    all_table_rows: List[List[str]] = []
    for block in blocks:
        if block["type"] == "table":
            header = block.get("header") or []
            if header and any(str(h).strip() for h in header):
                bold_rows.add(len(matrix))
                matrix.append([str(h) for h in header])
            for cells in block["rows"]:
                matrix.append(list(cells))
                all_table_rows.append(list(cells))
            # Leerzeile nur nach Tabellen - Absatz-fuer-Absatz-Streifen vermeiden.
            matrix.append([""])
        else:
            spans.add(len(matrix))
            matrix.append([block["text"]])
    while matrix and matrix[-1] == [""]:
        matrix.pop()
    width = max((len(r) for r in matrix), default=1)
    matrix = [row + [""] * (width - len(row)) for row in matrix]
    aligns = _bk_lmx_column_alignments(all_table_rows) if all_table_rows else []
    aligns = (aligns + ["left"] * width)[:width]
    return {"matrix": matrix, "spans": spans, "bold_rows": bold_rows,
            "aligns": aligns, "widths": None, "heights": None}


def _bk_lmx_write(self, path: str, fmt: str, blocks: List[Dict[str, Any]]) -> None:
    fmt = str(fmt).lower()
    if fmt == "docx":
        _bk_lmx_write_docx(path, blocks)
    elif fmt in ("odt", "xlsx", "ods"):
        layout = _bk_lmx_blocks_to_layout(blocks)
        matrix = layout["matrix"]
        source_layout = getattr(self, "_bk_lmx_source_layout", None) or {}
        source_widths = list(source_layout.get("widths") or []) if isinstance(source_layout, dict) else []
        matrix_width = max([len(row) for row in matrix] or [1])
        if len(source_widths) == matrix_width:
            # Die Bildgeometrie bleibt auch nach der Strukturierung durch das
            # LM verbindlich: breite Namensspalten und schmale Zahlenspalten
            # werden nicht wieder auf gleich breite Standardzellen reduziert.
            layout["widths"] = source_widths
        else:
            col_widths_fn = globals().get("_bk_simple_col_widths")
            if callable(col_widths_fn):
                layout["widths"] = col_widths_fn(matrix)
        layout["heights"] = [18.0] * max(1, len(matrix))
        writer = globals().get(f"_bk_write_simple_layout_{fmt}")
        if callable(writer):
            writer(path, layout, self)
            return
        # ODT besitzt keinen separaten layout-Wrapper; der Matrix-Writer
        # akzeptiert das Layout direkt und erzeugt bei breiten Tabellen
        # automatisch Querformat mit proportionalen Spaltenbreiten.
        writer = globals().get(f"_bk_write_simple_matrix_{fmt}")
        if not callable(writer):
            raise RuntimeError(_bk_lmx_tr(self, "lmx_err_writer_missing", "Matrix-Writer für {} nicht verfügbar.").format(fmt))
        try:
            writer(path, matrix, self, layout=layout)
        except TypeError:
            writer(path, matrix, self)
    else:
        raise RuntimeError(_bk_lmx_tr(self, "lmx_err_unknown_format", "Unbekanntes KI-Exportformat: {}").format(fmt))


# ---------------------------------------------------------------------------
# Ablaufsteuerung
# ---------------------------------------------------------------------------

def _bk_lmx_collect_lines(item) -> List[str]:
    """Zeilen fuer den KI-Export mit einem verbindlichen raeumlichen Raster.

    Leere Zellen sind fuer historische Doppeltabellen genauso wichtig wie der
    Text selbst. Die alte ``" | "``-Liste entfernte leere Positionen; dadurch
    schob das Modell die rechte Tabellenhaelfte nach links und erzeugte acht
    oder neun beliebige Spalten. Jetzt wird zuerst dieselbe deterministische
    Layoutanalyse wie beim normalen Export verwendet. Tabellenzeilen tragen
    ihre feste Spaltenzahl und jede leere Zelle einen ausdruecklichen Marker.
    """
    results = getattr(item, "results", None)
    lines: List[str] = []
    if results and len(results) >= 4 and results[3]:
        try:
            layout_fn = globals().get("_bk_layout_blocks_with_tables")
            row_text_fn = globals().get("_bk_layout_row_text")
            image_size = getattr(results[2], "size", None) if len(results) > 2 else None
            if callable(layout_fn):
                layout = layout_fn(results[3], image_size)
                spatial_fn = globals().get("_bk_simple_spatial_layout_from_item")
                if callable(spatial_fn):
                    try:
                        setattr(item, "_bk_lmx_source_layout", spatial_fn(item))
                    except Exception:
                        pass
                for block in layout.get("blocks") or []:
                    if block.get("type") == "table":
                        matrix = block.get("matrix") or []
                        width = max([len(row) for row in matrix] or [1])
                        for row in matrix:
                            values = []
                            for c in range(width):
                                value = str(row[c] if c < len(row) else "").strip()
                                values.append(value if value else _BK_LMX_EMPTY_CELL)
                            lines.append(f"[TABLE cols={width}] " + " | ".join(values))
                    else:
                        for row in block.get("rows") or []:
                            if callable(row_text_fn):
                                text = str(row_text_fn(row) or "").strip()
                            else:
                                text = " ".join(
                                    str(rec.get("text", "") or "").strip()
                                    for rec in row.get("items", []) if str(rec.get("text", "") or "").strip()
                                )
                            if text:
                                lines.append("[TEXT] " + text)
                if lines:
                    return lines
        except Exception:
            lines = []

        # Rueckfall fuer alte/ungeboxte Ergebnisse.
        for rec in results[3]:
            text = str(getattr(rec, "text", "") or "").strip()
            if text:
                lines.append("[TEXT] " + text)
    if not lines and results and results[0]:
        lines = ["[TEXT] " + line for line in str(results[0]).splitlines() if line.strip()]
    return lines


def _bk_lmx_current_item(self):
    # Wichtig: Die Methode heisst _current_task(), nicht _current_task_item().
    # Der alte (nirgends existierende) Name liess die Item-Aufloesung immer
    # scheitern, wodurch der KI-Export faelschlich "Keine erkannten Zeilen"
    # meldete, obwohl Zeilen vorhanden waren.
    for meth in ("_current_task", "_current_task_item"):
        fn = getattr(self, meth, None)
        if callable(fn):
            try:
                item = fn()
                if item is not None:
                    return item
            except Exception:
                pass
    for attr in ("current_item", "_current_item", "active_item"):
        item = getattr(self, attr, None)
        if item is not None:
            return item
    # Letzter Rueckfall: einziges bzw. erstes fertiges Element der Warteschlange.
    try:
        items = list(getattr(self, "queue_items", []) or [])
        done = [it for it in items if getattr(it, "results", None)]
        if len(done) == 1:
            return done[0]
        if done:
            return done[0]
        if len(items) == 1:
            return items[0]
    except Exception:
        pass
    return None


_BK_LMX_CHUNK_LINES = 40


def _bk_lmx_structured_blocks(self, lines: List[str], cancel_check=None, status_cb=None, fmt: str = "docx") -> List[Dict[str, Any]]:
    """Zeilen in Portionen strukturieren, die sicher ins Kontextfenster passen.

    Bei einem 8000er-Fenster passen Eingabe UND Antwort einer ganzen Seite
    nicht zusammen hinein - genau daran scheiterte der Export bisher (die
    Antwort wurde mittendrin abgeschnitten -> "Expecting ',' delimiter").
    Portionen von ~40 Zeilen halten beide Seiten klein; die Bloecke werden
    anschliessend zusammengefuehrt (gleich breite Tabellen an den
    Portionsgrenzen werden zu einer Tabelle verbunden)."""
    system = _bk_lmx_prompt(self, fmt, "system")
    template = _bk_lmx_prompt(self, fmt, "user")
    chunks = [lines[i:i + _BK_LMX_CHUNK_LINES] for i in range(0, len(lines), _BK_LMX_CHUNK_LINES)] or [[]]
    merged: List[Dict[str, Any]] = []
    for c_idx, chunk in enumerate(chunks):
        if callable(cancel_check) and cancel_check():
            raise RuntimeError(_bk_lmx_tr(self, "msg_ai_cancelled", "KI-Vorgang abgebrochen."))
        if callable(status_cb):
            status_cb(_bk_lmx_tr(
                self, "lmx_status_chunk",
                "Die lokale KI strukturiert das Dokument (Teil {}/{}) - bitte etwas Geduld …"
            ).format(c_idx + 1, len(chunks)))
        offset = c_idx * _BK_LMX_CHUNK_LINES
        user = template.format("\n".join(
            f"{offset + i:03d}: {line}" for i, line in enumerate(chunk)))
        # Fortsetzungs-Kontext: Ohne ihn wusste das Modell ab Teil 2 nicht,
        # dass ein Register fortgesetzt wird, und lieferte jede Zeile als
        # Einzelabsatz - genau die "kaputte zweite Haelfte". Die Spalten der
        # zuletzt erzeugten Tabelle werden deshalb explizit vorgegeben.
        if c_idx > 0:
            prev_table = next((b for b in reversed(merged) if b.get("type") == "table"), None)
            if prev_table is not None and prev_table.get("rows"):
                width = len(prev_table["rows"][0])
                header = prev_table.get("header") or []
                header_hint = " | ".join(str(h) for h in header if str(h).strip()) or \
                    _bk_lmx_tr(self, "lmx_hint_no_header", "wie zuvor")
                user += "\n\n" + _bk_lmx_tr(
                    self, "lmx_prompt_continuation",
                    "FORTSETZUNG: Dies ist Teil {} desselben Dokuments. Gleichartige "
                    "Registerzeilen gehören weiterhin in EINE Tabelle mit exakt {} "
                    "Spalten ({}). Gib KEIN header-Feld erneut aus und mache aus "
                    "Registerzeilen keine Absätze."
                ).format(c_idx + 1, width, header_hint)
        content = _bk_lmx_chat(self, system, user, cancel_check=cancel_check)
        data = _bk_lmx_extract_json(self, content)
        blocks = _bk_lmx_normalize_blocks(data)
        # Rueckfall-Erkennung: Liefert eine Folgeportion NUR Absaetze, obwohl
        # zuvor eine Tabelle lief und die Absaetze wie Datenzeilen aussehen
        # (Komma + Ziffer), einmal mit verschaerfter Anweisung wiederholen.
        if c_idx > 0 and blocks and not any(b.get("type") == "table" for b in blocks):
            paras = [b for b in blocks if b.get("type") == "paragraph"]
            datalike = [b for b in paras if re.search(r"\d", b.get("text", ""))
                        and "," in b.get("text", "")]
            prev_table_exists = any(b.get("type") == "table" for b in merged)
            if prev_table_exists and paras and len(datalike) / len(paras) >= 0.6:
                strict = user + "\n\n" + _bk_lmx_tr(
                    self, "lmx_prompt_strict_table",
                    "WICHTIG: Deine letzte Antwort bestand nur aus Absätzen. Das ist "
                    "falsch. Alle Registerzeilen MÜSSEN als rows einer table "
                    "ausgegeben werden, keine paragraph-Blöcke für Registerzeilen.")
                content = _bk_lmx_chat(self, system, strict, cancel_check=cancel_check)
                data = _bk_lmx_extract_json(self, content)
                retry_blocks = _bk_lmx_normalize_blocks(data)
                if any(b.get("type") == "table" for b in retry_blocks):
                    blocks = retry_blocks
        for block in blocks:
            prev = merged[-1] if merged else None
            if (prev is not None and block["type"] == "table" and prev.get("type") == "table"
                    and prev.get("rows") and block.get("rows")
                    and len(prev["rows"][0]) == len(block["rows"][0])
                    and not block.get("header")):
                prev["rows"].extend(block["rows"])
                continue
            merged.append(block)
    return merged


class _BKLMXExportWorker(QThread):
    finished_export = Signal(str)
    failed_export = Signal(str)
    status_changed = Signal(str)

    def __init__(self, window, lines: List[str], path: str, fmt: str):
        super().__init__(window)
        self._window = window
        self._lines = list(lines)
        self._path = str(path)
        self._fmt = str(fmt)
        self._cancelled = False
        # LM-Verbindungsdaten des Fensters spiegeln, damit die Chat-Helfer
        # aus dem Worker heraus identisch funktionieren.
        for attr in ("ai_endpoint", "lm_endpoint", "ai_model", "lm_model", "_tr"):
            try:
                setattr(self, attr, getattr(window, attr))
            except Exception:
                pass

    def cancel(self):
        self._cancelled = True
        try:
            self.requestInterruption()
        except Exception:
            pass

    def _is_cancelled(self):
        return self._cancelled or self.isInterruptionRequested()

    def run(self):
        try:
            blocks = _bk_lmx_structured_blocks(
                self._window, self._lines,
                cancel_check=self._is_cancelled,
                status_cb=lambda text: self.status_changed.emit(text),
                fmt=self._fmt)
            if self._is_cancelled():
                raise RuntimeError(_bk_lmx_tr(self._window, "msg_ai_cancelled", "KI-Vorgang abgebrochen."))
            if not blocks:
                raise RuntimeError(_bk_lmx_tr(self._window, "lmx_err_no_blocks",
                                              "Die KI-Antwort enthält kein gültiges 'blocks'-Feld."))
            self.status_changed.emit(_bk_lmx_tr(
                self._window, "lmx_status_writing", "Dokument wird geschrieben …"))
            _bk_lmx_write(self._window, self._path, self._fmt, blocks)
            self.finished_export.emit(self._path)
        except Exception as exc:
            self.failed_export.emit(str(exc))


def _bk_lmx_export(self, fmt: str) -> None:
    item = _bk_lmx_current_item(self)
    lines = _bk_lmx_collect_lines(item) if item is not None else []
    try:
        self._bk_lmx_source_layout = getattr(item, "_bk_lmx_source_layout", None) if item is not None else None
    except Exception:
        self._bk_lmx_source_layout = None
    if not lines:
        QMessageBox.information(
            self,
            _bk_lmx_tr(self, "info_title", "Hinweis"),
            _bk_lmx_tr(self, "lmx_err_no_lines",
                       "Keine erkannten Zeilen vorhanden. Bitte zuerst OCR ausführen."))
        return
    base = os.path.splitext(str(getattr(item, "display_name", "export")))[0]
    suggested = os.path.join(os.path.expanduser("~"), f"{base}_ki.{fmt}")
    path, _ = QFileDialog.getSaveFileName(
        self,
        _bk_lmx_tr(self, "lmx_save_title", "KI-Export speichern"),
        suggested,
        f"{fmt.upper()} (*.{fmt})")
    if not path:
        return
    if not path.lower().endswith(f".{fmt}"):
        path += f".{fmt}"

    # Der Export laeuft in einem Worker-Thread; ein Dialog mit Kreis-Lade-
    # Animation informiert, dass die lokale KI arbeitet. Die GUI bleibt
    # bedienbar und faellt nicht mehr in den "Reagiert nicht"-Zustand.
    dialog = ProgressStatusDialog(
        _bk_lmx_tr(self, "lmx_busy_title", "KI-Export"), self._tr, self)
    dialog.set_status(_bk_lmx_tr(
        self, "lmx_busy_message",
        "Die lokale KI erstellt gerade das Dokument. Das kann je nach Modell "
        "einige Minuten dauern - bitte etwas Geduld …"))
    worker = _BKLMXExportWorker(self, lines, path, fmt)
    self._bk_lmx_worker = worker

    def _cleanup():
        try:
            dialog.close()
        except Exception:
            pass
        try:
            self._bk_lmx_worker = None
        except Exception:
            pass

    def _on_finished(saved_path):
        _cleanup()
        try:
            self.status_bar.showMessage(_bk_lmx_tr(
                self, "msg_ai_export_done", "KI-Export abgeschlossen: {}").format(saved_path), 6000)
        except Exception:
            pass

    def _on_failed(message):
        _cleanup()
        text = str(message or "")
        cancelled_text = _bk_lmx_tr(self, "msg_ai_cancelled", "KI-Vorgang abgebrochen.")
        if cancelled_text and cancelled_text in text:
            try:
                self.status_bar.showMessage(cancelled_text, 4000)
            except Exception:
                pass
            return
        QMessageBox.warning(
            self,
            _bk_lmx_tr(self, "warn_title", "Warnung"),
            _bk_lmx_tr(self, "msg_ai_export_failed",
                       "KI-Export fehlgeschlagen:\n{}").format(text))

    worker.status_changed.connect(dialog.set_status)
    worker.finished_export.connect(_on_finished)
    worker.failed_export.connect(_on_failed)
    dialog.cancel_requested.connect(worker.cancel)
    dialog.show()
    worker.start()


# ---------------------------------------------------------------------------
# Menü-Verdrahtung
# ---------------------------------------------------------------------------

_BK_LMX_FORMAT_LABELS = {
    "docx": "Word (.docx)",
    "xlsx": "Excel (.xlsx)",
    "odt": "LibreOffice Writer (.odt)",
    "ods": "LibreOffice Calc (.ods)",
}


def _bk_lmx_menu_texts(self) -> None:
    menu = getattr(self, "_bk_lmx_menu", None)
    if menu is not None:
        menu.setTitle(_bk_lmx_tr(self, "menu_ai_export", "KI-Export (lokale KI)"))
    actions = getattr(self, "_bk_lmx_actions", {}) or {}
    for fmt, action in actions.items():
        action.setText(_bk_lmx_tr(self, f"menu_ai_export_{fmt}",
                                  _BK_LMX_FORMAT_LABELS.get(fmt, fmt.upper())))


def _bk_lmx_install_menu(self) -> None:
    export_menu = getattr(self, "export_menu", None)
    if export_menu is None or getattr(self, "_bk_lmx_menu", None) is not None:
        return
    menu = QMenu(export_menu)
    self._bk_lmx_menu = menu
    self._bk_lmx_actions = {}
    for fmt in _BK_LMX_FORMATS:
        action = QAction(menu)
        action.triggered.connect(lambda _checked=False, f=fmt: _bk_lmx_export(self, f))
        menu.addAction(action)
        self._bk_lmx_actions[fmt] = action
    export_menu.addSeparator()
    export_menu.addMenu(menu)
    _bk_lmx_menu_texts(self)


def _bk_lmx_init(self, *args, **kwargs) -> None:
    try:
        _bk_lmx_install_menu(self)
    except Exception:
        pass


def _bk_lmx_retranslate(self, *args, **kwargs) -> None:
    try:
        _bk_lmx_menu_texts(self)
    except Exception:
        pass


register_init_delta(_bk_lmx_init)
register_retranslate_delta(_bk_lmx_retranslate)

MainWindow._bk_lmx_export = _bk_lmx_export

__all__ = [
    "_bk_lmx_export",
    "_bk_lmx_chat",
    "_bk_lmx_prompt",
    "_bk_lmx_extract_json",
    "_bk_lmx_normalize_blocks",
    "_bk_lmx_write_docx",
]

register_globals("bk", globals(), __all__)
