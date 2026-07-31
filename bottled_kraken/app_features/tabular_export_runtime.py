from bottled_kraken.module_registry import register_globals, seed_globals
from bottled_kraken.common.chain_consolidation import register_render_handler, RENDER_NOT_HANDLED
seed_globals('bk', globals())
from bottled_kraken.common import (
    QFileDialog,
    QDialog,
    QDialogButtonBox,
    QLabel,
    QMessageBox,
    QRadioButton,
    QCheckBox,
    QHBoxLayout,
    QPushButton,
    QVBoxLayout,
    QApplication,
    TaskItem,
    _load_image_color,
    _page_to_data_url,
    csv,
    html,
    json,
    math,
    os,
    re,
)
import zipfile
import threading
import time
from bottled_kraken.main_window import MainWindow
from bottled_kraken.dialogs import BusyStatusDialog
try:
    from PySide6.QtCore import QEventLoop as _BK_QEventLoop
except Exception:
    _BK_QEventLoop = None
try:
    from PySide6.QtWidgets import QGridLayout, QGroupBox
except Exception:
    QGridLayout = None
    QGroupBox = None
from bottled_kraken.export_page_setup import (
    bk_get_export_orientation,
    bk_set_export_orientation,
    bk_resolve_landscape,
    bk_resolve_portrait,
    bk_page_size_cm,
    bk_page_size_inches,
    bk_xlsx_page_setup_xml,
    bk_odf_orientation_name,
)
from bottled_kraken.export_layout import (
    _clean_text,
    _content_bounds,
    _median_height,
    _group_rows,
    _page_size,
    _records_from_views,
    write_positioned_docx,
    write_positioned_odt,
    write_spatial_txt,
)

_BK_TABULAR_COLUMNS = [
    ("family_name", "export_column_family_name", "Familienname", 18.0, 2.8),
    ("given_names", "export_column_given_names", "Vorname(n)", 22.0, 3.4),
    ("relationship", "export_column_relationship", "Zusatz/Beziehung", 26.0, 4.2),
    ("age_original", "export_column_age", "Alter", 12.0, 2.2),
    ("date_original", "export_column_birth_date", "Geburtsdatum", 16.0, 2.8),
    ("year_resolved", "export_column_year", "Jahr", 12.0, 2.0),
    ("place_in_source", "export_column_place", "Ort", 20.0, 3.2),
    ("number", "export_column_number", "Nr. / Seitennr.", 14.0, 2.4),
    ("original_line", "export_column_original_line", "Originalzeile", 70.0, 9.5),
]
_BK_TABULAR_KEYS = [column[0] for column in _BK_TABULAR_COLUMNS]
_BK_TABULAR_DEFAULT_KEYS = list(_BK_TABULAR_KEYS)
_BK_TABULAR_HEADERS = [column[2] for column in _BK_TABULAR_COLUMNS]
_BK_TABULAR_COLUMN_BY_KEY = {column[0]: column for column in _BK_TABULAR_COLUMNS}
_BK_TEXT_LAYOUT_FMTS = {"txt", "text", "txt_plain", "docx", "word", "odt"}
_BK_TABLE_EXPORT_FMTS = {"csv", "json", "xlsx", "excel", "ods", "calc"}
_BK_DITTO_VALUES = {'"', "'", "„", "“", "”", "-\"-", "-„-", "-=-"}


def _bk_screen_available_geometry(widget=None):
    try:
        app = QApplication.instance()
    except Exception:
        app = None
    screen = None
    # Zuerst den Bildschirm des Elternfensters verwenden. Vor dem ersten
    # show() hat ein Dialog noch kein windowHandle und mapToGlobal liefert
    # Koordinaten nahe (0,0) - auf Multi-Monitor-Systemen wurde der Dialog
    # dadurch auf dem falschen (z. B. zweiten) Bildschirm platziert.
    try:
        parent = widget.parentWidget() if widget is not None and hasattr(widget, "parentWidget") else None
        if parent is not None:
            top = parent.window() if hasattr(parent, "window") else parent
            handle = top.windowHandle() if hasattr(top, "windowHandle") else None
            if handle is not None and handle.screen() is not None:
                screen = handle.screen()
            elif hasattr(QApplication, "screenAt"):
                screen = QApplication.screenAt(top.mapToGlobal(top.rect().center()))
    except Exception:
        screen = None
    try:
        if screen is None and widget is not None and hasattr(widget, "windowHandle"):
            handle = widget.windowHandle()
            if handle is not None:
                screen = handle.screen()
    except Exception:
        screen = None
    if screen is None:
        try:
            if widget is not None and hasattr(QApplication, "screenAt"):
                screen = QApplication.screenAt(widget.mapToGlobal(widget.rect().center()))
        except Exception:
            screen = None
    if screen is None:
        try:
            screen = app.primaryScreen() if app is not None else QApplication.primaryScreen()
        except Exception:
            screen = None
    try:
        return screen.availableGeometry() if screen is not None else None
    except Exception:
        return None


def _bk_keep_dialog_inside_screen(dialog, margin=12):
    geom = _bk_screen_available_geometry(dialog)
    if geom is None:
        return
    # Solange der Dialog noch nicht sichtbar war, wird er ueber dem
    # Elternfenster zentriert. Ohne diese Zentrierung landete er auf
    # Multi-Monitor-Systemen auf dem falschen Bildschirm.
    try:
        if not dialog.isVisible():
            parent = dialog.parentWidget()
            if parent is not None:
                top = parent.window()
                center = top.frameGeometry().center()
                size = dialog.frameGeometry().size()
                if size.width() <= 0 or size.height() <= 0:
                    size = dialog.sizeHint()
                dialog.move(int(center.x() - size.width() / 2), int(center.y() - size.height() / 2))
    except Exception:
        pass
    try:
        frame = dialog.frameGeometry()
        if frame.width() <= 0 or frame.height() <= 0:
            frame = dialog.geometry()
        left = geom.left() + int(margin)
        top = geom.top() + int(margin)
        right = geom.right() - int(margin)
        bottom = geom.bottom() - int(margin)
        x = frame.x()
        y = frame.y()
        if frame.width() >= max(1, right - left + 1):
            x = left
        else:
            if frame.left() < left:
                x += left - frame.left()
            if frame.right() > right:
                x -= frame.right() - right
        if frame.height() >= max(1, bottom - top + 1):
            y = top
        else:
            if frame.top() < top:
                y += top - frame.top()
            if frame.bottom() > bottom:
                y -= frame.bottom() - bottom
        dialog.move(max(left, int(x)), max(top, int(y)))
    except Exception:
        pass


def _bk_dialog_content_min_size(dialog):
    """Tatsaechlichen Mindestplatz des Dialoginhalts ermitteln.

    Frueher wurde der Dialog hart auf feste Groessen (z. B. 560x300)
    geklemmt. Sobald zusaetzliche Gruppen (etwa das Seitenformat)
    hinzukamen, passte der Inhalt nicht mehr und Qt quetschte die
    Darstellungs-Radiobuttons bis zur Unsichtbarkeit zusammen. Deshalb
    duerfen Minimal-/Maximalgroessen nie unter den Layout-Mindestbedarf
    fallen.
    """
    w = h = 0
    try:
        lay = dialog.layout()
        if lay is not None:
            ms = lay.totalMinimumSize()
            w = max(w, int(ms.width()))
            h = max(h, int(ms.height()))
            sh = lay.totalSizeHint()
            h = max(h, int(sh.height()))
            w = max(w, int(sh.width()))
    except Exception:
        pass
    try:
        sh = dialog.sizeHint()
        w = max(w, int(sh.width()))
        h = max(h, int(sh.height()))
    except Exception:
        pass
    return w, h


def _bk_resize_export_dialog_for_mode(dialog, detailed=False, txt_only=False):
    geom = _bk_screen_available_geometry(dialog)
    if txt_only or not detailed:
        target_w, target_h = (520, 220) if txt_only else (560, 300)
        need_w, need_h = _bk_dialog_content_min_size(dialog)
        target_w = max(target_w, need_w + 12)
        target_h = max(target_h, need_h + 12)
        if geom is not None:
            try:
                target_w = min(target_w, int(geom.width()) - 32)
                target_h = min(target_h, int(geom.height()) - 72)
            except Exception:
                pass
        try:
            dialog.setMinimumSize(target_w, target_h)
            dialog.setMaximumSize(16777215, 16777215)
            dialog.resize(target_w, target_h)
        except Exception:
            pass
        _bk_keep_dialog_inside_screen(dialog)
        return
    margin = 32
    max_w, max_h = 840, 720
    if geom is not None:
        try:
            # Die Höhe wird konservativer begrenzt, weil frameGeometry zusätzlich
            # Titelleiste/Fensterrahmen enthält. So bleibt der Dialog nach dem
            # Umschalten auf „Tabelle (erweitert)“ vollständig im sichtbaren Bereich.
            max_w = max(520, min(max_w, int(geom.width()) - margin))
            max_h = max(430, min(max_h, int(geom.height()) - 72))
        except Exception:
            pass
    need_w, need_h = _bk_dialog_content_min_size(dialog)
    max_w = max(max_w, min(need_w + 12, int(geom.width()) - 32) if geom is not None else need_w + 12)
    max_h = max(max_h, min(need_h + 12, int(geom.height()) - 72) if geom is not None else need_h + 12)
    min_w = min(max(760, need_w + 12), max_w)
    min_h = min(max(620, need_h + 12), max_h)
    try:
        dialog.setMinimumSize(min_w, min_h)
        dialog.setMaximumSize(16777215, 16777215)
        dialog.resize(max_w, max_h)
    except Exception:
        pass
    _bk_keep_dialog_inside_screen(dialog)


def _bk_registry_lookup(key):
    """Schluessel direkt in den Sprachdateien nachschlagen (ohne Fenster).

    Reihenfolge: Standardsprache, danach alle verfuegbaren Sprachen. Gibt
    "" zurueck, wenn der Schluessel in keiner Sprachdatei existiert.
    """
    try:
        from bottled_kraken.translation import translation as _bk_translation
    except Exception:
        return ""
    candidates = []
    try:
        default_lang = getattr(_bk_translation, "DEFAULT_LANGUAGE", None)
        if default_lang:
            candidates.append(default_lang)
    except Exception:
        pass
    try:
        candidates.extend(list(_bk_translation.available_languages()))
    except Exception:
        candidates.extend(["de", "en", "fr"])
    for lang in candidates:
        if not lang:
            continue
        try:
            value = _bk_translation.translate(lang, key)
        except Exception:
            continue
        if value and value != key:
            return value
    return ""


def _bk_tab_tr(window, key, fallback=""):
    try:
        value = window._tr(key)
        if value and value != key:
            return value
    except Exception:
        pass
    # Fehlt der Schluessel in der aktiven Sprache, zuerst die uebrigen
    # Sprachdateien befragen. Die fest im Code stehenden Fallback-Texte
    # kommen nur noch zum Zug, wenn ein Schluessel in KEINER Sprachdatei
    # existiert - dadurch erscheinen keine deutschen Texte mehr in der
    # englischen oder franzoesischen Oberflaeche.
    registry = _bk_registry_lookup(key)
    if registry:
        return registry
    return fallback or key


def _bk_tr_registry(window, key):
    """Uebersetzung ausschliesslich aus den Sprachdateien beziehen.

    Erst die aktive Sprache des Fensters, danach die uebrigen
    Sprachregistraturen. Es gibt bewusst KEINE fest im Quellcode
    hinterlegten Sprachtexte; schlimmstenfalls wird der Schluessel selbst
    zurueckgegeben.
    """
    try:
        value = window._tr(key)
        if value and value != key:
            return value
    except Exception:
        pass
    registry = _bk_registry_lookup(key)
    if registry:
        return registry
    return key


def _bk_add_export_orientation_group(self, dlg, layout):
    """Auswahlgruppe Hoch-/Querformat in einen Export-Dialog einfuegen.

    Die Gruppe persistiert die Auswahl selbst, sobald der Dialog mit OK
    bestaetigt wird (dlg.accepted). Dadurch funktioniert sie in allen
    Dialog-Generationen, ohne deren accept-Handler anzupassen. Bei
    Abbrechen bleibt die bisherige Einstellung unveraendert.
    """
    try:
        current = bk_get_export_orientation(self)
    except Exception:
        current = "auto"
    box = QGroupBox(_bk_tr_registry(self, "export_orientation_group"), dlg) if QGroupBox is not None else None
    row = QHBoxLayout(box) if box is not None else QHBoxLayout()
    row.setContentsMargins(10, 8, 10, 8)
    row.setSpacing(14)
    rb_auto = QRadioButton(_bk_tr_registry(self, "export_orientation_auto"), dlg)
    rb_portrait = QRadioButton(_bk_tr_registry(self, "export_orientation_portrait"), dlg)
    rb_landscape = QRadioButton(_bk_tr_registry(self, "export_orientation_landscape"), dlg)
    # Eigene Button-Gruppe: verhindert in jeder Dialog-Generation, dass die
    # Radiobuttons mit den Darstellungs-Radiobuttons exklusiv gekoppelt werden.
    try:
        from PySide6.QtWidgets import QButtonGroup as _BKQButtonGroup
        group = _BKQButtonGroup(dlg)
        group.setExclusive(True)
        for rb in (rb_auto, rb_portrait, rb_landscape):
            group.addButton(rb)
        dlg._bk_orientation_button_group = group
    except Exception:
        pass
    rb_auto.setChecked(current not in {"portrait", "landscape"})
    rb_portrait.setChecked(current == "portrait")
    rb_landscape.setChecked(current == "landscape")
    row.addWidget(rb_auto); row.addWidget(rb_portrait); row.addWidget(rb_landscape); row.addStretch(1)
    if box is not None:
        layout.addWidget(box)
    else:
        layout.addLayout(row)
    radios = {"auto": rb_auto, "portrait": rb_portrait, "landscape": rb_landscape, "box": box}

    def _persist_orientation():
        try:
            bk_set_export_orientation(_bk_orientation_from_radios(radios), self)
        except Exception:
            pass

    try:
        dlg.accepted.connect(_persist_orientation)
    except Exception:
        pass
    return radios


def _bk_orientation_from_radios(radios) -> str:
    try:
        if radios and radios.get("landscape") is not None and radios["landscape"].isChecked():
            return "landscape"
        if radios and radios.get("portrait") is not None and radios["portrait"].isChecked():
            return "portrait"
    except Exception:
        pass
    return "auto"


def _bk_tab_clean(value) -> str:
    return _clean_text(value)


def _bk_tab_is_ditto(value) -> bool:
    txt = _bk_tab_clean(value)
    return txt in _BK_DITTO_VALUES or txt.replace(" ", "") in _BK_DITTO_VALUES


def _bk_tab_is_separator(text: str) -> bool:
    txt = _bk_tab_clean(text)
    if not txt:
        return True
    if re.fullmatch(r"[-–—=•·*+<>_/\\|~\s]+", txt):
        return True
    if re.fullmatch(r"[-–—=]{1,4}[^A-Za-zÀ-ÿ0-9]{0,3}[-–—=]{1,4}", txt):
        return True
    return False


def _bk_tab_year_value(text: str) -> str:
    txt = _bk_tab_clean(text)
    if not txt:
        return ""
    # Jahresüberschriften stehen oft allein, aber mit Punkt, Strichen oder Schmuckzeichen.
    m = re.fullmatch(r"\s*[-–—=•·*]*\s*(1[5-9]\d{2}|20\d{2})\s*[.)]?\s*[-–—=•·*]*\s*", txt)
    return m.group(1) if m else ""


def _bk_tab_number_value(text: str) -> str:
    txt = _bk_tab_clean(text)
    if not txt or _bk_tab_year_value(txt):
        return ""
    m = re.fullmatch(r"\s*(\d{1,4})\s*[.)]?\s*", txt)
    return m.group(1) if m else ""


def _bk_tab_is_standalone_year(text: str) -> bool:
    return bool(_bk_tab_year_value(text))


def _bk_tab_is_standalone_number(text: str) -> bool:
    return bool(_bk_tab_number_value(text))


def _bk_tab_has_letters(text: str) -> bool:
    return bool(re.search(r"[A-Za-zÀ-ÿÄÖÜäöüß]", str(text or "")))


def _bk_tab_join_text_fragments(parts) -> str:
    text = " ".join(_bk_tab_clean(part) for part in (parts or []) if _bk_tab_clean(part))
    text = re.sub(r"\s+([,.;:!?])", r"\1", text)
    text = re.sub(r"([(])\s+", r"\1", text)
    text = re.sub(r"\s+([)])", r"\1", text)
    text = re.sub(r"\s+", " ", text).strip()
    return text


def _bk_tab_merge_visual_cluster(items):
    clean_items = [dict(item) for item in (items or []) if _bk_tab_clean(item.get("text", ""))]
    if not clean_items:
        return None
    clean_items.sort(key=lambda item: (float(item.get("x0", 0.0) or 0.0), int(item.get("index", 0) or 0)))
    if len(clean_items) == 1:
        return clean_items[0]
    texts = [item.get("text", "") for item in clean_items]
    bbox_items = [item for item in clean_items if item.get("bbox")]
    merged = dict(clean_items[0])
    merged["text"] = _bk_tab_join_text_fragments(texts)
    merged["index"] = min(int(item.get("index", 0) or 0) for item in clean_items)
    if bbox_items:
        x0 = min(float(item.get("x0", 0.0) or 0.0) for item in bbox_items)
        y0 = min(float(item.get("y0", 0.0) or 0.0) for item in bbox_items)
        x1 = max(float(item.get("x1", 0.0) or 0.0) for item in bbox_items)
        y1 = max(float(item.get("y1", 0.0) or 0.0) for item in bbox_items)
        merged.update({
            "bbox": (x0, y0, x1, y1),
            "x0": x0,
            "y0": y0,
            "x1": x1,
            "y1": y1,
            "cx": (x0 + x1) / 2.0,
            "cy": (y0 + y1) / 2.0,
            "w": x1 - x0,
            "h": y1 - y0,
        })
    return merged


def _bk_tab_merge_visual_row_records(records, page_width=0.0):
    records = [dict(record) for record in (records or []) if _bk_tab_clean(record.get("text", ""))]
    if not records:
        return []
    boxed_count = sum(1 for record in records if record.get("bbox"))
    if boxed_count < 2:
        return records
    median_height = _median_height(records)
    y_tolerance = max(3.0, min(28.0, median_height * 0.70))
    ordered = sorted(records, key=lambda record: (
        float(record.get("cy", record.get("y0", 0.0)) or 0.0),
        float(record.get("x0", 0.0) or 0.0),
        int(record.get("index", 0) or 0),
    ))
    visual_rows = []
    for record in ordered:
        cy = float(record.get("cy", record.get("y0", 0.0)) or 0.0)
        placed = False
        for row in visual_rows:
            if abs(cy - row["cy"]) <= y_tolerance:
                row["items"].append(record)
                row["cy"] = sum(float(item.get("cy", item.get("y0", 0.0)) or 0.0) for item in row["items"]) / max(1, len(row["items"]))
                placed = True
                break
        if not placed:
            visual_rows.append({"cy": cy, "items": [record]})
    split_gap = max(110.0, min(max(1.0, float(page_width or 0.0)) * 0.105, 260.0), median_height * 11.0)
    merged = []
    for row in visual_rows:
        items = sorted(row["items"], key=lambda item: (float(item.get("x0", 0.0) or 0.0), int(item.get("index", 0) or 0)))
        cluster = []
        prev_x1 = None
        for item in items:
            x0 = float(item.get("x0", 0.0) or 0.0)
            if cluster and prev_x1 is not None and (x0 - prev_x1) > split_gap:
                merged_item = _bk_tab_merge_visual_cluster(cluster)
                if merged_item:
                    merged.append(merged_item)
                cluster = []
            cluster.append(item)
            prev_x1 = max(float(item.get("x1", x0) or x0), prev_x1 if prev_x1 is not None else x0)
        merged_item = _bk_tab_merge_visual_cluster(cluster)
        if merged_item:
            merged.append(merged_item)
    merged.sort(key=lambda item: (float(item.get("y0", 0.0) or 0.0), float(item.get("x0", 0.0) or 0.0), int(item.get("index", 0) or 0)))
    return merged


def _bk_tab_expand_numeric_records(records):
    expanded = []
    med_h = _median_height(records) if records else 12.0
    for record in records or []:
        text = _bk_tab_clean(record.get("text", ""))
        if not text:
            continue
        parts = re.findall(r"\d{1,4}", text)
        if (
            record.get("bbox")
            and len(parts) >= 2
            and re.fullmatch(r"\s*\d{1,4}\.?(?:\s+\d{1,4}\.?)+\s*", text)
            and float(record.get("h", 0.0) or 0.0) > med_h * 1.45
        ):
            y0 = float(record.get("y0", 0.0) or 0.0)
            h = float(record.get("h", med_h) or med_h)
            step = h / max(1, len(parts))
            for idx, part in enumerate(parts):
                clone = dict(record)
                clone["text"] = part
                clone["y0"] = y0 + idx * step
                clone["y1"] = y0 + (idx + 1) * step
                clone["cy"] = (clone["y0"] + clone["y1"]) / 2.0
                clone["h"] = step
                expanded.append(clone)
        else:
            expanded.append(record)
    expanded.sort(key=lambda item: (float(item.get("y0", 0.0) or 0.0), float(item.get("x0", 0.0) or 0.0), int(item.get("index", 0) or 0)))
    return expanded


def _bk_tab_extract_age(text: str) -> str:
    m = re.search(r"\b\d{1,3}\s*(?:Jahre?|Jahr|J\.|Wochen?|W\.|Tage?|Monate?|Mon\.|Years?|Months?|Days?)\b", text, flags=re.IGNORECASE)
    return _bk_tab_clean(m.group(0)) if m else ""


def _bk_tab_extract_date(text: str) -> str:
    m = re.search(r"\b\d{1,2}\s*\.\s*(?:[IVXLCDM]{1,8}|\d{1,2})\s*\.?(?!\d)", text, flags=re.IGNORECASE)
    if not m:
        return ""
    value = re.sub(r"\s+", "", m.group(0))
    if not value.endswith("."):
        value += "."
    return value


def _bk_tab_extract_year_in_source(text: str) -> str:
    direct = _bk_tab_year_value(text)
    if direct:
        return direct
    m = re.search(r"\b(1[5-9]\d{2}|20\d{2})\b", text)
    if m:
        return m.group(1)
    if re.search(r"[„“”]\s*[,.;:]?", text):
        return "„"
    return ""


def _bk_tab_extract_number(text: str) -> str:
    txt = str(text or "")
    candidates = []
    age_match = re.search(r"\b(\d{1,3})\s*(?:Jahre?|Jahr|J\.|Wochen?|W\.|Tage?|Monate?|Mon\.|Years?|Months?|Days?)\b", txt, flags=re.IGNORECASE)
    age_span = age_match.span(1) if age_match else None
    for match in re.finditer(r"(?<![A-Za-zÀ-ÿ])\d{1,4}(?![A-Za-zÀ-ÿ])", txt):
        value = match.group(0)
        if len(value) == 4 and re.fullmatch(r"1[5-9]\d{2}|20\d{2}", value):
            continue
        if age_span and match.span() == age_span:
            continue
        candidates.append(value)
    return candidates[-1] if candidates else ""


def _bk_tab_extract_relationship(text: str) -> str:
    m = re.search(r"\(([^)]{1,120})\)", text)
    if m:
        return _bk_tab_clean(m.group(1))
    before_age = re.split(r"\b\d{1,3}\s*(?:Jahre?|Jahr|J\.|Wochen?|W\.|Tage?|Monate?|Mon\.)\b", text, maxsplit=1, flags=re.IGNORECASE)[0]
    if "," in before_age:
        rel = before_age.split(",", 1)[1]
        rel = rel.strip(" ,.;:-()")
        if re.search(r"\b(?:S\.?d\.?|T\.?d\.?|Weib|Wtwe|Witwe|Sohn|Tochter|Frau|Mann)\b", rel, flags=re.IGNORECASE):
            return _bk_tab_clean(rel)
    return ""


def _bk_tab_name_source_part(text: str) -> str:
    txt = _bk_tab_clean(text)
    if not txt:
        return ""
    txt = re.sub(r"\([^)]{0,140}\)", " ", txt)
    cut_points = []
    patterns = [
        r"\b\d{1,3}\s*(?:Jahre?|Jahr|J\.|Wochen?|W\.|Tage?|Monate?|Mon\.)\b",
        r"\b\d{1,2}\s*\.\s*(?:[IVXLCDM]{1,8}|\d{1,2})\s*\.?"]
    for pattern in patterns:
        m = re.search(pattern, txt, flags=re.IGNORECASE)
        if m:
            cut_points.append(m.start())
    y = re.search(r"\b(?:1[5-9]\d{2}|20\d{2})\b", txt)
    if y:
        cut_points.append(y.start())
    nr = re.search(r"\s+\d{1,4}\s*\.?$", txt)
    if nr:
        cut_points.append(nr.start())
    if cut_points:
        txt = txt[: min(cut_points)]
    if "," in txt:
        txt = txt.split(",", 1)[0]
    txt = txt.strip(" ,.;:-")
    return _bk_tab_clean(txt)


def _bk_tab_split_name(text: str):
    source = _bk_tab_name_source_part(text)
    if not source:
        return "", "", ""
    source = re.sub(r"^\d+\s+", "", source).strip()
    parts = source.split()
    if not parts:
        return "", "", ""
    family = parts[0]
    given = " ".join(parts[1:])
    return family, given, source


def _bk_tab_extract_place(text: str, year_text: str, number_text: str) -> str:
    txt = _bk_tab_clean(text)
    if not txt:
        return ""
    tmp = txt
    age = _bk_tab_extract_age(tmp)
    if age:
        tmp = tmp.replace(age, " ", 1)
    date = _bk_tab_extract_date(tmp)
    if date:
        tmp = re.sub(re.escape(date).replace("\\ ", r"\s*"), " ", tmp, count=1)
    if year_text and year_text != "„":
        tmp = re.sub(r"\b" + re.escape(str(year_text)) + r"\b", " ", tmp, count=1)
    if number_text:
        tmp = re.sub(r"\b" + re.escape(str(number_text)) + r"\b\s*\.?\s*$", " ", tmp, count=1)
    name_part = _bk_tab_name_source_part(txt)
    if name_part:
        tmp = tmp.replace(name_part, " ", 1)
    tmp = re.sub(r"\([^)]*\)", " ", tmp)
    tmp = re.sub(r"\b\d{1,3}\s*(?:Jahre?|Jahr|J\.|Wochen?|W\.|Tage?|Monate?|Mon\.)\b", " ", tmp, flags=re.IGNORECASE)
    tmp = re.sub(r"\b\d{1,2}\s*\.\s*(?:[IVXLCDM]{1,8}|\d{1,2})\s*\.??", " ", tmp, flags=re.IGNORECASE)
    tmp = re.sub(r"[„“”]", " ", tmp)
    tokens = [t.strip(" ,.;:-()") for t in tmp.split()]
    tokens = [t for t in tokens if t and not re.fullmatch(r"\d{1,4}", t)]
    place = " ".join(tokens).strip(" ,.;:-")
    if len(place) > 80:
        place = ""
    return place


def _bk_tab_nearest_context_year(record, years, page_width, median_height):
    if not years:
        return ""
    rx = float(record.get("cx", record.get("x0", 0.0)) or 0.0)
    ry = float(record.get("y0", 0.0) or 0.0)
    max_dx = max(page_width * 0.18, float(record.get("w", 0.0) or 0.0) * 1.3, 90.0)
    candidates = []
    for yrec in years:
        if float(yrec.get("y1", 0.0) or 0.0) > ry + median_height * 0.35:
            continue
        dx = abs(float(yrec.get("cx", yrec.get("x0", 0.0)) or 0.0) - rx)
        dy = ry - float(yrec.get("cy", yrec.get("y0", 0.0)) or 0.0)
        if dx <= max_dx and dy >= 0:
            candidates.append((dy, dx, yrec))
    if not candidates:
        # Fallback: latest year heading above the entry on the same half/third of the page.
        same_side = []
        for yrec in years:
            if float(yrec.get("y1", 0.0) or 0.0) > ry:
                continue
            dx = abs(float(yrec.get("cx", yrec.get("x0", 0.0)) or 0.0) - rx)
            dy = ry - float(yrec.get("cy", yrec.get("y0", 0.0)) or 0.0)
            if dx <= page_width * 0.32:
                same_side.append((dy, dx, yrec))
        candidates = same_side
    if not candidates:
        return ""
    candidates.sort(key=lambda item: (item[0], item[1]))
    return _bk_tab_year_value(candidates[0][2].get("text", "")) or _bk_tab_clean(candidates[0][2].get("text", ""))


def _bk_tab_matching_number(record, numbers, page_width, median_height):
    if not numbers:
        return ""
    ry = float(record.get("cy", record.get("y0", 0.0)) or 0.0)
    x1 = float(record.get("x1", record.get("x0", 0.0)) or 0.0)
    x0 = float(record.get("x0", 0.0) or 0.0)
    max_dx = max(page_width * 0.18, 140.0)
    candidates = []
    for nrec in numbers:
        nx0 = float(nrec.get("x0", 0.0) or 0.0)
        ncy = float(nrec.get("cy", nrec.get("y0", 0.0)) or 0.0)
        if nx0 < x0 - 5:
            continue
        dy = abs(ncy - ry)
        dx = max(0.0, nx0 - x1)
        if dy <= max(3.0, median_height * 0.85) and dx <= max_dx:
            candidates.append((dy, dx, nrec))
    if not candidates:
        return ""
    candidates.sort(key=lambda item: (item[0], item[1]))
    best = candidates[0][2]
    return _bk_tab_number_value(best.get("text", "")) or _bk_tab_clean(best.get("text", ""))


def _bk_tab_make_row_from_record(record, context_year="", right_number=""):
    original = _bk_tab_clean(record.get("text", ""))
    family, given, full_name = _bk_tab_split_name(original)
    if not full_name:
        return None
    age = _bk_tab_extract_age(original)
    date = _bk_tab_extract_date(original)
    year_in_source = _bk_tab_extract_year_in_source(original) or context_year
    number = _bk_tab_extract_number(original) or right_number
    relationship = _bk_tab_extract_relationship(original)
    place = _bk_tab_extract_place(original, year_in_source, number)
    return {
        "family_name": family,
        "given_names": given,
        "relationship": relationship,
        "age_original": age,
        "date_original": date,
        "year_resolved": "",
        "year_in_source": year_in_source,
        "place_resolved": "",
        "place_in_source": place,
        "number": number,
        "original_line": original,
        "full_name": full_name,
        "_source_y": float(record.get("y0", 0.0) or 0.0),
        "_source_x": float(record.get("x0", 0.0) or 0.0),
    }


def _bk_build_transcription_rows(record_views, image_size=None):
    raw_records = _records_from_views(record_views)
    if not raw_records:
        return []
    page_width, _page_height = _page_size(image_size, raw_records)
    records = _bk_tab_merge_visual_row_records(raw_records, page_width)
    records = _bk_tab_expand_numeric_records(records)
    if not records:
        return []
    page_width, _page_height = _page_size(image_size, records)
    median_height = _median_height(records)
    year_records = [r for r in records if _bk_tab_is_standalone_year(r.get("text", ""))]
    number_records = [r for r in records if _bk_tab_is_standalone_number(r.get("text", ""))]
    rows = []
    seen = set()
    for record in records:
        text = _bk_tab_clean(record.get("text", ""))
        if not text or _bk_tab_is_separator(text):
            continue
        if not _bk_tab_has_letters(text):
            continue
        # Skip labels or one-letter group headings like "U" unless they clearly contain a name.
        if len(text) <= 2 and len(text.split()) <= 1:
            continue
        context_year = _bk_tab_nearest_context_year(record, year_records, page_width, median_height)
        number = _bk_tab_matching_number(record, number_records, page_width, median_height)
        row = _bk_tab_make_row_from_record(record, context_year, number)
        if not row:
            continue
        key = (row.get("family_name", "").lower(), row.get("given_names", "").lower(), row.get("number", ""), row.get("original_line", "").lower())
        if key in seen:
            continue
        seen.add(key)
        rows.append(row)
    rows.sort(key=lambda item: (float(item.get("_source_x", 0.0) or 0.0), float(item.get("_source_y", 0.0) or 0.0)))
    # Preserve reading order better for multi-column pages: first by major column cluster, then y.
    if rows:
        xs = sorted(float(row.get("_source_x", 0.0) or 0.0) for row in rows)
        clusters = []
        for x in xs:
            if not clusters or abs(x - clusters[-1]) > max(120.0, page_width * 0.08):
                clusters.append(x)
            else:
                clusters[-1] = (clusters[-1] + x) / 2.0
        def row_sort_key(row):
            x = float(row.get("_source_x", 0.0) or 0.0)
            cluster = min(range(len(clusters)), key=lambda idx: abs(clusters[idx] - x)) if clusters else 0
            return (cluster, float(row.get("_source_y", 0.0) or 0.0), x)
        rows.sort(key=row_sort_key)
    last_year = ""
    last_place = ""
    out = []
    for idx, row in enumerate(rows, start=1):
        row = dict(row)
        ysrc = _bk_tab_clean(row.get("year_in_source", ""))
        if _bk_tab_is_ditto(ysrc):
            row["year_resolved"] = last_year
            row["year_in_source"] = "„"
        else:
            row["year_resolved"] = ysrc
            if ysrc:
                last_year = ysrc
        psrc = _bk_tab_clean(row.get("place_in_source", ""))
        if _bk_tab_is_ditto(psrc):
            row["place_resolved"] = last_place
            row["place_in_source"] = "„"
        else:
            row["place_resolved"] = psrc
            if psrc:
                last_place = psrc
        row["id"] = f"entry_{idx:04d}"
        row.pop("_source_x", None)
        row.pop("_source_y", None)
        out.append(row)
    return out






def _bk_tabular_columns_for_export(window=None, column_keys=None):
    keys = _bk_normalize_column_keys(column_keys)
    return [(key, _bk_tabular_column_title(window, key)) for key in keys]


def _bk_table_matrix_from_rows(rows, column_keys=None, window=None):
    columns = _bk_tabular_columns_for_export(window, column_keys)
    matrix = [[title for _key, title in columns]]
    for row in rows or []:
        matrix.append([_bk_tab_clean(row.get(key, "")) for key, _title in columns])
    return matrix


def _bk_table_dicts_from_rows(rows, column_keys=None):
    keys = _bk_normalize_column_keys(column_keys)
    table_rows = []
    for row in rows or []:
        rid = row.get("id")
        out = {"id": rid} if rid else {}
        for key in keys:
            out[key] = _bk_tab_clean(row.get(key, ""))
        table_rows.append(out)
    return table_rows


def _bk_write_transcription_csv(path, rows, column_keys=None, window=None):
    with open(path, "w", encoding="utf-8-sig", newline="") as handle:
        writer = csv.writer(handle, delimiter=",", quotechar='"', quoting=csv.QUOTE_MINIMAL, lineterminator="\r\n")
        writer.writerows(_bk_table_matrix_from_rows(rows, column_keys, window))

def _bk_xlsx_col_name(index):
    name = ""
    index = int(index)
    while index > 0:
        index, rem = divmod(index - 1, 26)
        name = chr(65 + rem) + name
    return name


def _bk_xml(value):
    return html.escape(str(value or ""), quote=True)


def _bk_write_transcription_xlsx(path, rows, column_keys=None, window=None):
    keys = _bk_normalize_column_keys(column_keys)
    matrix = _bk_table_matrix_from_rows(rows, keys, window)
    col_widths = [float(_BK_TABULAR_COLUMN_BY_KEY.get(key, (None, None, None, 16.0, 3.0))[3]) for key in keys]
    row_xml = []
    for r_idx, row in enumerate(matrix, start=1):
        cells = []
        for c_idx, value in enumerate(row, start=1):
            ref = f"{_bk_xlsx_col_name(c_idx)}{r_idx}"
            style = "1" if r_idx == 1 else "2"
            cells.append(f'<c r="{ref}" s="{style}" t="inlineStr"><is><t xml:space="preserve">{_bk_xml(value)}</t></is></c>')
        height = 22 if r_idx == 1 else 18
        row_xml.append(f'<row r="{r_idx}" ht="{height}" customHeight="1">{"".join(cells)}</row>')
    cols = ''.join(f'<col min="{i}" max="{i}" width="{col_widths[i-1]:.2f}" customWidth="1"/>' for i in range(1, len(col_widths) + 1))
    last_col = _bk_xlsx_col_name(max(1, len(col_widths)))
    sheet = ''.join([
        '<?xml version="1.0" encoding="UTF-8"?>',
        '<worksheet xmlns="http://schemas.openxmlformats.org/spreadsheetml/2006/main" xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships">',
        '<sheetViews><sheetView workbookViewId="0"><pane ySplit="1" topLeftCell="A2" activePane="bottomLeft" state="frozen"/></sheetView></sheetViews>',
        '<sheetFormatPr defaultRowHeight="18"/>',
        '<cols>', cols, '</cols>',
        '<sheetData>', ''.join(row_xml), '</sheetData>',
        '<autoFilter ref="A1:%s%d"/>' % (last_col, max(1, len(matrix))),
        '<pageMargins left="0.25" right="0.25" top="0.25" bottom="0.25" header="0" footer="0"/>',
        bk_xlsx_page_setup_xml(True, window),
        '</worksheet>',
    ])
    content_types = ''.join([
        '<?xml version="1.0" encoding="UTF-8"?>',
        '<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">',
        '<Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>',
        '<Default Extension="xml" ContentType="application/xml"/>',
        '<Override PartName="/xl/workbook.xml" ContentType="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet.main+xml"/>',
        '<Override PartName="/xl/worksheets/sheet1.xml" ContentType="application/vnd.openxmlformats-officedocument.spreadsheetml.worksheet+xml"/>',
        '<Override PartName="/xl/styles.xml" ContentType="application/vnd.openxmlformats-officedocument.spreadsheetml.styles+xml"/>',
        '<Override PartName="/docProps/core.xml" ContentType="application/vnd.openxmlformats-package.core-properties+xml"/>',
        '<Override PartName="/docProps/app.xml" ContentType="application/vnd.openxmlformats-officedocument.extended-properties+xml"/>',
        '</Types>',
    ])
    root_rels = ''.join([
        '<?xml version="1.0" encoding="UTF-8"?>',
        '<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships"><Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument" Target="xl/workbook.xml"/></Relationships>',
    ])
    workbook = ''.join([
        '<?xml version="1.0" encoding="UTF-8"?>',
        '<workbook xmlns="http://schemas.openxmlformats.org/spreadsheetml/2006/main" xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships"><sheets><sheet name="Transkription" sheetId="1" r:id="rId1"/></sheets></workbook>',
    ])
    workbook_rels = '<?xml version="1.0" encoding="UTF-8"?><Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships"><Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/worksheet" Target="worksheets/sheet1.xml"/><Relationship Id="rId2" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/styles" Target="styles.xml"/></Relationships>'
    styles = ''.join([
        '<?xml version="1.0" encoding="UTF-8"?>',
        '<styleSheet xmlns="http://schemas.openxmlformats.org/spreadsheetml/2006/main">',
        '<fonts count="3"><font><sz val="11"/><name val="Calibri"/></font><font><b/><sz val="10"/><name val="Calibri"/></font><font><sz val="9"/><name val="Calibri"/></font></fonts>',
        '<fills count="3"><fill><patternFill patternType="none"/></fill><fill><patternFill patternType="gray125"/></fill><fill><patternFill patternType="solid"><fgColor rgb="FFE9EEF6"/><bgColor indexed="64"/></patternFill></fill></fills>',
        '<borders count="2"><border><left/><right/><top/><bottom/><diagonal/></border><border><left style="thin"><color rgb="FFB7B7B7"/></left><right style="thin"><color rgb="FFB7B7B7"/></right><top style="thin"><color rgb="FFB7B7B7"/></top><bottom style="thin"><color rgb="FFB7B7B7"/></bottom><diagonal/></border></borders>',
        '<cellStyleXfs count="1"><xf numFmtId="0" fontId="0" fillId="0" borderId="0"/></cellStyleXfs>',
        '<cellXfs count="3"><xf numFmtId="0" fontId="0" fillId="0" borderId="0" xfId="0"/><xf numFmtId="0" fontId="1" fillId="2" borderId="1" xfId="0" applyFont="1" applyFill="1" applyBorder="1" applyAlignment="1"><alignment horizontal="left" vertical="center" wrapText="1"/></xf><xf numFmtId="0" fontId="2" fillId="0" borderId="1" xfId="0" applyFont="1" applyBorder="1" applyAlignment="1"><alignment horizontal="left" vertical="top" wrapText="1"/></xf></cellXfs>',
        '<cellStyles count="1"><cellStyle name="Normal" xfId="0" builtinId="0"/></cellStyles>',
        '</styleSheet>',
    ])
    core = '<?xml version="1.0" encoding="UTF-8"?><cp:coreProperties xmlns:cp="http://schemas.openxmlformats.org/package/2006/metadata/core-properties" xmlns:dc="http://purl.org/dc/elements/1.1/"><dc:creator>Bottled Kraken</dc:creator></cp:coreProperties>'
    app = '<?xml version="1.0" encoding="UTF-8"?><Properties xmlns="http://schemas.openxmlformats.org/officeDocument/2006/extended-properties"><Application>Bottled Kraken</Application></Properties>'
    with zipfile.ZipFile(path, "w", zipfile.ZIP_DEFLATED) as archive:
        archive.writestr("[Content_Types].xml", content_types)
        archive.writestr("_rels/.rels", root_rels)
        archive.writestr("xl/workbook.xml", workbook)
        archive.writestr("xl/_rels/workbook.xml.rels", workbook_rels)
        archive.writestr("xl/styles.xml", styles)
        archive.writestr("xl/worksheets/sheet1.xml", sheet)
        archive.writestr("docProps/core.xml", core)
        archive.writestr("docProps/app.xml", app)

def _bk_ods_text(value):
    return html.escape(str(value or ""), quote=False)


def _bk_write_transcription_ods(path, rows, column_keys=None, window=None):
    keys = _bk_normalize_column_keys(column_keys)
    matrix = _bk_table_matrix_from_rows(rows, keys, window)
    widths = [float(_BK_TABULAR_COLUMN_BY_KEY.get(key, (None, None, None, 16.0, 3.0))[4]) for key in keys]
    column_styles = []
    columns = []
    for idx, width in enumerate(widths, start=1):
        column_styles.append('<style:style style:name="co%d" style:family="table-column"><style:table-column-properties style:column-width="%.3fcm"/></style:style>' % (idx, width))
        columns.append('<table:table-column table:style-name="co%d"/>' % idx)
    table_rows = []
    for r_idx, row in enumerate(matrix, start=1):
        style = "ceHeader" if r_idx == 1 else "ceBody"
        cells = []
        for c_idx in range(len(widths)):
            text = row[c_idx] if c_idx < len(row) else ""
            cells.append('<table:table-cell table:style-name="%s" office:value-type="string"><text:p>%s</text:p></table:table-cell>' % (style, _bk_ods_text(text)))
        row_style = "roHeader" if r_idx == 1 else "roBody"
        table_rows.append('<table:table-row table:style-name="%s">%s</table:table-row>' % (row_style, ''.join(cells)))
    content = ''.join([
        '<?xml version="1.0" encoding="UTF-8"?>',
        '<office:document-content ',
        'xmlns:office="urn:oasis:names:tc:opendocument:xmlns:office:1.0" ',
        'xmlns:style="urn:oasis:names:tc:opendocument:xmlns:style:1.0" ',
        'xmlns:text="urn:oasis:names:tc:opendocument:xmlns:text:1.0" ',
        'xmlns:table="urn:oasis:names:tc:opendocument:xmlns:table:1.0" ',
        'xmlns:fo="urn:oasis:names:tc:opendocument:xmlns:xsl-fo-compatible:1.0" ',
        'xmlns:svg="urn:oasis:names:tc:opendocument:xmlns:svg-compatible:1.0" ',
        'office:version="1.2">',
        '<office:font-face-decls><style:font-face style:name="Calibri" svg:font-family="Calibri"/></office:font-face-decls>',
        '<office:automatic-styles>',
        ''.join(column_styles),
        '<style:style style:name="roHeader" style:family="table-row"><style:table-row-properties style:row-height="0.65cm"/></style:style>',
        '<style:style style:name="roBody" style:family="table-row"><style:table-row-properties style:row-height="0.52cm" fo:break-before="auto"/></style:style>',
        '<style:style style:name="ceHeader" style:family="table-cell"><style:table-cell-properties fo:border="0.05pt solid #808080" fo:background-color="#E9EEF6" fo:padding="0.06cm"/><style:text-properties fo:font-size="10pt" fo:font-weight="bold" style:font-name="Calibri"/><style:paragraph-properties fo:margin-top="0cm" fo:margin-bottom="0cm"/></style:style>',
        '<style:style style:name="ceBody" style:family="table-cell"><style:table-cell-properties fo:border="0.05pt solid #B7B7B7" fo:padding="0.06cm"/><style:text-properties fo:font-size="9pt" style:font-name="Calibri"/><style:paragraph-properties fo:margin-top="0cm" fo:margin-bottom="0cm"/></style:style>',
        '</office:automatic-styles>',
        '<office:body><office:spreadsheet><table:table table:name="Transkription">',
        ''.join(columns), ''.join(table_rows),
        '</table:table></office:spreadsheet></office:body></office:document-content>',
    ])
    styles = '<?xml version="1.0" encoding="UTF-8"?><office:document-styles xmlns:office="urn:oasis:names:tc:opendocument:xmlns:office:1.0" xmlns:style="urn:oasis:names:tc:opendocument:xmlns:style:1.0" xmlns:fo="urn:oasis:names:tc:opendocument:xmlns:xsl-fo-compatible:1.0" office:version="1.2"><office:styles><style:default-style style:family="table-cell"><style:text-properties fo:font-size="9pt" style:font-name="Calibri"/></style:default-style></office:styles></office:document-styles>'
    meta = '<?xml version="1.0" encoding="UTF-8"?><office:document-meta xmlns:office="urn:oasis:names:tc:opendocument:xmlns:office:1.0" xmlns:meta="urn:oasis:names:tc:opendocument:xmlns:meta:1.0" office:version="1.2"><office:meta><meta:generator>Bottled Kraken</meta:generator></office:meta></office:document-meta>'
    manifest = '<?xml version="1.0" encoding="UTF-8"?><manifest:manifest xmlns:manifest="urn:oasis:names:tc:opendocument:xmlns:manifest:1.0" manifest:version="1.2"><manifest:file-entry manifest:full-path="/" manifest:media-type="application/vnd.oasis.opendocument.spreadsheet"/><manifest:file-entry manifest:full-path="content.xml" manifest:media-type="text/xml"/><manifest:file-entry manifest:full-path="styles.xml" manifest:media-type="text/xml"/><manifest:file-entry manifest:full-path="meta.xml" manifest:media-type="text/xml"/></manifest:manifest>'
    with zipfile.ZipFile(path, "w") as archive:
        info = zipfile.ZipInfo("mimetype")
        info.date_time = (2020, 1, 1, 0, 0, 0)
        info.compress_type = zipfile.ZIP_STORED
        archive.writestr(info, "application/vnd.oasis.opendocument.spreadsheet")
        for name, data in (("content.xml", content), ("styles.xml", styles), ("meta.xml", meta), ("META-INF/manifest.xml", manifest)):
            zi = zipfile.ZipInfo(name)
            zi.date_time = (2020, 1, 1, 0, 0, 0)
            zi.compress_type = zipfile.ZIP_DEFLATED
            archive.writestr(zi, data.encode("utf-8"))

def _bk_sqlite_tabular_payload(task, rows, column_keys=None, window=None):
    source_path = str(getattr(task, "path", "") or "")
    title = os.path.basename(source_path)
    keys = _bk_normalize_column_keys(column_keys)
    selected_columns = [
        {"key": key, "title": _bk_tabular_column_title(window, key)}
        for key in keys
    ]
    table_rows = _bk_table_dicts_from_rows(rows, keys)
    years = []
    year_seen = set()
    for row in rows or []:
        year = _bk_tab_clean(row.get("year_resolved", ""))
        if re.fullmatch(r"1[5-9]\d{2}|20\d{2}", year) and year not in year_seen:
            year_seen.add(year)
            years.append({"id": f"year_{year}", "year": int(year), "context": "source_heading_or_line"})
    return {
        "schema": "bottled_kraken.sqlite_json.tabular.v2",
        "database_hint": "sqlite",
        "columns": [col["title"] for col in selected_columns],
        "column_definitions": selected_columns,
        "tables": {
            "documents": [{"id": 1, "source_path": source_path, "title": title}],
            "transcription_entries": table_rows,
            "years": years,
        },
    }


def _bk_write_transcription_json(path, task, rows, column_keys=None, window=None):
    payload = _bk_sqlite_tabular_payload(task, rows, column_keys, window)
    with open(path, "w", encoding="utf-8") as handle:
        json.dump(payload, handle, ensure_ascii=False, indent=2)

def _bk_plain_lines(record_views):
    return [_bk_tab_clean(getattr(rv, "text", "")) for rv in (record_views or []) if _bk_tab_clean(getattr(rv, "text", ""))]


def _bk_write_lines_txt(path, record_views):
    with open(path, "w", encoding="utf-8", newline="") as handle:
        for line in _bk_plain_lines(record_views):
            handle.write(line + "\n")


def _bk_write_lines_docx(path, record_views):
    try:
        from docx import Document
        from docx.shared import Pt
    except Exception as exc:
        raise RuntimeError(_bk_registry_lookup("err_no_docx_package_short") or "python-docx") from exc
    doc = Document()
    try:
        normal = doc.styles["Normal"]
        normal.font.name = "Arial"
        normal.font.size = Pt(10)
    except Exception:
        pass
    for line in _bk_plain_lines(record_views):
        p = doc.add_paragraph(line)
        try:
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
        except Exception:
            pass
    doc.save(path)


def _bk_write_lines_odt(path, record_views):
    body = []
    for line in _bk_plain_lines(record_views):
        body.append('<text:p text:style-name="P1">%s</text:p>' % _bk_ods_text(line))
    content = ''.join([
        '<?xml version="1.0" encoding="UTF-8"?>',
        '<office:document-content xmlns:office="urn:oasis:names:tc:opendocument:xmlns:office:1.0" xmlns:style="urn:oasis:names:tc:opendocument:xmlns:style:1.0" xmlns:text="urn:oasis:names:tc:opendocument:xmlns:text:1.0" xmlns:fo="urn:oasis:names:tc:opendocument:xmlns:xsl-fo-compatible:1.0" office:version="1.2">',
        '<office:automatic-styles><style:style style:name="P1" style:family="paragraph"><style:paragraph-properties fo:margin-top="0cm" fo:margin-bottom="0cm"/><style:text-properties fo:font-size="10pt" style:font-name="Arial"/></style:style></office:automatic-styles>',
        '<office:body><office:text>', ''.join(body), '</office:text></office:body></office:document-content>',
    ])
    _odt_page_w, _odt_page_h = bk_page_size_cm(False, None)
    styles = '<?xml version="1.0" encoding="UTF-8"?><office:document-styles xmlns:office="urn:oasis:names:tc:opendocument:xmlns:office:1.0" xmlns:style="urn:oasis:names:tc:opendocument:xmlns:style:1.0" xmlns:text="urn:oasis:names:tc:opendocument:xmlns:text:1.0" xmlns:fo="urn:oasis:names:tc:opendocument:xmlns:xsl-fo-compatible:1.0" xmlns:svg="urn:oasis:names:tc:opendocument:xmlns:svg-compatible:1.0" office:version="1.2"><office:font-face-decls><style:font-face style:name="Arial" svg:font-family="Arial"/></office:font-face-decls><office:styles><style:default-style style:family="paragraph"><style:text-properties fo:font-size="10pt" style:font-name="Arial"/></style:default-style></office:styles><office:automatic-styles><style:page-layout style:name="pm1"><style:page-layout-properties fo:page-width="%.1fcm" fo:page-height="%.1fcm" style:print-orientation="%s" fo:margin-top="1.5cm" fo:margin-bottom="1.5cm" fo:margin-left="1.5cm" fo:margin-right="1.5cm"/></style:page-layout></office:automatic-styles><office:master-styles><style:master-page style:name="Standard" style:page-layout-name="pm1"/></office:master-styles></office:document-styles>' % (_odt_page_w, _odt_page_h, bk_odf_orientation_name(False, None))
    meta = '<?xml version="1.0" encoding="UTF-8"?><office:document-meta xmlns:office="urn:oasis:names:tc:opendocument:xmlns:office:1.0" xmlns:meta="urn:oasis:names:tc:opendocument:xmlns:meta:1.0" office:version="1.2"><office:meta><meta:generator>Bottled Kraken</meta:generator></office:meta></office:document-meta>'
    settings = '<?xml version="1.0" encoding="UTF-8"?><office:document-settings xmlns:office="urn:oasis:names:tc:opendocument:xmlns:office:1.0" office:version="1.2"><office:settings/></office:document-settings>'
    manifest = '<?xml version="1.0" encoding="UTF-8"?><manifest:manifest xmlns:manifest="urn:oasis:names:tc:opendocument:xmlns:manifest:1.0" manifest:version="1.2"><manifest:file-entry manifest:full-path="/" manifest:media-type="application/vnd.oasis.opendocument.text"/><manifest:file-entry manifest:full-path="content.xml" manifest:media-type="text/xml"/><manifest:file-entry manifest:full-path="styles.xml" manifest:media-type="text/xml"/><manifest:file-entry manifest:full-path="meta.xml" manifest:media-type="text/xml"/><manifest:file-entry manifest:full-path="settings.xml" manifest:media-type="text/xml"/></manifest:manifest>'
    with zipfile.ZipFile(path, "w") as archive:
        info = zipfile.ZipInfo("mimetype")
        info.date_time = (2020, 1, 1, 0, 0, 0)
        info.compress_type = zipfile.ZIP_STORED
        archive.writestr(info, "application/vnd.oasis.opendocument.text")
        for name, data in (("content.xml", content), ("styles.xml", styles), ("meta.xml", meta), ("settings.xml", settings), ("META-INF/manifest.xml", manifest)):
            zi = zipfile.ZipInfo(name)
            zi.date_time = (2020, 1, 1, 0, 0, 0)
            zi.compress_type = zipfile.ZIP_DEFLATED
            archive.writestr(zi, data.encode("utf-8"))


def _bk_write_table_txt(path, rows, column_keys=None, window=None):
    with open(path, "w", encoding="utf-8", newline="") as handle:
        writer = csv.writer(handle, delimiter="\t", quotechar='"', quoting=csv.QUOTE_MINIMAL, lineterminator="\n")
        writer.writerows(_bk_table_matrix_from_rows(rows, column_keys, window))


def _bk_write_table_docx(path, rows, column_keys=None, window=None):
    try:
        from docx import Document
        from docx.shared import Pt
    except Exception as exc:
        raise RuntimeError(_bk_registry_lookup("err_no_docx_package_short") or "python-docx") from exc
    matrix = _bk_table_matrix_from_rows(rows, column_keys, window)
    doc = Document()
    try:
        section = doc.sections[0]
        # Bisher war hier immer Querformat erzwungen. Jetzt entscheidet die
        # im Export-Dialog gewaehlte Ausrichtung (Standard weiterhin quer).
        use_landscape = bool(bk_resolve_landscape(True, window))
        if use_landscape:
            section.orientation = 1
            section.page_width, section.page_height = section.page_height, section.page_width
        normal = doc.styles["Normal"]
        normal.font.name = "Arial"
        normal.font.size = Pt(8)
    except Exception:
        pass
    if not matrix:
        doc.save(path)
        return
    table = doc.add_table(rows=1, cols=len(matrix[0]))
    try:
        table.style = "Table Grid"
        table.autofit = True
    except Exception:
        pass
    hdr = table.rows[0].cells
    for idx, value in enumerate(matrix[0]):
        hdr[idx].text = str(value or "")
        try:
            for p in hdr[idx].paragraphs:
                for run in p.runs:
                    run.bold = True
                    run.font.size = Pt(8)
        except Exception:
            pass
    for row in matrix[1:]:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = str(value or "")
            try:
                for p in cells[idx].paragraphs:
                    for run in p.runs:
                        run.font.size = Pt(8)
            except Exception:
                pass
    doc.save(path)




def _bk_load_saved_column_keys(window):
    raw = None
    try:
        settings = getattr(window, "settings", None)
        if settings is not None:
            raw = settings.value("export/table_column_keys", "", str)
    except Exception:
        raw = None
    keys = []
    if raw:
        try:
            parsed = json.loads(str(raw))
            if isinstance(parsed, list):
                keys = [str(x) for x in parsed]
        except Exception:
            keys = [part.strip() for part in str(raw).split(",") if part.strip()]
    if not keys:
        keys = getattr(window, "_bk_export_selected_column_keys", None) or _BK_TABULAR_DEFAULT_KEYS
    return _bk_normalize_column_keys(keys)




def _bk_current_column_keys_for_render(window):
    keys = getattr(window, "_bk_export_current_column_keys", None)
    if keys:
        return _bk_normalize_column_keys(keys)
    return _bk_load_saved_column_keys(window)


def _bk_column_choice_dialog(self, fmt=None, include_text_modes=False):
    dlg = QDialog(self)
    dlg.setWindowTitle(_bk_tab_tr(self, "export_text_layout_title" if include_text_modes else "export_table_columns_title"))
    dlg.setModal(True)

    layout = QVBoxLayout(dlg)
    layout.setContentsMargins(16, 14, 16, 14)
    layout.setSpacing(10)

    intro_key = "export_text_layout_intro_extended" if include_text_modes else "export_table_columns_intro"
    label = QLabel(_bk_tab_tr(self, intro_key), dlg)
    label.setWordWrap(True)
    layout.addWidget(label)

    rb_original = rb_lines = rb_table = None
    mode_box = None
    if include_text_modes:
        if QGroupBox is not None:
            mode_box = QGroupBox(_bk_tab_tr(self, "export_layout_mode_group"), dlg)
            mode_layout = QVBoxLayout(mode_box)
            mode_layout.setContentsMargins(12, 10, 12, 10)
            mode_layout.setSpacing(6)
        else:
            mode_layout = QVBoxLayout()
            mode_layout.setSpacing(6)
        rb_original = QRadioButton(_bk_tab_tr(self, "export_text_layout_original"), dlg)
        rb_lines = QRadioButton(_bk_tab_tr(self, "export_text_layout_lines"), dlg)
        rb_table = QRadioButton(_bk_tab_tr(self, "export_text_layout_table"), dlg)
        mode = str(getattr(self, "_bk_export_text_layout_mode", "original") or "original").lower()
        rb_original.setChecked(mode not in {"lines", "table"})
        rb_lines.setChecked(mode == "lines")
        rb_table.setChecked(mode == "table")
        mode_layout.addWidget(rb_original)
        mode_layout.addWidget(rb_lines)
        mode_layout.addWidget(rb_table)
        if mode_box is not None:
            layout.addWidget(mode_box)
        else:
            layout.addLayout(mode_layout)

    _bk_add_export_orientation_group(self, dlg, layout)

    selected_keys = _bk_load_saved_column_keys_for_dialog(self)
    checkboxes = {}

    if QGroupBox is not None:
        columns_box = QGroupBox(_bk_tab_tr(self, "export_table_columns_label"), dlg)
        columns_layout = QVBoxLayout(columns_box)
        columns_layout.setContentsMargins(12, 10, 12, 10)
        columns_layout.setSpacing(8)
    else:
        columns_box = None
        columns_layout = QVBoxLayout()
        columns_layout.setSpacing(8)
        columns_layout.addWidget(QLabel(_bk_tab_tr(self, "export_table_columns_label"), dlg))

    if QGridLayout is not None:
        checkbox_grid = QGridLayout()
        checkbox_grid.setContentsMargins(0, 0, 0, 0)
        checkbox_grid.setHorizontalSpacing(22)
        checkbox_grid.setVerticalSpacing(6)
        grid_cols = 3
        for idx, key in enumerate(_BK_TABULAR_KEYS):
            cb = QCheckBox(_bk_tabular_column_title(self, key), dlg)
            cb.setChecked(key in selected_keys)
            checkboxes[key] = cb
            checkbox_grid.addWidget(cb, idx // grid_cols, idx % grid_cols)
        columns_layout.addLayout(checkbox_grid)
    else:
        checkbox_layout = QVBoxLayout()
        checkbox_layout.setSpacing(4)
        for key in _BK_TABULAR_KEYS:
            cb = QCheckBox(_bk_tabular_column_title(self, key), dlg)
            cb.setChecked(key in selected_keys)
            checkboxes[key] = cb
            checkbox_layout.addWidget(cb)
        columns_layout.addLayout(checkbox_layout)

    quick_row = QHBoxLayout()
    quick_row.setSpacing(8)
    btn_all = QPushButton(_bk_tab_tr(self, "export_table_columns_all"), dlg)
    btn_none = QPushButton(_bk_tab_tr(self, "export_table_columns_none_button"), dlg)
    btn_remember = QPushButton(_bk_tab_tr(self, "export_table_columns_remember"), dlg)
    quick_row.addWidget(btn_all)
    quick_row.addWidget(btn_none)
    quick_row.addWidget(btn_remember)
    quick_row.addStretch(1)
    columns_layout.addLayout(quick_row)

    if columns_box is not None:
        layout.addWidget(columns_box)
    else:
        layout.addLayout(columns_layout)

    def current_checked_keys():
        return [key for key, cb in checkboxes.items() if cb.isChecked()]

    def set_all():
        for cb in checkboxes.values():
            cb.setChecked(True)

    def set_none():
        for cb in checkboxes.values():
            cb.setChecked(False)

    result = {"mode": "table", "columns": selected_keys, "remembered": False}

    def remember_selection():
        keys = current_checked_keys()
        if not keys:
            QMessageBox.warning(dlg, _bk_tab_tr(self, "warn_title"), _bk_tab_tr(self, "export_table_columns_none"))
            return
        result["remembered"] = True
        result["columns"] = _bk_save_column_keys(self, keys)

    btn_all.clicked.connect(set_all)
    btn_none.clicked.connect(set_none)
    btn_remember.clicked.connect(remember_selection)

    def sync_columns_enabled():
        enabled = True
        if include_text_modes and rb_table is not None:
            enabled = bool(rb_table.isChecked())
        if columns_box is not None:
            columns_box.setEnabled(enabled)
        else:
            for cb in checkboxes.values():
                cb.setEnabled(enabled)
            btn_all.setEnabled(enabled)
            btn_none.setEnabled(enabled)
            btn_remember.setEnabled(enabled)

    if include_text_modes:
        try:
            rb_original.toggled.connect(sync_columns_enabled)
            rb_lines.toggled.connect(sync_columns_enabled)
            rb_table.toggled.connect(sync_columns_enabled)
        except Exception:
            pass
        sync_columns_enabled()

    buttons = QDialogButtonBox(QDialogButtonBox.Ok | QDialogButtonBox.Cancel, dlg)
    try:
        buttons.button(QDialogButtonBox.Ok).setText(_bk_tab_tr(self, "btn_ok"))
        buttons.button(QDialogButtonBox.Cancel).setText(_bk_tab_tr(self, "btn_cancel"))
    except Exception:
        pass

    def cancel_dialog():
        result["cancelled"] = True
        try:
            dlg.done(QDialog.Rejected)
        except Exception:
            try:
                dlg.reject()
            except Exception:
                pass

    def accept_checked():
        result["cancelled"] = False
        mode = "table"
        if include_text_modes:
            if rb_lines and rb_lines.isChecked():
                mode = "lines"
            elif rb_original and rb_original.isChecked():
                mode = "original"
            else:
                mode = "table"
        keys = current_checked_keys()
        if mode == "table" and not keys:
            QMessageBox.warning(dlg, _bk_tab_tr(self, "warn_title"), _bk_tab_tr(self, "export_table_columns_none"))
            return
        result["mode"] = mode
        result["columns"] = _bk_normalize_column_keys(keys)
        dlg.accept()

    buttons.accepted.connect(accept_checked)
    buttons.rejected.connect(cancel_dialog)
    layout.addWidget(buttons)

    if include_text_modes:
        dlg.setMinimumSize(560, 410)
        try:
            dlg.resize(620, 450)
        except Exception:
            pass
    else:
        dlg.setMinimumSize(560, 300)
        try:
            dlg.resize(620, 330)
        except Exception:
            pass

    try:
        _exec_result = dlg.exec()
    except RuntimeError:
        result["cancelled"] = True
        return None
    except Exception:
        result["cancelled"] = True
        return None
    if _exec_result != QDialog.Accepted:
        result["cancelled"] = True
        return None
    return result

def _bk_text_layout_choice_dialog(self, fmt):
    result = _bk_column_choice_dialog(self, fmt, include_text_modes=True)
    if not result:
        return None
    self._bk_export_current_column_keys = result.get("columns") or list(_BK_TABULAR_DEFAULT_KEYS)
    if result.get("remembered"):
        self._bk_export_selected_column_keys = result.get("columns") or list(_BK_TABULAR_DEFAULT_KEYS)
    return result.get("mode", "original")



try:
    _BK_TABULAR_PREV_EXPORT_SINGLE = MainWindow._export_single_interactive
except Exception:
    _BK_TABULAR_PREV_EXPORT_SINGLE = None
try:
    _BK_TABULAR_PREV_EXPORT_BATCH = MainWindow._export_batch
except Exception:
    _BK_TABULAR_PREV_EXPORT_BATCH = None


def _bk_tabular_render_file(self, path: str, fmt: str, item: TaskItem):
    fmt_l = str(fmt or "").lower().lstrip(".")
    if not item or not getattr(item, "results", None):
        return RENDER_NOT_HANDLED
    _text, _kr, pil_image, record_views = item.results
    try:
        export_image = _load_image_color(item.path)
    except Exception:
        export_image = pil_image
    image_size = getattr(export_image, "size", None) or getattr(pil_image, "size", None)
    column_keys = _bk_current_column_keys_for_render(self)
    if fmt_l in _BK_TABLE_EXPORT_FMTS:
        rows = _bk_build_transcription_rows(record_views, image_size)
        if fmt_l == "csv":
            return _bk_write_transcription_csv(path, rows, column_keys, self)
        if fmt_l == "json":
            return _bk_write_transcription_json(path, item, rows, column_keys, self)
        if fmt_l in {"xlsx", "excel"}:
            return _bk_write_transcription_xlsx(path, rows, column_keys, self)
        if fmt_l in {"ods", "calc"}:
            return _bk_write_transcription_ods(path, rows, column_keys, self)
    if fmt_l in _BK_TEXT_LAYOUT_FMTS:
        mode = str(getattr(self, "_bk_export_text_layout_mode", "original") or "original").lower()
        if mode == "table":
            rows = _bk_build_transcription_rows(record_views, image_size)
            if fmt_l in {"txt", "text", "txt_plain"}:
                return _bk_write_table_txt(path, rows, column_keys, self)
            if fmt_l in {"docx", "word"}:
                return _bk_write_table_docx(path, rows, column_keys, self)
            if fmt_l == "odt":
                return _bk_write_table_odt(path, rows, column_keys, self)
        if mode == "lines":
            if fmt_l in {"txt", "text", "txt_plain"}:
                return _bk_write_lines_txt(path, record_views)
            if fmt_l in {"docx", "word"}:
                return _bk_write_lines_docx(path, record_views)
            if fmt_l == "odt":
                return _bk_write_lines_odt(path, record_views)
        if fmt_l in {"txt", "text", "txt_plain"}:
            return write_spatial_txt(path, record_views, image_size)
        if fmt_l in {"docx", "word"}:
            return write_positioned_docx(path, item, export_image, record_views)
        if fmt_l == "odt":
            return write_positioned_odt(path, item, export_image, record_views)
    return RENDER_NOT_HANDLED






def _bk_export_sqlite_json_tabular(self):
    task = _bk_fix36_current_task(self) if callable(globals().get("_bk_fix36_current_task")) else None
    if not task or not getattr(task, "results", None):
        QMessageBox.information(self, _bk_tab_tr(self, "info_title"), _bk_tab_tr(self, "warn_no_ocr_results"))
        return
    try:
        _txt, _kr, pil_image, recs = task.results
        try:
            export_image = _load_image_color(task.path)
            image_size = export_image.size
        except Exception:
            image_size = getattr(pil_image, "size", None)
        rows = _bk_build_transcription_rows(recs, image_size)
    except Exception:
        rows = []
    if not rows:
        QMessageBox.information(self, _bk_tab_tr(self, "info_title"), _bk_tab_tr(self, "warn_no_exportable_person_entries"))
        return
    result = _bk_column_choice_dialog(self, "sqlite-json", include_text_modes=False)
    if result is None:
        return
    column_keys = result.get("columns") or list(_BK_TABULAR_DEFAULT_KEYS)
    self._bk_export_current_column_keys = column_keys
    if result.get("remembered"):
        self._bk_export_selected_column_keys = column_keys
    start_dir = getattr(self, "current_export_dir", "") or os.path.dirname(getattr(task, "path", "") or "") or os.getcwd()
    default_name = os.path.splitext(os.path.basename(getattr(task, "path", "bottled_kraken")))[0] + "_sqlite.json"
    path, _filter = QFileDialog.getSaveFileName(
        self,
        _bk_tab_tr(self, "dlg_sqlite_json_title"),
        os.path.join(start_dir, default_name),
        _bk_tab_tr(self, "filter_json_files"),
    )
    if not path:
        return
    if not path.lower().endswith(".json"):
        path += ".json"
    _bk_write_transcription_json(path, task, rows, column_keys, self)
    try:
        self.current_export_dir = os.path.dirname(path)
        self.status_bar.showMessage(_bk_tab_tr(self, "msg_sqlite_export_done").format(os.path.basename(path)), 5000)
    except Exception:
        pass


try:
    register_render_handler(_bk_tabular_render_file)
    MainWindow._bk_export_sqlite_json = _bk_export_sqlite_json_tabular
    MainWindow.bk_export_sqlite_persons = _bk_export_sqlite_json_tabular
except Exception:
    pass

__all__ = [
    '_BK_TABULAR_COLUMNS',
    '_BK_TABULAR_DEFAULT_KEYS',
    '_BK_TABULAR_HEADERS',
    '_BK_TABULAR_KEYS',
    '_BK_TABLE_EXPORT_FMTS',
    '_BK_TEXT_LAYOUT_FMTS',
    '_bk_build_transcription_rows',
    '_bk_column_choice_dialog',
    '_bk_current_column_keys_for_render',
    '_bk_load_saved_column_keys',
    '_bk_save_column_keys',
    '_bk_tab_merge_visual_row_records',
    '_bk_export_sqlite_json_tabular',
    '_bk_plain_lines',
    '_bk_sqlite_tabular_payload',
    '_bk_tabular_render_file',
    '_bk_text_layout_choice_dialog',
    '_bk_write_lines_docx',
    '_bk_write_lines_odt',
    '_bk_write_lines_txt',
    '_bk_write_table_docx',
    '_bk_write_table_odt',
    '_bk_write_table_txt',
    '_bk_write_transcription_csv',
    '_bk_write_transcription_json',
    '_bk_write_transcription_ods',
    '_bk_write_transcription_xlsx',
]
register_globals('bk', globals(), __all__)

# ---------------------------------------------------------------------------
# Manuelle Tabellenbereiche / Spaltenzonen
# ---------------------------------------------------------------------------
try:
    from PySide6.QtCore import QRectF, Qt, QTimer
    from PySide6.QtGui import QColor, QBrush, QPen, QPixmap
    from PySide6.QtWidgets import (
        QAbstractItemView,
        QComboBox,
        QFormLayout,
        QGraphicsRectItem,
        QGraphicsScene,
        QGraphicsView,
        QLineEdit,
        QSpinBox,
        QTableWidget,
        QTableWidgetItem,
        QHeaderView,
    )
except Exception:
    QRectF = None
    QColor = None
    QBrush = None
    QPen = None
    QPixmap = None
    QGraphicsRectItem = None
    QGraphicsScene = None
    QGraphicsView = object
    QLineEdit = None
    QSpinBox = None
    QComboBox = None
    QFormLayout = None
    QTableWidget = None
    QTableWidgetItem = None
    QHeaderView = None
    QAbstractItemView = None
try:
    from PIL.ImageQt import ImageQt as _BK_ImageQt
except Exception:
    _BK_ImageQt = None

_BK_ZONE_TYPES = ("data", "number", "year", "ignore")
_BK_ZONE_TYPE_TRANSLATION = {
    "data": ("export_zone_type_data", "Datenbereich"),
    "number": ("export_zone_type_number", "Nummern-/Seitennummernspalte"),
    "year": ("export_zone_type_year", "Überschriften/Jahre"),
    "ignore": ("export_zone_type_ignore", "Ignorieren"),
}






def _bk_get_task_export_zones(task) -> list:
    zones = []
    for idx, zone in enumerate(getattr(task, "export_zones", []) or []):
        clean = _bk_export_clean_zone(zone, idx)
        if clean:
            zones.append(clean)
    zones.sort(key=lambda z: (int(z.get("order", 0) or 0), float(z.get("x0", 0.0)), float(z.get("y0", 0.0))))
    for idx, zone in enumerate(zones):
        zone["order"] = idx
    return zones


def _bk_set_task_export_zones(task, zones) -> list:
    clean = []
    for idx, zone in enumerate(zones or []):
        item = _bk_export_clean_zone(zone, idx)
        if item:
            item["order"] = len(clean)
            clean.append(item)
    try:
        setattr(task, "export_zones", clean)
        task.edited = True
    except Exception:
        pass
    return clean



def _bk_zone_records(records, zone):
    return [record for record in (records or []) if _bk_zone_contains_record(zone, record)]


def _bk_zone_is_ignored(record, zones) -> bool:
    return any(_bk_zone_contains_record(zone, record) for zone in zones if zone.get("type") == "ignore")












def _bk_open_export_zones_dialog(window, task):
    if task is None:
        try:
            task = window._current_task()
        except Exception:
            task = None
    if task is None:
        QMessageBox.warning(window, _bk_tab_tr(window, "warn_title"), _bk_tab_tr(window, "warn_need_done_for_ai"))
        return None
    path = getattr(task, "path", "")
    if not path or not os.path.exists(path):
        QMessageBox.warning(window, _bk_tab_tr(window, "warn_title"), _bk_tab_tr(window, "warn_project_file_missing").format(path))
        return None
    dlg = _BKExportZonesDialog(window, task, path, _bk_get_task_export_zones(task))
    try:
        # Der Bereichsdialog ist ein echter Unterdialog des Export-Workflows.
        # Er muss vor dem Export-Darstellungsdialog bleiben und darf nicht
        # von diesem überdeckt werden, wenn der Benutzer zwischen Fenstern wechselt.
        dlg.setWindowModality(Qt.ApplicationModal)
        dlg.setWindowFlag(Qt.WindowStaysOnTopHint, True)
        dlg.setWindowFlag(Qt.WindowContextHelpButtonHint, False)
    except Exception:
        pass
    try:
        dlg.showMaximized()
        dlg.raise_()
        dlg.activateWindow()
        QApplication.processEvents()
    except Exception:
        pass
    if dlg.exec() != QDialog.Accepted:
        return None
    zones = _bk_set_task_export_zones(task, dlg.zones())
    return zones


def _bk_zone_column_choice_dialog(self, fmt=None, include_text_modes=False):
    dlg = QDialog(self)
    try:
        dlg.setAttribute(Qt.WA_DeleteOnClose, False)
        dlg.setWindowModality(Qt.ApplicationModal)
        dlg.setWindowFlag(Qt.WindowContextHelpButtonHint, False)
    except Exception:
        pass
    dlg.setWindowTitle(_bk_tab_tr(self, "export_text_layout_title" if include_text_modes else "export_table_columns_title"))
    dlg.setModal(True)

    layout = QVBoxLayout(dlg)
    layout.setContentsMargins(16, 14, 16, 14)
    layout.setSpacing(10)

    intro_key = "export_text_layout_intro_extended" if include_text_modes else "export_table_columns_intro"
    label = QLabel(_bk_tab_tr(self, intro_key), dlg)
    label.setWordWrap(True)
    layout.addWidget(label)

    rb_original = rb_lines = rb_table = None
    mode_box = None
    if include_text_modes:
        if QGroupBox is not None:
            mode_box = QGroupBox(_bk_tab_tr(self, "export_layout_mode_group"), dlg)
            mode_layout = QVBoxLayout(mode_box)
            mode_layout.setContentsMargins(12, 10, 12, 10)
            mode_layout.setSpacing(6)
        else:
            mode_layout = QVBoxLayout()
            mode_layout.setSpacing(6)
        rb_original = QRadioButton(_bk_tab_tr(self, "export_text_layout_original"), dlg)
        rb_lines = QRadioButton(_bk_tab_tr(self, "export_text_layout_lines"), dlg)
        rb_table = QRadioButton(_bk_tab_tr(self, "export_text_layout_table"), dlg)
        mode = str(getattr(self, "_bk_export_text_layout_mode", "original") or "original").lower()
        rb_original.setChecked(mode not in {"lines", "table"})
        rb_lines.setChecked(mode == "lines")
        rb_table.setChecked(mode == "table")
        mode_layout.addWidget(rb_original)
        mode_layout.addWidget(rb_lines)
        mode_layout.addWidget(rb_table)
        if mode_box is not None:
            layout.addWidget(mode_box)
        else:
            layout.addLayout(mode_layout)

    selected_keys = _bk_load_saved_column_keys(self)
    checkboxes = {}

    if QGroupBox is not None:
        columns_box = QGroupBox(_bk_tab_tr(self, "export_table_columns_label"), dlg)
        columns_layout = QVBoxLayout(columns_box)
        columns_layout.setContentsMargins(12, 10, 12, 10)
        columns_layout.setSpacing(8)
    else:
        columns_box = None
        columns_layout = QVBoxLayout()
        columns_layout.setSpacing(8)
        columns_layout.addWidget(QLabel(_bk_tab_tr(self, "export_table_columns_label"), dlg))

    if QGridLayout is not None:
        checkbox_grid = QGridLayout()
        checkbox_grid.setContentsMargins(0, 0, 0, 0)
        checkbox_grid.setHorizontalSpacing(22)
        checkbox_grid.setVerticalSpacing(6)
        grid_cols = 3
        for idx, key in enumerate(_BK_TABULAR_KEYS):
            cb = QCheckBox(_bk_tabular_column_title(self, key), dlg)
            cb.setChecked(key in selected_keys)
            checkboxes[key] = cb
            checkbox_grid.addWidget(cb, idx // grid_cols, idx % grid_cols)
        columns_layout.addLayout(checkbox_grid)
    else:
        checkbox_layout = QVBoxLayout()
        checkbox_layout.setSpacing(4)
        for key in _BK_TABULAR_KEYS:
            cb = QCheckBox(_bk_tabular_column_title(self, key), dlg)
            cb.setChecked(key in selected_keys)
            checkboxes[key] = cb
            checkbox_layout.addWidget(cb)
        columns_layout.addLayout(checkbox_layout)

    zone_row = QHBoxLayout()
    zone_row.setSpacing(8)
    cb_zones = QCheckBox(_bk_tab_tr(self, "export_table_use_zones"), dlg)
    cb_zones.setChecked(bool(getattr(self, "_bk_export_use_zones", False)))
    btn_zones = QPushButton(_bk_tab_tr(self, "export_table_define_zones"), dlg)
    zone_row.addWidget(cb_zones, 1)
    zone_row.addWidget(btn_zones)
    columns_layout.addLayout(zone_row)

    quick_row = QHBoxLayout()
    quick_row.setSpacing(8)
    btn_all = QPushButton(_bk_tab_tr(self, "export_table_columns_all"), dlg)
    btn_none = QPushButton(_bk_tab_tr(self, "export_table_columns_none_button"), dlg)
    btn_remember = QPushButton(_bk_tab_tr(self, "export_table_columns_remember"), dlg)
    quick_row.addWidget(btn_all)
    quick_row.addWidget(btn_none)
    quick_row.addWidget(btn_remember)
    quick_row.addStretch(1)
    columns_layout.addLayout(quick_row)

    if columns_box is not None:
        layout.addWidget(columns_box)
    else:
        layout.addLayout(columns_layout)

    def current_checked_keys():
        return [key for key, cb in checkboxes.items() if cb.isChecked()]

    def set_all():
        for cb in checkboxes.values():
            cb.setChecked(True)

    def set_none():
        for cb in checkboxes.values():
            cb.setChecked(False)

    result = {"mode": "table", "columns": selected_keys, "remembered": False, "use_zones": cb_zones.isChecked(), "cancelled": False}

    def remember_selection():
        keys = current_checked_keys()
        if not keys:
            QMessageBox.warning(dlg, _bk_tab_tr(self, "warn_title"), _bk_tab_tr(self, "export_table_columns_none"))
            return
        result["remembered"] = True
        result["columns"] = _bk_save_column_keys(self, keys)
        try:
            self.settings.setValue("export/table_use_zones", bool(cb_zones.isChecked()))
            self.settings.sync()
        except Exception:
            pass

    def define_zones():
        task = None
        try:
            task = self._current_task()
        except Exception:
            task = None
        zones = _bk_open_export_zones_dialog(self, task)
        if zones is not None:
            cb_zones.setChecked(bool(zones))

    btn_all.clicked.connect(set_all)
    btn_none.clicked.connect(set_none)
    btn_remember.clicked.connect(remember_selection)
    btn_zones.clicked.connect(define_zones)

    def sync_columns_enabled():
        enabled = True
        if include_text_modes and rb_table is not None:
            enabled = bool(rb_table.isChecked())
        if columns_box is not None:
            columns_box.setEnabled(enabled)
        else:
            for cb in checkboxes.values():
                cb.setEnabled(enabled)
            btn_all.setEnabled(enabled)
            btn_none.setEnabled(enabled)
            btn_remember.setEnabled(enabled)
            cb_zones.setEnabled(enabled)
            btn_zones.setEnabled(enabled)

    if include_text_modes:
        try:
            rb_original.toggled.connect(sync_columns_enabled)
            rb_lines.toggled.connect(sync_columns_enabled)
            rb_table.toggled.connect(sync_columns_enabled)
        except Exception:
            pass
        sync_columns_enabled()

    buttons = QDialogButtonBox(QDialogButtonBox.Ok | QDialogButtonBox.Cancel, dlg)
    try:
        buttons.button(QDialogButtonBox.Ok).setText(_bk_tab_tr(self, "btn_ok"))
        buttons.button(QDialogButtonBox.Cancel).setText(_bk_tab_tr(self, "btn_cancel"))
    except Exception:
        pass

    def cancel_dialog():
        result["cancelled"] = True
        try:
            dlg.done(QDialog.Rejected)
        except Exception:
            try:
                dlg.reject()
            except Exception:
                pass

    def accept_checked():
        result["cancelled"] = False
        mode = "table"
        if include_text_modes:
            if rb_lines and rb_lines.isChecked():
                mode = "lines"
            elif rb_original and rb_original.isChecked():
                mode = "original"
            else:
                mode = "table"
        keys = current_checked_keys()
        if mode == "table" and not keys:
            QMessageBox.warning(dlg, _bk_tab_tr(self, "warn_title"), _bk_tab_tr(self, "export_table_columns_none"))
            return
        result["mode"] = mode
        result["columns"] = _bk_normalize_column_keys(keys)
        result["use_zones"] = bool(cb_zones.isChecked())
        try:
            self._bk_export_use_zones = result["use_zones"]
        except Exception:
            pass
        dlg.accept()

    buttons.accepted.connect(accept_checked)
    buttons.rejected.connect(cancel_dialog)
    layout.addWidget(buttons)

    if include_text_modes:
        dlg.setMinimumSize(610, 470)
        try:
            dlg.resize(680, 500)
        except Exception:
            pass
    else:
        dlg.setMinimumSize(610, 360)
        try:
            dlg.resize(680, 390)
        except Exception:
            pass

    if dlg.exec() != QDialog.Accepted:
        return None
    return result


try:
    _bk_column_choice_dialog = _bk_zone_column_choice_dialog
except Exception:
    pass



def _bk_zone_rows_for_item(item, image_size, use_zones):
    _text, _kr, _pil, record_views = item.results
    zones = _bk_get_task_export_zones(item) if use_zones else []
    if zones:
        return _bk_build_transcription_rows_with_zones(record_views, image_size, zones)
    return _bk_build_transcription_rows(record_views, image_size)






def _bk_zone_export_sqlite_json(self):
    task = _bk_fix36_current_task(self) if callable(globals().get("_bk_fix36_current_task")) else None
    if not task or not getattr(task, "results", None):
        QMessageBox.information(self, _bk_tab_tr(self, "info_title"), _bk_tab_tr(self, "warn_no_ocr_results"))
        return
    result = _bk_column_choice_dialog(self, "sqlite-json", include_text_modes=False)
    if result is None:
        return
    column_keys = result.get("columns") or list(_BK_TABULAR_DEFAULT_KEYS)
    self._bk_export_current_column_keys = column_keys
    self._bk_export_use_zones = bool(result.get("use_zones", False))
    if result.get("remembered"):
        self._bk_export_selected_column_keys = column_keys
    try:
        _txt, _kr, pil_image, recs = task.results
        try:
            export_image = _load_image_color(task.path)
            image_size = export_image.size
        except Exception:
            image_size = getattr(pil_image, "size", None)
        rows = _bk_zone_rows_for_item(task, image_size, self._bk_export_use_zones)
    except Exception:
        rows = []
    if not rows:
        QMessageBox.information(self, _bk_tab_tr(self, "info_title"), _bk_tab_tr(self, "warn_no_exportable_person_entries"))
        return
    start_dir = getattr(self, "current_export_dir", "") or os.path.dirname(getattr(task, "path", "") or "") or os.getcwd()
    default_name = os.path.splitext(os.path.basename(getattr(task, "path", "bottled_kraken")))[0] + "_sqlite.json"
    path, _filter = QFileDialog.getSaveFileName(
        self,
        _bk_tab_tr(self, "dlg_sqlite_json_title"),
        os.path.join(start_dir, default_name),
        _bk_tab_tr(self, "filter_json_files"),
    )
    if not path:
        return
    if not path.lower().endswith(".json"):
        path += ".json"
    _bk_write_transcription_json(path, task, rows, column_keys, self)
    try:
        self.current_export_dir = os.path.dirname(path)
        self.status_bar.showMessage(_bk_tab_tr(self, "msg_sqlite_export_done").format(os.path.basename(path)), 5000)
    except Exception:
        pass


try:
    _BK_ZONE_PREV_TASK_TO_DICT = MainWindow._task_to_dict
except Exception:
    _BK_ZONE_PREV_TASK_TO_DICT = None
try:
    _BK_ZONE_PREV_TASK_FROM_DICT = MainWindow._task_from_dict
except Exception:
    _BK_ZONE_PREV_TASK_FROM_DICT = None


def _bk_zone_task_to_dict(self, task):
    payload = _BK_ZONE_PREV_TASK_TO_DICT(self, task) if callable(_BK_ZONE_PREV_TASK_TO_DICT) else {}
    try:
        zones = _bk_get_task_export_zones(task)
        if zones:
            payload["export_zones"] = zones
    except Exception:
        pass
    return payload


def _bk_zone_task_from_dict(self, data):
    task = _BK_ZONE_PREV_TASK_FROM_DICT(self, data) if callable(_BK_ZONE_PREV_TASK_FROM_DICT) else None
    try:
        if task is not None and isinstance(data, dict):
            _bk_set_task_export_zones(task, data.get("export_zones") or [])
    except Exception:
        pass
    return task


try:
    MainWindow._bk_export_sqlite_json = _bk_zone_export_sqlite_json
    MainWindow.bk_export_sqlite_persons = _bk_zone_export_sqlite_json
    MainWindow._task_to_dict = _bk_zone_task_to_dict
    MainWindow._task_from_dict = _bk_zone_task_from_dict
except Exception:
    pass

try:
    __all__.extend([
        '_bk_build_transcription_rows_with_zones',
        '_bk_export_zone_title',
        '_bk_get_task_export_zones',
        '_bk_open_export_zones_dialog',
        '_bk_set_task_export_zones',
        '_bk_zone_column_choice_dialog',
        '_bk_zone_export_sqlite_json',
    ])
    register_globals('bk', globals(), __all__)
except Exception:
    pass

# ---------------------------------------------------------------------------
# Tabellenbereiche: direkte Bearbeitung per Maus + fachliche Datentypen
# ---------------------------------------------------------------------------
try:
    from PySide6.QtCore import QPointF
    from PySide6.QtGui import QPainter
except Exception:
    QPointF = None
    QPainter = None

_BK_ZONE_TYPES = ("heading", "names", "age", "years", "places", "page_numbers", "other", "ignore")
_BK_ZONE_TYPE_TRANSLATION = {
    "heading": ("export_zone_type_heading", "Überschrift"),
    "names": ("export_zone_type_names", "Namen"),
    "age": ("export_zone_type_age", "Alter"),
    "years": ("export_zone_type_years", "Jahre"),
    "places": ("export_zone_type_places", "Orte"),
    "page_numbers": ("export_zone_type_page_numbers", "Seitenzahlen"),
    "other": ("export_zone_type_other", "Sonstiges"),
    "ignore": ("export_zone_type_ignore", "Ignorieren"),
}
_BK_ZONE_LEGACY_TYPES = {
    "data": "names",
    "number": "page_numbers",
    "year": "years",
}







def _bk_zone_age_value(text: str) -> str:
    value = _bk_tab_extract_age(text or "")
    if value:
        return value
    txt = _bk_tab_clean(text)
    if re.fullmatch(r"\d{1,3}\s*(?:Jahre?|Jahr|J\.?|Monate?|Mon\.?|Tage?|Wochen?)\.?", txt, flags=re.IGNORECASE):
        return txt
    return ""


def _bk_zone_year_value(text: str) -> str:
    return _bk_tab_year_value(text) or _bk_tab_extract_year_in_source(text)


def _bk_zone_place_value(text: str) -> str:
    txt = _bk_tab_clean(text)
    if not txt or not _bk_tab_has_letters(txt):
        return ""
    if len(txt) > 90:
        return ""
    return txt.strip(" ,.;:-")


def _bk_zone_page_number_value(text: str) -> str:
    return _bk_tab_number_value(text)






class _BKExportZoneRectItem(QGraphicsRectItem if QGraphicsRectItem is not None else object):
    HANDLE_SIZE = 10.0

    def __init__(self, dialog, row, rect, fill, stroke):
        super().__init__(rect)
        self._dialog = dialog
        self._row = row
        self._mode = None
        self._start_scene = None
        self._start_rect = None
        self.setPen(QPen(stroke, 2))
        self.setBrush(QBrush(fill))
        self.setZValue(20 + row)
        try:
            self.setAcceptHoverEvents(True)
            self.setAcceptedMouseButtons(Qt.LeftButton)
        except Exception:
            pass

    def _handle_at(self, pos):
        rect = self.rect()
        s = max(6.0, min(18.0, self.HANDLE_SIZE))
        x = float(pos.x())
        y = float(pos.y())
        left, right = rect.left(), rect.right()
        top, bottom = rect.top(), rect.bottom()
        near_left = abs(x - left) <= s
        near_right = abs(x - right) <= s
        near_top = abs(y - top) <= s
        near_bottom = abs(y - bottom) <= s
        if near_left and near_top:
            return "tl"
        if near_right and near_top:
            return "tr"
        if near_left and near_bottom:
            return "bl"
        if near_right and near_bottom:
            return "br"
        if near_left:
            return "l"
        if near_right:
            return "r"
        if near_top:
            return "t"
        if near_bottom:
            return "b"
        return "move"

    def mousePressEvent(self, event):
        self._dialog.select_zone(self._row)
        self._mode = self._handle_at(event.pos())
        self._start_scene = event.scenePos()
        self._start_rect = QRectF(self.rect())
        event.accept()

    def mouseMoveEvent(self, event):
        if self._mode is None or self._start_scene is None or self._start_rect is None:
            event.accept()
            return
        dx = float(event.scenePos().x() - self._start_scene.x())
        dy = float(event.scenePos().y() - self._start_scene.y())
        rect = QRectF(self._start_rect)
        mode = self._mode
        if mode == "move":
            rect.translate(dx, dy)
        else:
            if "l" in mode:
                rect.setLeft(rect.left() + dx)
            if "r" in mode:
                rect.setRight(rect.right() + dx)
            if "t" in mode:
                rect.setTop(rect.top() + dy)
            if "b" in mode:
                rect.setBottom(rect.bottom() + dy)
        rect = self._dialog.clamp_rect(rect)
        self.setRect(rect)
        self._dialog.set_zone_rect(self._row, rect, redraw=False)
        event.accept()

    def mouseReleaseEvent(self, event):
        self._mode = None
        self._start_scene = None
        self._start_rect = None
        self._dialog.redraw_zones()
        event.accept()

    def paint(self, painter, option, widget=None):
        super().paint(painter, option, widget)
        try:
            rect = self.rect()
            s = self.HANDLE_SIZE
            painter.save()
            painter.setPen(QPen(QColor(0, 0, 0), 1))
            painter.setBrush(QBrush(QColor(255, 255, 255)))
            for x, y in (
                (rect.left(), rect.top()),
                (rect.right(), rect.top()),
                (rect.left(), rect.bottom()),
                (rect.right(), rect.bottom()),
            ):
                painter.drawRect(QRectF(x - s / 2.0, y - s / 2.0, s, s))
            painter.restore()
        except Exception:
            pass


class _BKExportZoneGraphicsView(QGraphicsView):
    def __init__(self, parent_dialog):
        super().__init__(parent_dialog)
        self._dialog = parent_dialog
        self._start = None
        self._rubber = None
        self.setMouseTracking(True)
        try:
            self.setDragMode(QGraphicsView.NoDrag)
        except Exception:
            pass

    def _zone_item_at(self, event):
        try:
            pos = self.mapToScene(event.position().toPoint() if hasattr(event, "position") else event.pos())
            items = self.scene().items(pos)
            for item in items:
                if isinstance(item, _BKExportZoneRectItem):
                    return item
        except Exception:
            pass
        return None

    def mousePressEvent(self, event):
        if event.button() == Qt.LeftButton and self._zone_item_at(event) is None:
            self._start = self.mapToScene(event.position().toPoint() if hasattr(event, "position") else event.pos())
            if self._rubber is not None:
                try:
                    self.scene().removeItem(self._rubber)
                except Exception:
                    pass
                self._rubber = None
            self._rubber = QGraphicsRectItem(QRectF(self._start, self._start))
            self._rubber.setPen(QPen(QColor(40, 130, 255), 2, Qt.DashLine))
            self._rubber.setBrush(QBrush(QColor(40, 130, 255, 35)))
            self.scene().addItem(self._rubber)
            event.accept()
            return
        super().mousePressEvent(event)

    def mouseMoveEvent(self, event):
        if self._start is not None and self._rubber is not None:
            pos = self.mapToScene(event.position().toPoint() if hasattr(event, "position") else event.pos())
            self._rubber.setRect(QRectF(self._start, pos).normalized())
            event.accept()
            return
        super().mouseMoveEvent(event)

    def mouseReleaseEvent(self, event):
        if self._start is not None and self._rubber is not None:
            rect = self._rubber.rect().normalized()
            try:
                self.scene().removeItem(self._rubber)
            except Exception:
                pass
            self._rubber = None
            self._start = None
            if rect.width() >= 6 and rect.height() >= 6:
                self._dialog.add_zone_from_rect(rect)
            event.accept()
            return
        super().mouseReleaseEvent(event)

    def resizeEvent(self, event):
        super().resizeEvent(event)
        try:
            self.fitInView(self.sceneRect(), Qt.KeepAspectRatio)
        except Exception:
            pass


class _BKExportZonesDialog(QDialog):
    def __init__(self, window, task, image_path, zones=None):
        super().__init__(window)
        self._window = window
        self._task = task
        self._image_path = image_path
        self._zones = [_bk_export_clean_zone(z, i) for i, z in enumerate(zones or [])]
        self._zones = [z for z in self._zones if z]
        self._items = []
        self.setWindowTitle(_bk_tab_tr(window, "export_zones_title"))
        self.setModal(True)
        self.resize(1480, 940)
        try:
            self.setWindowState(self.windowState() | Qt.WindowMaximized)
        except Exception:
            pass

        root = QHBoxLayout(self)
        root.setContentsMargins(14, 14, 14, 14)
        root.setSpacing(12)

        self.scene = QGraphicsScene(self)
        self.view = _BKExportZoneGraphicsView(self)
        self.view.setScene(self.scene)
        pixmap = QPixmap(str(image_path or "")) if QPixmap is not None else None
        self._image_w = 0
        self._image_h = 0
        if pixmap is not None and not pixmap.isNull():
            self._image_w = int(pixmap.width())
            self._image_h = int(pixmap.height())
            self.scene.addPixmap(pixmap)
            self.scene.setSceneRect(QRectF(0, 0, pixmap.width(), pixmap.height()))
        root.addWidget(self.view, 3)

        side = QVBoxLayout()
        side.setSpacing(8)
        intro = QLabel(_bk_tab_tr(window, "export_zones_intro"), self)
        intro.setWordWrap(True)
        side.addWidget(intro)

        self.table = QTableWidget(0, 3, self)
        self.table.setMinimumWidth(720)
        self.table.setHorizontalHeaderLabels([
            _bk_tab_tr(window, "export_zones_col_name"),
            _bk_tab_tr(window, "export_zones_col_type"),
            _bk_tab_tr(window, "export_zones_col_rect"),
        ])
        try:
            self.table.setSelectionBehavior(QAbstractItemView.SelectRows)
            self.table.setSelectionMode(QAbstractItemView.SingleSelection)
            self.table.horizontalHeader().setStretchLastSection(False)
            self.table.setColumnWidth(0, 230)
            self.table.setColumnWidth(1, 260)
            self.table.setColumnWidth(2, 210)
        except Exception:
            pass
        side.addWidget(self.table, 1)

        button_row1 = QHBoxLayout()
        self.btn_up = QPushButton(_bk_tab_tr(window, "export_zones_up"), self)
        self.btn_down = QPushButton(_bk_tab_tr(window, "export_zones_down"), self)
        button_row1.addWidget(self.btn_up)
        button_row1.addWidget(self.btn_down)
        side.addLayout(button_row1)

        button_row2 = QHBoxLayout()
        self.btn_delete = QPushButton(_bk_tab_tr(window, "export_zones_delete"), self)
        self.btn_clear = QPushButton(_bk_tab_tr(window, "export_zones_clear"), self)
        button_row2.addWidget(self.btn_delete)
        button_row2.addWidget(self.btn_clear)
        side.addLayout(button_row2)

        buttons = QDialogButtonBox(QDialogButtonBox.Ok | QDialogButtonBox.Cancel, self)
        try:
            buttons.button(QDialogButtonBox.Ok).setText(_bk_tab_tr(window, "btn_ok"))
            buttons.button(QDialogButtonBox.Cancel).setText(_bk_tab_tr(window, "btn_cancel"))
        except Exception:
            pass
        buttons.accepted.connect(self.accept)
        buttons.rejected.connect(self.reject)
        side.addWidget(buttons)
        root.addLayout(side, 2)

        self.btn_delete.clicked.connect(self.delete_selected)
        self.btn_clear.clicked.connect(self.clear_zones)
        self.btn_up.clicked.connect(lambda: self.move_selected(-1))
        self.btn_down.clicked.connect(lambda: self.move_selected(1))
        self.table.itemChanged.connect(self._table_item_changed)
        self.table.itemDoubleClicked.connect(self._table_item_double_clicked)
        try:
            self.table.itemSelectionChanged.connect(self._selected_zone_changed)
        except Exception:
            pass
        self.refresh_table()

    def clamp_rect(self, rect):
        rect = QRectF(rect).normalized()
        min_size = 3.0
        if rect.width() < min_size:
            rect.setWidth(min_size)
        if rect.height() < min_size:
            rect.setHeight(min_size)
        max_w = float(getattr(self, "_image_w", 0) or 0)
        max_h = float(getattr(self, "_image_h", 0) or 0)
        if max_w > 0:
            if rect.left() < 0:
                rect.translate(-rect.left(), 0)
            if rect.right() > max_w:
                rect.translate(max_w - rect.right(), 0)
            rect.setLeft(max(0.0, rect.left()))
            rect.setRight(min(max_w, max(rect.left() + min_size, rect.right())))
        if max_h > 0:
            if rect.top() < 0:
                rect.translate(0, -rect.top())
            if rect.bottom() > max_h:
                rect.translate(0, max_h - rect.bottom())
            rect.setTop(max(0.0, rect.top()))
            rect.setBottom(min(max_h, max(rect.top() + min_size, rect.bottom())))
        return rect.normalized()

    def _zone_type_combo(self, row, zone_type):
        combo = QComboBox(self.table)
        for value in _BK_ZONE_TYPES:
            combo.addItem(_bk_export_zone_title(self._window, value), value)
        idx = combo.findData(zone_type if zone_type in _BK_ZONE_TYPES else "names")
        combo.setCurrentIndex(max(0, idx))
        combo.currentIndexChanged.connect(lambda _=0, r=row, c=combo: self._combo_changed(r, c))
        return combo

    def _combo_changed(self, row, combo):
        if 0 <= row < len(self._zones):
            self._zones[row]["type"] = str(combo.currentData() or "names")
            self.redraw_zones()

    def _table_item_changed(self, item):
        if item is None:
            return
        row = item.row()
        if not (0 <= row < len(self._zones)):
            return
        if item.column() == 0:
            self._zones[row]["name"] = str(item.text() or "").strip() or f"Bereich {row + 1}"
            self.redraw_zones()
        elif item.column() == 2:
            rect = self.parse_rect_text(str(item.text() or ""), self._zones[row])
            if rect is None:
                item.setText(self._format_zone_rect(self._zones[row]))
                return
            self.set_zone_rect(row, rect, redraw=True)

    def _table_item_double_clicked(self, item):
        if item is not None and item.column() == 2:
            self.table.editItem(item)

    def parse_rect_text(self, text, fallback_zone=None):
        nums = re.findall(r"-?\d+(?:[.,]\d+)?", str(text or ""))
        if len(nums) < 4:
            return None
        try:
            x = float(nums[0].replace(",", "."))
            y = float(nums[1].replace(",", "."))
            w = float(nums[2].replace(",", "."))
            h = float(nums[3].replace(",", "."))
        except Exception:
            return None
        if w < 0 or h < 0:
            return None
        return self.clamp_rect(QRectF(x, y, max(3.0, w), max(3.0, h)))

    def add_zone_from_rect(self, rect):
        zone = _bk_export_clean_zone({
            "x0": rect.left(),
            "y0": rect.top(),
            "x1": rect.right(),
            "y1": rect.bottom(),
            "type": "names",
            "name": _bk_tab_tr(self._window, "export_zone_default_name").format(len(self._zones) + 1),
            "order": len(self._zones),
        }, len(self._zones))
        if zone:
            self._zones.append(zone)
            self.refresh_table(select_row=len(self._zones) - 1)

    def _format_zone_rect(self, zone):
        return f'{int(round(zone["x0"]))} / {int(round(zone["y0"]))} / {int(round(zone["x1"] - zone["x0"]))} × {int(round(zone["y1"] - zone["y0"]))}'

    def set_zone_rect(self, row, rect, redraw=True):
        if not (0 <= row < len(self._zones)):
            return
        rect = self.clamp_rect(rect)
        self._zones[row].update({
            "x0": float(rect.left()),
            "y0": float(rect.top()),
            "x1": float(rect.right()),
            "y1": float(rect.bottom()),
        })
        try:
            self.table.blockSignals(True)
            item = self.table.item(row, 2)
            if item is not None:
                item.setText(self._format_zone_rect(self._zones[row]))
        finally:
            try:
                self.table.blockSignals(False)
            except Exception:
                pass
        if redraw:
            self.redraw_zones()

    def select_zone(self, row):
        if 0 <= row < len(self._zones):
            try:
                self.table.selectRow(row)
            except Exception:
                pass
            self.redraw_zones()

    def _selected_zone_changed(self):
        self.redraw_zones()

    def refresh_table(self, select_row=-1):
        self.table.blockSignals(True)
        self.table.setRowCount(0)
        for row, zone in enumerate(self._zones):
            zone["order"] = row
            self.table.insertRow(row)
            name_item = QTableWidgetItem(str(zone.get("name", f"Bereich {row + 1}")))
            rect_item = QTableWidgetItem(self._format_zone_rect(zone))
            rect_item.setFlags(Qt.ItemIsEnabled | Qt.ItemIsSelectable | Qt.ItemIsEditable)
            self.table.setItem(row, 0, name_item)
            self.table.setCellWidget(row, 1, self._zone_type_combo(row, zone.get("type", "names")))
            self.table.setItem(row, 2, rect_item)
        self.table.blockSignals(False)
        if 0 <= select_row < self.table.rowCount():
            self.table.selectRow(select_row)
        self.redraw_zones()

    def redraw_zones(self):
        for item in list(self._items):
            try:
                self.scene.removeItem(item)
            except Exception:
                pass
        self._items = []
        colors = {
            "heading": QColor(0, 170, 80, 90),
            "names": QColor(40, 130, 255, 85),
            "age": QColor(130, 90, 210, 80),
            "years": QColor(255, 135, 0, 85),
            "places": QColor(0, 150, 170, 80),
            "page_numbers": QColor(255, 205, 0, 95),
            "other": QColor(120, 120, 120, 70),
            "ignore": QColor(220, 20, 20, 65),
        }
        pens = {
            "heading": QColor(0, 145, 80),
            "names": QColor(40, 130, 255),
            "age": QColor(110, 70, 180),
            "years": QColor(215, 115, 0),
            "places": QColor(0, 125, 150),
            "page_numbers": QColor(210, 150, 0),
            "other": QColor(100, 100, 100),
            "ignore": QColor(210, 20, 20),
        }
        selected = self.selected_row()
        for row, zone in enumerate(self._zones):
            rect = QRectF(float(zone["x0"]), float(zone["y0"]), float(zone["x1"] - zone["x0"]), float(zone["y1"] - zone["y0"]))
            ztype = zone.get("type", "names")
            item = _BKExportZoneRectItem(self, row, rect, colors.get(ztype, colors["names"]), pens.get(ztype, pens["names"]))
            if row == selected:
                item.setPen(QPen(pens.get(ztype, pens["names"]), 4))
                item.setZValue(100 + row)
            self.scene.addItem(item)
            self._items.append(item)

    def selected_row(self):
        try:
            rows = self.table.selectionModel().selectedRows()
            if rows:
                return rows[0].row()
        except Exception:
            pass
        return -1

    def delete_selected(self):
        row = self.selected_row()
        if 0 <= row < len(self._zones):
            self._zones.pop(row)
            self.refresh_table(select_row=min(row, len(self._zones) - 1))

    def clear_zones(self):
        self._zones = []
        self.refresh_table()

    def move_selected(self, delta):
        row = self.selected_row()
        new_row = row + int(delta)
        if 0 <= row < len(self._zones) and 0 <= new_row < len(self._zones):
            self._zones[row], self._zones[new_row] = self._zones[new_row], self._zones[row]
            self.refresh_table(select_row=new_row)

    def zones(self):
        out = []
        for row, zone in enumerate(self._zones):
            clean = _bk_export_clean_zone(zone, row)
            if clean:
                clean["order"] = len(out)
                out.append(clean)
        return out

try:
    __all__.extend([
        '_BKExportZoneRectItem',
        '_bk_zone_age_value',
        '_bk_zone_page_number_value',
        '_bk_zone_place_value',
        '_bk_zone_year_value',
    ])
    register_globals('bk', globals(), sorted(set(__all__)))
except Exception:
    pass

# ---------------------------------------------------------------------------
# Tabellenbereiche: Spalten-Datentypen wie Exportspalten + echte Griffpunkte
# ---------------------------------------------------------------------------
_BK_ZONE_TYPES = tuple(_BK_TABULAR_KEYS + ["ignore"])
_BK_ZONE_LEGACY_TYPES = {
    "data": "original_line",
    "names": "original_line",
    "other": "original_line",
    "heading": "year_resolved",
    "year": "year_resolved",
    "years": "year_resolved",
    "age": "age_original",
    "places": "place_in_source",
    "page_numbers": "number",
    "number": "number",
}






def _bk_zone_contains_record(zone, record) -> bool:
    if not zone or not record.get("bbox"):
        return False
    try:
        zx0, zy0, zx1, zy1 = float(zone["x0"]), float(zone["y0"]), float(zone["x1"]), float(zone["y1"])
        rx0 = float(record.get("x0", 0.0) or 0.0)
        ry0 = float(record.get("y0", 0.0) or 0.0)
        rx1 = float(record.get("x1", rx0) or rx0)
        ry1 = float(record.get("y1", ry0) or ry0)
    except Exception:
        return False
    cx = (rx0 + rx1) / 2.0
    cy = (ry0 + ry1) / 2.0
    if zx0 <= cx <= zx1 and zy0 <= cy <= zy1:
        return True
    ix0, iy0 = max(zx0, rx0), max(zy0, ry0)
    ix1, iy1 = min(zx1, rx1), min(zy1, ry1)
    if ix1 <= ix0 or iy1 <= iy0:
        return False
    # Sobald eine Overlay-Box sichtbar in einem Bereich liegt, zählt sie zu diesem Bereich.
    return ((ix1 - ix0) * (iy1 - iy0)) >= 1.0








# Griffpunkt-Fix: Bei Mausklick auf einen Bereich darf der Bereich nicht sofort neu gezeichnet werden,
# sonst wird das gerade angeklickte QGraphicsItem entfernt, bevor Drag/Resize beginnen kann.
def _bk_zone_rect_mouse_press_final(self, event):
    try:
        self._dialog._suppress_zone_redraw = True
        self._dialog.table.selectRow(self._row)
    except Exception:
        pass
    finally:
        try:
            self._dialog._suppress_zone_redraw = False
        except Exception:
            pass
    self._mode = self._handle_at(event.pos())
    self._start_scene = event.scenePos()
    self._start_rect = QRectF(self.rect())
    event.accept()


def _bk_zone_selected_changed_final(self):
    if bool(getattr(self, "_suppress_zone_redraw", False)):
        return
    self.redraw_zones()


def _bk_zone_dialog_init_final(self, window, task, image_path, zones=None):
    QDialog.__init__(self, window)
    self._window = window
    self._task = task
    self._image_path = image_path
    self._zones = [_bk_export_clean_zone(z, i) for i, z in enumerate(zones or [])]
    self._zones = [z for z in self._zones if z]
    self._items = []
    self._suppress_zone_redraw = False
    self.setWindowTitle(_bk_tab_tr(window, "export_zones_title"))
    self.setModal(True)
    self.resize(1520, 940)
    try:
        self.setWindowState(self.windowState() | Qt.WindowMaximized)
    except Exception:
        pass

    root = QHBoxLayout(self)
    root.setContentsMargins(14, 14, 14, 14)
    root.setSpacing(12)

    self.scene = QGraphicsScene(self)
    self.view = _BKExportZoneGraphicsView(self)
    self.view.setScene(self.scene)
    pixmap = QPixmap(str(image_path or "")) if QPixmap is not None else None
    self._image_w = 0
    self._image_h = 0
    if pixmap is not None and not pixmap.isNull():
        self._image_w = int(pixmap.width())
        self._image_h = int(pixmap.height())
        self.scene.addPixmap(pixmap)
        self.scene.setSceneRect(QRectF(0, 0, pixmap.width(), pixmap.height()))
    root.addWidget(self.view, 3)

    side = QVBoxLayout()
    side.setSpacing(8)
    intro = QLabel(_bk_tab_tr(
        window,
        "export_zones_intro_column_types",
    ), self)
    intro.setWordWrap(True)
    side.addWidget(intro)

    self.table = QTableWidget(0, 3, self)
    self.table.setMinimumWidth(760)
    self.table.setHorizontalHeaderLabels([
        _bk_tab_tr(window, "export_zones_col_name"),
        _bk_tab_tr(window, "export_zones_col_type"),
        _bk_tab_tr(window, "export_zones_col_rect"),
    ])
    try:
        self.table.setSelectionBehavior(QAbstractItemView.SelectRows)
        self.table.setSelectionMode(QAbstractItemView.SingleSelection)
        self.table.horizontalHeader().setStretchLastSection(False)
        self.table.setColumnWidth(0, 230)
        self.table.setColumnWidth(1, 260)
        self.table.setColumnWidth(2, 250)
    except Exception:
        pass
    side.addWidget(self.table, 1)

    button_row1 = QHBoxLayout()
    self.btn_up = QPushButton(_bk_tab_tr(window, "export_zones_up"), self)
    self.btn_down = QPushButton(_bk_tab_tr(window, "export_zones_down"), self)
    button_row1.addWidget(self.btn_up)
    button_row1.addWidget(self.btn_down)
    side.addLayout(button_row1)

    button_row2 = QHBoxLayout()
    self.btn_delete = QPushButton(_bk_tab_tr(window, "export_zones_delete"), self)
    self.btn_clear = QPushButton(_bk_tab_tr(window, "export_zones_clear"), self)
    button_row2.addWidget(self.btn_delete)
    button_row2.addWidget(self.btn_clear)
    side.addLayout(button_row2)

    buttons = QDialogButtonBox(QDialogButtonBox.Ok | QDialogButtonBox.Cancel, self)
    try:
        buttons.button(QDialogButtonBox.Ok).setText(_bk_tab_tr(window, "btn_ok"))
        buttons.button(QDialogButtonBox.Cancel).setText(_bk_tab_tr(window, "btn_cancel"))
    except Exception:
        pass
    buttons.accepted.connect(self.accept)
    buttons.rejected.connect(self.reject)
    side.addWidget(buttons)
    root.addLayout(side, 2)

    self.btn_delete.clicked.connect(self.delete_selected)
    self.btn_clear.clicked.connect(self.clear_zones)
    self.btn_up.clicked.connect(lambda: self.move_selected(-1))
    self.btn_down.clicked.connect(lambda: self.move_selected(1))
    self.table.itemChanged.connect(self._table_item_changed)
    self.table.itemDoubleClicked.connect(self._table_item_double_clicked)
    try:
        self.table.itemSelectionChanged.connect(self._selected_zone_changed)
    except Exception:
        pass
    self.refresh_table()




def _bk_zone_add_from_rect_final(self, rect):
    zone = _bk_export_clean_zone({
        "x0": rect.left(),
        "y0": rect.top(),
        "x1": rect.right(),
        "y1": rect.bottom(),
        "type": "original_line",
        "name": _bk_tab_tr(self._window, "export_zone_default_name").format(len(self._zones) + 1),
        "order": len(self._zones),
    }, len(self._zones))
    if zone:
        self._zones.append(zone)
        self.refresh_table(select_row=len(self._zones) - 1)


def _bk_zone_redraw_final(self):
    for item in list(self._items):
        try:
            self.scene.removeItem(item)
        except Exception:
            pass
    self._items = []
    colors = {
        "family_name": QColor(40, 130, 255, 80),
        "given_names": QColor(90, 165, 255, 80),
        "relationship": QColor(130, 90, 210, 80),
        "age_original": QColor(165, 90, 210, 80),
        "date_original": QColor(255, 140, 0, 85),
        "year_resolved": QColor(0, 170, 80, 85),
        "place_in_source": QColor(0, 150, 170, 80),
        "number": QColor(255, 205, 0, 95),
        "original_line": QColor(40, 130, 255, 85),
        "ignore": QColor(220, 20, 20, 65),
    }
    pens = {
        "family_name": QColor(40, 130, 255),
        "given_names": QColor(80, 145, 235),
        "relationship": QColor(110, 70, 180),
        "age_original": QColor(145, 70, 180),
        "date_original": QColor(215, 115, 0),
        "year_resolved": QColor(0, 145, 80),
        "place_in_source": QColor(0, 125, 150),
        "number": QColor(210, 150, 0),
        "original_line": QColor(40, 130, 255),
        "ignore": QColor(210, 20, 20),
    }
    selected = self.selected_row()
    for row, zone in enumerate(self._zones):
        rect = QRectF(float(zone["x0"]), float(zone["y0"]), float(zone["x1"] - zone["x0"]), float(zone["y1"] - zone["y0"]))
        ztype = zone.get("type", "original_line")
        item = _BKExportZoneRectItem(self, row, rect, colors.get(ztype, colors["original_line"]), pens.get(ztype, pens["original_line"]))
        if row == selected:
            item.setPen(QPen(pens.get(ztype, pens["original_line"]), 4))
            item.setZValue(100 + row)
        self.scene.addItem(item)
        self._items.append(item)


def _bk_zone_refresh_table_final(self, select_row=-1):
    self.table.blockSignals(True)
    self.table.setRowCount(0)
    for row, zone in enumerate(self._zones):
        zone["order"] = row
        self.table.insertRow(row)
        name_item = QTableWidgetItem(str(zone.get("name", f"Bereich {row + 1}")))
        rect_item = QTableWidgetItem(self._format_zone_rect(zone))
        rect_item.setFlags(Qt.ItemIsEnabled | Qt.ItemIsSelectable | Qt.ItemIsEditable)
        self.table.setItem(row, 0, name_item)
        self.table.setCellWidget(row, 1, self._zone_type_combo(row, zone.get("type", "original_line")))
        self.table.setItem(row, 2, rect_item)
    self.table.blockSignals(False)
    if 0 <= select_row < self.table.rowCount():
        self.table.selectRow(select_row)
    self.redraw_zones()


def _bk_zone_table_item_changed_final(self, item):
    if item is None:
        return
    row = item.row()
    if not (0 <= row < len(self._zones)):
        return
    if item.column() == 0:
        self._zones[row]["name"] = str(item.text() or "").strip() or f"Bereich {row + 1}"
        self.redraw_zones()
    elif item.column() == 2:
        rect = self.parse_rect_text(str(item.text() or ""), self._zones[row])
        if rect is None:
            item.setText(self._format_zone_rect(self._zones[row]))
            return
        self.set_zone_rect(row, rect, redraw=True)


def _bk_zone_column_choice_dialog_final(self, fmt=None, include_text_modes=False):
    dlg = QDialog(self)
    dlg.setWindowTitle(_bk_tab_tr(self, "export_text_layout_title" if include_text_modes else "export_table_columns_title"))
    dlg.setModal(True)
    layout = QVBoxLayout(dlg)
    layout.setContentsMargins(16, 14, 16, 14)
    layout.setSpacing(10)

    intro_key = "export_text_layout_intro_extended" if include_text_modes else "export_table_columns_intro"
    label = QLabel(_bk_tab_tr(self, intro_key), dlg)
    label.setWordWrap(True)
    layout.addWidget(label)

    rb_original = rb_lines = rb_table = None
    if include_text_modes:
        if QGroupBox is not None:
            mode_box = QGroupBox(_bk_tab_tr(self, "export_layout_mode_group"), dlg)
            mode_layout = QVBoxLayout(mode_box)
            mode_layout.setContentsMargins(12, 10, 12, 10)
        else:
            mode_box = None
            mode_layout = QVBoxLayout()
        rb_original = QRadioButton(_bk_tab_tr(self, "export_text_layout_original"), dlg)
        rb_lines = QRadioButton(_bk_tab_tr(self, "export_text_layout_lines"), dlg)
        rb_table = QRadioButton(_bk_tab_tr(self, "export_text_layout_table"), dlg)
        mode = str(getattr(self, "_bk_export_text_layout_mode", "original") or "original").lower()
        rb_original.setChecked(mode not in {"lines", "table"})
        rb_lines.setChecked(mode == "lines")
        rb_table.setChecked(mode == "table")
        mode_layout.addWidget(rb_original)
        mode_layout.addWidget(rb_lines)
        mode_layout.addWidget(rb_table)
        if mode_box is not None:
            layout.addWidget(mode_box)
        else:
            layout.addLayout(mode_layout)

    selected_keys = _bk_load_saved_column_keys(self)
    checkboxes = {}
    if QGroupBox is not None:
        columns_box = QGroupBox(_bk_tab_tr(self, "export_table_columns_label"), dlg)
        columns_layout = QVBoxLayout(columns_box)
        columns_layout.setContentsMargins(12, 10, 12, 10)
    else:
        columns_box = None
        columns_layout = QVBoxLayout()
        columns_layout.addWidget(QLabel(_bk_tab_tr(self, "export_table_columns_label"), dlg))

    if QGridLayout is not None:
        checkbox_grid = QGridLayout()
        checkbox_grid.setContentsMargins(0, 0, 0, 0)
        checkbox_grid.setHorizontalSpacing(22)
        checkbox_grid.setVerticalSpacing(6)
        for idx, key in enumerate(_BK_TABULAR_KEYS):
            cb = QCheckBox(_bk_tabular_column_title(self, key), dlg)
            cb.setChecked(key in selected_keys)
            checkboxes[key] = cb
            checkbox_grid.addWidget(cb, idx // 3, idx % 3)
        columns_layout.addLayout(checkbox_grid)
    else:
        for key in _BK_TABULAR_KEYS:
            cb = QCheckBox(_bk_tabular_column_title(self, key), dlg)
            cb.setChecked(key in selected_keys)
            checkboxes[key] = cb
            columns_layout.addWidget(cb)

    quick_row = QHBoxLayout()
    btn_all = QPushButton(_bk_tab_tr(self, "export_table_columns_all"), dlg)
    btn_none = QPushButton(_bk_tab_tr(self, "export_table_columns_none_button"), dlg)
    btn_remember = QPushButton(_bk_tab_tr(self, "export_table_columns_remember"), dlg)
    quick_row.addWidget(btn_all)
    quick_row.addWidget(btn_none)
    quick_row.addWidget(btn_remember)
    quick_row.addStretch(1)
    columns_layout.addLayout(quick_row)
    if columns_box is not None:
        layout.addWidget(columns_box)
    else:
        layout.addLayout(columns_layout)

    # Unabhängig von der Tabellen-Darstellung: Bereiche können immer definiert werden.
    zone_box = QGroupBox(_bk_tab_tr(self, "export_zones_group"), dlg) if QGroupBox is not None else None
    zone_layout = QHBoxLayout(zone_box) if zone_box is not None else QHBoxLayout()
    cb_zones = QCheckBox(_bk_tab_tr(self, "export_table_use_zones"), dlg)
    try:
        remembered_zones = self.settings.value("export/table_use_zones", bool(getattr(self, "_bk_export_use_zones", False)), type=bool)
    except Exception:
        remembered_zones = bool(getattr(self, "_bk_export_use_zones", False))
    cb_zones.setChecked(bool(remembered_zones))
    btn_zones = QPushButton(_bk_tab_tr(self, "export_table_define_zones"), dlg)
    zone_layout.addWidget(cb_zones, 1)
    zone_layout.addWidget(btn_zones)
    if zone_box is not None:
        layout.addWidget(zone_box)
    else:
        layout.addLayout(zone_layout)

    def current_checked_keys():
        return [key for key, cb in checkboxes.items() if cb.isChecked()]

    def set_all():
        for cb in checkboxes.values():
            cb.setChecked(True)

    def set_none():
        for cb in checkboxes.values():
            cb.setChecked(False)

    result = {"mode": "table", "columns": selected_keys, "remembered": False, "use_zones": bool(cb_zones.isChecked())}

    def remember_selection():
        keys = current_checked_keys()
        if not keys:
            QMessageBox.warning(dlg, _bk_tab_tr(self, "warn_title"), _bk_tab_tr(self, "export_table_columns_none"))
            return
        result["remembered"] = True
        result["columns"] = _bk_save_column_keys(self, keys)
        try:
            self.settings.setValue("export/table_use_zones", bool(cb_zones.isChecked()))
            self.settings.sync()
        except Exception:
            pass

    def define_zones():
        try:
            task = self._current_task()
        except Exception:
            task = None
        was_enabled = True
        try:
            was_enabled = dlg.isEnabled()
            dlg.setEnabled(False)
            # Nicht verstecken: sonst kann der Exportdialog nach dem Bereichsdialog
            # verloren gehen und der eigentliche Export wirkt "verschwunden".
            dlg.lower()
            QApplication.processEvents()
        except Exception:
            pass
        try:
            zones = _bk_open_export_zones_dialog(self, task)
        finally:
            try:
                dlg.setEnabled(was_enabled)
                dlg.show()
                dlg.raise_()
                dlg.activateWindow()
                QApplication.processEvents()
            except Exception:
                pass
        if zones is not None:
            cb_zones.setChecked(bool(zones))

    btn_all.clicked.connect(set_all)
    btn_none.clicked.connect(set_none)
    btn_remember.clicked.connect(remember_selection)
    btn_zones.clicked.connect(define_zones)

    def sync_columns_enabled():
        enabled = True
        if include_text_modes and rb_table is not None:
            enabled = bool(rb_table.isChecked())
        target = columns_box if columns_box is not None else None
        if target is not None:
            target.setEnabled(enabled)
        else:
            for cb in checkboxes.values():
                cb.setEnabled(enabled)
            btn_all.setEnabled(enabled)
            btn_none.setEnabled(enabled)
            btn_remember.setEnabled(enabled)
        # cb_zones und btn_zones bleiben bewusst immer aktiv.

    if include_text_modes:
        try:
            rb_original.toggled.connect(sync_columns_enabled)
            rb_lines.toggled.connect(sync_columns_enabled)
            rb_table.toggled.connect(sync_columns_enabled)
        except Exception:
            pass
        sync_columns_enabled()

    buttons = QDialogButtonBox(QDialogButtonBox.Ok | QDialogButtonBox.Cancel, dlg)
    try:
        buttons.button(QDialogButtonBox.Ok).setText(_bk_tab_tr(self, "btn_ok"))
        buttons.button(QDialogButtonBox.Cancel).setText(_bk_tab_tr(self, "btn_cancel"))
    except Exception:
        pass

    def cancel_dialog():
        result["cancelled"] = True
        try:
            dlg.done(QDialog.Rejected)
        except Exception:
            try:
                dlg.reject()
            except Exception:
                pass

    def accept_checked():
        result["cancelled"] = False
        mode = "table"
        if include_text_modes:
            if rb_lines and rb_lines.isChecked():
                mode = "lines"
            elif rb_original and rb_original.isChecked():
                mode = "original"
        keys = current_checked_keys()
        if mode == "table" and not keys:
            QMessageBox.warning(dlg, _bk_tab_tr(self, "warn_title"), _bk_tab_tr(self, "export_table_columns_none"))
            return
        result["mode"] = mode
        result["columns"] = _bk_normalize_column_keys(keys)
        result["use_zones"] = bool(cb_zones.isChecked())
        try:
            self._bk_export_use_zones = result["use_zones"]
        except Exception:
            pass
        dlg.accept()

    buttons.accepted.connect(accept_checked)
    buttons.rejected.connect(cancel_dialog)
    layout.addWidget(buttons)
    dlg.setMinimumSize(620, 430 if include_text_modes else 360)
    try:
        dlg.resize(700, 500 if include_text_modes else 410)
    except Exception:
        pass
    if dlg.exec() != QDialog.Accepted:
        return None
    return result


try:
    _BKExportZoneRectItem.mousePressEvent = _bk_zone_rect_mouse_press_final
    _BKExportZonesDialog.__init__ = _bk_zone_dialog_init_final
    _BKExportZonesDialog.add_zone_from_rect = _bk_zone_add_from_rect_final
    _BKExportZonesDialog.redraw_zones = _bk_zone_redraw_final
    _BKExportZonesDialog.refresh_table = _bk_zone_refresh_table_final
    _BKExportZonesDialog._table_item_changed = _bk_zone_table_item_changed_final
    _BKExportZonesDialog._selected_zone_changed = _bk_zone_selected_changed_final
    _bk_column_choice_dialog = _bk_zone_column_choice_dialog_final
except Exception:
    pass

try:
    __all__.extend([
        '_bk_zone_column_choice_dialog_final',
        '_bk_zone_value_for_column',
        '_bk_zone_rect_mouse_press_final',
    ])
    register_globals('bk', globals(), sorted(set(__all__)))
except Exception:
    pass

# ---------------------------------------------------------------------------
# Tabellenbereiche: erweiterte genealogische Datentypen + globale sensible Bereiche
# ---------------------------------------------------------------------------
_BK_EXTENDED_TABULAR_COLUMNS = [
    ("given_names", "export_column_given_names", "Vornamen", 22.0, 3.4),
    ("family_name", "export_column_family_name", "Familiennamen", 18.0, 2.8),
    ("full_name", "export_column_full_name", "Vor- und Familiennamen", 28.0, 4.6),
    ("middle_names", "export_column_middle_names", "Zweit- und Drittnamen", 24.0, 4.0),
    ("days", "export_column_days", "Tage", 10.0, 1.8),
    ("months", "export_column_months", "Monate", 10.0, 1.8),
    ("year_resolved", "export_column_year", "Jahre", 12.0, 2.0),
    ("date_original", "export_column_birth_date", "Geburtsdatum", 16.0, 2.8),
    ("marriage_date", "export_column_marriage_date", "Hochzeitsdatum", 17.0, 3.0),
    ("death_date", "export_column_death_date", "Sterbedatum", 16.0, 2.8),
    ("place_in_source", "export_column_place", "Orte", 20.0, 3.2),
    ("residence_place", "export_column_residence_place", "Wohnorte", 18.0, 3.0),
    ("birth_place", "export_column_birth_place", "Geburtsorte", 18.0, 3.0),
    ("death_place", "export_column_death_place", "Sterbeorte", 18.0, 3.0),
    ("street", "export_column_street", "Straßen", 18.0, 3.0),
    ("house_number", "export_column_house_number", "Hausnummern", 12.0, 2.0),
    ("postal_code", "export_column_postal_code", "Postleitzahlen", 14.0, 2.4),
    ("occupation", "export_column_occupation", "Berufe", 18.0, 3.0),
    ("spouse", "export_column_spouse", "Ehepartner", 20.0, 3.2),
    ("partner", "export_column_partner", "Lebenspartner", 20.0, 3.2),
    ("children", "export_column_children", "Kinder", 20.0, 3.2),
    ("grandchildren", "export_column_grandchildren", "Enkel", 20.0, 3.2),
    ("grandparents", "export_column_grandparents", "Großeltern", 20.0, 3.2),
    ("great_grandparents", "export_column_great_grandparents", "Ur-Großeltern", 22.0, 3.6),
    ("great_great_grandparents", "export_column_great_great_grandparents", "Ur-Ur-Großeltern", 24.0, 4.0),
    ("relationship", "export_column_relationship", "Zusatz", 24.0, 4.0),
    ("other", "export_column_other", "Sonstiges", 24.0, 4.0),
    ("unknown", "export_column_unknown", "Unbekannt", 20.0, 3.2),
    ("number", "export_column_number", "Seitenzahlen", 14.0, 2.4),
    ("original_line", "export_column_original_line", "Originalzeile", 70.0, 9.5),
]
_BK_TABULAR_COLUMNS = list(_BK_EXTENDED_TABULAR_COLUMNS)
_BK_TABULAR_KEYS = [column[0] for column in _BK_TABULAR_COLUMNS]
_BK_TABULAR_DEFAULT_KEYS = [
    "family_name",
    "given_names",
    "relationship",
    "age_original",
    "date_original",
    "year_resolved",
    "place_in_source",
    "number",
    "original_line",
]
# Alter bleibt als Exportspalte und Datentyp erhalten, auch wenn die erweiterten
# Datentypen die Anzeige neu sortieren.
if "age_original" not in [c[0] for c in _BK_TABULAR_COLUMNS]:
    _BK_TABULAR_COLUMNS.insert(4, ("age_original", "export_column_age", "Alter", 12.0, 2.2))
    _BK_TABULAR_KEYS = [column[0] for column in _BK_TABULAR_COLUMNS]
_BK_TABULAR_HEADERS = [column[2] for column in _BK_TABULAR_COLUMNS]
_BK_TABULAR_COLUMN_BY_KEY = {column[0]: column for column in _BK_TABULAR_COLUMNS}
_BK_ZONE_TYPES = tuple(_BK_TABULAR_KEYS + ["ignore"])
_BK_ZONE_LEGACY_TYPES = {
    "data": "original_line",
    "names": "full_name",
    "heading": "year_resolved",
    "year": "year_resolved",
    "years": "year_resolved",
    "age": "age_original",
    "places": "place_in_source",
    "page_numbers": "number",
    "number": "number",
    "other": "other",
}






def _bk_zone_date_parts(text: str):
    value = _bk_tab_extract_date(text or "")
    if not value:
        return "", ""
    m = re.match(r"\s*(\d{1,2})\s*\.\s*([IVXLCDM]{1,8}|\d{1,2})", value, flags=re.IGNORECASE)
    if not m:
        return "", ""
    return m.group(1), m.group(2).upper()


def _bk_zone_middle_names(text: str) -> str:
    _family, given, _full = _bk_tab_split_name(text or "")
    parts = [p for p in re.split(r"\s+", given.strip()) if p]
    return " ".join(parts[1:]) if len(parts) > 1 else ""


def _bk_zone_house_number(text: str) -> str:
    txt = _bk_tab_clean(text)
    m = re.search(r"\b\d{1,4}\s*[A-Za-z]?\b", txt)
    return _bk_tab_clean(m.group(0)) if m else ""


def _bk_zone_postal_code(text: str) -> str:
    txt = _bk_tab_clean(text)
    m = re.search(r"\b\d{4,5}\b", txt)
    return m.group(0) if m else ""








def _bk_zone_apply_name_values(row_values, key, value):
    value = _bk_tab_clean(value)
    if not value:
        return
    if key == "full_name":
        row_values["full_name"] = value
        fam, given, _full = _bk_tab_split_name(value)
        if fam and not row_values.get("family_name"):
            row_values["family_name"] = fam
        if given and not row_values.get("given_names"):
            row_values["given_names"] = given
        if not row_values.get("middle_names"):
            row_values["middle_names"] = _bk_zone_middle_names(value)
    elif key == "family_name":
        row_values["family_name"] = value
    elif key == "given_names":
        row_values["given_names"] = value
        if not row_values.get("middle_names"):
            parts = [p for p in value.split() if p]
            if len(parts) > 1:
                row_values["middle_names"] = " ".join(parts[1:])






def _bk_zone_rect_expanded_bounding_rect(self):
    try:
        margin = float(getattr(self, "HANDLE_SIZE", 10.0)) + 4.0
        return self.rect().adjusted(-margin, -margin, margin, margin)
    except Exception:
        return self.rect()


def _bk_zone_rect_expanded_shape(self):
    try:
        from PySide6.QtGui import QPainterPath
        path = QPainterPath()
        path.addRect(_bk_zone_rect_expanded_bounding_rect(self))
        return path
    except Exception:
        try:
            return self._bk_original_shape()
        except Exception:
            return None


def _bk_zone_rect_mouse_press_resizable(self, event):
    try:
        self._dialog._suppress_zone_redraw = True
        self._dialog.table.blockSignals(True)
        self._dialog.table.selectRow(self._row)
        self._dialog.table.blockSignals(False)
    except Exception:
        try:
            self._dialog.table.blockSignals(False)
        except Exception:
            pass
    finally:
        try:
            self._dialog._suppress_zone_redraw = False
        except Exception:
            pass
    self._mode = self._handle_at(event.pos())
    self._start_scene = event.scenePos()
    self._start_rect = QRectF(self.rect())
    event.accept()


def _bk_export_load_global_zones(window):
    try:
        raw = window.settings.value("export/global_zones", "", str)
        data = json.loads(raw) if raw else []
    except Exception:
        data = []
    zones = [_bk_export_clean_zone(zone, idx) for idx, zone in enumerate(data or [])]
    return [zone for zone in zones if zone]


def _bk_export_save_global_zones(window, zones):
    clean = [_bk_export_clean_zone(zone, idx) for idx, zone in enumerate(zones or [])]
    clean = [zone for zone in clean if zone]
    try:
        window.settings.setValue("export/global_zones", json.dumps(clean, ensure_ascii=False))
        window.settings.setValue("export/remember_zones", bool(clean))
        window.settings.sync()
    except Exception:
        pass
    return clean


def _bk_export_global_zones_enabled(window) -> bool:
    try:
        return bool(window.settings.value("export/remember_zones", False, type=bool))
    except Exception:
        return False


def _bk_effective_export_zones(window, task, use_zones):
    if not use_zones:
        return []
    if _bk_export_global_zones_enabled(window):
        global_zones = _bk_export_load_global_zones(window)
        if global_zones:
            return global_zones
    return _bk_get_task_export_zones(task)


def _bk_open_export_zones_dialog_global(window, task):
    if task is None:
        try:
            task = window._current_task()
        except Exception:
            task = None
    if task is None:
        QMessageBox.warning(window, _bk_tab_tr(window, "warn_title"), _bk_tab_tr(window, "warn_need_done_for_ai"))
        return None
    path = getattr(task, "path", "")
    if not path or not os.path.exists(path):
        QMessageBox.warning(window, _bk_tab_tr(window, "warn_title"), _bk_tab_tr(window, "warn_project_file_missing").format(path))
        return None
    task_zones = _bk_get_task_export_zones(task)
    initial_zones = task_zones or _bk_export_load_global_zones(window)
    dlg = _BKExportZonesDialog(window, task, path, initial_zones)
    try:
        remember = bool(_bk_export_global_zones_enabled(window))
        dlg.cb_remember_zones.setChecked(remember)
    except Exception:
        pass
    try:
        dlg.showMaximized()
    except Exception:
        pass
    if dlg.exec() != QDialog.Accepted:
        return None
    zones = _bk_set_task_export_zones(task, dlg.zones())
    try:
        if bool(dlg.cb_remember_zones.isChecked()):
            _bk_export_save_global_zones(window, zones)
        else:
            window.settings.setValue("export/remember_zones", False)
            window.settings.sync()
    except Exception:
        pass
    return zones


def _bk_zone_dialog_init_genealogy(self, window, task, image_path, zones=None):
    _bk_zone_dialog_init_final(self, window, task, image_path, zones)
    try:
        self.cb_remember_zones = QCheckBox(_bk_tab_tr(window, "export_zones_remember"), self)
        layout = self.layout()
        side_layout = layout.itemAt(1).layout() if layout and layout.count() > 1 else None
        if side_layout is not None:
            # Vor den OK/Abbrechen-Buttons einfügen.
            insert_at = max(0, side_layout.count() - 1)
            side_layout.insertWidget(insert_at, self.cb_remember_zones)
    except Exception:
        pass


def _bk_zone_rows_for_item_global(window, item, image_size, use_zones):
    _text, _kr, _pil, record_views = item.results
    zones = _bk_effective_export_zones(window, item, use_zones)
    if zones:
        return _bk_build_transcription_rows_with_zones(record_views, image_size, zones)
    rows = _bk_build_transcription_rows(record_views, image_size)
    # Neue Spalten bleiben bei klassischen Exporten leer, ältere Spalten werden beibehalten.
    for row in rows:
        if "full_name" not in row:
            row["full_name"] = _bk_tab_join_text_fragments([row.get("family_name", ""), row.get("given_names", "")])
        if "middle_names" not in row:
            row["middle_names"] = _bk_zone_middle_names(row.get("full_name", ""))
        if row.get("date_original"):
            row.setdefault("days", _bk_zone_date_parts(row.get("date_original", ""))[0])
            row.setdefault("months", _bk_zone_date_parts(row.get("date_original", ""))[1])
    return rows




def _bk_zone_export_sqlite_json_global(self):
    task = _bk_fix36_current_task(self) if callable(globals().get("_bk_fix36_current_task")) else None
    if not task or not getattr(task, "results", None):
        QMessageBox.information(self, _bk_tab_tr(self, "info_title"), _bk_tab_tr(self, "warn_no_ocr_results"))
        return
    result = _bk_column_choice_dialog(self, "sqlite-json", include_text_modes=False)
    if result is None:
        return
    column_keys = result.get("columns") or list(_BK_TABULAR_DEFAULT_KEYS)
    self._bk_export_current_column_keys = column_keys
    self._bk_export_use_zones = bool(result.get("use_zones", False))
    if result.get("remembered"):
        self._bk_export_selected_column_keys = column_keys
    try:
        _txt, _kr, pil_image, recs = task.results
        try:
            export_image = _load_image_color(task.path)
            image_size = export_image.size
        except Exception:
            image_size = getattr(pil_image, "size", None)
        rows = _bk_zone_rows_for_item_global(self, task, image_size, self._bk_export_use_zones)
    except Exception:
        rows = []
    if not rows:
        QMessageBox.information(self, _bk_tab_tr(self, "info_title"), _bk_tab_tr(self, "warn_no_exportable_person_entries"))
        return
    start_dir = getattr(self, "current_export_dir", "") or os.path.dirname(getattr(task, "path", "") or "") or os.getcwd()
    default_name = os.path.splitext(os.path.basename(getattr(task, "path", "bottled_kraken")))[0] + "_sqlite.json"
    path, _filter = QFileDialog.getSaveFileName(
        self,
        _bk_tab_tr(self, "dlg_sqlite_json_title"),
        os.path.join(start_dir, default_name),
        _bk_tab_tr(self, "filter_json_files"),
    )
    if not path:
        return
    if not path.lower().endswith(".json"):
        path += ".json"
    _bk_write_transcription_json(path, task, rows, column_keys, self)
    try:
        self.current_export_dir = os.path.dirname(path)
        self.status_bar.showMessage(_bk_tab_tr(self, "msg_sqlite_export_done").format(os.path.basename(path)), 5000)
    except Exception:
        pass

try:
    if not hasattr(_BKExportZoneRectItem, "_bk_original_bounding_rect"):
        _BKExportZoneRectItem._bk_original_bounding_rect = _BKExportZoneRectItem.boundingRect
    if not hasattr(_BKExportZoneRectItem, "_bk_original_shape"):
        _BKExportZoneRectItem._bk_original_shape = _BKExportZoneRectItem.shape
    _BKExportZoneRectItem.boundingRect = _bk_zone_rect_expanded_bounding_rect
    _BKExportZoneRectItem.shape = _bk_zone_rect_expanded_shape
    _BKExportZoneRectItem.mousePressEvent = _bk_zone_rect_mouse_press_resizable
    _BKExportZonesDialog.__init__ = _bk_zone_dialog_init_genealogy
    _BKExportZonesDialog.add_zone_from_rect = _bk_zone_add_from_rect_final
    _BKExportZonesDialog.redraw_zones = _bk_zone_redraw_final
    _BKExportZonesDialog.refresh_table = _bk_zone_refresh_table_final
    _BKExportZonesDialog._table_item_changed = _bk_zone_table_item_changed_final
    _BKExportZonesDialog._selected_zone_changed = _bk_zone_selected_changed_final
    _bk_open_export_zones_dialog = _bk_open_export_zones_dialog_global
    MainWindow._bk_export_sqlite_json = _bk_zone_export_sqlite_json_global
    MainWindow.bk_export_sqlite_persons = _bk_zone_export_sqlite_json_global
except Exception:
    pass

try:
    __all__.extend([
        '_bk_export_global_zones_enabled',
        '_bk_export_load_global_zones',
        '_bk_export_save_global_zones',
        '_bk_open_export_zones_dialog_global',
        '_bk_zone_rows_for_item_global',
        '_bk_zone_value_for_column',
    ])
    register_globals('bk', globals(), sorted(set(__all__)))
except Exception:
    pass

# ---------------------------------------------------------------------------
# Exportdialog/Bereichslogik: höhere Spaltenauswahl, spaltenweise Sortierung,
# robuste Bereichsgriffe und zonenabhängige Exportspalten
# ---------------------------------------------------------------------------

def _bk_zone_grid_position_columnwise(index: int, total: int, columns: int = 3):
    columns = max(1, int(columns or 1))
    rows_per_col = int(math.ceil(max(1, int(total or 0)) / float(columns)))
    return index % rows_per_col, index // rows_per_col


def _bk_zone_column_choice_dialog_tall(self, fmt=None, include_text_modes=False):
    dlg = QDialog(self)
    dlg.setWindowTitle(_bk_tab_tr(self, "export_text_layout_title" if include_text_modes else "export_table_columns_title"))
    dlg.setModal(True)
    layout = QVBoxLayout(dlg)
    layout.setContentsMargins(18, 16, 18, 16)
    layout.setSpacing(12)

    intro_key = "export_text_layout_intro_extended" if include_text_modes else "export_table_columns_intro"
    label = QLabel(_bk_tab_tr(self, intro_key), dlg)
    label.setWordWrap(True)
    layout.addWidget(label)

    rb_original = rb_lines = rb_table = None
    if include_text_modes:
        if QGroupBox is not None:
            mode_box = QGroupBox(_bk_tab_tr(self, "export_layout_mode_group"), dlg)
            mode_layout = QVBoxLayout(mode_box)
            mode_layout.setContentsMargins(12, 10, 12, 10)
            mode_layout.setSpacing(6)
        else:
            mode_box = None
            mode_layout = QVBoxLayout()
            mode_layout.setSpacing(6)
        rb_original = QRadioButton(_bk_tab_tr(self, "export_text_layout_original"), dlg)
        rb_lines = QRadioButton(_bk_tab_tr(self, "export_text_layout_lines"), dlg)
        rb_table = QRadioButton(_bk_tab_tr(self, "export_text_layout_table"), dlg)
        mode = str(getattr(self, "_bk_export_text_layout_mode", "original") or "original").lower()
        rb_original.setChecked(mode not in {"lines", "table"})
        rb_lines.setChecked(mode == "lines")
        rb_table.setChecked(mode == "table")
        for rb in (rb_original, rb_lines, rb_table):
            mode_layout.addWidget(rb)
        if mode_box is not None:
            layout.addWidget(mode_box)
        else:
            layout.addLayout(mode_layout)

    selected_keys = _bk_load_saved_column_keys(self)
    checkboxes = {}
    if QGroupBox is not None:
        columns_box = QGroupBox(_bk_tab_tr(self, "export_table_columns_label"), dlg)
        columns_layout = QVBoxLayout(columns_box)
        columns_layout.setContentsMargins(12, 10, 12, 10)
        columns_layout.setSpacing(8)
    else:
        columns_box = None
        columns_layout = QVBoxLayout()
        columns_layout.setSpacing(8)
        columns_layout.addWidget(QLabel(_bk_tab_tr(self, "export_table_columns_label"), dlg))

    if QGridLayout is not None:
        checkbox_grid = QGridLayout()
        checkbox_grid.setContentsMargins(0, 0, 0, 0)
        checkbox_grid.setHorizontalSpacing(26)
        checkbox_grid.setVerticalSpacing(7)
        total = len(_BK_TABULAR_KEYS)
        for idx, key in enumerate(_BK_TABULAR_KEYS):
            cb = QCheckBox(_bk_tabular_column_title(self, key), dlg)
            cb.setChecked(key in selected_keys)
            checkboxes[key] = cb
            row, col = _bk_zone_grid_position_columnwise(idx, total, 3)
            checkbox_grid.addWidget(cb, row, col)
        columns_layout.addLayout(checkbox_grid)
    else:
        for key in _BK_TABULAR_KEYS:
            cb = QCheckBox(_bk_tabular_column_title(self, key), dlg)
            cb.setChecked(key in selected_keys)
            checkboxes[key] = cb
            columns_layout.addWidget(cb)

    quick_row = QHBoxLayout()
    quick_row.setSpacing(8)
    btn_all = QPushButton(_bk_tab_tr(self, "export_table_columns_all"), dlg)
    btn_none = QPushButton(_bk_tab_tr(self, "export_table_columns_none_button"), dlg)
    btn_remember = QPushButton(_bk_tab_tr(self, "export_table_columns_remember"), dlg)
    quick_row.addWidget(btn_all)
    quick_row.addWidget(btn_none)
    quick_row.addWidget(btn_remember)
    quick_row.addStretch(1)
    columns_layout.addLayout(quick_row)

    if columns_box is not None:
        layout.addWidget(columns_box)
    else:
        layout.addLayout(columns_layout)

    zone_box = QGroupBox(_bk_tab_tr(self, "export_zones_group"), dlg) if QGroupBox is not None else None
    zone_layout = QHBoxLayout(zone_box) if zone_box is not None else QHBoxLayout()
    zone_layout.setSpacing(8)
    cb_zones = QCheckBox(_bk_tab_tr(self, "export_table_use_zones"), dlg)
    try:
        remembered_zones = self.settings.value("export/table_use_zones", bool(getattr(self, "_bk_export_use_zones", False)), type=bool)
    except Exception:
        remembered_zones = bool(getattr(self, "_bk_export_use_zones", False))
    cb_zones.setChecked(bool(remembered_zones))
    btn_zones = QPushButton(_bk_tab_tr(self, "export_table_define_zones"), dlg)
    zone_layout.addWidget(cb_zones, 1)
    zone_layout.addWidget(btn_zones)
    if zone_box is not None:
        layout.addWidget(zone_box)
    else:
        layout.addLayout(zone_layout)

    def current_checked_keys():
        return [key for key, cb in checkboxes.items() if cb.isChecked()]

    def set_all():
        for cb in checkboxes.values():
            cb.setChecked(True)

    def set_none():
        for cb in checkboxes.values():
            cb.setChecked(False)

    result = {"mode": "table", "columns": selected_keys, "remembered": False, "use_zones": bool(cb_zones.isChecked())}

    def remember_selection():
        keys = current_checked_keys()
        if not keys:
            QMessageBox.warning(dlg, _bk_tab_tr(self, "warn_title"), _bk_tab_tr(self, "export_table_columns_none"))
            return
        result["remembered"] = True
        result["columns"] = _bk_save_column_keys(self, keys)
        try:
            self.settings.setValue("export/table_use_zones", bool(cb_zones.isChecked()))
            self.settings.sync()
        except Exception:
            pass

    def define_zones():
        try:
            task = self._current_task()
        except Exception:
            task = None
        was_enabled = True
        try:
            was_enabled = dlg.isEnabled()
            dlg.setEnabled(False)
            # Nicht verstecken: sonst kann der Exportdialog nach dem Bereichsdialog
            # verloren gehen und der eigentliche Export wirkt "verschwunden".
            dlg.lower()
            QApplication.processEvents()
        except Exception:
            pass
        try:
            zones = _bk_open_export_zones_dialog(self, task)
        finally:
            try:
                dlg.setEnabled(was_enabled)
                dlg.show()
                dlg.raise_()
                dlg.activateWindow()
                QApplication.processEvents()
            except Exception:
                pass
        if zones is not None:
            cb_zones.setChecked(bool(zones))

    btn_all.clicked.connect(set_all)
    btn_none.clicked.connect(set_none)
    btn_remember.clicked.connect(remember_selection)
    btn_zones.clicked.connect(define_zones)

    def sync_columns_enabled():
        enabled = True
        if include_text_modes and rb_table is not None:
            enabled = bool(rb_table.isChecked())
        if columns_box is not None:
            columns_box.setEnabled(enabled)
        else:
            for cb in checkboxes.values():
                cb.setEnabled(enabled)
            btn_all.setEnabled(enabled)
            btn_none.setEnabled(enabled)
            btn_remember.setEnabled(enabled)
        # Bereichsauswahl bleibt bewusst unabhängig von der Darstellungs-/Tabellenoption.
        cb_zones.setEnabled(True)
        btn_zones.setEnabled(True)

    if include_text_modes:
        try:
            rb_original.toggled.connect(sync_columns_enabled)
            rb_lines.toggled.connect(sync_columns_enabled)
            rb_table.toggled.connect(sync_columns_enabled)
        except Exception:
            pass
        sync_columns_enabled()

    buttons = QDialogButtonBox(QDialogButtonBox.Ok | QDialogButtonBox.Cancel, dlg)
    try:
        buttons.button(QDialogButtonBox.Ok).setText(_bk_tab_tr(self, "btn_ok"))
        buttons.button(QDialogButtonBox.Cancel).setText(_bk_tab_tr(self, "btn_cancel"))
    except Exception:
        pass

    def cancel_dialog():
        result["cancelled"] = True
        try:
            dlg.done(QDialog.Rejected)
        except Exception:
            try:
                dlg.reject()
            except Exception:
                pass

    def accept_checked():
        result["cancelled"] = False
        mode = "table"
        if include_text_modes:
            if rb_lines and rb_lines.isChecked():
                mode = "lines"
            elif rb_original and rb_original.isChecked():
                mode = "original"
        keys = current_checked_keys()
        if mode == "table" and not keys:
            QMessageBox.warning(dlg, _bk_tab_tr(self, "warn_title"), _bk_tab_tr(self, "export_table_columns_none"))
            return
        result["mode"] = mode
        result["columns"] = _bk_normalize_column_keys(keys)
        result["use_zones"] = bool(cb_zones.isChecked())
        try:
            self._bk_export_use_zones = result["use_zones"]
        except Exception:
            pass
        dlg.accept()

    def cancel_dialog():
        result["cancelled"] = True
        try:
            dlg.done(QDialog.Rejected)
        except Exception:
            dlg.reject()

    try:
        cancel_btn = buttons.button(QDialogButtonBox.Cancel)
        if cancel_btn is not None:
            cancel_btn.clicked.connect(cancel_dialog)
        else:
            buttons.rejected.connect(cancel_dialog)
    except Exception:
        buttons.rejected.connect(cancel_dialog)
    buttons.accepted.connect(accept_checked)

    def _close_event(event):
        result["cancelled"] = True
        try:
            dlg.done(QDialog.Rejected)
        except Exception:
            try:
                dlg.reject()
            except Exception:
                pass
        try:
            event.accept()
        except Exception:
            pass

    try:
        dlg.closeEvent = _close_event
    except Exception:
        pass
    layout.addWidget(buttons)

    dlg.setMinimumSize(720, 760 if include_text_modes else 690)
    try:
        dlg.resize(760, 820 if include_text_modes else 740)
    except Exception:
        pass
    if dlg.exec() != QDialog.Accepted:
        return None
    return result


def _bk_zone_handle_mode_for_scene_rect(rect, scene_pos, tolerance=26.0):
    x = float(scene_pos.x())
    y = float(scene_pos.y())
    left, right = float(rect.left()), float(rect.right())
    top, bottom = float(rect.top()), float(rect.bottom())
    near_left = abs(x - left) <= tolerance and (top - tolerance) <= y <= (bottom + tolerance)
    near_right = abs(x - right) <= tolerance and (top - tolerance) <= y <= (bottom + tolerance)
    near_top = abs(y - top) <= tolerance and (left - tolerance) <= x <= (right + tolerance)
    near_bottom = abs(y - bottom) <= tolerance and (left - tolerance) <= x <= (right + tolerance)
    if near_left and near_top:
        return "tl"
    if near_right and near_top:
        return "tr"
    if near_left and near_bottom:
        return "bl"
    if near_right and near_bottom:
        return "br"
    if near_left:
        return "l"
    if near_right:
        return "r"
    if near_top:
        return "t"
    if near_bottom:
        return "b"
    if rect.adjusted(-4, -4, 4, 4).contains(scene_pos):
        return "move"
    return ""


def _bk_zone_find_at_scene_pos(dialog, scene_pos):
    zones = getattr(dialog, "_zones", []) or []
    for row in range(len(zones) - 1, -1, -1):
        zone = zones[row]
        rect = QRectF(float(zone["x0"]), float(zone["y0"]), float(zone["x1"] - zone["x0"]), float(zone["y1"] - zone["y0"])).normalized()
        mode = _bk_zone_handle_mode_for_scene_rect(rect, scene_pos)
        if mode:
            return row, mode, rect
    return -1, "", None


def _bk_zone_view_scene_pos(self, event):
    return self.mapToScene(event.position().toPoint() if hasattr(event, "position") else event.pos())


def _bk_zone_view_mouse_press_precise(self, event):
    if event.button() == Qt.LeftButton:
        pos = _bk_zone_view_scene_pos(self, event)
        row, mode, rect = _bk_zone_find_at_scene_pos(self._dialog, pos)
        if row >= 0:
            try:
                self._dialog._suppress_zone_redraw = True
                self._dialog.table.blockSignals(True)
                self._dialog.table.selectRow(row)
                self._dialog.table.blockSignals(False)
            except Exception:
                try:
                    self._dialog.table.blockSignals(False)
                except Exception:
                    pass
            finally:
                self._dialog._suppress_zone_redraw = False
            self._active_zone_row = row
            self._active_zone_mode = mode
            self._active_zone_start = pos
            self._active_zone_rect = QRectF(rect)
            event.accept()
            return
        self._start = pos
        if self._rubber is not None:
            try:
                self.scene().removeItem(self._rubber)
            except Exception:
                pass
            self._rubber = None
        self._rubber = QGraphicsRectItem(QRectF(self._start, self._start))
        self._rubber.setPen(QPen(QColor(40, 130, 255), 2, Qt.DashLine))
        self._rubber.setBrush(QBrush(QColor(40, 130, 255, 35)))
        self.scene().addItem(self._rubber)
        event.accept()
        return
    return QGraphicsView.mousePressEvent(self, event)


def _bk_zone_rect_from_drag(start_rect, start_pos, current_pos, mode):
    dx = float(current_pos.x() - start_pos.x())
    dy = float(current_pos.y() - start_pos.y())
    rect = QRectF(start_rect)
    if mode == "move":
        rect.translate(dx, dy)
    else:
        if "l" in mode:
            rect.setLeft(rect.left() + dx)
        if "r" in mode:
            rect.setRight(rect.right() + dx)
        if "t" in mode:
            rect.setTop(rect.top() + dy)
        if "b" in mode:
            rect.setBottom(rect.bottom() + dy)
    return rect.normalized()


def _bk_zone_view_mouse_move_precise(self, event):
    if getattr(self, "_active_zone_row", -1) >= 0:
        pos = _bk_zone_view_scene_pos(self, event)
        rect = _bk_zone_rect_from_drag(self._active_zone_rect, self._active_zone_start, pos, self._active_zone_mode)
        rect = self._dialog.clamp_rect(rect)
        self._dialog.set_zone_rect(int(self._active_zone_row), rect, redraw=True)
        event.accept()
        return
    if getattr(self, "_start", None) is not None and getattr(self, "_rubber", None) is not None:
        pos = _bk_zone_view_scene_pos(self, event)
        self._rubber.setRect(QRectF(self._start, pos).normalized())
        event.accept()
        return
    return QGraphicsView.mouseMoveEvent(self, event)


def _bk_zone_view_mouse_release_precise(self, event):
    if getattr(self, "_active_zone_row", -1) >= 0:
        self._active_zone_row = -1
        self._active_zone_mode = ""
        self._active_zone_start = None
        self._active_zone_rect = None
        self._dialog.redraw_zones()
        event.accept()
        return
    if getattr(self, "_start", None) is not None and getattr(self, "_rubber", None) is not None:
        rect = self._rubber.rect().normalized()
        try:
            self.scene().removeItem(self._rubber)
        except Exception:
            pass
        self._rubber = None
        self._start = None
        if rect.width() >= 6 and rect.height() >= 6:
            self._dialog.add_zone_from_rect(rect)
        event.accept()
        return
    return QGraphicsView.mouseReleaseEvent(self, event)


def _bk_zone_rect_handle_at_large(self, pos):
    return _bk_zone_handle_mode_for_scene_rect(self.rect(), pos, tolerance=24.0) or "move"


def _bk_zone_rect_mouse_press_no_redraw(self, event):
    try:
        self._dialog._suppress_zone_redraw = True
        self._dialog.table.blockSignals(True)
        self._dialog.table.selectRow(self._row)
        self._dialog.table.blockSignals(False)
    except Exception:
        try:
            self._dialog.table.blockSignals(False)
        except Exception:
            pass
    finally:
        try:
            self._dialog._suppress_zone_redraw = False
        except Exception:
            pass
    self._mode = self._handle_at(event.pos())
    self._start_scene = event.scenePos()
    self._start_rect = QRectF(self.rect())
    event.accept()


def _bk_zone_filter_columns_for_active_zones(window, item, column_keys, use_zones):
    keys = _bk_normalize_column_keys(column_keys)
    if not use_zones or item is None:
        return keys
    try:
        zones = _bk_effective_export_zones(window, item, True)
    except Exception:
        zones = []
    zone_types = []
    for zone in zones or []:
        clean = _bk_export_clean_zone(zone, len(zone_types))
        if not clean:
            continue
        ztype = clean.get("type")
        if ztype in _BK_TABULAR_COLUMN_BY_KEY and ztype not in zone_types:
            zone_types.append(ztype)
    if not zone_types:
        return keys
    filtered = [key for key in keys if key in zone_types]
    return filtered or keys


def _bk_zone_tabular_render_file_filtered(self, path: str, fmt: str, item: TaskItem):
    fmt_l = str(fmt or "").lower().lstrip(".")
    if item and getattr(item, "results", None) and (fmt_l in _BK_TABLE_EXPORT_FMTS or fmt_l in _BK_TEXT_LAYOUT_FMTS):
        _text, _kr, pil_image, record_views = item.results
        try:
            export_image = _load_image_color(item.path)
        except Exception:
            export_image = pil_image
        image_size = getattr(export_image, "size", None) or getattr(pil_image, "size", None)
        use_zones = bool(getattr(self, "_bk_export_use_zones", False))
        column_keys = _bk_zone_filter_columns_for_active_zones(self, item, _bk_current_column_keys_for_render(self), use_zones)
        if fmt_l in _BK_TABLE_EXPORT_FMTS:
            rows = _bk_zone_rows_for_item_global(self, item, image_size, use_zones)
            if fmt_l == "csv":
                return _bk_write_transcription_csv(path, rows, column_keys, self)
            if fmt_l == "json":
                return _bk_write_transcription_json(path, item, rows, column_keys, self)
            if fmt_l in {"xlsx", "excel"}:
                return _bk_write_transcription_xlsx(path, rows, column_keys, self)
            if fmt_l in {"ods", "calc"}:
                return _bk_write_transcription_ods(path, rows, column_keys, self)
        if fmt_l in _BK_TEXT_LAYOUT_FMTS:
            mode = str(getattr(self, "_bk_export_text_layout_mode", "original") or "original").lower()
            if mode == "table":
                rows = _bk_zone_rows_for_item_global(self, item, image_size, use_zones)
                if fmt_l in {"txt", "text", "txt_plain"}:
                    return _bk_write_table_txt(path, rows, column_keys, self)
                if fmt_l in {"docx", "word"}:
                    return _bk_write_table_docx(path, rows, column_keys, self)
                if fmt_l == "odt":
                    return _bk_write_table_odt(path, rows, column_keys, self)
    return RENDER_NOT_HANDLED


def _bk_zone_export_sqlite_json_filtered(self):
    task = _bk_fix36_current_task(self) if callable(globals().get("_bk_fix36_current_task")) else None
    if not task or not getattr(task, "results", None):
        QMessageBox.information(self, _bk_tab_tr(self, "info_title"), _bk_tab_tr(self, "warn_no_ocr_results"))
        return
    result = _bk_column_choice_dialog(self, "sqlite-json", include_text_modes=False)
    if result is None:
        return
    self._bk_export_use_zones = bool(result.get("use_zones", False))
    column_keys = _bk_zone_filter_columns_for_active_zones(self, task, result.get("columns") or list(_BK_TABULAR_DEFAULT_KEYS), self._bk_export_use_zones)
    self._bk_export_current_column_keys = column_keys
    if result.get("remembered"):
        self._bk_export_selected_column_keys = result.get("columns") or column_keys
    try:
        _txt, _kr, pil_image, recs = task.results
        try:
            export_image = _load_image_color(task.path)
            image_size = export_image.size
        except Exception:
            image_size = getattr(pil_image, "size", None)
        rows = _bk_zone_rows_for_item_global(self, task, image_size, self._bk_export_use_zones)
    except Exception:
        rows = []
    if not rows:
        QMessageBox.information(self, _bk_tab_tr(self, "info_title"), _bk_tab_tr(self, "warn_no_exportable_person_entries"))
        return
    start_dir = getattr(self, "current_export_dir", "") or os.path.dirname(getattr(task, "path", "") or "") or os.getcwd()
    default_name = os.path.splitext(os.path.basename(getattr(task, "path", "bottled_kraken")))[0] + "_sqlite.json"
    path, _filter = QFileDialog.getSaveFileName(
        self,
        _bk_tab_tr(self, "dlg_sqlite_json_title"),
        os.path.join(start_dir, default_name),
        _bk_tab_tr(self, "filter_json_files"),
    )
    if not path:
        return
    if not path.lower().endswith(".json"):
        path += ".json"
    _bk_write_transcription_json(path, task, rows, column_keys, self)
    try:
        self.current_export_dir = os.path.dirname(path)
        self.status_bar.showMessage(_bk_tab_tr(self, "msg_sqlite_export_done").format(os.path.basename(path)), 5000)
    except Exception:
        pass

try:
    _bk_column_choice_dialog = _bk_zone_column_choice_dialog_tall
    _BKExportZoneGraphicsView.mousePressEvent = _bk_zone_view_mouse_press_precise
    _BKExportZoneGraphicsView.mouseMoveEvent = _bk_zone_view_mouse_move_precise
    _BKExportZoneGraphicsView.mouseReleaseEvent = _bk_zone_view_mouse_release_precise
    _BKExportZoneRectItem.HANDLE_SIZE = 18.0
    _BKExportZoneRectItem._handle_at = _bk_zone_rect_handle_at_large
    _BKExportZoneRectItem.mousePressEvent = _bk_zone_rect_mouse_press_no_redraw
    register_render_handler(_bk_zone_tabular_render_file_filtered)
    MainWindow._bk_export_sqlite_json = _bk_zone_export_sqlite_json_filtered
    MainWindow.bk_export_sqlite_persons = _bk_zone_export_sqlite_json_filtered
except Exception:
    pass

try:
    __all__.extend([
        '_bk_zone_column_choice_dialog_tall',
        '_bk_zone_filter_columns_for_active_zones',
        '_bk_zone_tabular_render_file_filtered',
        '_bk_zone_export_sqlite_json_filtered',
    ])
    register_globals('bk', globals(), sorted(set(__all__)))
except Exception:
    pass

# ---------------------------------------------------------------------------
# KI-Unterstützung für sensible Tabellenbereiche
# ---------------------------------------------------------------------------
try:
    from bottled_kraken.common import _extract_json_payload, _force_text
    from bottled_kraken.workers import AIRevisionWorker
except Exception:
    _extract_json_payload = None
    _force_text = str
    AIRevisionWorker = None


def _bk_export_zones_ai_enabled(window) -> bool:
    try:
        return bool(window.settings.value("export/zones_ai_support", False, type=bool))
    except Exception:
        return False


def _bk_export_set_zones_ai_enabled(window, enabled: bool):
    try:
        window.settings.setValue("export/zones_ai_support", bool(enabled))
        window.settings.sync()
    except Exception:
        pass


def _bk_ai_zone_records_payload(raw_records, zones, window):
    payload = []
    for zone in zones or []:
        ztype = str(zone.get("type", "") or "")
        if ztype == "ignore":
            continue
        records = _bk_zone_records(raw_records, zone)
        records = _bk_tab_merge_visual_row_records(
            records,
            max(1.0, float(zone.get("x1", 0.0)) - float(zone.get("x0", 0.0)))
        )
        lines = []
        for rec in sorted(records, key=lambda r: (float(r.get("y0", 0.0) or 0.0), float(r.get("x0", 0.0) or 0.0))):
            text = _bk_tab_clean(rec.get("text", ""))
            if not text:
                continue
            lines.append({
                "x": int(round(float(rec.get("x0", 0.0) or 0.0))),
                "y": int(round(float(rec.get("y0", 0.0) or 0.0))),
                "text": text,
            })
        if not lines:
            continue
        payload.append({
            "name": str(zone.get("name", "") or ""),
            "type": ztype,
            "label": _bk_export_zone_title(window, ztype),
            "rect": [
                int(round(float(zone.get("x0", 0.0) or 0.0))),
                int(round(float(zone.get("y0", 0.0) or 0.0))),
                int(round(float(zone.get("x1", 0.0) or 0.0))),
                int(round(float(zone.get("y1", 0.0) or 0.0))),
            ],
            "lines": lines,
        })
    return payload


def _bk_extract_ai_rows_json(text: str):
    raw = str(text or "").strip()
    if not raw:
        return []
    raw = re.sub(r"^```(?:json)?\s*", "", raw, flags=re.IGNORECASE).strip()
    raw = re.sub(r"\s*```$", "", raw).strip()
    data = None
    try:
        data = json.loads(raw)
    except Exception:
        if callable(_extract_json_payload):
            try:
                data = _extract_json_payload(raw)
            except Exception:
                data = None
    if isinstance(data, list):
        rows = data
    elif isinstance(data, dict):
        rows = data.get("rows") or data.get("entries") or data.get("data") or []
    else:
        rows = []
    if not isinstance(rows, list):
        return []
    out = []
    for row in rows:
        if not isinstance(row, dict):
            continue
        item = {}
        for key in _BK_TABULAR_KEYS:
            value = row.get(key, "")
            if value is None:
                value = ""
            if isinstance(value, (list, tuple)):
                value = "; ".join(_bk_tab_clean(v) for v in value if _bk_tab_clean(v))
            else:
                value = _bk_tab_clean(value)
            item[key] = value
        if any(item.get(k) for k in _BK_TABULAR_KEYS):
            out.append(item)
    return out




def _bk_build_transcription_rows_with_zones_ai(window, item, image_size=None, zones=None):
    if AIRevisionWorker is None:
        return []
    if not item or not getattr(item, "results", None):
        return []
    try:
        model_id = window._resolve_ai_model_id()
    except Exception:
        model_id = (getattr(window, "ai_model_id", "") or "").strip()
    if not model_id:
        _bk_export_zones_ai_warn(window, "model", "")
        return []
    try:
        _text, _kr, _pil, record_views = item.results
    except Exception:
        return []
    raw_records = _records_from_views(record_views)
    raw_records = [record for record in raw_records if not _bk_zone_is_ignored(record, zones)]
    raw_records = _bk_tab_expand_numeric_records(raw_records)
    zone_payload = _bk_ai_zone_records_payload(raw_records, zones, window)
    if not zone_payload:
        return []
    column_keys = []
    for zone in zones or []:
        ztype = str(zone.get("type", "") or "")
        if ztype in _BK_TABULAR_COLUMN_BY_KEY and ztype not in column_keys:
            column_keys.append(ztype)
    if not column_keys:
        return []
    allowed = {key: _bk_tabular_column_title(window, key) for key in column_keys}
    system_prompt = _bk_tab_tr(window, "ai_prompt_export_zones_system")
    user_prompt = _bk_tab_tr(window, "ai_prompt_export_zones_user").format(
        ", ".join(column_keys),
        json.dumps(allowed, ensure_ascii=False),
        json.dumps({"zones": zones or [], "candidates": zone_payload}, ensure_ascii=False),
    )
    try:
        worker = AIRevisionWorker(
            path=getattr(item, "path", "") or "",
            recs=[],
            lm_model=model_id,
            endpoint=getattr(window, "ai_endpoint", "http://127.0.0.1:1234/v1/chat/completions"),
            enable_thinking=False,
            tr_func=getattr(window, "_tr", None),
            temperature=0.1,
            top_p=0.8,
            max_tokens=max(1800, min(9000, 350 + 180 * max(1, len(raw_records)))),
        )
        schema = {
            "type": "json_schema",
            "json_schema": {
                "name": "table_rows",
                "schema": {
                    "type": "object",
                    "properties": {
                        "rows": {
                            "type": "array",
                            "items": {
                                "type": "object",
                                "additionalProperties": {"type": "string"},
                            },
                        }
                    },
                    "required": ["rows"],
                    "additionalProperties": False,
                },
            },
        }
        payload = {
            "model": model_id,
            "messages": [
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt},
            ],
            **worker._build_sampling_payload(response_format=schema, override_max_tokens=worker.max_tokens),
        }
        payload["enable_thinking"] = False
        payload["reasoning"] = {"enabled": False}
        data = worker._post_json(payload)
        content = worker._extract_message_content(data)
        rows = _bk_extract_ai_rows_json(content)
        return _bk_finish_ai_zone_rows(rows)
    except Exception as exc:
        try:
            print(f"BK export zones AI fallback: {exc}")
        except Exception:
            pass
        return []


try:
    _BK_PREV_ZONE_ROWS_FOR_ITEM_GLOBAL_AI = _bk_zone_rows_for_item_global
except Exception:
    _BK_PREV_ZONE_ROWS_FOR_ITEM_GLOBAL_AI = None


def _bk_zone_rows_for_item_global_ai(window, item, image_size, use_zones):
    zones = _bk_effective_export_zones(window, item, use_zones)
    if zones and _bk_export_zones_ai_enabled(window):
        ai_rows = _bk_build_transcription_rows_with_zones_ai(window, item, image_size, zones)
        if ai_rows:
            return ai_rows
    if callable(_BK_PREV_ZONE_ROWS_FOR_ITEM_GLOBAL_AI):
        return _BK_PREV_ZONE_ROWS_FOR_ITEM_GLOBAL_AI(window, item, image_size, use_zones)
    _text, _kr, _pil, record_views = item.results
    return _bk_build_transcription_rows_with_zones(record_views, image_size, zones) if zones else _bk_build_transcription_rows(record_views, image_size)



try:
    _BK_PREV_EXPORT_ZONES_INIT_AI = _BKExportZonesDialog.__init__
except Exception:
    _BK_PREV_EXPORT_ZONES_INIT_AI = None


def _bk_zone_dialog_init_ai_options(self, window, task, image_path, zones=None):
    if callable(_BK_PREV_EXPORT_ZONES_INIT_AI):
        _BK_PREV_EXPORT_ZONES_INIT_AI(self, window, task, image_path, zones)
    try:
        self.cb_ai_zones = QCheckBox(_bk_tab_tr(window, "export_zones_ai_support"), self)
        self.cb_ai_zones.setChecked(_bk_export_zones_ai_enabled(window))
        layout = self.layout()
        side_layout = layout.itemAt(1).layout() if layout and layout.count() > 1 else None
        old_cb = getattr(self, "cb_remember_zones", None)
        if side_layout is not None:
            insert_at = max(0, side_layout.count() - 1)
            for idx in range(side_layout.count()):
                item = side_layout.itemAt(idx)
                if item and item.widget() is old_cb:
                    taken = side_layout.takeAt(idx)
                    if taken and taken.widget():
                        taken.widget().setParent(None)
                    insert_at = idx
                    break
            row = QHBoxLayout()
            if old_cb is not None:
                row.addWidget(old_cb)
            row.addWidget(self.cb_ai_zones)
            row.addStretch(1)
            side_layout.insertLayout(insert_at, row)
    except Exception:
        pass


def _bk_open_export_zones_dialog_ai(window, task):
    if task is None:
        try:
            task = window._current_task()
        except Exception:
            task = None
    if task is None:
        QMessageBox.warning(window, _bk_tab_tr(window, "warn_title"), _bk_tab_tr(window, "warn_need_done_for_ai"))
        return None
    path = getattr(task, "path", "")
    if not path or not os.path.exists(path):
        QMessageBox.warning(window, _bk_tab_tr(window, "warn_title"), _bk_tab_tr(window, "warn_project_file_missing").format(path))
        return None
    task_zones = _bk_get_task_export_zones(task)
    initial_zones = task_zones or _bk_export_load_global_zones(window)
    dlg = _BKExportZonesDialog(window, task, path, initial_zones)
    try:
        dlg.cb_remember_zones.setChecked(bool(_bk_export_global_zones_enabled(window)))
    except Exception:
        pass
    try:
        dlg.cb_ai_zones.setChecked(bool(_bk_export_zones_ai_enabled(window)))
    except Exception:
        pass
    try:
        dlg.showMaximized()
    except Exception:
        pass
    if dlg.exec() != QDialog.Accepted:
        return None
    zones = _bk_set_task_export_zones(task, dlg.zones())
    try:
        if bool(dlg.cb_remember_zones.isChecked()):
            _bk_export_save_global_zones(window, zones)
        else:
            window.settings.setValue("export/remember_zones", False)
        _bk_export_set_zones_ai_enabled(window, bool(getattr(dlg, "cb_ai_zones", None) and dlg.cb_ai_zones.isChecked()))
        window.settings.sync()
    except Exception:
        pass
    return zones

try:
    _BKExportZonesDialog.__init__ = _bk_zone_dialog_init_ai_options
    _bk_open_export_zones_dialog = _bk_open_export_zones_dialog_ai
    _bk_zone_rows_for_item_global = _bk_zone_rows_for_item_global_ai
except Exception:
    pass

try:
    __all__.extend([
        '_bk_export_zones_ai_enabled',
        '_bk_export_set_zones_ai_enabled',
        '_bk_build_transcription_rows_with_zones_ai',
        '_bk_zone_rows_for_item_global_ai',
        '_bk_open_export_zones_dialog_ai',
        '_bk_zone_dialog_init_ai_options',
    ])
    register_globals('bk', globals(), sorted(set(__all__)))
except Exception:
    pass

# ---------------------------------------------------------------------------
# Fix: KI-Unterstützung bei sensiblen Bereichen muss bei ODT/DOCX/TXT eine
# echte Tabelle erzeugen. Sonst fiel der Export bei zuletzt gewähltem
# räumlichen Layout wieder auf den Positions-/Original-Renderer zurück.
# ---------------------------------------------------------------------------


def _bk_zone_ai_text_export_should_be_table(window, item, use_zones: bool) -> bool:
    if not use_zones:
        return False
    try:
        if not _bk_export_zones_ai_enabled(window):
            return False
    except Exception:
        return False
    try:
        return bool(_bk_effective_export_zones(window, item, True))
    except Exception:
        return False


def _bk_zone_tabular_render_file_ai_text_table(self, path: str, fmt: str, item: TaskItem):
    fmt_l = str(fmt or "").lower().lstrip(".")
    if item and getattr(item, "results", None) and fmt_l in _BK_TEXT_LAYOUT_FMTS:
        try:
            _text, _kr, pil_image, _record_views = item.results
        except Exception:
            pil_image = None
        try:
            export_image = _load_image_color(item.path)
        except Exception:
            export_image = pil_image
        image_size = getattr(export_image, "size", None) or getattr(pil_image, "size", None)
        use_zones = bool(getattr(self, "_bk_export_use_zones", False))
        mode = str(getattr(self, "_bk_export_text_layout_mode", "original") or "original").lower()
        force_table = _bk_zone_ai_text_export_should_be_table(self, item, use_zones)
        if mode == "table" or force_table:
            column_keys = _bk_zone_filter_columns_for_active_zones(
                self,
                item,
                _bk_current_column_keys_for_render(self),
                use_zones,
            )
            rows = _bk_zone_rows_for_item_global(self, item, image_size, use_zones)
            if fmt_l in {"txt", "text", "txt_plain"}:
                return _bk_write_table_txt(path, rows, column_keys, self)
            if fmt_l in {"docx", "word"}:
                return _bk_write_table_docx(path, rows, column_keys, self)
            if fmt_l == "odt":
                return _bk_write_table_odt(path, rows, column_keys, self)
    return RENDER_NOT_HANDLED


register_render_handler(_bk_zone_tabular_render_file_ai_text_table)

try:
    __all__.extend([
        '_bk_zone_ai_text_export_should_be_table',
        '_bk_zone_tabular_render_file_ai_text_table',
    ])
    register_globals('bk', globals(), sorted(set(__all__)))
except Exception:
    pass

# ---------------------------------------------------------------------------
# Fix: KI-Unterstützung für sensible Bereiche mit kompakten, gechunkten Daten.
# Lokale Modelle mit n_ctx=8192 dürfen nicht die vollständigen Bereichslisten
# aller OCR-Boxen bekommen. Die KI erhält deshalb nur vorgruppierte Zeilen-
# kandidaten und wird in kleinen Blöcken aufgerufen.
# ---------------------------------------------------------------------------

def _bk_ai_short_text(value, limit=140):
    text = _bk_tab_clean(value)
    if len(text) <= int(limit):
        return text
    return text[: int(limit) - 1].rstrip() + "…"


def _bk_ai_zone_column_keys(zones):
    keys = []
    for zone in zones or []:
        try:
            clean = _bk_export_clean_zone(zone, len(keys))
        except Exception:
            clean = None
        if not clean:
            continue
        key = str(clean.get("type", "") or "")
        if key == "ignore":
            continue
        if key in _BK_TABULAR_COLUMN_BY_KEY and key not in keys:
            keys.append(key)
    return keys



def _bk_ai_group_page_records(raw_records):
    """Gruppiert alle Overlay-Boxen der Seite in visuelle Zeilen.

    Der frühere Toleranzbereich war für eng gesetzte Registerseiten zu groß und
    konnte Nachbarzeilen zusammenziehen. Für die KI-Unterstützung ist das
    besonders schädlich, weil das Modell dann mehrere Personen in einer
    Tabellenzeile zusammenfasst. Deshalb wird hier deutlich strenger gruppiert
    und immer die nächstliegende vorhandene Zeilengruppe gewählt.
    """
    records = [r for r in (raw_records or []) if _bk_tab_clean(r.get("text", ""))]
    if not records:
        return []
    median_height = max(4.0, _median_height(records))
    y_tol = max(3.0, min(7.0, median_height * 0.38))
    ordered = sorted(records, key=lambda r: (
        float(r.get("cy", r.get("y0", 0.0)) or 0.0),
        float(r.get("x0", 0.0) or 0.0),
        int(r.get("index", 0) or 0),
    ))
    groups = []
    for rec in ordered:
        cy = float(rec.get("cy", rec.get("y0", 0.0)) or 0.0)
        best = None
        best_dist = 10 ** 9
        for group in groups[-6:]:
            dist = abs(cy - group["cy"])
            if dist <= y_tol and dist < best_dist:
                best = group
                best_dist = dist
        if best is not None:
            best["records"].append(rec)
            n = len(best["records"])
            best["cy"] = ((best["cy"] * (n - 1)) + cy) / n
        else:
            groups.append({"cy": cy, "records": [rec]})
    out = []
    for group in groups:
        group_records = sorted(group["records"], key=lambda r: float(r.get("x0", 0.0) or 0.0))
        if any(not _bk_tab_is_separator(r.get("text", "")) for r in group_records):
            out.append(group_records)
    return out




def _bk_ai_zone_chunk_candidates(candidates, max_chars=3600, max_rows=6):
    chunks = []
    current = []
    current_chars = 0
    for candidate in candidates or []:
        try:
            size = len(json.dumps(candidate, ensure_ascii=False))
        except Exception:
            size = 200
        if current and (len(current) >= int(max_rows) or current_chars + size > int(max_chars)):
            chunks.append(current)
            current = []
            current_chars = 0
        current.append(candidate)
        current_chars += size
    if current:
        chunks.append(current)
    return chunks





def _bk_ai_zone_page_data_url(item):
    path = str(getattr(item, "path", "") or "")
    if not path:
        return ""
    try:
        return _page_to_data_url(path, max_side=1400, image_format="JPEG", jpeg_quality=78)
    except Exception:
        try:
            return _page_to_data_url(path, max_side=1100, image_format="JPEG", jpeg_quality=72)
        except Exception:
            return ""


def _bk_ai_compact_json(obj, limit=26000):
    raw = json.dumps(obj, ensure_ascii=False, separators=(",", ":"))
    if len(raw) <= int(limit):
        return raw
    # Notbremse für lokale 8k-Kontexte: zuerst Overlay-Kontext einkürzen,
    # Zonen und Kandidaten bleiben erhalten.
    try:
        obj = dict(obj)
        overlays = list(obj.get("overlay_boxes") or [])
        while overlays and len(json.dumps(obj, ensure_ascii=False, separators=(",", ":"))) > int(limit):
            overlays = overlays[: max(1, int(len(overlays) * 0.80))]
            obj["overlay_boxes"] = overlays
            obj["overlay_boxes_truncated"] = True
        return json.dumps(obj, ensure_ascii=False, separators=(",", ":"))
    except Exception:
        return raw[: int(limit)]




def _bk_ai_zone_post_json_with_events(worker, payload, window=None, dlg=None):
    result = {"data": None, "error": None}
    def _target():
        try:
            result["data"] = worker._post_json(payload)
        except Exception as exc:
            result["error"] = exc
    thread = threading.Thread(target=_target, daemon=True)
    thread.start()
    while thread.is_alive():
        try:
            if dlg is not None:
                spinner = getattr(dlg, "spinner", None)
                if spinner is not None and hasattr(spinner, "_advance"):
                    spinner._advance()
                dlg.repaint()
            if _BK_QEventLoop is not None:
                QApplication.processEvents(_BK_QEventLoop.AllEvents, 50)
            else:
                QApplication.processEvents()
        except Exception:
            try:
                QApplication.processEvents()
            except Exception:
                pass
        time.sleep(0.02)
    if result["error"] is not None:
        raise result["error"]
    return result["data"]


def _bk_ai_zone_force_no_think(text: str) -> str:
    """Qwen-/Reasoning-Modelle sollen für diesen Export direkt JSON liefern.
    /no_think wird am Anfang und am Ende gesetzt, weil manche lokale
    Chat-Templates nur einen der beiden Textteile zuverlässig beachten.
    """
    value = str(text or "").strip()
    directive = (
        "/no_think\n"
        "Keine Analyse. Kein Schritt-für-Schritt-Denken. Keine Aufzählung der Verarbeitung. "
        "Gib sofort ausschließlich das finale JSON zurück.\n"
    )
    if not value:
        return directive + "/no_think"
    cleaned = re.sub(r"(?im)^\s*/no_think\s*", "", value).strip()
    return directive + cleaned + "\n/no_think"


def _bk_ai_zone_output_tokens_for_chunk(token_limit: int, chunk_len: int) -> int:
    """Antwortlimit für Tabellen-JSON.
    Die Einstellung in LM-Optionen ist das harte Antwortlimit für diesen Task.
    Sie darf nicht anhand der Chunkgröße auf wenige Tausend Tokens gekappt werden,
    weil Qwen-Reasoning-Modelle sonst vor der eigentlichen JSON-Ausgabe im
    reasoning_content stehen bleiben.
    """
    try:
        limit = int(token_limit)
    except Exception:
        limit = 900
    # Die Kandidaten werden in kleine Chunks zerlegt. Für das reine Zuordnen
    # von Overlay-Texten zu Spalten reichen kurze Antworten; große Limits
    # fördern bei Qwen/Reasoning-Modellen nur lange Denktexte.
    return max(256, min(1000, limit))


def _bk_ai_zone_message_text_from_response(worker, data: dict) -> str:
    """Extrahiert JSON auch dann, wenn lokale Reasoning-Backends es ungewöhnlich ablegen."""
    try:
        return worker._extract_message_content(data)
    except Exception as exc:
        choices = data.get("choices") if isinstance(data, dict) else None
        choice0 = choices[0] if isinstance(choices, list) and choices else {}
        message = choice0.get("message", {}) if isinstance(choice0, dict) else {}
        finish_reason = str(choice0.get("finish_reason", "") if isinstance(choice0, dict) else "")
        for value in (
            message.get("content") if isinstance(message, dict) else None,
            message.get("reasoning_content") if isinstance(message, dict) else None,
            choice0.get("content") if isinstance(choice0, dict) else None,
            choice0.get("text") if isinstance(choice0, dict) else None,
        ):
            if isinstance(value, str) and value.strip():
                raw = value.strip()
                rows = _bk_extract_ai_rows_json(raw)
                if rows:
                    return raw
        details = str(exc)
        if finish_reason:
            details = f"{details} (finish_reason={finish_reason})"
        raise RuntimeError(details)


def _bk_ai_zone_call(window, worker, model_id, prompt, max_tokens, page_data_url="", dlg=None):
    # System-Prompt ausschliesslich aus den Sprachdateien; ohne Treffer wird
    # ein sprachneutraler Minimal-Systemtext verwendet.
    system_text = _bk_tr_registry(window, "ai_prompt_export_zones_system")
    if not system_text or system_text == "ai_prompt_export_zones_system":
        system_text = 'JSON only: {"rows":[{...}]}'

    system_text = _bk_ai_zone_force_no_think(system_text)
    user_content_text = _bk_ai_zone_force_no_think(prompt)
    user_message = {"role": "user", "content": user_content_text}
    if page_data_url:
        user_message = {
            "role": "user",
            "content": [
                {"type": "text", "text": user_content_text},
                {"type": "image_url", "image_url": {"url": page_data_url}},
            ],
        }
    base_payload = {
        "model": model_id,
        "messages": [
            {"role": "system", "content": system_text},
            user_message,
        ],
        **worker._build_sampling_payload(response_format=None, override_max_tokens=max_tokens),
    }
    base_payload["temperature"] = 0.03
    base_payload["top_p"] = 0.8
    base_payload["top_k"] = 10
    base_payload["enable_thinking"] = False
    base_payload["thinking"] = False
    base_payload["reasoning"] = {"enabled": False, "effort": "none"}
    base_payload["reasoning_effort"] = "none"
    base_payload["chat_template_kwargs"] = {"enable_thinking": False, "thinking": False}

    # Erster Versuch: mit Bild, wenn verfügbar. Falls das lokale Modell/Backend keine
    # Vision-Nachrichten unterstützt, folgt automatisch ein reiner Textversuch.
    attempts = []
    if page_data_url:
        attempts.append(("vision", True))
    attempts.append(("text", False))

    last_error = None
    for _mode, use_image in attempts:
        for response_format in ({"type": "json_object"}, None):
            payload = dict(base_payload)
            if not use_image:
                payload["messages"] = [
                    {"role": "system", "content": system_text},
                    {"role": "user", "content": user_content_text},
                ]
            if response_format is not None:
                payload["response_format"] = response_format
            try:
                data = _bk_ai_zone_post_json_with_events(worker, payload, window, dlg)
                content = _bk_ai_zone_message_text_from_response(worker, data)
                rows = _bk_extract_ai_rows_json(content)
                if rows:
                    return rows
            except Exception as exc:
                last_error = exc
                message = str(exc).lower()
                can_retry = (
                    response_format is not None and (
                        "response_format" in message
                        or "json" in message
                        or "context" in message
                        or "n_ctx" in message
                        or "tokens" in message
                        or "400" in message
                    )
                )
                vision_failed = use_image and (
                    "image" in message
                    or "vision" in message
                    or "content" in message
                    or "unsupported" in message
                    or "invalid" in message
                    or "400" in message
                )
                if can_retry or vision_failed:
                    continue
                raise
    if last_error is not None:
        raise last_error
    return []




def _bk_export_zones_ai_busy_dialog(window, chunk_count=0):
    title = _bk_tab_tr(window, "export_zones_ai_busy_title")
    message = _bk_tab_tr(
        window,
        "export_zones_ai_busy_message",
    )
    dlg = BusyStatusDialog(title, message, getattr(window, "_tr", None), window)
    try:
        dlg.btn_cancel.hide()
    except Exception:
        pass
    if chunk_count and chunk_count > 1:
        try:
            dlg.set_status(message + "\n\n" + _bk_tab_tr(window, "export_zones_ai_busy_chunks"))
        except Exception:
            pass
    return dlg


def _bk_export_zones_ai_set_busy_status(window, dlg, current_chunk, total_chunks):
    if dlg is None:
        return
    base = _bk_tab_tr(
        window,
        "export_zones_ai_busy_message",
    )
    if total_chunks > 1:
        suffix = _bk_tab_tr(window, "export_zones_ai_busy_progress").format(int(current_chunk), int(total_chunks))
        text = base + "\n\n" + suffix
    else:
        text = base
    try:
        dlg.set_status(text)
    except Exception:
        pass
    try:
        QApplication.processEvents()
    except Exception:
        pass


def _bk_export_zones_ai_warning_text(window, reason: str, details: str = "") -> str:
    reason_l = str(reason or "").strip().lower()
    details_s = str(details or "").strip()
    if "reasoning" in reason_l or "finish_reason=length" in details_s.lower():
        base = _bk_tab_tr(
            window,
            "export_zones_ai_warn_reasoning",
        )
    elif "context" in reason_l or "n_ctx" in details_s.lower() or "context" in details_s.lower():
        base = _bk_tab_tr(
            window,
            "export_zones_ai_warn_context",
        )
    elif "token" in reason_l or "tokens" in details_s.lower() or "length" in details_s.lower():
        base = _bk_tab_tr(
            window,
            "export_zones_ai_warn_token",
        )
    elif "model" in reason_l:
        base = _bk_tab_tr(
            window,
            "export_zones_ai_warn_model",
        )
    elif "too_few_rows" in reason_l:
        base = _bk_tab_tr(
            window,
            "export_zones_ai_warn_too_few_rows",
        )
    elif "no_rows" in reason_l:
        base = _bk_tab_tr(
            window,
            "export_zones_ai_warn_no_rows",
        )
    else:
        base = _bk_tab_tr(
            window,
            "export_zones_ai_warn_general",
        )
    fallback = _bk_tab_tr(
        window,
        "export_zones_ai_warn_fallback",
    )
    if details_s:
        details_s = details_s[:1200]
        return f"{base}\n\n{fallback}\n\n{details_s}"
    return f"{base}\n\n{fallback}"


def _bk_export_zones_ai_warn(window, reason: str, details: str = ""):
    try:
        # Nicht mehrfach dieselbe Fehlermeldung innerhalb desselben Exports anzeigen.
        key = (str(reason or ""), str(details or "")[:180])
        seen = getattr(window, "_bk_export_zones_ai_warning_seen", set())
        if key in seen:
            return
        seen = set(seen)
        seen.add(key)
        setattr(window, "_bk_export_zones_ai_warning_seen", seen)
        QMessageBox.warning(
            window,
            _bk_tab_tr(window, "export_zones_ai_warn_title"),
            _bk_export_zones_ai_warning_text(window, reason, details),
        )
    except Exception:
        try:
            print("BK export zones AI warning:", reason, details)
        except Exception:
            pass


def _bk_ai_row_value_count(value: str, pattern: str) -> int:
    try:
        return len(re.findall(pattern, str(value or ""), flags=re.IGNORECASE))
    except Exception:
        return 0


def _bk_ai_rows_look_merged(rows, candidates_count: int = 0) -> bool:
    rows = list(rows or [])
    if not rows:
        return True
    if candidates_count and candidates_count >= 8 and len(rows) < max(4, int(candidates_count * 0.35)):
        return True
    suspicious = 0
    for row in rows:
        if not isinstance(row, dict):
            continue
        age = str(row.get("age_original", "") or "")
        date = str(row.get("date_original", "") or "")
        year = str(row.get("year_resolved", "") or "")
        number = str(row.get("number", "") or "")
        # Mehrere Alters-, Datums-, Jahres- oder Seitenzahlen in einer Zelle sind
        # ein starkes Zeichen, dass Nachbarzeilen zusammengezogen wurden.
        if _bk_ai_row_value_count(age, r"\b\d{1,3}\s*(?:Jahr|Jahre|Wochen?|Monate?|Tage?)\b") > 1:
            suspicious += 1
        if _bk_ai_row_value_count(date, r"\b\d{1,2}\s*\.\s*(?:[IVXLCDM]{1,8}|\d{1,2})") > 1:
            suspicious += 1
        if _bk_ai_row_value_count(year, r"\b1[5-9]\d{2}\b|\b20\d{2}\b") > 1:
            suspicious += 1
        if _bk_ai_row_value_count(number, r"\b\d{1,4}\b") > 2:
            suspicious += 1
    return suspicious >= max(2, int(len(rows) * 0.35))



def _bk_ai_zone_free_only_keys():
    return {"unknown", "heading", "subheading", "original_line"}


def _bk_ai_zone_has_real_data(cells, column_keys):
    cells = cells or {}
    data_keys = [key for key in (column_keys or []) if key not in _bk_ai_zone_free_only_keys()]
    if not data_keys:
        return any(_bk_tab_clean(cells.get(key, "")) for key in (column_keys or []))
    return any(_bk_tab_clean(cells.get(key, "")) for key in data_keys)


def _bk_ai_zone_row_has_real_data(row, column_keys):
    row = row or {}
    data_keys = [key for key in (column_keys or []) if key not in _bk_ai_zone_free_only_keys()]
    if not data_keys:
        return any(_bk_tab_clean(row.get(key, "")) for key in (column_keys or []))
    return any(_bk_tab_clean(row.get(key, "")) for key in data_keys)


def _bk_ai_zone_row_signature(row, column_keys):
    values = []
    for key in (column_keys or _BK_TABULAR_KEYS):
        if key in {"id"}:
            continue
        val = _bk_tab_clean((row or {}).get(key, "")).casefold()
        if val:
            values.append((key, val))
    return tuple(values)



def _bk_build_transcription_rows_with_zones_ai_compact(window, item, image_size=None, zones=None):
    if AIRevisionWorker is None:
        _bk_export_zones_ai_warn(window, "model")
        return []
    if not item or not getattr(item, "results", None):
        _bk_export_zones_ai_warn(window, "no_rows")
        return []
    try:
        model_id = window._resolve_ai_model_id()
    except Exception:
        model_id = (getattr(window, "ai_model_id", "") or "").strip()
    if not model_id:
        _bk_export_zones_ai_warn(window, "model")
        return []
    zones = [_bk_export_clean_zone(zone, idx) for idx, zone in enumerate(zones or [])]
    zones = [zone for zone in zones if zone]
    column_keys = _bk_ai_zone_column_keys(zones)
    if not column_keys:
        return []
    try:
        _text, _kr, _pil, record_views = item.results
    except Exception:
        return []
    raw_records = _records_from_views(record_views)
    raw_records = [record for record in raw_records if not _bk_zone_is_ignored(record, zones)]
    raw_records = _bk_tab_expand_numeric_records(raw_records)
    if not raw_records:
        return []
    candidates = _bk_ai_zone_candidate_rows(raw_records, zones, column_keys)
    if not candidates:
        _bk_export_zones_ai_warn(window, "no_rows")
        return []
    allowed = {key: _bk_tabular_column_title(window, key) for key in column_keys}
    context_payload = _bk_ai_zone_context_payload(window, raw_records, zones, image_size=image_size)
    page_data_url = _bk_ai_zone_page_data_url(item)
    try:
        token_limit = _bk_export_zones_ai_max_tokens(window, 900)
        worker = AIRevisionWorker(
            path=getattr(item, "path", "") or "",
            recs=[],
            lm_model=model_id,
            endpoint=getattr(window, "ai_endpoint", "http://127.0.0.1:1234/v1/chat/completions"),
            enable_thinking=False,
            tr_func=getattr(window, "_tr", None),
            temperature=0.05,
            top_p=0.8,
            max_tokens=token_limit,
        )
        chunks = _bk_ai_zone_chunk_candidates(candidates)
        busy_dlg = _bk_export_zones_ai_busy_dialog(window, len(chunks))
        try:
            busy_dlg.show()
            QApplication.processEvents()
        except Exception:
            busy_dlg = None
        rows = []
        try:
            for idx, chunk in enumerate(chunks, start=1):
                _bk_export_zones_ai_set_busy_status(window, busy_dlg, idx, len(chunks))
                prompt = _bk_ai_zone_prompt(window, column_keys, allowed, chunk, context_payload=context_payload)
                # Antwortlimit aus der LM-Option verwenden. Nicht wieder hart auf 500-1400
                # Tokens kappen, sonst liefern Qwen-/Reasoning-Modelle nur reasoning_content
                # und werden vor dem eigentlichen JSON abgeschnitten.
                max_tokens = _bk_ai_zone_output_tokens_for_chunk(token_limit, len(chunk))
                chunk_rows = _bk_ai_zone_call(window, worker, model_id, prompt, max_tokens, page_data_url=page_data_url, dlg=busy_dlg)
                rows.extend(_bk_ai_zone_sanitize_rows_for_chunk(chunk_rows, chunk, column_keys))
                try:
                    QApplication.processEvents()
                except Exception:
                    pass
        finally:
            try:
                if busy_dlg is not None:
                    busy_dlg.close()
                    busy_dlg.deleteLater()
            except Exception:
                pass
        finished = _bk_finish_ai_zone_rows(rows)
        if _bk_ai_rows_look_merged(finished, len(candidates)):
            _bk_export_zones_ai_warn(window, "too_few_rows", f"AI rows: {len(finished)} | candidates: {len(candidates)}")
            return []
        return finished
    except Exception as exc:
        try:
            print(f"BK export zones AI fallback: {exc}")
        except Exception:
            pass
        _bk_export_zones_ai_warn(window, "error", str(exc))
        return []


try:
    _bk_build_transcription_rows_with_zones_ai = _bk_build_transcription_rows_with_zones_ai_compact
except Exception:
    pass

try:
    __all__.extend([
        '_bk_ai_zone_candidate_rows',
        '_bk_ai_zone_chunk_candidates',
        '_bk_ai_zone_prompt',
        '_bk_build_transcription_rows_with_zones_ai_compact',
    ])
    register_globals('bk', globals(), sorted(set(__all__)))
except Exception:
    pass


# ---------------------------------------------------------------------------
# Präzisere Griffpunkte für Exportbereiche:
# Resize startet nur noch direkt auf den sichtbaren Eckgriffen. Klicks auf
# Kanten oder im Bereich selbst verschieben/selektieren den Bereich.
# ---------------------------------------------------------------------------

def _bk_zone_corner_handle_hit(rect, scene_pos, handle_size=None):
    try:
        size = float(handle_size if handle_size is not None else getattr(_BKExportZoneRectItem, "HANDLE_SIZE", 18.0))
    except Exception:
        size = 18.0
    half = max(4.0, min(10.0, size / 2.0))
    x = float(scene_pos.x())
    y = float(scene_pos.y())
    corners = (
        ("tl", float(rect.left()), float(rect.top())),
        ("tr", float(rect.right()), float(rect.top())),
        ("bl", float(rect.left()), float(rect.bottom())),
        ("br", float(rect.right()), float(rect.bottom())),
    )
    for mode, cx, cy in corners:
        if (cx - half) <= x <= (cx + half) and (cy - half) <= y <= (cy + half):
            return mode
    return ""


def _bk_zone_handle_mode_for_scene_rect_precise(rect, scene_pos, tolerance=None):
    mode = _bk_zone_corner_handle_hit(rect, scene_pos)
    if mode:
        return mode
    try:
        if rect.normalized().contains(scene_pos):
            return "move"
    except Exception:
        pass
    return ""


def _bk_zone_rect_handle_at_corners_only(self, pos):
    rect = self.rect().normalized()
    mode = _bk_zone_corner_handle_hit(rect, pos)
    if mode:
        return mode
    try:
        if rect.contains(pos):
            return "move"
    except Exception:
        pass
    return "move"


try:
    _bk_zone_handle_mode_for_scene_rect = _bk_zone_handle_mode_for_scene_rect_precise
    _BKExportZoneRectItem._handle_at = _bk_zone_rect_handle_at_corners_only
except Exception:
    pass

try:
    __all__.extend([
        '_bk_zone_corner_handle_hit',
        '_bk_zone_handle_mode_for_scene_rect_precise',
        '_bk_zone_rect_handle_at_corners_only',
    ])
    register_globals('bk', globals(), sorted(set(__all__)))
except Exception:
    pass


# Finaler KI-Fallback-Pfad: verhindert, dass alte gestapelte Wrapper nach einem
# KI-Fehler wieder eine frühere, fehlerhafte KI-Variante aufrufen.
def _bk_zone_rows_for_item_global_ai_warn_final(window, item, image_size, use_zones):
    zones = _bk_effective_export_zones(window, item, use_zones)
    try:
        setattr(window, "_bk_export_zones_ai_warning_seen", set())
    except Exception:
        pass
    if zones and _bk_export_zones_ai_enabled(window):
        ai_rows = _bk_build_transcription_rows_with_zones_ai(window, item, image_size, zones)
        if ai_rows:
            return ai_rows
    try:
        _text, _kr, _pil, record_views = item.results
    except Exception:
        return []
    return _bk_build_transcription_rows_with_zones(record_views, image_size, zones) if zones else _bk_build_transcription_rows(record_views, image_size)

try:
    _bk_zone_rows_for_item_global = _bk_zone_rows_for_item_global_ai_warn_final
except Exception:
    pass


# ---------------------------------------------------------------------------
# Final fix: Export-KI muss gesetzte Auswahlbereiche strikt respektieren.
# Die Bereiche sind keine Full-Page-Spaltenprojektion mehr, sondern begrenzen
# die zu exportierenden Overlay-Boxen. Das Seitenbild bleibt nur Hintergrund-
# kontext. Zusätzlich: Überschrift/Teil-Überschrift als Spalten/Datentypen.
# ---------------------------------------------------------------------------
_BK_HEADING_COLUMNS = [
    ("heading", "export_column_heading", "Überschrift", 24.0, 4.0),
    ("subheading", "export_column_subheading", "Teil-Überschrift", 26.0, 4.2),
]
try:
    _existing_keys = [c[0] for c in _BK_TABULAR_COLUMNS]
    _insert_at = 0
    for _col in reversed(_BK_HEADING_COLUMNS):
        if _col[0] not in _existing_keys:
            _BK_TABULAR_COLUMNS.insert(_insert_at, _col)
    _BK_TABULAR_KEYS = [column[0] for column in _BK_TABULAR_COLUMNS]
    _BK_TABULAR_HEADERS = [column[2] for column in _BK_TABULAR_COLUMNS]
    _BK_TABULAR_COLUMN_BY_KEY = {column[0]: column for column in _BK_TABULAR_COLUMNS}
    _BK_ZONE_TYPES = tuple(_BK_TABULAR_KEYS + ["ignore"])
    _BK_ZONE_LEGACY_TYPES.update({
        "heading": "heading",
        "subheading": "subheading",
        "year": "year_resolved",
        "years": "year_resolved",
    })
except Exception:
    pass


def _bk_zone_value_for_column(column_key: str, text: str) -> str:
    text = _bk_tab_clean(text)
    if not text:
        return ""
    if column_key in {"heading", "subheading", "unknown", "other", "original_line"}:
        # Diese Spalten sind bewusst freie Textfelder. Bei "Unbekannt" darf nie
        # automatisch das Label "Unbekannt" geschrieben werden, sondern nur der
        # echte Text aus dem Auswahlbereich.
        return text
    if column_key == "family_name":
        fam, _given, _full = _bk_tab_split_name(text)
        return fam or text.strip(" ,.;:-")
    if column_key == "given_names":
        _fam, given, _full = _bk_tab_split_name(text)
        return given or text.strip(" ,.;:-")
    if column_key == "full_name":
        fam, given, full = _bk_tab_split_name(text)
        return full or _bk_tab_join_text_fragments([fam, given]) or text.strip(" ,.;:-")
    if column_key == "middle_names":
        return _bk_zone_middle_names(text) or text.strip(" ,.;:-")
    if column_key == "relationship":
        return _bk_tab_extract_relationship(text) or text.strip(" ,.;:-")
    if column_key == "age_original":
        return _bk_zone_age_value(text) or text.strip(" ,.;:-")
    if column_key == "days":
        day, _month = _bk_zone_date_parts(text)
        return day or text.strip(" ,.;:-")
    if column_key == "months":
        _day, month = _bk_zone_date_parts(text)
        return month or text.strip(" ,.;:-")
    if column_key in {"date_original", "marriage_date", "death_date"}:
        return _bk_tab_extract_date(text) or text.strip(" ,.;:-")
    if column_key == "year_resolved":
        return _bk_zone_year_value(text) or text.strip(" ,.;:-")
    if column_key in {"place_in_source", "residence_place", "birth_place", "death_place", "street"}:
        return _bk_zone_place_value(text) or text.strip(" ,.;:-")
    if column_key == "house_number":
        return _bk_zone_house_number(text) or text.strip(" ,.;:-")
    if column_key == "postal_code":
        return _bk_zone_postal_code(text) or text.strip(" ,.;:-")
    if column_key == "number":
        return _bk_zone_page_number_value(text) or text.strip(" ,.;:-")
    if column_key in {
        "occupation",
        "spouse",
        "partner",
        "children",
        "grandchildren",
        "grandparents",
        "great_grandparents",
        "great_great_grandparents",
    }:
        return text
    return text


def _bk_zone_join_original_line(row_values, fallback=""):
    parts = []
    for key in _BK_TABULAR_KEYS:
        if key in {"original_line", "unknown", "heading", "subheading"}:
            continue
        value = _bk_tab_clean(row_values.get(key, ""))
        if value:
            parts.append(value)
    joined = _bk_tab_join_text_fragments(parts)
    return joined or _bk_tab_clean(fallback)


def _bk_ai_record_zone_match(record, zones):
    """Findet den besten echten Bereichstreffer einer Overlay-Box.

    Anders als die alte KI-Projektion auf die ganze Seite zählt hier nur eine
    echte Überlappung mit dem gezeichneten Auswahlbereich. Dadurch exportiert die
    KI nicht mehr die komplette Seite, wenn oben nur Beispielbereiche gesetzt
    wurden.
    """
    try:
        rx0 = float(record.get("x0", 0.0) or 0.0)
        ry0 = float(record.get("y0", 0.0) or 0.0)
        rx1 = float(record.get("x1", rx0) or rx0)
        ry1 = float(record.get("y1", ry0) or ry0)
        rw = max(1.0, rx1 - rx0)
        rh = max(1.0, ry1 - ry0)
        rcx = (rx0 + rx1) / 2.0
        rcy = (ry0 + ry1) / 2.0
    except Exception:
        return "", 0.0
    best = (0.0, "")
    for zone in zones or []:
        ztype = str(zone.get("type", "") or "")
        if not ztype or ztype == "ignore" or ztype not in _BK_TABULAR_COLUMN_BY_KEY:
            continue
        try:
            zx0 = float(zone.get("x0", 0.0) or 0.0)
            zy0 = float(zone.get("y0", 0.0) or 0.0)
            zx1 = float(zone.get("x1", 0.0) or 0.0)
            zy1 = float(zone.get("y1", 0.0) or 0.0)
        except Exception:
            continue
        if zx1 < zx0:
            zx0, zx1 = zx1, zx0
        if zy1 < zy0:
            zy0, zy1 = zy1, zy0
        ix0, iy0 = max(rx0, zx0), max(ry0, zy0)
        ix1, iy1 = min(rx1, zx1), min(ry1, zy1)
        overlap_area = max(0.0, ix1 - ix0) * max(0.0, iy1 - iy0)
        area_ratio = overlap_area / max(1.0, rw * rh)
        center_inside = zx0 <= rcx <= zx1 and zy0 <= rcy <= zy1
        score = area_ratio + (0.60 if center_inside else 0.0)
        if score > best[0]:
            best = (score, ztype)
    # Eine sehr kleine Berührung am Rand reicht nicht; Mittelpunkt oder sichtbare
    # Fläche im Bereich müssen plausibel sein.
    return (best[1], best[0]) if best[0] >= 0.12 else ("", 0.0)




def _bk_ai_zone_context_payload(window, raw_records, zones, image_size=None, max_boxes=220):
    """Kontext für die Export-KI.

    Zonen werden vollständig übergeben. Overlay-Boxen werden aber auf echte
    Bereichstreffer begrenzt, damit das Modell nicht wieder die ganze Seite als
    Exportauftrag interpretiert.
    """
    clean_zones = []
    zone_rows = []
    for idx, zone in enumerate(zones or [], start=1):
        clean = _bk_export_clean_zone(zone, idx - 1)
        if not clean or clean.get("type") == "ignore":
            continue
        clean_zones.append(clean)
        zone_rows.append({
            "id": idx,
            "type": str(clean.get("type", "")),
            "label": _bk_export_zone_title(window, str(clean.get("type", ""))),
            "rect": [
                int(round(float(clean.get("x0", 0.0) or 0.0))),
                int(round(float(clean.get("y0", 0.0) or 0.0))),
                int(round(float(clean.get("x1", 0.0) or 0.0))),
                int(round(float(clean.get("y1", 0.0) or 0.0))),
            ],
        })
    overlay_rows = []
    for idx, rec in enumerate(sorted(raw_records or [], key=lambda r: (
        float(r.get("y0", 0.0) or 0.0),
        float(r.get("x0", 0.0) or 0.0),
        int(r.get("index", 0) or 0),
    )), start=1):
        text = _bk_ai_short_text(rec.get("text", ""), 90)
        if not text:
            continue
        zkey, _score = _bk_ai_record_zone_match(rec, clean_zones)
        if not zkey:
            continue
        overlay_rows.append({
            "i": idx,
            "type_hint": zkey,
            "x": int(round(float(rec.get("x0", 0.0) or 0.0))),
            "y": int(round(float(rec.get("y0", 0.0) or 0.0))),
            "w": int(round(float(rec.get("w", (float(rec.get("x1", 0.0) or 0.0) - float(rec.get("x0", 0.0) or 0.0))) or 0.0))),
            "h": int(round(float(rec.get("h", (float(rec.get("y1", 0.0) or 0.0) - float(rec.get("y0", 0.0) or 0.0))) or 0.0))),
            "text": text,
        })
        if len(overlay_rows) >= int(max_boxes):
            break
    payload = {
        "image_size": list(image_size or []) if image_size else [],
        "zones": zone_rows,
        "overlay_boxes": overlay_rows,
    }
    if len(overlay_rows) >= int(max_boxes):
        payload["overlay_boxes_truncated"] = True
    return payload


def _bk_ai_zone_prompt(window, column_keys, allowed_titles, candidates, context_payload=None):
    compact_columns = {key: allowed_titles.get(key, key) for key in column_keys}
    column_text = ", ".join(column_keys)
    payload = {
        "columns": compact_columns,
        "zones": (context_payload or {}).get("zones", []),
        "overlay_boxes_in_selected_zones_only": (context_payload or {}).get("overlay_boxes", []),
        "candidates_from_selected_zones_only": candidates,
    }
    if (context_payload or {}).get("image_size"):
        payload["image_size"] = (context_payload or {}).get("image_size")
    payload_json = _bk_ai_compact_json(payload, limit=16000)
    # Der Prompt kommt ausschliesslich aus den Sprachdateien (de/en/fr).
    # Es gibt keinen fest einkodierten Sprachtext mehr; sollten die
    # Sprachdateien unvollstaendig sein, wird ein sprachneutraler
    # Datenprompt aus Schluesseln, Spalten und Kontext gebaut.
    template = _bk_tr_registry(window, "ai_prompt_export_zones_user_compact")
    prompt = None
    if template and template != "ai_prompt_export_zones_user_compact":
        try:
            prompt = template.format(column_text, json.dumps(compact_columns, ensure_ascii=False), payload_json)
        except Exception:
            prompt = None
    if prompt is None:
        prompt = (
            "/no_think\n"
            "JSON: {\"rows\":[{...}]}\n"
            "keys: " + column_text + "\n"
            "columns: " + json.dumps(compact_columns, ensure_ascii=False) + "\n"
            "context:\n" + payload_json + "\n/no_think"
        )
    return _bk_ai_zone_orientation_hint(window, prompt)


def _bk_ai_zone_orientation_hint(window, prompt: str) -> str:
    """Bei explizit gewaehltem Hoch-/Querformat einen Hinweis anfuegen.

    Der Hinweistext stammt vollstaendig aus den Sprachdateien; ist der
    Schluessel nirgends vorhanden, wird der Prompt unveraendert gelassen.
    """
    try:
        orientation = bk_get_export_orientation(window)
    except Exception:
        orientation = "auto"
    if orientation not in {"portrait", "landscape"}:
        return prompt
    key = "ai_prompt_export_orientation_portrait" if orientation == "portrait" else "ai_prompt_export_orientation_landscape"
    hint = _bk_tr_registry(window, key)
    if not hint or hint == key:
        return prompt
    if prompt.rstrip().endswith("/no_think"):
        stripped = prompt.rstrip()
        return stripped[: -len("/no_think")].rstrip("\n") + "\n" + hint + "\n/no_think"
    return prompt + "\n" + hint


def _bk_finish_ai_zone_rows_base(rows):
    out = []
    seen = set()
    for row in rows or []:
        item = {key: _bk_tab_clean(row.get(key, "")) for key in _BK_TABULAR_KEYS}
        # Niemals Spaltenlabels als Inhalt übernehmen.
        for key, bad_values in {
            "unknown": {"unbekannt", "unknown", "inconnu"},
            "heading": {"überschrift", "heading", "titre"},
            "subheading": {"teil-überschrift", "teilüberschrift", "subheading", "sous-titre"},
        }.items():
            if item.get(key, "").strip().casefold() in bad_values:
                item[key] = ""
        if not item.get("full_name"):
            item["full_name"] = _bk_tab_join_text_fragments([item.get("family_name", ""), item.get("given_names", "")])
        if not item.get("middle_names") and item.get("full_name"):
            item["middle_names"] = _bk_zone_middle_names(item.get("full_name", ""))
        if item.get("date_original"):
            item["days"] = item.get("days") or _bk_zone_date_parts(item.get("date_original", ""))[0]
            item["months"] = item.get("months") or _bk_zone_date_parts(item.get("date_original", ""))[1]
        if not item.get("original_line"):
            item["original_line"] = _bk_zone_join_original_line(item)
        signature = tuple(_bk_tab_clean(item.get(k, "")).casefold() for k in _BK_TABULAR_KEYS if k not in {"id"})
        if signature in seen:
            continue
        seen.add(signature)
        item["year_in_source"] = item.get("year_resolved", "")
        item["place_resolved"] = item.get("place_in_source", "")
        item["id"] = f"entry_{len(out) + 1:04d}"
        out.append(item)
    return out

try:
    __all__.extend([
        '_bk_ai_record_zone_match',
        '_bk_ai_zone_candidate_rows',
        '_bk_ai_zone_context_payload',
        '_bk_ai_zone_prompt',
        '_bk_finish_ai_zone_rows',
    ])
    register_globals('bk', globals(), sorted(set(__all__)))
except Exception:
    pass


# ---------------------------------------------------------------------------
# Fix: Exportdialog-Abbruch darf keine nachgelagerten Exportdialoge starten.
# ---------------------------------------------------------------------------
def _bk_tabular_export_single_interactive_safe_cancel(self, item: TaskItem, fmt: str):
    fmt_l = str(fmt or "").lower().lstrip(".")
    if fmt_l in _BK_TABLE_EXPORT_FMTS:
        result = _bk_column_choice_dialog(self, fmt_l, include_text_modes=False)
        if result is None or result.get("cancelled"):
            return None
        self._bk_export_current_column_keys = result.get("columns") or list(_BK_TABULAR_DEFAULT_KEYS)
        self._bk_export_use_zones = bool(result.get("use_zones", False))
        if result.get("remembered"):
            self._bk_export_selected_column_keys = result.get("columns") or list(_BK_TABULAR_DEFAULT_KEYS)
    elif fmt_l in _BK_TEXT_LAYOUT_FMTS:
        result = _bk_column_choice_dialog(self, fmt_l, include_text_modes=True)
        if result is None or result.get("cancelled"):
            return None
        self._bk_export_text_layout_mode = str(result.get("mode") or "original")
        self._bk_export_current_column_keys = result.get("columns") or list(_BK_TABULAR_DEFAULT_KEYS)
        self._bk_export_use_zones = bool(result.get("use_zones", False))
        if result.get("remembered"):
            self._bk_export_selected_column_keys = result.get("columns") or list(_BK_TABULAR_DEFAULT_KEYS)
    if callable(_BK_TABULAR_PREV_EXPORT_SINGLE):
        return _BK_TABULAR_PREV_EXPORT_SINGLE(self, item, fmt)
    return None


def _bk_tabular_export_batch_safe_cancel(self, items, fmt: str):
    fmt_l = str(fmt or "").lower().lstrip(".")
    if fmt_l in _BK_TABLE_EXPORT_FMTS:
        result = _bk_column_choice_dialog(self, fmt_l, include_text_modes=False)
        if result is None or result.get("cancelled"):
            return None
        self._bk_export_current_column_keys = result.get("columns") or list(_BK_TABULAR_DEFAULT_KEYS)
        self._bk_export_use_zones = bool(result.get("use_zones", False))
        if result.get("remembered"):
            self._bk_export_selected_column_keys = result.get("columns") or list(_BK_TABULAR_DEFAULT_KEYS)
    elif fmt_l in _BK_TEXT_LAYOUT_FMTS:
        result = _bk_column_choice_dialog(self, fmt_l, include_text_modes=True)
        if result is None or result.get("cancelled"):
            return None
        self._bk_export_text_layout_mode = str(result.get("mode") or "original")
        self._bk_export_current_column_keys = result.get("columns") or list(_BK_TABULAR_DEFAULT_KEYS)
        self._bk_export_use_zones = bool(result.get("use_zones", False))
        if result.get("remembered"):
            self._bk_export_selected_column_keys = result.get("columns") or list(_BK_TABULAR_DEFAULT_KEYS)
    if callable(_BK_TABULAR_PREV_EXPORT_BATCH):
        return _BK_TABULAR_PREV_EXPORT_BATCH(self, items, fmt)
    return None

try:
    MainWindow._export_single_interactive = _bk_tabular_export_single_interactive_safe_cancel
    MainWindow._export_batch = _bk_tabular_export_batch_safe_cancel
except Exception:
    pass

try:
    __all__.extend([
        '_bk_ai_zone_has_real_data',
        '_bk_ai_zone_sanitize_rows_for_chunk',
        '_bk_tabular_export_single_interactive_safe_cancel',
        '_bk_tabular_export_batch_safe_cancel',
    ])
    register_globals('bk', globals(), sorted(set(__all__)))
except Exception:
    pass


# ---------------------------------------------------------------------------
# Final fix: stabiler Exportdialog + Überschrift/Teil-Überschrift als Kontext.
# ---------------------------------------------------------------------------





try:
    __all__.extend([
        '_bk_ai_context_records_value',
        '_bk_ai_zone_candidate_rows',
        '_bk_ai_zone_sanitize_rows_for_chunk',
    ])
    register_globals('bk', globals(), sorted(set(__all__)))
except Exception:
    pass

# ---------------------------------------------------------------------------
# Wiederherstellung aus gutem Zwischenstand:
# Die Zeilen-/Spalten-Befüllung der Export-KI nutzt wieder die bewährte
# Kandidaten- und Sanitize-Logik des funktionierenden Zwischenstands.
# Neuere Dialog-/UI-Fixes bleiben davor erhalten.
# ---------------------------------------------------------------------------
# ---------------------------------------------------------------------------
# Final fix: stabiler Exportdialog + Überschrift/Teil-Überschrift als Kontext.
# ---------------------------------------------------------------------------
def _bk_ai_context_records_value(records, row_y, median_height=12.0):
    records = list(records or [])
    if not records:
        return ""
    try:
        row_y = float(row_y or 0.0)
    except Exception:
        row_y = 0.0
    candidates = []
    for rec in records:
        try:
            cy = float(rec.get("cy", rec.get("y0", 0.0)) or 0.0)
        except Exception:
            cy = 0.0
        if cy <= row_y + max(2.0, float(median_height or 12.0) * 0.35):
            candidates.append((cy, rec))
    if not candidates:
        return ""
    candidates.sort(key=lambda item: item[0])
    best_y = candidates[-1][0]
    same_line = [
        rec for cy, rec in candidates
        if abs(cy - best_y) <= max(3.0, float(median_height or 12.0) * 0.65)
    ]
    same_line.sort(key=lambda rec: float(rec.get("x0", 0.0) or 0.0))
    return _bk_tab_join_text_fragments([
        _bk_tab_clean(rec.get("text", "")) for rec in same_line
        if _bk_tab_clean(rec.get("text", ""))
    ])


def _bk_ai_zone_candidate_rows_base(raw_records, zones, column_keys):
    """KI-Kandidaten nur aus Auswahlbereichen; Überschriften werden als Kontext übernommen."""
    if not raw_records or not zones or not column_keys:
        return []
    zones = [_bk_export_clean_zone(zone, idx) for idx, zone in enumerate(zones or [])]
    zones = [z for z in zones if z and z.get("type") != "ignore"]
    if not zones:
        return []

    matched = []
    context_records = {"heading": [], "subheading": []}
    for rec in raw_records or []:
        text = _bk_tab_clean(rec.get("text", ""))
        if not text or _bk_tab_is_separator(text):
            continue
        key, score = _bk_ai_record_zone_match(rec, zones)
        if not key:
            continue
        item = dict(rec)
        item["_zone_key"] = key
        item["_zone_score"] = score
        if key in {"heading", "subheading"}:
            context_records.setdefault(key, []).append(item)
            continue
        if key in column_keys:
            matched.append(item)

    if not matched:
        return []

    try:
        median_height = _median_height(raw_records)
    except Exception:
        median_height = 12.0

    out = []
    for group in _bk_ai_group_page_records(matched):
        fragments = []
        cells = {}
        all_text_parts = []
        y_values = []
        for rec in group:
            text = _bk_tab_clean(rec.get("text", ""))
            if not text or _bk_tab_is_separator(text):
                continue
            zone_key = str(rec.get("_zone_key", "") or "")
            if not zone_key or zone_key not in column_keys:
                continue
            x0 = int(round(float(rec.get("x0", 0.0) or 0.0)))
            y0 = int(round(float(rec.get("y0", 0.0) or 0.0)))
            y_values.append(float(rec.get("cy", rec.get("y0", 0.0)) or 0.0))
            fragments.append({"x": x0, "y": y0, "type_hint": zone_key, "text": _bk_ai_short_text(text, 110)})
            all_text_parts.append(text)
            value = _bk_zone_value_for_column(zone_key, text) or text
            if value:
                current = cells.get(zone_key, "")
                cells[zone_key] = _bk_tab_join_text_fragments([current, _bk_ai_short_text(value, 140)]) if current else _bk_ai_short_text(value, 140)

        if not fragments or not cells:
            continue
        if not _bk_ai_zone_has_real_data(cells, column_keys):
            continue

        row_y = int(round(sum(y_values) / max(1, len(y_values)))) if y_values else 0

        # Überschrift und Teil-Überschrift sind Kontextspalten:
        # Sie erzeugen keine eigenen Datenzeilen, sondern werden für die
        # darunterliegenden Registereinträge übernommen.
        for ctx_key in ("heading", "subheading"):
            if ctx_key in column_keys and not _bk_tab_clean(cells.get(ctx_key, "")):
                ctx_value = _bk_ai_context_records_value(context_records.get(ctx_key), row_y, median_height)
                if ctx_value:
                    cells[ctx_key] = _bk_ai_short_text(ctx_value, 160)

        original_line = _bk_tab_join_text_fragments(all_text_parts)
        if "original_line" in column_keys and original_line:
            cells["original_line"] = _bk_ai_short_text(original_line, 240)
        out.append({
            "n": len(out) + 1,
            "y": row_y,
            "fragments": fragments,
            "cells": cells,
            "line_text": _bk_ai_short_text(original_line, 260),
        })
    return out


def _bk_ai_zone_sanitize_rows_for_chunk_base(rows, chunk, column_keys):
    """Begrenzt KI-Ausgaben auf Kandidaten und ergänzt Kontextspalten aus candidates."""
    max_rows = max(0, len(chunk or []))
    cleaned = []
    seen = set()
    for row in rows or []:
        if not isinstance(row, dict):
            continue
        item = {key: _bk_tab_clean(row.get(key, "")) for key in _BK_TABULAR_KEYS}

        # Candidate-Kontext zurückschreiben, falls das Modell Überschrift /
        # Teil-Überschrift auslässt. Dadurch stimmen diese Spalten auch bei
        # guter KI-Zeilenzerlegung zuverlässig.
        cand = (chunk or [])[len(cleaned)] if len(cleaned) < len(chunk or []) else {}
        cand_cells = cand.get("cells", {}) if isinstance(cand, dict) else {}
        for ctx_key in ("heading", "subheading"):
            if ctx_key in (column_keys or []) and not _bk_tab_clean(item.get(ctx_key, "")):
                cand_value = _bk_tab_clean(cand_cells.get(ctx_key, ""))
                if cand_value:
                    item[ctx_key] = cand_value

        for key, bad_values in {
            "unknown": {"unbekannt", "unknown", "inconnu"},
            "heading": {"überschrift", "heading", "titre"},
            "subheading": {"teil-überschrift", "teilüberschrift", "subheading", "sous-titre"},
        }.items():
            if item.get(key, "").strip().casefold() in bad_values:
                item[key] = ""
        if not _bk_ai_zone_row_has_real_data(item, column_keys):
            continue
        sig = _bk_ai_zone_row_signature(item, column_keys)
        if not sig or sig in seen:
            continue
        seen.add(sig)
        cleaned.append(item)
        if max_rows and len(cleaned) >= max_rows:
            break
    return cleaned

try:
    __all__.extend([
        '_bk_ai_context_records_value',
        '_bk_ai_zone_candidate_rows',
        '_bk_ai_zone_sanitize_rows_for_chunk',
    ])
    register_globals('bk', globals(), sorted(set(__all__)))
except Exception:
    pass

# ---------------------------------------------------------------------------
# Finaler Feinschliff: Überschrift/Teil-Überschrift nur einmal ausgeben.
# Die bewährte LM-Zeilen-/Spaltenlogik bleibt unverändert; nur wiederholte
# Kontextwerte werden vor und nach der KI-Antwort ausgedünnt.
# ---------------------------------------------------------------------------
def _bk_ai_heading_once_key(value: str) -> str:
    txt = _bk_tab_clean(value)
    if not txt:
        return ""
    return re.sub(r"\s+", " ", txt).strip().casefold()


def _bk_ai_keep_heading_values_once_in_candidates(candidates):
    seen = {"heading": set(), "subheading": set()}
    out = []
    for cand in candidates or []:
        if not isinstance(cand, dict):
            out.append(cand)
            continue
        cand = dict(cand)
        cells = dict(cand.get("cells", {}) or {})
        for ctx_key in ("heading", "subheading"):
            marker = _bk_ai_heading_once_key(cells.get(ctx_key, ""))
            if not marker:
                continue
            if marker in seen[ctx_key]:
                cells[ctx_key] = ""
            else:
                seen[ctx_key].add(marker)
        cand["cells"] = cells
        out.append(cand)
    return out


def _bk_ai_keep_heading_values_once_in_rows(rows):
    seen = {"heading": set(), "subheading": set()}
    out = []
    for row in rows or []:
        if not isinstance(row, dict):
            continue
        item = dict(row)
        for ctx_key in ("heading", "subheading"):
            marker = _bk_ai_heading_once_key(item.get(ctx_key, ""))
            if not marker:
                continue
            if marker in seen[ctx_key]:
                item[ctx_key] = ""
            else:
                seen[ctx_key].add(marker)
        out.append(item)
    return out




def _bk_ai_zone_candidate_rows(raw_records, zones, column_keys):
    candidates = _bk_ai_zone_candidate_rows_base(raw_records, zones, column_keys)
    return _bk_ai_keep_heading_values_once_in_candidates(candidates)


def _bk_ai_zone_sanitize_rows_for_chunk(rows, chunk, column_keys):
    cleaned = _bk_ai_zone_sanitize_rows_for_chunk_base(rows, chunk, column_keys)
    # Innerhalb eines Chunks keine wiederholten Überschriften/Teil-Überschriften
    # übernehmen. Chunk-übergreifend greift zusätzlich _bk_finish_ai_zone_rows.
    return _bk_ai_keep_heading_values_once_in_rows(cleaned)


def _bk_finish_ai_zone_rows(rows):
    finished = _bk_finish_ai_zone_rows_base(rows)
    return _bk_ai_keep_heading_values_once_in_rows(finished)


try:
    __all__.extend([
        '_bk_ai_heading_once_key',
        '_bk_ai_keep_heading_values_once_in_candidates',
        '_bk_ai_keep_heading_values_once_in_rows',
        '_bk_ai_zone_candidate_rows',
        '_bk_ai_zone_sanitize_rows_for_chunk',
        '_bk_finish_ai_zone_rows',
    ])
    register_globals('bk', globals(), sorted(set(__all__)))
except Exception:
    pass

# ---------------------------------------------------------------------------
# Heuristik-Fix für sensible Bereiche ohne KI-Unterstützung.
# Die funktionierende KI-Logik bleibt unverändert. Ohne KI werden breite
# OCR-Zeilen nicht mehr als Spaltentext mehrfach zusammengeführt, sondern als
# eine Registerzeile interpretiert und anschließend feldweise ergänzt.
# ---------------------------------------------------------------------------
def _bk_heuristic_zone_overlap_info(record, zone):
    try:
        rx0 = float(record.get("x0", 0.0) or 0.0)
        ry0 = float(record.get("y0", 0.0) or 0.0)
        rx1 = float(record.get("x1", rx0) or rx0)
        ry1 = float(record.get("y1", ry0) or ry0)
        zx0 = float(zone.get("x0", 0.0) or 0.0)
        zy0 = float(zone.get("y0", 0.0) or 0.0)
        zx1 = float(zone.get("x1", 0.0) or 0.0)
        zy1 = float(zone.get("y1", 0.0) or 0.0)
    except Exception:
        return 0.0, 0.0, False
    if rx1 < rx0:
        rx0, rx1 = rx1, rx0
    if ry1 < ry0:
        ry0, ry1 = ry1, ry0
    if zx1 < zx0:
        zx0, zx1 = zx1, zx0
    if zy1 < zy0:
        zy0, zy1 = zy1, zy0
    ix0, iy0 = max(rx0, zx0), max(ry0, zy0)
    ix1, iy1 = min(rx1, zx1), min(ry1, zy1)
    iw, ih = max(0.0, ix1 - ix0), max(0.0, iy1 - iy0)
    if iw <= 0.0 or ih <= 0.0:
        return 0.0, 0.0, False
    rec_area = max(1.0, (rx1 - rx0) * (ry1 - ry0))
    zone_area = max(1.0, (zx1 - zx0) * (zy1 - zy0))
    cx, cy = (rx0 + rx1) / 2.0, (ry0 + ry1) / 2.0
    center_inside = zx0 <= cx <= zx1 and zy0 <= cy <= zy1
    return (iw * ih) / rec_area, (iw * ih) / zone_area, center_inside


def _bk_heuristic_record_zone_matches(record, zones, min_ratio=0.035):
    matches = []
    for zone in zones or []:
        ztype = str(zone.get("type", "") or "")
        if not ztype or ztype == "ignore" or ztype not in _BK_TABULAR_COLUMN_BY_KEY:
            continue
        rec_ratio, zone_ratio, center_inside = _bk_heuristic_zone_overlap_info(record, zone)
        if rec_ratio <= 0.0 and zone_ratio <= 0.0 and not center_inside:
            continue
        score = rec_ratio + (0.65 if center_inside else 0.0) + min(0.25, zone_ratio * 0.04)
        if rec_ratio >= float(min_ratio) or center_inside or zone_ratio >= 0.015:
            matches.append((score, ztype, zone))
    matches.sort(key=lambda item: item[0], reverse=True)
    return matches


def _bk_heuristic_context_matches(record, zones, ctx_key):
    out = []
    for zone in zones or []:
        if zone.get("type") != ctx_key:
            continue
        rec_ratio, zone_ratio, center_inside = _bk_heuristic_zone_overlap_info(record, zone)
        # Kontextbereiche müssen deutlich getroffen werden. Sonst werden breite
        # normale OCR-Zeilen fälschlich als Überschrift übernommen.
        if center_inside or zone_ratio >= 0.18 or rec_ratio >= 0.20:
            out.append((zone_ratio + rec_ratio + (1.0 if center_inside else 0.0), zone))
    out.sort(key=lambda item: item[0], reverse=True)
    return out


def _bk_heuristic_record_is_wide_line(record, page_width, matched_types):
    try:
        width = float(record.get("x1", 0.0) or 0.0) - float(record.get("x0", 0.0) or 0.0)
    except Exception:
        width = 0.0
    meaningful = [t for t in set(matched_types or []) if t not in {"heading", "subheading", "ignore"}]
    return width >= max(150.0, float(page_width or 0.0) * 0.32) or len(meaningful) >= 3


def _bk_heuristic_put_value(row_values, key, value):
    key = str(key or "")
    value = _bk_tab_clean(value)
    if not key or not value:
        return
    if key in {"heading", "subheading"}:
        # Kontextspalten werden separat genau einmal gesetzt.
        return
    if key in {"family_name", "given_names", "full_name"}:
        _bk_zone_apply_name_values(row_values, key, value)
        return
    current = _bk_tab_clean(row_values.get(key, ""))
    if not current:
        row_values[key] = value
    elif value.casefold() not in [p.casefold() for p in re.split(r"\s+", current) if p] and value.casefold() not in current.casefold():
        row_values[key] = _bk_tab_join_text_fragments([current, value])


def _bk_heuristic_apply_parsed_line(row_values, parsed):
    parsed = parsed or {}
    for key in ("family_name", "given_names", "relationship", "age_original", "date_original", "year_resolved", "number", "original_line"):
        value = _bk_tab_clean(parsed.get(key, ""))
        if value and not _bk_tab_clean(row_values.get(key, "")):
            row_values[key] = value
    place = _bk_tab_clean(parsed.get("place_in_source", "") or parsed.get("place_resolved", ""))
    if place:
        for place_key in ("birth_place", "place_in_source", "residence_place", "death_place"):
            if place_key in _BK_TABULAR_KEYS and not _bk_tab_clean(row_values.get(place_key, "")):
                row_values[place_key] = place
    full = _bk_tab_join_text_fragments([row_values.get("family_name", ""), row_values.get("given_names", "")])
    if full and "full_name" in _BK_TABULAR_KEYS and not row_values.get("full_name"):
        row_values["full_name"] = full
    if row_values.get("date_original"):
        day, month = _bk_zone_date_parts(row_values["date_original"])
        if day and "days" in _BK_TABULAR_KEYS and not row_values.get("days"):
            row_values["days"] = day
        if month and "months" in _BK_TABULAR_KEYS and not row_values.get("months"):
            row_values["months"] = month


def _bk_heuristic_context_events(raw_records, zones):
    events = []
    seen = set()
    for ctx_key in ("heading", "subheading"):
        for rec in raw_records or []:
            text = _bk_tab_clean(rec.get("text", ""))
            if not text or _bk_tab_is_separator(text):
                continue
            if not _bk_heuristic_context_matches(rec, zones, ctx_key):
                continue
            marker = (ctx_key, round(float(rec.get("cy", rec.get("y0", 0.0)) or 0.0), 2), text.casefold())
            if marker in seen:
                continue
            seen.add(marker)
            events.append({
                "key": ctx_key,
                "value": text,
                "y": float(rec.get("cy", rec.get("y0", 0.0)) or 0.0),
                "x": float(rec.get("x0", 0.0) or 0.0),
            })
    events.sort(key=lambda item: (float(item.get("y", 0.0) or 0.0), float(item.get("x", 0.0) or 0.0)))
    return events


def _bk_heuristic_apply_context_once(rows, events):
    if not rows or not events:
        return rows
    used = set()
    for event in events:
        key = str(event.get("key", "") or "")
        value = _bk_tab_clean(event.get("value", ""))
        if key not in {"heading", "subheading"} or key not in _BK_TABULAR_KEYS or not value:
            continue
        sig = (key, value.casefold())
        if sig in used:
            continue
        ey = float(event.get("y", 0.0) or 0.0)
        target = None
        for row in rows:
            if float(row.get("_source_y", 0.0) or 0.0) >= ey - 1.0:
                target = row
                break
        if target is None:
            continue
        if not _bk_tab_clean(target.get(key, "")):
            target[key] = value
            used.add(sig)
    # Sicherheit: Wiederholte gleiche Überschrift/Teil-Überschrift entfernen.
    seen = {"heading": set(), "subheading": set()}
    for row in rows:
        for key in ("heading", "subheading"):
            value = _bk_tab_clean(row.get(key, ""))
            if not value:
                continue
            marker = value.casefold()
            if marker in seen[key]:
                row[key] = ""
            else:
                seen[key].add(marker)
    return rows


def _bk_build_transcription_rows_with_zones(record_views, image_size=None, zones=None):
    zones = [_bk_export_clean_zone(zone, idx) for idx, zone in enumerate(zones or [])]
    zones = [zone for zone in zones if zone]
    zones.sort(key=lambda z: (int(z.get("order", 0) or 0), float(z.get("x0", 0.0)), float(z.get("y0", 0.0))))
    if not zones:
        return _bk_build_transcription_rows(record_views, image_size)
    raw_records = _records_from_views(record_views)
    if not raw_records:
        return []
    raw_records = [record for record in raw_records if not _bk_zone_is_ignored(record, zones)]
    raw_records = _bk_tab_expand_numeric_records(raw_records)
    if not raw_records:
        return []
    page_width, _page_height = _page_size(image_size, raw_records)

    data_records = []
    data_seen = set()
    for rec in raw_records:
        text = _bk_tab_clean(rec.get("text", ""))
        if not text or _bk_tab_is_separator(text):
            continue
        matches = _bk_heuristic_record_zone_matches(rec, zones)
        data_matches = [m for m in matches if m[1] not in {"heading", "subheading", "ignore"}]
        if not data_matches:
            continue
        item = dict(rec)
        item["_zone_matches"] = data_matches
        key = (
            round(float(item.get("x0", 0.0) or 0.0), 2),
            round(float(item.get("y0", 0.0) or 0.0), 2),
            text.casefold(),
        )
        if key in data_seen:
            continue
        data_seen.add(key)
        data_records.append(item)
    if not data_records:
        return _bk_build_transcription_rows(record_views, image_size)

    rows = []
    seen_rows = set()
    for group in _bk_ai_group_page_records(data_records):
        row_values = {key: "" for key in _BK_TABULAR_KEYS}
        group = sorted(group, key=lambda r: float(r.get("x0", 0.0) or 0.0))
        y_values = []
        original_parts = []
        for rec in group:
            text = _bk_tab_clean(rec.get("text", ""))
            if not text or _bk_tab_is_separator(text):
                continue
            y_values.append(float(rec.get("cy", rec.get("y0", 0.0)) or 0.0))
            matches = list(rec.get("_zone_matches") or _bk_heuristic_record_zone_matches(rec, zones))
            matched_types = [m[1] for m in matches]
            if _bk_heuristic_record_is_wide_line(rec, page_width, matched_types):
                parsed = _bk_tab_make_row_from_record(rec, "", "") or {}
                _bk_heuristic_apply_parsed_line(row_values, parsed)
                original_parts.append(text)
                # Falls eine breite Zeile zusätzlich in expliziten Spaltenbereichen liegt,
                # dürfen Spezialspalten wie Unbekannt/Sonstiges ihren echten Bereichstext behalten.
                for _score, ztype, _zone in matches:
                    if ztype in {"unknown", "other"}:
                        _bk_heuristic_put_value(row_values, ztype, _bk_zone_value_for_column(ztype, text))
                continue
            for _score, ztype, _zone in matches:
                if ztype in {"heading", "subheading", "ignore"}:
                    continue
                value = _bk_zone_value_for_column(ztype, text)
                _bk_heuristic_put_value(row_values, ztype, value)
            original_parts.append(text)

        if not any(_bk_tab_clean(row_values.get(key, "")) for key in _BK_TABULAR_KEYS if key not in {"heading", "subheading", "unknown", "original_line"}):
            # Nur freie Kontext-/Unbekannt-Spalten sind keine echte Registerzeile.
            continue
        if not row_values.get("original_line"):
            row_values["original_line"] = _bk_zone_join_original_line(row_values, _bk_tab_join_text_fragments(original_parts))
        if not row_values.get("full_name") and "full_name" in _BK_TABULAR_KEYS:
            full = _bk_tab_join_text_fragments([row_values.get("family_name", ""), row_values.get("given_names", "")])
            if full:
                row_values["full_name"] = full
        if row_values.get("date_original"):
            day, month = _bk_zone_date_parts(row_values["date_original"])
            if day and "days" in _BK_TABULAR_KEYS and not row_values.get("days"):
                row_values["days"] = day
            if month and "months" in _BK_TABULAR_KEYS and not row_values.get("months"):
                row_values["months"] = month
        if row_values.get("birth_place") and not row_values.get("place_in_source"):
            row_values["place_in_source"] = row_values["birth_place"]
        if row_values.get("place_in_source") and not row_values.get("birth_place") and "birth_place" in _BK_TABULAR_KEYS:
            row_values["birth_place"] = row_values["place_in_source"]

        row = {key: row_values.get(key, "") for key in _BK_TABULAR_KEYS}
        row["year_in_source"] = row.get("year_resolved", "")
        row["place_resolved"] = row.get("place_in_source", "")
        row["_source_y"] = sum(y_values) / max(1, len(y_values)) if y_values else 0.0
        row["_source_x"] = min([float(r.get("x0", 0.0) or 0.0) for r in group] or [0.0])
        sig = tuple(_bk_tab_clean(row.get(key, "")).casefold() for key in ("family_name", "given_names", "full_name", "age_original", "date_original", "number", "original_line"))
        if sig in seen_rows:
            continue
        seen_rows.add(sig)
        rows.append(row)

    rows.sort(key=lambda row: (float(row.get("_source_y", 0.0) or 0.0), float(row.get("_source_x", 0.0) or 0.0)))
    rows = _bk_heuristic_apply_context_once(rows, _bk_heuristic_context_events(raw_records, zones))

    out = []
    last_year = ""
    last_place = ""
    for idx, row in enumerate(rows, start=1):
        row = dict(row)
        ysrc = _bk_tab_clean(row.get("year_in_source", row.get("year_resolved", "")))
        if _bk_tab_is_ditto(ysrc):
            row["year_resolved"] = last_year
            row["year_in_source"] = "„"
        elif ysrc:
            row["year_resolved"] = ysrc
            row["year_in_source"] = ysrc
            last_year = ysrc
        psrc = _bk_tab_clean(row.get("place_in_source", ""))
        if _bk_tab_is_ditto(psrc):
            row["place_resolved"] = last_place
            row["place_in_source"] = "„"
        elif psrc:
            row["place_resolved"] = psrc
            last_place = psrc
        row["id"] = f"entry_{idx:04d}"
        row.pop("_source_x", None)
        row.pop("_source_y", None)
        out.append(row)
    return out

try:
    __all__.extend([
        '_bk_heuristic_zone_overlap_info',
        '_bk_heuristic_record_zone_matches',
        '_bk_heuristic_context_matches',
        '_bk_heuristic_record_is_wide_line',
        '_bk_heuristic_apply_context_once',
    ])
    register_globals('bk', globals(), sorted(set(__all__)))
except Exception:
    pass


# ---------------------------------------------------------------------------
# Simple Tabellenexport: rohe Overlay-Boxen als einfache Layout-Tabelle
# ---------------------------------------------------------------------------
_BK_SIMPLE_TABLE_FMTS = {"docx", "word", "odt", "xlsx", "excel", "ods", "calc"}
_BK_SIMPLE_TABLE_SPREADSHEET_FMTS = {"xlsx", "excel", "ods", "calc"}
_BK_SIMPLE_TABLE_TEXT_FMTS = {"docx", "word", "odt"}


def _bk_simple_record_text(record):
    return _bk_tab_clean(record.get("text", ""))


def _bk_simple_table_records(record_views):
    records = []
    for rec in _records_from_views(record_views):
        text = _bk_simple_record_text(rec)
        if not text or _bk_tab_is_separator(text):
            continue
        try:
            x0 = float(rec.get("x0", rec.get("cx", 0.0)) or 0.0)
            y0 = float(rec.get("y0", rec.get("cy", 0.0)) or 0.0)
            x1 = float(rec.get("x1", x0 + rec.get("w", 1.0)) or x0 + 1.0)
            y1 = float(rec.get("y1", y0 + rec.get("h", 1.0)) or y0 + 1.0)
        except Exception:
            continue
        if x1 < x0:
            x0, x1 = x1, x0
        if y1 < y0:
            y0, y1 = y1, y0
        records.append({
            "text": text,
            "x0": x0,
            "y0": y0,
            "x1": x1,
            "y1": y1,
            "cx": (x0 + x1) / 2.0,
            "cy": (y0 + y1) / 2.0,
            "w": max(1.0, x1 - x0),
            "h": max(1.0, y1 - y0),
        })
    return records


def _bk_simple_group_rows(records):
    """Gruppiert Overlay-Boxen für Tabelle (einfach) in visuelle Zeilen.

    Wichtig: Diese Funktion darf keine benachbarten Registerzeilen per
    Kettenreaktion zusammenziehen. Genau das war der Grund, warum mehrere
    Overlay-Boxen in einer Tabellenzelle landeten. Deshalb wird hier nur nach
    stabilen y-Mittelpunkten gruppiert; die wachsende Gruppenhöhe wird nicht als
    zusätzlicher Trefferbereich benutzt.
    """
    if not records:
        return []
    heights = sorted(max(1.0, float(r.get("h", 1.0) or 1.0)) for r in records)
    med_h = heights[len(heights) // 2] if heights else 12.0
    # Sehr enge Registerzeilen brauchen eine kleine Toleranz. Größere Werte
    # lassen zwei benachbarte OCR-Zeilen zu einer Tabellenzeile verschmelzen.
    y_tol = max(2.5, min(7.0, med_h * 0.38))
    groups = []
    for rec in sorted(records, key=lambda r: (float(r.get("cy", 0.0) or 0.0), float(r.get("x0", 0.0) or 0.0))):
        cy = float(rec.get("cy", 0.0) or 0.0)
        y0 = float(rec.get("y0", cy) or cy)
        best = None
        best_delta = None
        for group in groups:
            gy = float(group.get("cy", 0.0) or 0.0)
            gy0 = float(group.get("y0", gy) or gy)
            delta = min(abs(cy - gy), abs(y0 - gy0))
            if delta <= y_tol and (best_delta is None or delta < best_delta):
                best = group
                best_delta = delta
        if best is None:
            groups.append({
                "cy": cy,
                "y0": y0,
                "records": [rec],
            })
            continue
        best["records"].append(rec)
        n = len(best["records"])
        # Stabiler Mittelwert; keine Ausweitung über y1, damit keine Ketten-
        # Verschmelzung über nahe Folgezeilen entsteht.
        best["cy"] = (float(best["cy"]) * (n - 1) + cy) / n
        best["y0"] = (float(best["y0"]) * (n - 1) + y0) / n
    groups.sort(key=lambda g: (float(g.get("cy", 0.0)), min(float(r.get("x0", 0.0)) for r in g.get("records", []) or [{"x0": 0.0}])))
    return [sorted(g.get("records", []) or [], key=lambda r: (float(r.get("x0", 0.0) or 0.0), float(r.get("cy", 0.0) or 0.0))) for g in groups]





def _bk_simple_normalize_token_text(text):
    value = _bk_tab_clean(text)
    if not value:
        return ""
    # Bei OCR-Zeilen kleben Datums-/Jahres-/Ortsangaben oft direkt zusammen.
    # Für den einfachen Tabellenexport werden solche harten Übergänge nur als
    # Wortgrenzen behandelt, nicht inhaltlich interpretiert.
    value = re.sub(r",(?=\()", ", ", value)
    value = re.sub(r"(?<=\))(?=\d)", " ", value)
    value = re.sub(r"(?<=[A-Za-zÀ-ÿÄÖÜäöüß])(?=\d{1,2}[.])", " ", value)
    value = re.sub(r"(?<=\d)[.](?=[A-Za-zÀ-ÿÄÖÜäöüß])", ". ", value)
    value = re.sub(r"(?<=\d)(?=[A-Za-zÀ-ÿÄÖÜäöüß])", " ", value)
    value = re.sub(r"\s+", " ", value).strip()
    return value







def _bk_simple_col_widths(matrix):
    if not matrix:
        return [12.0]
    count = max(1, max(len(row) for row in matrix))
    widths = []
    for c in range(count):
        vals = [str(row[c] if c < len(row) else "") for row in matrix]
        longest = max([len(v) for v in vals] or [8])
        widths.append(max(8.0, min(34.0, longest * 0.95 + 2.0)))
    return widths


def _bk_write_simple_matrix_docx(path, matrix, window=None, layout=None):
    try:
        from docx import Document
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except Exception as exc:
        raise RuntimeError(_bk_registry_lookup("err_no_docx_package_short") or "python-docx") from exc
    spans = set((layout or {}).get("spans") or set())
    aligns = list((layout or {}).get("aligns") or [])
    doc = Document()
    try:
        section = doc.sections[0]
        # Bisher war hier immer Querformat erzwungen. Jetzt entscheidet die
        # im Export-Dialog gewaehlte Ausrichtung (Standard weiterhin quer).
        use_landscape = bool(bk_resolve_landscape(True, window))
        if use_landscape:
            section.orientation = 1
            section.page_width, section.page_height = section.page_height, section.page_width
        normal = doc.styles["Normal"]
        normal.font.name = "Arial"
        normal.font.size = Pt(10)
    except Exception:
        pass
    if not matrix:
        doc.save(path)
        return
    col_count = max(1, max(len(row) for row in matrix))
    table = doc.add_table(rows=0, cols=col_count)
    try:
        table.style = "Table Grid"
        table.autofit = True
    except Exception:
        pass
    for r_i, row_values in enumerate(matrix):
        cells = table.add_row().cells
        if r_i in spans and col_count > 1:
            # Abschnittszeile: alle Zellen verbinden, fett und zentriert.
            try:
                merged = cells[0]
                for extra in cells[1:]:
                    merged = merged.merge(extra)
                merged.text = str(row_values[0] if row_values else "")
                for p in merged.paragraphs:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in p.runs:
                        run.bold = True
                        run.font.size = Pt(10)
                continue
            except Exception:
                pass
        for idx in range(col_count):
            cells[idx].text = str(row_values[idx] if idx < len(row_values) else "")
            try:
                for p in cells[idx].paragraphs:
                    if idx < len(aligns) and aligns[idx] == "right":
                        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    for run in p.runs:
                        run.font.size = Pt(10)
            except Exception:
                pass
    doc.save(path)



def _bk_write_simple_matrix_odf(path, matrix, mimetype, text_document=True, layout=None, window=None):
    matrix = matrix or []
    layout = layout or {}
    spans = set(layout.get("spans") or set())
    aligns = list(layout.get("aligns") or [])
    bold_rows = set(layout.get("bold_rows") or set())
    col_count = max(1, max([len(row) for row in matrix] or [1]))

    # Bei raeumlichen Exporten sind die aus den Bildkoordinaten ermittelten
    # Breiten verbindlich. Der alte Writer berechnete sie erneut aus der
    # Textlaenge und machte dadurch beide Namensspalten viel zu schmal.
    source_widths = list(layout.get("widths") or _bk_simple_col_widths(matrix))
    source_widths += [12.0] * max(0, col_count - len(source_widths))
    widths = [max(1.15, min(7.5, float(source_widths[i]) / 6.2))
              for i in range(col_count)]

    # Breite historische Tabellen werden in Writer automatisch auf einer
    # A4-Querformatseite angelegt. Die Spalten werden proportional skaliert,
    # nicht einzeln hart abgeschnitten.
    landscape_auto = bool(col_count >= 5 or sum(widths) > 18.0)
    landscape = bool(bk_resolve_landscape(landscape_auto, window))
    usable_cm = 28.2 if landscape else 19.5
    total_cm = sum(widths)
    if text_document and total_cm > usable_cm:
        factor = usable_cm / max(0.1, total_cm)
        widths = [max(0.95, width * factor) for width in widths]

    font_pt = 8.4 if col_count >= 7 else (9.0 if col_count >= 5 else 10.0)
    column_styles = []
    columns = []
    for idx in range(1, col_count + 1):
        width = widths[idx - 1] if idx - 1 < len(widths) else 2.2
        column_styles.append(
            '<style:style style:name="co%d" style:family="table-column">'
            '<style:table-column-properties style:column-width="%.3fcm"/>'
            '</style:style>' % (idx, width)
        )
        columns.append('<table:table-column table:style-name="co%d"/>' % idx)

    non_span_rows = [idx for idx in range(len(matrix)) if idx not in spans]
    first_data = non_span_rows[0] if non_span_rows else -1
    last_data = non_span_rows[-1] if non_span_rows else -1
    table_rows = []
    for r_i, row in enumerate(matrix):
        cells = []
        if r_i in spans and col_count > 1:
            text = row[0] if row else ""
            cells.append(
                '<table:table-cell table:style-name="ceSection" office:value-type="string" '
                'table:number-columns-spanned="%d"><text:p>%s</text:p></table:table-cell>'
                % (col_count, _bk_ods_text(text))
            )
            cells.append('<table:covered-table-cell table:number-columns-repeated="%d"/>' % (col_count - 1))
        else:
            edge = "Only" if r_i == first_data == last_data else (
                "Top" if r_i == first_data else ("Bottom" if r_i == last_data else "")
            )
            for c_idx in range(col_count):
                text = row[c_idx] if c_idx < len(row) else ""
                if r_i in bold_rows:
                    style = "ceHead"
                else:
                    prefix = "ceRight" if c_idx < len(aligns) and aligns[c_idx] == "right" else "ceBody"
                    style = prefix + edge
                cells.append(
                    '<table:table-cell table:style-name="%s" office:value-type="string">'
                    '<text:p>%s</text:p></table:table-cell>' % (style, _bk_ods_text(text))
                )
        table_rows.append('<table:table-row>%s</table:table-row>' % ''.join(cells))

    def _cell_style(name, align="start", top=False, bottom=False, bold=False):
        borders = [
            'fo:border-left="0.05pt solid #8A8A8A"',
            'fo:border-right="0.05pt solid #8A8A8A"',
            'fo:border-top="%s"' % ('0.5pt solid #555555' if top else 'none'),
            'fo:border-bottom="%s"' % ('0.5pt solid #555555' if bottom else 'none'),
        ]
        weight = ' fo:font-weight="bold"' if bold else ''
        return (
            '<style:style style:name="%s" style:family="table-cell">'
            '<style:table-cell-properties %s fo:padding="0.035cm" '
            'style:vertical-align="middle" fo:wrap-option="wrap"/>'
            '<style:text-properties fo:font-size="%.1fpt"%s style:font-name="Arial"/>'
            '<style:paragraph-properties fo:text-align="%s" fo:margin-top="0cm" '
            'fo:margin-bottom="0cm"/>'
            '</style:style>'
        ) % (name, ' '.join(borders), font_pt, weight, align)

    cell_styles = [
        _cell_style("ceBody"),
        _cell_style("ceBodyTop", top=True),
        _cell_style("ceBodyBottom", bottom=True),
        _cell_style("ceBodyOnly", top=True, bottom=True),
        _cell_style("ceRight", align="end"),
        _cell_style("ceRightTop", align="end", top=True),
        _cell_style("ceRightBottom", align="end", bottom=True),
        _cell_style("ceRightOnly", align="end", top=True, bottom=True),
        _cell_style("ceHead", top=True, bottom=True, bold=True),
        '<style:style style:name="ceSection" style:family="table-cell">'
        '<style:table-cell-properties fo:border="none" fo:padding="0.05cm" '
        'style:vertical-align="middle" fo:wrap-option="wrap"/>'
        '<style:text-properties fo:font-size="%.1fpt" fo:font-weight="bold" style:font-name="Arial"/>'
        '<style:paragraph-properties fo:text-align="center" fo:margin-top="0.08cm" '
        'fo:margin-bottom="0.03cm"/>'
        '</style:style>' % max(font_pt, 9.0),
    ]

    body_tag = 'office:text' if text_document else 'office:spreadsheet'
    content = ''.join([
        '<?xml version="1.0" encoding="UTF-8"?>',
        '<office:document-content xmlns:office="urn:oasis:names:tc:opendocument:xmlns:office:1.0" '
        'xmlns:style="urn:oasis:names:tc:opendocument:xmlns:style:1.0" '
        'xmlns:text="urn:oasis:names:tc:opendocument:xmlns:text:1.0" '
        'xmlns:table="urn:oasis:names:tc:opendocument:xmlns:table:1.0" '
        'xmlns:fo="urn:oasis:names:tc:opendocument:xmlns:xsl-fo-compatible:1.0" '
        'xmlns:svg="urn:oasis:names:tc:opendocument:xmlns:svg-compatible:1.0" office:version="1.2">',
        '<office:font-face-decls><style:font-face style:name="Arial" svg:font-family="Arial"/></office:font-face-decls>',
        '<office:automatic-styles>', ''.join(column_styles), ''.join(cell_styles),
        '</office:automatic-styles>',
        '<office:body><', body_tag, '><table:table table:name="Transkription">',
        ''.join(columns), ''.join(table_rows), '</table:table></', body_tag,
        '></office:body></office:document-content>',
    ])

    page_w, page_h = ((29.7, 21.0) if landscape else (21.0, 29.7))
    orientation = "landscape" if landscape else "portrait"
    styles = ''.join([
        '<?xml version="1.0" encoding="UTF-8"?>',
        '<office:document-styles xmlns:office="urn:oasis:names:tc:opendocument:xmlns:office:1.0" '
        'xmlns:style="urn:oasis:names:tc:opendocument:xmlns:style:1.0" '
        'xmlns:text="urn:oasis:names:tc:opendocument:xmlns:text:1.0" '
        'xmlns:fo="urn:oasis:names:tc:opendocument:xmlns:xsl-fo-compatible:1.0" '
        'xmlns:svg="urn:oasis:names:tc:opendocument:xmlns:svg-compatible:1.0" office:version="1.2">',
        '<office:font-face-decls><style:font-face style:name="Arial" svg:font-family="Arial"/></office:font-face-decls>',
        '<office:styles><style:default-style style:family="paragraph">'
        '<style:text-properties fo:font-size="%.1fpt" style:font-name="Arial"/>'
        '</style:default-style></office:styles>' % font_pt,
        '<office:automatic-styles><style:page-layout style:name="pm1">'
        '<style:page-layout-properties fo:page-width="%.1fcm" fo:page-height="%.1fcm" '
        'style:print-orientation="%s" fo:margin-top="0.65cm" fo:margin-bottom="0.65cm" '
        'fo:margin-left="0.65cm" fo:margin-right="0.65cm"/>'
        '</style:page-layout></office:automatic-styles>' % (page_w, page_h, orientation),
        '<office:master-styles><style:master-page style:name="Standard" style:page-layout-name="pm1"/>'
        '</office:master-styles></office:document-styles>',
    ])
    meta = '<?xml version="1.0" encoding="UTF-8"?><office:document-meta xmlns:office="urn:oasis:names:tc:opendocument:xmlns:office:1.0" xmlns:meta="urn:oasis:names:tc:opendocument:xmlns:meta:1.0" office:version="1.2"><office:meta><meta:generator>Bottled Kraken</meta:generator></office:meta></office:document-meta>'
    settings = '<?xml version="1.0" encoding="UTF-8"?><office:document-settings xmlns:office="urn:oasis:names:tc:opendocument:xmlns:office:1.0" office:version="1.2"><office:settings/></office:document-settings>'
    manifest = '<?xml version="1.0" encoding="UTF-8"?><manifest:manifest xmlns:manifest="urn:oasis:names:tc:opendocument:xmlns:manifest:1.0" manifest:version="1.2"><manifest:file-entry manifest:full-path="/" manifest:media-type="%s"/><manifest:file-entry manifest:full-path="content.xml" manifest:media-type="text/xml"/><manifest:file-entry manifest:full-path="styles.xml" manifest:media-type="text/xml"/><manifest:file-entry manifest:full-path="meta.xml" manifest:media-type="text/xml"/><manifest:file-entry manifest:full-path="settings.xml" manifest:media-type="text/xml"/></manifest:manifest>' % mimetype
    with zipfile.ZipFile(path, "w") as archive:
        info = zipfile.ZipInfo("mimetype")
        info.date_time = (2020, 1, 1, 0, 0, 0)
        info.compress_type = zipfile.ZIP_STORED
        archive.writestr(info, mimetype)
        for name, data in (("content.xml", content), ("styles.xml", styles), ("meta.xml", meta),
                           ("settings.xml", settings), ("META-INF/manifest.xml", manifest)):
            zi = zipfile.ZipInfo(name)
            zi.date_time = (2020, 1, 1, 0, 0, 0)
            zi.compress_type = zipfile.ZIP_DEFLATED
            archive.writestr(zi, data.encode("utf-8"))


def _bk_write_simple_matrix_odt(path, matrix, window=None, layout=None):
    return _bk_write_simple_matrix_odf(path, matrix, "application/vnd.oasis.opendocument.text", True, layout=layout, window=window)


def _bk_write_simple_matrix_ods(path, matrix, window=None, layout=None):
    return _bk_write_simple_matrix_odf(path, matrix, "application/vnd.oasis.opendocument.spreadsheet", False, layout=layout, window=window)


try:
    _BK_SIMPLE_PREV_COLUMN_CHOICE_DIALOG = _bk_column_choice_dialog
except Exception:
    _BK_SIMPLE_PREV_COLUMN_CHOICE_DIALOG = None


def _bk_column_choice_dialog_simple_table(self, fmt=None, include_text_modes=False):
    fmt_l = str(fmt or "").lower().lstrip(".")
    supports_simple = fmt_l in _BK_SIMPLE_TABLE_FMTS
    if not supports_simple:
        if callable(_BK_SIMPLE_PREV_COLUMN_CHOICE_DIALOG):
            return _BK_SIMPLE_PREV_COLUMN_CHOICE_DIALOG(self, fmt, include_text_modes)
        return None

    dlg = QDialog(self)
    dlg.setWindowTitle(_bk_tab_tr(self, "export_text_layout_title" if include_text_modes else "export_table_columns_title"))
    dlg.setModal(True)
    layout = QVBoxLayout(dlg)
    layout.setContentsMargins(16, 14, 16, 14)
    layout.setSpacing(10)
    intro_key = "export_text_layout_intro_extended" if include_text_modes else "export_table_columns_intro"
    label = QLabel(_bk_tab_tr(self, intro_key), dlg)
    label.setWordWrap(True)
    layout.addWidget(label)

    rb_original = rb_lines = None
    rb_simple = QRadioButton(_bk_tab_tr(self, "export_text_layout_table_simple"), dlg)
    rb_table = QRadioButton(_bk_tab_tr(self, "export_text_layout_table"), dlg)
    if QGroupBox is not None:
        mode_box = QGroupBox(_bk_tab_tr(self, "export_layout_mode_group"), dlg)
        mode_layout = QVBoxLayout(mode_box)
        mode_layout.setContentsMargins(12, 10, 12, 10)
    else:
        mode_box = None
        mode_layout = QVBoxLayout()
    if include_text_modes:
        rb_original = QRadioButton(_bk_tab_tr(self, "export_text_layout_original"), dlg)
        rb_lines = QRadioButton(_bk_tab_tr(self, "export_text_layout_lines"), dlg)
        mode = str(getattr(self, "_bk_export_text_layout_mode", "original") or "original").lower()
        rb_original.setChecked(mode not in {"lines", "table", "table_simple"})
        rb_lines.setChecked(mode == "lines")
        rb_table.setChecked(mode == "table")
        rb_simple.setChecked(mode == "table_simple")
        mode_layout.addWidget(rb_original)
        mode_layout.addWidget(rb_lines)
    else:
        # Für reine Spreadsheet-Exporte ist die einfache ABBYY-ähnliche Tabelle
        # der Ausgangszustand. Erst wenn der Benutzer in der laufenden Sitzung
        # bewusst auf erweitert umstellt, bleibt erweitert aktiv.
        current_mode = getattr(self, "_bk_export_table_mode", None)
        if current_mode is None:
            table_mode = "simple" if fmt_l in _BK_SIMPLE_TABLE_SPREADSHEET_FMTS else "table"
        else:
            table_mode = str(current_mode or "table").lower()
        rb_simple.setChecked(table_mode == "simple")
        rb_table.setChecked(table_mode != "simple")
    mode_layout.addWidget(rb_simple)
    mode_layout.addWidget(rb_table)
    if mode_box is not None:
        layout.addWidget(mode_box)
    else:
        layout.addLayout(mode_layout)

    selected_keys = _bk_load_saved_column_keys_for_dialog(self)
    checkboxes = {}
    columns_box = QGroupBox(_bk_tab_tr(self, "export_table_columns_label"), dlg) if QGroupBox is not None else None
    columns_layout = QVBoxLayout(columns_box) if columns_box is not None else QVBoxLayout()
    if columns_box is not None:
        columns_layout.setContentsMargins(12, 10, 12, 10)
    else:
        columns_layout.addWidget(QLabel(_bk_tab_tr(self, "export_table_columns_label"), dlg))
    if QGridLayout is not None:
        grid = QGridLayout()
        grid.setContentsMargins(0, 0, 0, 0)
        grid.setHorizontalSpacing(22)
        grid.setVerticalSpacing(6)
        for idx, key in enumerate(_BK_TABULAR_KEYS):
            cb = QCheckBox(_bk_tabular_column_title(self, key), dlg)
            cb.setChecked(key in selected_keys)
            checkboxes[key] = cb
            grid.addWidget(cb, idx // 3, idx % 3)
        columns_layout.addLayout(grid)
    else:
        for key in _BK_TABULAR_KEYS:
            cb = QCheckBox(_bk_tabular_column_title(self, key), dlg)
            cb.setChecked(key in selected_keys)
            checkboxes[key] = cb
            columns_layout.addWidget(cb)
    quick_row = QHBoxLayout()
    btn_all = QPushButton(_bk_tab_tr(self, "export_table_columns_all"), dlg)
    btn_none = QPushButton(_bk_tab_tr(self, "export_table_columns_none_button"), dlg)
    btn_remember = QPushButton(_bk_tab_tr(self, "export_table_columns_remember"), dlg)
    quick_row.addWidget(btn_all)
    quick_row.addWidget(btn_none)
    quick_row.addWidget(btn_remember)
    quick_row.addStretch(1)
    columns_layout.addLayout(quick_row)
    if columns_box is not None:
        layout.addWidget(columns_box)
    else:
        layout.addLayout(columns_layout)

    zone_box = QGroupBox(_bk_tab_tr(self, "export_zones_group"), dlg) if QGroupBox is not None else None
    zone_layout = QHBoxLayout(zone_box) if zone_box is not None else QHBoxLayout()
    cb_zones = QCheckBox(_bk_tab_tr(self, "export_table_use_zones"), dlg)
    try:
        remembered_zones = self.settings.value("export/table_use_zones", bool(getattr(self, "_bk_export_use_zones", False)), type=bool)
    except Exception:
        remembered_zones = bool(getattr(self, "_bk_export_use_zones", False))
    cb_zones.setChecked(bool(remembered_zones))
    btn_zones = QPushButton(_bk_tab_tr(self, "export_table_define_zones"), dlg)
    zone_layout.addWidget(cb_zones, 1)
    zone_layout.addWidget(btn_zones)
    if zone_box is not None:
        layout.addWidget(zone_box)
    else:
        layout.addLayout(zone_layout)

    hint = QLabel(_bk_tab_tr(self, "export_table_simple_hint"), dlg)
    hint.setWordWrap(True)
    layout.addWidget(hint)

    def current_checked_keys():
        return [key for key, cb in checkboxes.items() if cb.isChecked()]

    def set_all():
        for cb in checkboxes.values():
            cb.setChecked(True)

    def set_none():
        for cb in checkboxes.values():
            cb.setChecked(False)

    result = {"mode": "table", "columns": selected_keys, "remembered": False, "use_zones": cb_zones.isChecked(), "cancelled": False}

    def remember_selection():
        keys = current_checked_keys()
        if not keys and not rb_simple.isChecked():
            QMessageBox.warning(dlg, _bk_tab_tr(self, "warn_title"), _bk_tab_tr(self, "export_table_columns_none"))
            return
        result["remembered"] = True
        if keys:
            result["columns"] = _bk_save_column_keys(self, keys)
        try:
            self.settings.setValue("export/table_use_zones", bool(cb_zones.isChecked()))
            self.settings.sync()
        except Exception:
            pass

    def define_zones():
        task = None
        try:
            task = self._current_task()
        except Exception:
            task = None
        zones = _bk_open_export_zones_dialog(self, task)
        if zones is not None:
            cb_zones.setChecked(bool(zones))

    btn_all.clicked.connect(set_all)
    btn_none.clicked.connect(set_none)
    btn_remember.clicked.connect(remember_selection)
    btn_zones.clicked.connect(define_zones)

    def sync_enabled():
        detailed = rb_table.isChecked()
        if columns_box is not None:
            columns_box.setVisible(detailed)
            columns_box.setEnabled(detailed)
        else:
            for cb in checkboxes.values():
                cb.setVisible(detailed)
                cb.setEnabled(detailed)
            btn_all.setVisible(detailed)
            btn_none.setVisible(detailed)
            btn_remember.setVisible(detailed)
            btn_all.setEnabled(detailed)
            btn_none.setEnabled(detailed)
            btn_remember.setEnabled(detailed)
        if zone_box is not None:
            zone_box.setVisible(detailed)
            zone_box.setEnabled(detailed)
        cb_zones.setEnabled(detailed)
        btn_zones.setEnabled(detailed)
        hint.setVisible(rb_simple.isChecked())
        try:
            dlg.adjustSize()
        except Exception:
            pass

    for rb in (rb_original, rb_lines, rb_table, rb_simple):
        if rb is not None:
            try:
                rb.toggled.connect(sync_enabled)
            except Exception:
                pass
    sync_enabled()

    buttons = QDialogButtonBox(QDialogButtonBox.Ok | QDialogButtonBox.Cancel, dlg)
    try:
        buttons.button(QDialogButtonBox.Ok).setText(_bk_tab_tr(self, "btn_ok"))
        buttons.button(QDialogButtonBox.Cancel).setText(_bk_tab_tr(self, "btn_cancel"))
    except Exception:
        pass

    def cancel_dialog():
        result["cancelled"] = True
        dlg.done(QDialog.Rejected)

    def accept_checked():
        result["cancelled"] = False
        if rb_simple.isChecked():
            mode = "table_simple"
            keys = []
        elif include_text_modes and rb_lines is not None and rb_lines.isChecked():
            mode = "lines"
            keys = []
        elif include_text_modes and rb_original is not None and rb_original.isChecked():
            mode = "original"
            keys = []
        else:
            mode = "table"
            keys = current_checked_keys()
            if not keys:
                QMessageBox.warning(dlg, _bk_tab_tr(self, "warn_title"), _bk_tab_tr(self, "export_table_columns_none"))
                return
        result["mode"] = mode
        result["columns"] = _bk_normalize_column_keys(keys)
        result["use_zones"] = bool(cb_zones.isChecked()) if mode == "table" else False
        try:
            self._bk_export_use_zones = result["use_zones"]
        except Exception:
            pass
        dlg.accept()

    buttons.accepted.connect(accept_checked)
    buttons.rejected.connect(cancel_dialog)
    layout.addWidget(buttons)
    dlg.setMinimumSize(690, 560 if include_text_modes else 500)
    try:
        dlg.resize(760, 640 if include_text_modes else 560)
    except Exception:
        pass
    try:
        exec_result = dlg.exec()
    except Exception:
        result["cancelled"] = True
        return None
    if exec_result != QDialog.Accepted:
        result["cancelled"] = True
        return None
    return result








try:
    _BK_SIMPLE_PREV_EXPORT_SINGLE = MainWindow._export_single_interactive
except Exception:
    _BK_SIMPLE_PREV_EXPORT_SINGLE = None
try:
    _BK_SIMPLE_PREV_EXPORT_BATCH = MainWindow._export_batch
except Exception:
    _BK_SIMPLE_PREV_EXPORT_BATCH = None






try:
    _bk_column_choice_dialog = _bk_column_choice_dialog_simple_table
except Exception:
    pass

try:
    __all__.extend([
        '_bk_write_simple_matrix_docx',
        '_bk_write_simple_matrix_odt',
        '_bk_write_simple_matrix_ods',
        '_bk_column_choice_dialog_simple_table',
    ])
    register_globals('bk', globals(), sorted(set(__all__)))
except Exception:
    pass


# ---------------------------------------------------------------------------
# Simple Tabellenexport: kompakter Dialog, sinnvolle Feldaufteilung und Vorschau
# ---------------------------------------------------------------------------
_BK_SIMPLE_PREVIEW_COLUMNS = [
    ("heading", "export_column_heading", "Überschrift"),
    ("subheading", "export_column_subheading", "Teil-Überschrift"),
    ("family_name", "export_column_family_names", "Familienname"),
    ("given_names", "export_column_given_names_multi", "Vorname(n)"),
    ("extra", "export_column_extra", "Zusatz"),
    ("age", "export_column_age", "Alter"),
    ("date_year", "export_column_date_year", "Datum/Jahr"),
    ("place", "export_column_birth_places", "Ort"),
    ("number", "export_column_page_numbers", "Seitenzahl"),
    ("original", "export_column_original_line", "Originalzeile"),
]
_BK_SIMPLE_PREVIEW_COLUMN_KEYS = [key for key, _tr, _fb in _BK_SIMPLE_PREVIEW_COLUMNS]


def _bk_simple_column_title(window, key):
    for col_key, tr_key, fallback in _BK_SIMPLE_PREVIEW_COLUMNS:
        if col_key == key:
            return _bk_tab_tr(window, tr_key, fallback)
    return str(key or "")


def _bk_simple_join_clean(parts):
    out = []
    for part in parts or []:
        value = _bk_tab_clean(part)
        if value:
            out.append(value)
    return " ".join(out).strip()


def _bk_simple_strip_punct(value):
    return _bk_tab_clean(value).strip(" ,;:.·•|/-–—")


def _bk_simple_row_raw_text(row_records):
    return _bk_tab_join_text_fragments([
        _bk_simple_record_text(rec)
        for rec in sorted(row_records or [], key=lambda r: (float(r.get("x0", 0.0) or 0.0), float(r.get("y0", 0.0) or 0.0)))
    ])


def _bk_simple_is_page_heading(text):
    value = _bk_tab_clean(text).lower()
    return bool(re.fullmatch(r"[-–—=\s]*seite\s*[-–—=\s]*\d+\s*[-–—=\s]*", value))


def _bk_simple_is_subheading(text):
    value = _bk_tab_clean(text)
    if not value:
        return False
    if len(value) <= 80 and re.search(r"[A-Za-zÀ-ÿÄÖÜäöüß]", value) and not re.search(r"\d{1,2}\s*(?:Jahre?|Wochen?|Monate?|Stunden?|Tag|Tage)\b", value, re.I):
        if value.endswith("!") or re.search(r"\b(?:unter|weiter|weitersuchen|fortsetzung|forts\.)\b", value, re.I):
            return True
    return False


def _bk_simple_extract_last_number(value):
    text = _bk_tab_clean(value)
    m = re.search(r"(?:^|\s)(\d{1,4})\s*[.)]?\s*$", text)
    if not m:
        return "", text
    number = m.group(1)
    start = m.start(1)
    rest = text[:start].rstrip(" ,;:.")
    # Vierstellige Jahreszahlen am Ende sind eher Datum/Jahr als Seitenzahl.
    try:
        n = int(number)
        if 1500 <= n <= 2099 and not re.search(r"[A-Za-zÀ-ÿÄÖÜäöüß][.,]\s*$", rest):
            return "", text
    except Exception:
        pass
    return number, rest


def _bk_simple_parse_register_text(text):
    raw = _bk_tab_clean(text)
    cells = {key: "" for key in _BK_SIMPLE_PREVIEW_COLUMN_KEYS}
    cells["original"] = raw
    if not raw:
        return cells
    if _bk_simple_is_page_heading(raw):
        cells["heading"] = raw
        cells["original"] = raw
        return cells
    if _bk_simple_is_subheading(raw):
        cells["subheading"] = raw
        cells["original"] = raw
        return cells

    work = _bk_simple_normalize_token_text(raw)
    number, work_wo_number = _bk_simple_extract_last_number(work)
    cells["number"] = number
    work = work_wo_number

    # Zusatzinformationen in Klammern bleiben als ein Feld zusammen.
    extras = re.findall(r"\([^)]*\)", work)
    if extras:
        cells["extra"] = _bk_simple_join_clean(extras)
        for extra in extras:
            work = work.replace(extra, " ")

    # Alter erkennen und entfernen.
    age_match = re.search(r"\b(\d{1,3}\s*(?:Jahre?|Jahr|Wochen?|Woche|Monate?|Monat|Stunden?|Stunde|Tag|Tage))\b\.?,?", work, re.I)
    if age_match:
        cells["age"] = _bk_tab_clean(age_match.group(1))
        work = (work[:age_match.start()] + " " + work[age_match.end():]).strip()

    # Datums-/Jahresfragmente. Römische Monatszahlen und vierstellige Jahre bleiben zusammen.
    # Entfernung über die Original-Treffer-Spannen: Das erneute Suchen des
    # bereinigten Fragments ("21.V.") scheiterte, wenn die Vorlage Leerzeichen
    # enthielt ("21. V.") - dann blieben Reste wie "21 V" im Namen kleben.
    date_parts = []
    spans = []
    for m in re.finditer(r"\b\d{1,2}\s*[.]\s*(?:[IVXLCDMivxlcdm]{1,6}|[0-9]{1,2})\s*[.]?", work):
        part = _bk_tab_clean(m.group(0)).replace(" ", "")
        if part:
            date_parts.append(part)
            spans.append(m.span())
    for m in re.finditer(r"\b(?:1[5-9]\d{2}|20\d{2})\b", work):
        part = m.group(0)
        if part not in date_parts:
            date_parts.append(part)
            spans.append(m.span())
    if date_parts:
        cells["date_year"] = _bk_simple_join_clean(date_parts)
        merged_spans = []
        for start, end in sorted(spans):
            if merged_spans and start <= merged_spans[-1][1]:
                merged_spans[-1] = (merged_spans[-1][0], max(end, merged_spans[-1][1]))
            else:
                merged_spans.append((start, end))
        for start, end in reversed(merged_spans):
            work = work[:start] + " " + work[end:]

    # Übrig bleiben meist Name und Ort. Kommas trennen Namensteil zuverlässiger als Punkte.
    work = re.sub(r"\s+", " ", work).strip(" ,;:.")
    # Entferne typische Wiederholungszeichen/Ditto-Markierungen nicht als Inhalt.
    work = re.sub(r"(?:^|\s)[\"'„“”]+(?:\s|$)", " ", work).strip()

    place = ""
    # Ortsangabe ist häufig das letzte Segment nach Datum/Jahr/Alter, besonders nach Punkt.
    # Wenn mehrere alphabetische Segmente übrig sind, bleibt vorn der Name und hinten der Ort.
    segments = [_bk_simple_strip_punct(seg) for seg in re.split(r"\s{2,}|\s*[.]\s+", work) if _bk_simple_strip_punct(seg)]
    name_part = work
    if len(segments) >= 2:
        # Von hinten den ersten plausiblen Orts-Kandidaten suchen: grossgeschrieben
        # und mindestens 3 Zeichen. Streu-Segmente (einzelne Buchstaben aus
        # OCR-Rauschen wie "v") werden uebersprungen und verworfen.
        idx = len(segments) - 1
        while idx >= 1:
            candidate = segments[idx]
            if (re.match(r"^[A-ZÄÖÜ]", candidate) and len(candidate) >= 3
                    and re.search(r"[A-Za-zÀ-ÿÄÖÜäöüß]", candidate) and len(candidate) <= 32):
                place = candidate
                name_part = " ".join(segments[:idx]).strip()
                break
            if len(candidate) <= 2:
                idx -= 1
                continue
            break
        else:
            name_part = " ".join(segments).strip()
    else:
        # Fallback: letztes alphabetisches Wort nach einem Jahr/Datum-Trenner wurde oben oft übriggelassen.
        m = re.search(r"(.+?)\s+([A-ZÄÖÜ][A-Za-zÀ-ÿÄÖÜäöüß.\-]{2,})$", work)
        if m and cells.get("age") and (cells.get("date_year") or cells.get("number")):
            name_part = m.group(1).strip()
            place = m.group(2).strip()

    cells["place"] = _bk_simple_strip_punct(place)
    name_part = _bk_simple_strip_punct(name_part)
    # Komma-Variante der Vorlage ("1763,Kahlenberg."): Nach der Datums-Entfernung
    # haengt der Ort dann per Komma am Namen. Nur abtrennen, wenn Datums-/Nummern-
    # Kontext existiert und der Kandidat ein einzelnes grossgeschriebenes Wort ist.
    if not cells["place"] and (cells.get("date_year") or cells.get("number")):
        m = re.match(r"^(.*?),\s*([A-ZÄÖÜ][A-Za-zÀ-ÿÄÖÜäöüß.\-]{2,})$", name_part)
        if m and " " not in m.group(2):
            name_part = m.group(1).strip()
            cells["place"] = _bk_simple_strip_punct(m.group(2))
    # Falls der Ort nicht zuverlässig abgetrennt werden konnte, stört das nicht; Originalzeile bleibt erhalten.
    if name_part:
        # erstes Wort = Familienname, Rest = Vornamen. Das ist für Registerseiten im Stil ABBYY lesbarer als wortweise Spalten.
        name_part = name_part.replace(",", " ")
        name_tokens = [tok for tok in name_part.split() if tok]
        if name_tokens:
            cells["family_name"] = _bk_simple_strip_punct(name_tokens[0])
            cells["given_names"] = _bk_simple_strip_punct(" ".join(name_tokens[1:]))
    return cells


def _bk_simple_table_matrix_from_records_semantic(records, window=None):
    records = [rec for rec in (records or []) if _bk_tab_clean(rec.get("text", "")) and not _bk_tab_is_separator(rec.get("text", ""))]
    if not records:
        return []
    rows = _bk_simple_group_rows(records)
    parsed_rows = []
    active_heading = ""
    active_subheading = ""
    heading_used = False
    subheading_used = False
    for row_records in rows:
        raw = _bk_simple_row_raw_text(row_records)
        cells = _bk_simple_parse_register_text(raw)
        if cells.get("heading") and not any(cells.get(k) for k in ("family_name", "given_names", "age", "date_year", "place", "number")):
            active_heading = cells.get("heading", "")
            heading_used = False
            continue
        if cells.get("subheading") and not any(cells.get(k) for k in ("family_name", "given_names", "age", "date_year", "place", "number")):
            active_subheading = cells.get("subheading", "")
            subheading_used = False
            continue
        if not any(cells.get(k) for k in ("family_name", "given_names", "age", "date_year", "place", "number", "extra")):
            continue
        if active_heading and not heading_used:
            cells["heading"] = active_heading
            heading_used = True
        else:
            cells["heading"] = ""
        if active_subheading and not subheading_used:
            cells["subheading"] = active_subheading
            subheading_used = True
        else:
            cells["subheading"] = ""
        parsed_rows.append(cells)
    if not parsed_rows:
        return []
    # Nur Spalten ausgeben, die mindestens einen Wert enthalten. Originalzeile bleibt als Kontrollspalte erhalten.
    keys = [key for key in _BK_SIMPLE_PREVIEW_COLUMN_KEYS if key != "original"]
    keep = [key for key in keys if any(_bk_tab_clean(row.get(key, "")) for row in parsed_rows)]
    if "original" not in keep:
        keep.append("original")
    header = [_bk_simple_column_title(window, key) if window is not None else key for key in keep]
    matrix = [header]
    for row in parsed_rows:
        matrix.append([row.get(key, "") for key in keep])
    return matrix


class _BKSimplePreviewRectItem(QGraphicsRectItem if QGraphicsRectItem is not None else object):
    HANDLE_SIZE = 8.0

    def __init__(self, dialog, row, rect):
        super().__init__(rect)
        self._dialog = dialog
        self._row = row
        self._mode = None
        self._start_scene = None
        self._start_rect = None
        try:
            self.setPen(QPen(QColor(30, 120, 255), 2))
            self.setBrush(QBrush(QColor(30, 120, 255, 35)))
            self.setZValue(30 + row)
            self.setAcceptHoverEvents(True)
            self.setAcceptedMouseButtons(Qt.LeftButton)
        except Exception:
            pass

    def _handle_at(self, pos):
        rect = self.rect()
        s = self.HANDLE_SIZE + 2.0
        x, y = float(pos.x()), float(pos.y())
        checks = (
            ("tl", rect.left(), rect.top()),
            ("tr", rect.right(), rect.top()),
            ("bl", rect.left(), rect.bottom()),
            ("br", rect.right(), rect.bottom()),
        )
        for mode, hx, hy in checks:
            if abs(x - hx) <= s and abs(y - hy) <= s:
                return mode
        return "move"

    def mousePressEvent(self, event):
        self._dialog.select_box(self._row)
        self._mode = self._handle_at(event.pos())
        self._start_scene = event.scenePos()
        self._start_rect = QRectF(self.rect())
        event.accept()

    def mouseMoveEvent(self, event):
        if self._mode is None or self._start_scene is None or self._start_rect is None:
            event.accept(); return
        dx = float(event.scenePos().x() - self._start_scene.x())
        dy = float(event.scenePos().y() - self._start_scene.y())
        rect = QRectF(self._start_rect)
        if self._mode == "move":
            rect.translate(dx, dy)
        else:
            if "l" in self._mode: rect.setLeft(rect.left() + dx)
            if "r" in self._mode: rect.setRight(rect.right() + dx)
            if "t" in self._mode: rect.setTop(rect.top() + dy)
            if "b" in self._mode: rect.setBottom(rect.bottom() + dy)
        rect = self._dialog.clamp_rect(rect)
        self.setRect(rect)
        self._dialog.set_box_rect(self._row, rect, redraw=False)
        event.accept()

    def mouseReleaseEvent(self, event):
        self._mode = None
        self._start_scene = None
        self._start_rect = None
        self._dialog.redraw_boxes()
        event.accept()

    def paint(self, painter, option, widget=None):
        super().paint(painter, option, widget)
        try:
            rect = self.rect(); s = self.HANDLE_SIZE
            painter.save()
            painter.setPen(QPen(QColor(0, 0, 0), 1))
            painter.setBrush(QBrush(QColor(255, 255, 255)))
            for x, y in ((rect.left(), rect.top()), (rect.right(), rect.top()), (rect.left(), rect.bottom()), (rect.right(), rect.bottom())):
                painter.drawRect(QRectF(x - s / 2.0, y - s / 2.0, s, s))
            painter.restore()
        except Exception:
            pass


def _bk_simple_preview_pixmap(item):
    """Return a full-size preview pixmap. Prefer the already loaded PIL page.
    QPixmap(path) may load only a tiny embedded thumbnail for some source images.
    """
    try:
        if item is not None and getattr(item, "results", None):
            _text, _kr, pil_image, _record_views = item.results
            if pil_image is not None and _BK_ImageQt is not None and QPixmap is not None:
                try:
                    image = pil_image.convert("RGB")
                    qimg = _BK_ImageQt(image)
                    pix = QPixmap.fromImage(qimg)
                    if pix is not None and not pix.isNull() and pix.width() > 50 and pix.height() > 50:
                        return pix
                except Exception:
                    pass
    except Exception:
        pass
    try:
        path = str(getattr(item, "path", "") or "")
        pix = QPixmap(path) if QPixmap is not None and path else None
        if pix is not None and not pix.isNull():
            return pix
    except Exception:
        pass
    try:
        path = str(getattr(item, "path", "") or "")
        if path and _BK_ImageQt is not None and QPixmap is not None:
            image = _load_image_color(path).convert("RGB")
            qimg = _BK_ImageQt(image)
            pix = QPixmap.fromImage(qimg)
            if pix is not None and not pix.isNull():
                return pix
    except Exception:
        pass
    return None


class _BKSimplePreviewDialog(QDialog):
    def __init__(self, window, item, records):
        super().__init__(window)
        self._window = window
        self._item = item
        self._records = [dict(rec) for rec in records or []]
        self._items = []
        self._image_w = 0
        self._image_h = 0
        self.setWindowTitle(_bk_tab_tr(window, "export_simple_preview_title"))
        self.setModal(True)
        self.resize(1240, 820)
        root = QHBoxLayout(self)
        root.setContentsMargins(12, 12, 12, 12)
        root.setSpacing(10)
        self.scene = QGraphicsScene(self)
        self.view = QGraphicsView(self)
        self.view.setScene(self.scene)
        try:
            self.view.setDragMode(QGraphicsView.NoDrag)
        except Exception:
            pass
        self._pixmap_item = None
        pixmap = _bk_simple_preview_pixmap(item)
        if pixmap is not None and not pixmap.isNull():
            self._image_w = int(pixmap.width()); self._image_h = int(pixmap.height())
            self._pixmap_item = self.scene.addPixmap(pixmap)
            self.scene.setSceneRect(QRectF(0, 0, pixmap.width(), pixmap.height()))
        root.addWidget(self.view, 3)
        side = QVBoxLayout(); side.setSpacing(8)
        intro = QLabel(_bk_tab_tr(window, "export_simple_preview_intro"), self)
        intro.setWordWrap(True)
        side.addWidget(intro)
        self.table = QTableWidget(0, 5, self)
        self.table.setHorizontalHeaderLabels([
            _bk_tab_tr(window, "export_simple_preview_col_text"), "X", "Y",
            _bk_tab_tr(window, "export_simple_preview_col_width"),
            _bk_tab_tr(window, "export_simple_preview_col_height"),
        ])
        try:
            self.table.setSelectionBehavior(QAbstractItemView.SelectRows)
            self.table.setSelectionMode(QAbstractItemView.SingleSelection)
            self.table.setColumnWidth(0, 320)
            for col in range(1, 5):
                self.table.setColumnWidth(col, 70)
            if QHeaderView is not None:
                self.table.horizontalHeader().setStretchLastSection(False)
        except Exception:
            pass
        side.addWidget(self.table, 1)
        btn_row = QHBoxLayout()
        self.btn_delete = QPushButton(_bk_tab_tr(window, "export_simple_preview_delete"), self)
        self.btn_reset = QPushButton(_bk_tab_tr(window, "export_simple_preview_reset"), self)
        btn_row.addWidget(self.btn_delete); btn_row.addWidget(self.btn_reset)
        side.addLayout(btn_row)
        buttons = QDialogButtonBox(QDialogButtonBox.Ok | QDialogButtonBox.Cancel, self)
        try:
            buttons.button(QDialogButtonBox.Ok).setText(_bk_tab_tr(window, "btn_ok"))
            buttons.button(QDialogButtonBox.Cancel).setText(_bk_tab_tr(window, "btn_cancel"))
        except Exception:
            pass
        side.addWidget(buttons)
        root.addLayout(side, 2)
        self.btn_delete.clicked.connect(self.delete_selected)
        self.btn_reset.clicked.connect(self.reset_boxes)
        self.table.itemChanged.connect(self._table_item_changed)
        try:
            self.table.itemSelectionChanged.connect(self._selection_changed)
        except Exception:
            pass
        buttons.accepted.connect(self.accept)
        buttons.rejected.connect(self.reject)
        self.refresh_table()
        self._fit_preview_pending = True
        self._schedule_fit_preview()

    def _schedule_fit_preview(self):
        if self._image_w <= 0 or self._image_h <= 0:
            return
        if QTimer is not None:
            try:
                QTimer.singleShot(0, self._fit_preview)
                QTimer.singleShot(80, self._fit_preview)
                return
            except Exception:
                pass
        self._fit_preview()

    def _fit_preview(self):
        if self._image_w <= 0 or self._image_h <= 0:
            return
        try:
            self.view.resetTransform()
            self.view.fitInView(QRectF(0, 0, float(self._image_w), float(self._image_h)), Qt.KeepAspectRatio)
            self._fit_preview_pending = False
        except Exception:
            pass

    def showEvent(self, event):
        try:
            super().showEvent(event)
        except Exception:
            pass
        self._schedule_fit_preview()

    def resizeEvent(self, event):
        try:
            super().resizeEvent(event)
        except Exception:
            pass
        if getattr(self, "_fit_preview_pending", True):
            self._schedule_fit_preview()

    def clamp_rect(self, rect):
        rect = QRectF(rect).normalized()
        min_size = 3.0
        if rect.width() < min_size: rect.setWidth(min_size)
        if rect.height() < min_size: rect.setHeight(min_size)
        if self._image_w > 0:
            if rect.left() < 0: rect.translate(-rect.left(), 0)
            if rect.right() > self._image_w: rect.translate(self._image_w - rect.right(), 0)
            rect.setLeft(max(0.0, rect.left()))
            rect.setRight(min(float(self._image_w), max(rect.left() + min_size, rect.right())))
        if self._image_h > 0:
            if rect.top() < 0: rect.translate(0, -rect.top())
            if rect.bottom() > self._image_h: rect.translate(0, self._image_h - rect.bottom())
            rect.setTop(max(0.0, rect.top()))
            rect.setBottom(min(float(self._image_h), max(rect.top() + min_size, rect.bottom())))
        return rect.normalized()

    def _format_rect_value(self, row, col):
        rec = self._records[row]
        if col == 1: return int(round(float(rec.get("x0", 0))))
        if col == 2: return int(round(float(rec.get("y0", 0))))
        if col == 3: return int(round(float(rec.get("x1", 0)) - float(rec.get("x0", 0))))
        if col == 4: return int(round(float(rec.get("y1", 0)) - float(rec.get("y0", 0))))
        return ""

    def refresh_table(self, select_row=-1):
        try: self.table.blockSignals(True)
        except Exception: pass
        self.table.setRowCount(0)
        for row, rec in enumerate(self._records):
            self.table.insertRow(row)
            self.table.setItem(row, 0, QTableWidgetItem(str(rec.get("text", ""))))
            for col in range(1, 5):
                self.table.setItem(row, col, QTableWidgetItem(str(self._format_rect_value(row, col))))
        try: self.table.blockSignals(False)
        except Exception: pass
        if 0 <= select_row < self.table.rowCount():
            self.table.selectRow(select_row)
        self.redraw_boxes()

    def redraw_boxes(self):
        for item in list(self._items):
            try: self.scene.removeItem(item)
            except Exception: pass
        self._items = []
        selected = -1
        try:
            indexes = self.table.selectionModel().selectedRows()
            if indexes: selected = indexes[0].row()
        except Exception:
            pass
        for row, rec in enumerate(self._records):
            try:
                rect = QRectF(float(rec.get("x0", 0)), float(rec.get("y0", 0)), max(3.0, float(rec.get("x1", 0)) - float(rec.get("x0", 0))), max(3.0, float(rec.get("y1", 0)) - float(rec.get("y0", 0))))
                item = _BKSimplePreviewRectItem(self, row, rect)
                if row == selected:
                    item.setPen(QPen(QColor(255, 120, 0), 3))
                    item.setBrush(QBrush(QColor(255, 160, 0, 45)))
                self.scene.addItem(item)
                self._items.append(item)
            except Exception:
                pass

    def select_box(self, row):
        try: self.table.selectRow(row)
        except Exception: pass
        self.redraw_boxes()

    def set_box_rect(self, row, rect, redraw=True):
        if not (0 <= row < len(self._records)): return
        rect = self.clamp_rect(rect)
        self._records[row].update({"x0": float(rect.left()), "y0": float(rect.top()), "x1": float(rect.right()), "y1": float(rect.bottom()), "cx": (rect.left()+rect.right())/2.0, "cy": (rect.top()+rect.bottom())/2.0, "w": max(1.0, rect.width()), "h": max(1.0, rect.height())})
        try:
            self.table.blockSignals(True)
            for col in range(1, 5):
                item = self.table.item(row, col)
                if item is not None: item.setText(str(self._format_rect_value(row, col)))
        finally:
            try: self.table.blockSignals(False)
            except Exception: pass
        if redraw: self.redraw_boxes()

    def _table_item_changed(self, item):
        if item is None: return
        row = item.row(); col = item.column()
        if not (0 <= row < len(self._records)): return
        if col == 0:
            self._records[row]["text"] = str(item.text() or "")
            return
        try:
            nums = [float(self.table.item(row, c).text().replace(",", ".")) for c in range(1, 5)]
            x, y, w, h = nums
            self.set_box_rect(row, QRectF(x, y, max(3.0, w), max(3.0, h)), redraw=True)
        except Exception:
            self.refresh_table(select_row=row)

    def _selection_changed(self):
        self.redraw_boxes()

    def delete_selected(self):
        rows = []
        try: rows = sorted([idx.row() for idx in self.table.selectionModel().selectedRows()], reverse=True)
        except Exception: rows = []
        for row in rows:
            if 0 <= row < len(self._records):
                del self._records[row]
        self.refresh_table(select_row=min(rows[-1] if rows else 0, len(self._records)-1))

    def reset_boxes(self):
        try:
            _text, _kr, _pil, record_views = self._item.results
            self._records = _bk_simple_table_records(record_views)
        except Exception:
            pass
        self.refresh_table()

    def records(self):
        return [dict(rec) for rec in self._records if _bk_tab_clean(rec.get("text", ""))]


def _bk_open_simple_table_preview(window, item):
    if QGraphicsScene is None or QTableWidget is None or not item or not getattr(item, "results", None):
        return None
    try:
        _text, _kr, _pil, record_views = item.results
        records = _bk_simple_table_records(record_views)
    except Exception:
        return None
    dlg = _BKSimplePreviewDialog(window, item, records)
    try:
        result = dlg.exec()
    except Exception:
        return records
    if result != QDialog.Accepted:
        return None
    return dlg.records()


def _bk_simple_table_matrix_for_item_semantic(window, item):
    try:
        preview_path = str(getattr(window, "_bk_simple_preview_item_path", "") or "")
        item_path = str(getattr(item, "path", "") or "")
        if preview_path and item_path and preview_path == item_path:
            records = getattr(window, "_bk_simple_preview_records", None)
            if records is not None:
                return _bk_simple_table_matrix_from_records_semantic(records, window)
    except Exception:
        pass
    try:
        _text, _kr, _pil, record_views = item.results
    except Exception:
        return []
    return _bk_simple_table_matrix_from_records_semantic(_bk_simple_table_records(record_views), window)


# Vorherige einfache Exportfunktion überschreiben, ohne die erweiterte Tabellenlogik zu ändern.
def _bk_simple_table_matrix_for_item(item):
    try:
        window = globals().get("_bk_simple_current_window")
        if window is not None:
            return _bk_simple_table_matrix_for_item_semantic(window, item)
    except Exception:
        pass
    try:
        _text, _kr, _pil, record_views = item.results
    except Exception:
        return []
    return _bk_simple_table_matrix_from_records_semantic(_bk_simple_table_records(record_views), None)






def _bk_column_choice_dialog_simple_table_compact(self, fmt=None, include_text_modes=False):
    fmt_l = str(fmt or "").lower().lstrip(".")
    supports_simple = fmt_l in _BK_SIMPLE_TABLE_FMTS
    if not supports_simple:
        if callable(_BK_SIMPLE_PREV_COLUMN_CHOICE_DIALOG):
            return _BK_SIMPLE_PREV_COLUMN_CHOICE_DIALOG(self, fmt, include_text_modes)
        return None
    dlg = QDialog(self)
    dlg.setWindowTitle(_bk_tab_tr(self, "export_text_layout_title" if include_text_modes else "export_table_columns_title"))
    dlg.setModal(True)
    layout = QVBoxLayout(dlg)
    layout.setContentsMargins(14, 12, 14, 12)
    layout.setSpacing(8)
    intro_key = "export_text_layout_intro_extended" if include_text_modes else "export_table_columns_intro"
    label = QLabel(_bk_tab_tr(self, intro_key), dlg)
    label.setWordWrap(True)
    layout.addWidget(label)
    rb_original = rb_lines = None
    rb_simple = QRadioButton(_bk_tab_tr(self, "export_text_layout_table_simple"), dlg)
    rb_table = QRadioButton(_bk_tab_tr(self, "export_text_layout_table"), dlg)
    mode_box = QGroupBox(_bk_tab_tr(self, "export_layout_mode_group"), dlg) if QGroupBox is not None else None
    mode_layout = QVBoxLayout(mode_box) if mode_box is not None else QVBoxLayout()
    mode_layout.setContentsMargins(10, 8, 10, 8)
    if include_text_modes:
        rb_original = QRadioButton(_bk_tab_tr(self, "export_text_layout_original"), dlg)
        rb_lines = QRadioButton(_bk_tab_tr(self, "export_text_layout_lines"), dlg)
        mode = str(getattr(self, "_bk_export_text_layout_mode", "original") or "original").lower()
        rb_original.setChecked(mode not in {"lines", "table", "table_simple"})
        rb_lines.setChecked(mode == "lines")
        rb_simple.setChecked(mode == "table_simple")
        rb_table.setChecked(mode == "table")
        mode_layout.addWidget(rb_original); mode_layout.addWidget(rb_lines)
    else:
        # Für reine Spreadsheet-Exporte ist die einfache ABBYY-ähnliche Tabelle
        # der Ausgangszustand. Erst wenn der Benutzer in der laufenden Sitzung
        # bewusst auf erweitert umstellt, bleibt erweitert aktiv.
        current_mode = getattr(self, "_bk_export_table_mode", None)
        if current_mode is None:
            table_mode = "simple" if fmt_l in _BK_SIMPLE_TABLE_SPREADSHEET_FMTS else "table"
        else:
            table_mode = str(current_mode or "table").lower()
        rb_simple.setChecked(table_mode == "simple")
        rb_table.setChecked(table_mode != "simple")
    mode_layout.addWidget(rb_simple); mode_layout.addWidget(rb_table)
    if mode_box is not None: layout.addWidget(mode_box)
    else: layout.addLayout(mode_layout)

    _bk_add_export_orientation_group(self, dlg, layout)

    selected_keys = _bk_load_saved_column_keys_for_dialog(self)
    checkboxes = {}
    columns_box = QGroupBox(_bk_tab_tr(self, "export_table_columns_label"), dlg) if QGroupBox is not None else None
    columns_layout = QVBoxLayout(columns_box) if columns_box is not None else QVBoxLayout()
    if QGridLayout is not None:
        grid = QGridLayout(); grid.setHorizontalSpacing(18); grid.setVerticalSpacing(4)
        for idx, key in enumerate(_BK_TABULAR_KEYS):
            cb = QCheckBox(_bk_tabular_column_title(self, key), dlg); cb.setChecked(key in selected_keys); checkboxes[key]=cb
            grid.addWidget(cb, idx % 8, idx // 8)
        columns_layout.addLayout(grid)
    else:
        for key in _BK_TABULAR_KEYS:
            cb = QCheckBox(_bk_tabular_column_title(self, key), dlg); cb.setChecked(key in selected_keys); checkboxes[key]=cb; columns_layout.addWidget(cb)
    quick_row = QHBoxLayout()
    btn_all = QPushButton(_bk_tab_tr(self, "export_table_columns_all"), dlg)
    btn_none = QPushButton(_bk_tab_tr(self, "export_table_columns_none_button"), dlg)
    btn_remember = QPushButton(_bk_tab_tr(self, "export_table_columns_remember"), dlg)
    for btn in (btn_all, btn_none, btn_remember): quick_row.addWidget(btn)
    quick_row.addStretch(1); columns_layout.addLayout(quick_row)
    if columns_box is not None: layout.addWidget(columns_box)
    else: layout.addLayout(columns_layout)

    zone_box = QGroupBox(_bk_tab_tr(self, "export_zones_group"), dlg) if QGroupBox is not None else None
    zone_layout = QHBoxLayout(zone_box) if zone_box is not None else QHBoxLayout()
    cb_zones = QCheckBox(_bk_tab_tr(self, "export_table_use_zones"), dlg)
    try: remembered_zones = self.settings.value("export/table_use_zones", bool(getattr(self, "_bk_export_use_zones", False)), type=bool)
    except Exception: remembered_zones = bool(getattr(self, "_bk_export_use_zones", False))
    cb_zones.setChecked(bool(remembered_zones))
    btn_zones = QPushButton(_bk_tab_tr(self, "export_table_define_zones"), dlg)
    zone_layout.addWidget(cb_zones, 1); zone_layout.addWidget(btn_zones)
    if zone_box is not None: layout.addWidget(zone_box)
    else: layout.addLayout(zone_layout)
    hint = QLabel(_bk_tab_tr(self, "export_table_simple_hint"), dlg)
    hint.setWordWrap(True); layout.addWidget(hint)

    def current_checked_keys(): return [key for key, cb in checkboxes.items() if cb.isChecked()]
    def set_all():
        for cb in checkboxes.values(): cb.setChecked(True)
    def set_none():
        for cb in checkboxes.values(): cb.setChecked(False)
    result = {"mode": "table", "columns": selected_keys, "remembered": False, "use_zones": cb_zones.isChecked(), "cancelled": False}
    def remember_selection():
        keys = current_checked_keys()
        if not keys and not rb_simple.isChecked():
            QMessageBox.warning(dlg, _bk_tab_tr(self, "warn_title"), _bk_tab_tr(self, "export_table_columns_none")); return
        result["remembered"] = True
        if keys: result["columns"] = _bk_save_column_keys(self, keys)
        try: self.settings.setValue("export/table_use_zones", bool(cb_zones.isChecked())); self.settings.sync()
        except Exception: pass
    def define_zones():
        try: task = self._current_task()
        except Exception: task = None
        zones = _bk_open_export_zones_dialog(self, task)
        if zones is not None: cb_zones.setChecked(bool(zones))
    btn_all.clicked.connect(set_all); btn_none.clicked.connect(set_none); btn_remember.clicked.connect(remember_selection); btn_zones.clicked.connect(define_zones)
    def sync_enabled():
        detailed = rb_table.isChecked()
        if columns_box is not None: columns_box.setVisible(detailed); columns_box.setEnabled(detailed)
        if zone_box is not None: zone_box.setVisible(detailed); zone_box.setEnabled(detailed)
        cb_zones.setEnabled(detailed); btn_zones.setEnabled(detailed); hint.setVisible(rb_simple.isChecked())
        try:
            if detailed:
                dlg.resize(720, 680 if include_text_modes else 610)
            else:
                dlg.resize(640, 360 if include_text_modes else 320)
            dlg.adjustSize()
        except Exception: pass
    for rb in (rb_original, rb_lines, rb_layout_tables, rb_simple, rb_table):
        if rb is not None:
            try: rb.toggled.connect(sync_enabled)
            except Exception: pass
    buttons = QDialogButtonBox(QDialogButtonBox.Ok | QDialogButtonBox.Cancel, dlg)
    try:
        buttons.button(QDialogButtonBox.Ok).setText(_bk_tab_tr(self, "btn_ok")); buttons.button(QDialogButtonBox.Cancel).setText(_bk_tab_tr(self, "btn_cancel"))
    except Exception: pass
    def cancel_dialog(): result["cancelled"] = True; dlg.done(QDialog.Rejected)
    def accept_checked():
        result["cancelled"] = False
        if rb_simple.isChecked(): mode="table_simple"; keys=[]
        elif include_text_modes and rb_lines is not None and rb_lines.isChecked(): mode="lines"; keys=[]
        elif include_text_modes and rb_original is not None and rb_original.isChecked(): mode="original"; keys=[]
        else:
            mode="table"; keys=current_checked_keys()
            if not keys:
                QMessageBox.warning(dlg, _bk_tab_tr(self, "warn_title"), _bk_tab_tr(self, "export_table_columns_none")); return
        result["mode"] = mode; result["columns"] = _bk_normalize_column_keys(keys); result["use_zones"] = bool(cb_zones.isChecked()) if mode == "table" else False
        try: self._bk_export_use_zones = result["use_zones"]
        except Exception: pass
        dlg.accept()
    buttons.accepted.connect(accept_checked); buttons.rejected.connect(cancel_dialog); layout.addWidget(buttons)
    sync_enabled()
    try: exec_result = dlg.exec()
    except Exception: result["cancelled"] = True; return None
    if exec_result != QDialog.Accepted: result["cancelled"] = True; return None
    return result




try:
    _bk_column_choice_dialog = _bk_column_choice_dialog_simple_table_compact
except Exception:
    pass
try:
    __all__.extend([
        '_bk_simple_table_matrix_from_records_semantic',
        '_BKSimplePreviewDialog',
        '_bk_open_simple_table_preview',
        '_bk_column_choice_dialog_simple_table_compact',
    ])
    register_globals('bk', globals(), sorted(set(__all__)))
except Exception:
    pass

# ---------------------------------------------------------------------------
# Simple Tabellenexport: keine Vorschau, räumliche Overlay-Box-Tabelle
# ---------------------------------------------------------------------------
def _bk_semantic_blocks_from_item(self, item):
    """Deterministische Variante des KI-Exports: erkannte Zeilen in derselben
    Blockstruktur (heading / paragraph / table mit header) aufbereiten - ohne
    KI, in garantiert originaler Reihenfolge. Registerzeilen werden ueber den
    vorhandenen semantischen Parser in Name | Zusatz | Alter | Datum/Jahr |
    Ort | Nr. zerlegt. Liefert None, wenn die Seite nicht registerartig genug
    ist (dann greift das raeumliche Layout als Rueckfall)."""
    try:
        record_views = item.results[3] if item is not None and getattr(item, "results", None) else []
    except Exception:
        record_views = []
    records = _bk_simple_table_records(record_views)
    if not records:
        return None
    row_groups = _bk_simple_group_rows(records)
    if not row_groups:
        return None
    header = [
        _bk_simple_column_title(self, "family_name") or "Name",
        _bk_simple_column_title(self, "extra") or "Zusatz",
        _bk_simple_column_title(self, "age") or "Alter",
        _bk_simple_column_title(self, "date_year") or "Datum/Jahr",
        _bk_simple_column_title(self, "place") or "Ort",
        _bk_simple_column_title(self, "number") or "Nr.",
    ]
    blocks = []
    open_rows = []
    data_rows = 0
    register_rows = 0

    def _close_table():
        nonlocal open_rows
        if open_rows:
            blocks.append({"type": "table", "header": list(header), "rows": open_rows})
            open_rows = []

    for row in row_groups:
        raw = _bk_simple_row_raw_text(row)
        if not _bk_tab_clean(raw):
            continue
        cells = _bk_simple_parse_register_text(raw)
        if cells.get("heading"):
            _close_table()
            blocks.append({"type": "heading", "text": cells["heading"]})
            continue
        if cells.get("subheading"):
            _close_table()
            blocks.append({"type": "paragraph", "text": cells["subheading"]})
            continue
        data_rows += 1
        name = _bk_tab_clean(" ".join(p for p in (cells.get("family_name", ""), cells.get("given_names", "")) if p))
        filled = [v for v in (name, cells.get("extra"), cells.get("age"),
                              cells.get("date_year"), cells.get("place"), cells.get("number")) if _bk_tab_clean(v or "")]
        if name and len(filled) >= 2:
            register_rows += 1
            open_rows.append([
                name,
                _bk_tab_clean(cells.get("extra", "")),
                _bk_tab_clean(cells.get("age", "")),
                _bk_tab_clean(cells.get("date_year", "")),
                _bk_tab_clean(cells.get("place", "")),
                _bk_tab_clean(cells.get("number", "")),
            ])
        else:
            _close_table()
            blocks.append({"type": "paragraph", "text": _bk_tab_clean(raw)})
    _close_table()
    if data_rows == 0 or register_rows / max(1, data_rows) < 0.6:
        return None
    return blocks


def _bk_semantic_blocks_write(self, path, fmt, item):
    """Versucht den KI-Stil deterministisch; True bei Erfolg, False -> Rueckfall."""
    blocks = _bk_semantic_blocks_from_item(self, item)
    if not blocks:
        return False
    writer = globals().get("_bk_lmx_write")
    if not callable(writer):
        return False
    try:
        writer(self, path, str(fmt).lower(), blocks)
        return True
    except Exception:
        return False


def _bk_simple_spatial_layout_from_item(item):
    try:
        _text, _kr, _pil, record_views = item.results
    except Exception:
        return {"matrix": [], "widths": [12.0], "heights": [18.0]}
    records = _bk_simple_table_records(record_views)
    return _bk_simple_spatial_layout_from_records(records)














try:
    register_globals('bk', globals(), sorted(set(__all__)))
except Exception:
    pass

# ---------------------------------------------------------------------------
# Tabelle (einfach): ABBYY-ähnliches festes Spaltengitter für XLSX/ODS
# ---------------------------------------------------------------------------


# ---------------------------------------------------------------------------
# Final export UX + simple-table grid + custom extended columns/datentypen
# ---------------------------------------------------------------------------
try:
    from PySide6.QtWidgets import QLineEdit as _BK_QLineEdit
except Exception:
    _BK_QLineEdit = None

# CSV kann ebenfalls eine einfache Overlay-Tabelle bekommen; TXT bleibt ohne Tabellenmodus.
try:
    _BK_SIMPLE_TABLE_SPREADSHEET_FMTS = set(_BK_SIMPLE_TABLE_SPREADSHEET_FMTS) | {"csv"}
    _BK_SIMPLE_TABLE_FMTS = set(_BK_SIMPLE_TABLE_FMTS) | {"csv"}
except Exception:
    pass


def _bk_custom_column_slug(title: str) -> str:
    txt = _bk_tab_clean(title).lower()
    txt = txt.replace("ä", "ae").replace("ö", "oe").replace("ü", "ue").replace("ß", "ss")
    txt = re.sub(r"[^a-z0-9]+", "_", txt).strip("_")
    return txt or "spalte"


def _bk_load_custom_columns(window=None):
    raw = ""
    try:
        settings = getattr(window, "settings", None)
        if settings is not None:
            raw = settings.value("export/custom_tabular_columns", "", str)
    except Exception:
        raw = ""
    items = []
    if raw:
        try:
            parsed = json.loads(str(raw))
            if isinstance(parsed, list):
                for entry in parsed:
                    if not isinstance(entry, dict):
                        continue
                    key = str(entry.get("key", "") or "").strip()
                    title = _bk_tab_clean(entry.get("title", ""))
                    if key.startswith("custom_") and title:
                        items.append({"key": key, "title": title})
        except Exception:
            items = []
    return items


def _bk_save_custom_columns(window, items):
    clean = []
    seen = set()
    for entry in items or []:
        key = str(entry.get("key", "") or "").strip()
        title = _bk_tab_clean(entry.get("title", ""))
        if not key.startswith("custom_") or not title or key in seen:
            continue
        clean.append({"key": key, "title": title})
        seen.add(key)
    try:
        settings = getattr(window, "settings", None)
        if settings is not None:
            settings.setValue("export/custom_tabular_columns", json.dumps(clean, ensure_ascii=False))
            settings.sync()
    except Exception:
        pass
    return clean


def _bk_ensure_custom_columns(window=None):
    global _BK_TABULAR_COLUMNS, _BK_TABULAR_KEYS, _BK_TABULAR_HEADERS, _BK_TABULAR_COLUMN_BY_KEY, _BK_ZONE_TYPES
    base = [col for col in _BK_TABULAR_COLUMNS if not str(col[0]).startswith("custom_")]
    custom = []
    for entry in _bk_load_custom_columns(window):
        key = str(entry.get("key", ""))
        title = _bk_tab_clean(entry.get("title", ""))
        if key and title:
            custom.append((key, key, title, 18.0, 3.0))
    _BK_TABULAR_COLUMNS = base + custom
    _BK_TABULAR_KEYS = [col[0] for col in _BK_TABULAR_COLUMNS]
    _BK_TABULAR_HEADERS = [col[2] for col in _BK_TABULAR_COLUMNS]
    _BK_TABULAR_COLUMN_BY_KEY = {col[0]: col for col in _BK_TABULAR_COLUMNS}
    _BK_ZONE_TYPES = tuple(_BK_TABULAR_KEYS + ["ignore"])
    return list(_BK_TABULAR_KEYS)


def _bk_tabular_column_title(window, key):
    _bk_ensure_custom_columns(window)
    column = _BK_TABULAR_COLUMN_BY_KEY.get(str(key))
    if not column:
        return str(key or "")
    if str(column[1]).startswith("custom_"):
        return str(column[2])
    return _bk_tab_tr(window, column[1], column[2])


def _bk_normalize_column_keys(column_keys=None):
    _bk_ensure_custom_columns(None)
    keys = []
    source = column_keys if column_keys is not None else []
    for key in source:
        key = str(key or "")
        if key in _BK_TABULAR_COLUMN_BY_KEY and key not in keys:
            keys.append(key)
    return keys


def _bk_load_saved_column_keys_for_dialog(window):
    # Ausgangszustand: keine Spalten vorausgewählt. Gespeichert wird erst ab dieser Version.
    _bk_ensure_custom_columns(window)
    raw = None
    user_saved = False
    try:
        settings = getattr(window, "settings", None)
        if settings is not None:
            user_saved = bool(settings.value("export/table_column_keys_user_saved_v3", False, type=bool))
            if user_saved:
                raw = settings.value("export/table_column_keys_v3", "", str)
    except Exception:
        raw = None
        user_saved = False
    if raw:
        try:
            parsed = json.loads(str(raw))
            if isinstance(parsed, list):
                return _bk_normalize_column_keys([str(x) for x in parsed])
        except Exception:
            return _bk_normalize_column_keys([p.strip() for p in str(raw).split(",") if p.strip()])
    return []


def _bk_save_column_keys(window, keys):
    _bk_ensure_custom_columns(window)
    normalized = _bk_normalize_column_keys(keys)
    try:
        setattr(window, "_bk_export_selected_column_keys", list(normalized))
        settings = getattr(window, "settings", None)
        if settings is not None:
            settings.setValue("export/table_column_keys_v3", json.dumps(list(normalized), ensure_ascii=False))
            settings.setValue("export/table_column_keys_user_saved_v3", True)
            settings.sync()
    except Exception:
        pass
    return normalized


def _bk_export_zone_title(window, zone_type: str) -> str:
    _bk_ensure_custom_columns(window)
    zone_type = _BK_ZONE_LEGACY_TYPES.get(str(zone_type or "full_name"), str(zone_type or "full_name"))
    if zone_type == "ignore":
        return _bk_tab_tr(window, "export_zone_type_ignore")
    if zone_type in _BK_TABULAR_COLUMN_BY_KEY:
        return _bk_tabular_column_title(window, zone_type)
    return str(zone_type or "")


def _bk_export_clean_zone(zone, order: int = 0):
    _bk_ensure_custom_columns(None)
    if not isinstance(zone, dict):
        return None
    try:
        x0 = float(zone.get("x0", 0.0)); y0 = float(zone.get("y0", 0.0)); x1 = float(zone.get("x1", 0.0)); y1 = float(zone.get("y1", 0.0))
    except Exception:
        return None
    if x1 < x0: x0, x1 = x1, x0
    if y1 < y0: y0, y1 = y1, y0
    if x1 - x0 < 3 or y1 - y0 < 3:
        return None
    zone_type = str(zone.get("type", "full_name") or "full_name")
    zone_type = _BK_ZONE_LEGACY_TYPES.get(zone_type, zone_type)
    if zone_type not in _BK_ZONE_TYPES:
        zone_type = "full_name" if "full_name" in _BK_ZONE_TYPES else (_BK_TABULAR_KEYS[0] if _BK_TABULAR_KEYS else "original_line")
    name = str(zone.get("name", "") or "").strip() or f"Bereich {int(order) + 1}"
    try: order_value = int(zone.get("order", order) or order)
    except Exception: order_value = int(order or 0)
    return {"x0": x0, "y0": y0, "x1": x1, "y1": y1, "type": zone_type, "name": name, "order": order_value}


def _bk_zone_combo_final(self, row, zone_type):
    _bk_ensure_custom_columns(self._window)
    combo = QComboBox(self.table)
    for value in _BK_ZONE_TYPES:
        combo.addItem(_bk_export_zone_title(self._window, value), value)
    fallback = zone_type if zone_type in _BK_ZONE_TYPES else ("original_line" if "original_line" in _BK_ZONE_TYPES else (_BK_TABULAR_KEYS[0] if _BK_TABULAR_KEYS else ""))
    idx = combo.findData(fallback)
    combo.setCurrentIndex(max(0, idx))
    combo.currentIndexChanged.connect(lambda _=0, r=row, c=combo: self._combo_changed(r, c))
    return combo

try:
    _BKExportZonesDialog._zone_type_combo = _bk_zone_combo_final
except Exception:
    pass


def _bk_export_zones_ai_max_tokens(window, default=300):
    try:
        return max(64, min(300, int(window._lm_token_limit("export_zones_ai"))))
    except Exception:
        try:
            limits = getattr(window, "lm_token_limits", {}) or {}
            return max(64, min(300, int(limits.get("export_zones_ai", default))))
        except Exception:
            return 300





def _bk_column_choice_dialog_final(self, fmt=None, include_text_modes=False):
    _bk_ensure_custom_columns(self)
    fmt_l = str(fmt or "").lower().lstrip(".")
    txt_only = fmt_l in {"txt", "text", "txt_plain"}
    supports_simple = fmt_l in _BK_SIMPLE_TABLE_FMTS and not txt_only
    supports_extended = not txt_only
    if not supports_simple and not supports_extended and not include_text_modes:
        if callable(_BK_SIMPLE_PREV_COLUMN_CHOICE_DIALOG):
            return _BK_SIMPLE_PREV_COLUMN_CHOICE_DIALOG(self, fmt, include_text_modes)
        return None
    dlg = QDialog(self)
    dlg.setWindowTitle(_bk_tab_tr(self, "export_text_layout_title" if include_text_modes else "export_table_columns_title"))
    dlg.setModal(True)
    layout = QVBoxLayout(dlg)
    layout.setContentsMargins(14, 12, 14, 12)
    layout.setSpacing(8)
    intro_key = "export_text_layout_intro_extended" if include_text_modes else "export_table_columns_intro"
    label = QLabel(_bk_tab_tr(self, intro_key), dlg)
    label.setWordWrap(True)
    layout.addWidget(label)
    rb_original = rb_lines = rb_layout_tables = rb_simple = rb_table = None
    mode_box = QGroupBox(_bk_tab_tr(self, "export_layout_mode_group"), dlg) if QGroupBox is not None else None
    mode_layout = QVBoxLayout(mode_box) if mode_box is not None else QVBoxLayout()
    mode_layout.setContentsMargins(10, 8, 10, 8)
    if include_text_modes:
        rb_original = QRadioButton(_bk_tab_tr(self, "export_text_layout_original"), dlg)
        rb_lines = QRadioButton(_bk_tab_tr(self, "export_text_layout_lines"), dlg)
        mode = str(getattr(self, "_bk_export_text_layout_mode", "original") or "original").lower()
        rb_original.setChecked(mode not in {"lines", "table", "table_simple", "layout_tables"} or txt_only)
        rb_lines.setChecked(mode == "lines")
        mode_layout.addWidget(rb_original); mode_layout.addWidget(rb_lines)
    if supports_simple:
        rb_simple = QRadioButton(_bk_tab_tr(self, "export_text_layout_table_simple"), dlg)
        mode_layout.addWidget(rb_simple)
    if supports_extended:
        rb_table = QRadioButton(_bk_tab_tr(self, "export_text_layout_table"), dlg)
        mode_layout.addWidget(rb_table)
    if not include_text_modes and rb_simple is not None:
        rb_simple.setChecked(True)
    elif not include_text_modes and rb_table is not None:
        rb_table.setChecked(True)
    elif include_text_modes and not txt_only:
        mode = str(getattr(self, "_bk_export_text_layout_mode", "original") or "original").lower()
        if rb_layout_tables is not None: rb_layout_tables.setChecked(mode == "layout_tables")
        if rb_simple is not None: rb_simple.setChecked(mode == "table_simple")
        if rb_table is not None: rb_table.setChecked(mode == "table")
    if mode_box is not None: layout.addWidget(mode_box)
    else: layout.addLayout(mode_layout)

    _bk_add_export_orientation_group(self, dlg, layout)

    selected_keys = _bk_load_saved_column_keys_for_dialog(self)
    checkboxes = {}
    columns_box = QGroupBox(_bk_tab_tr(self, "export_table_columns_label"), dlg) if QGroupBox is not None else None
    columns_layout = QVBoxLayout(columns_box) if columns_box is not None else QVBoxLayout()
    if QGridLayout is not None:
        grid = QGridLayout(); grid.setHorizontalSpacing(18); grid.setVerticalSpacing(4)
        for idx, key in enumerate(_BK_TABULAR_KEYS):
            cb = QCheckBox(_bk_tabular_column_title(self, key), dlg); cb.setChecked(key in selected_keys); checkboxes[key] = cb
            grid.addWidget(cb, idx % 8, idx // 8)
        columns_layout.addLayout(grid)
    else:
        for key in _BK_TABULAR_KEYS:
            cb = QCheckBox(_bk_tabular_column_title(self, key), dlg); cb.setChecked(key in selected_keys); checkboxes[key] = cb; columns_layout.addWidget(cb)
    if _BK_QLineEdit is not None:
        custom_row = QHBoxLayout()
        custom_label = QLabel(_bk_tab_tr(self, "export_table_custom_column_label"), dlg)
        custom_edit = _BK_QLineEdit(dlg)
        custom_edit.setPlaceholderText(_bk_tab_tr(self, "export_table_custom_column_placeholder"))
        btn_add_custom = QPushButton(_bk_tab_tr(self, "export_table_custom_column_add"), dlg)
        custom_row.addWidget(custom_label); custom_row.addWidget(custom_edit, 1); custom_row.addWidget(btn_add_custom)
        columns_layout.addLayout(custom_row)
    else:
        custom_edit = None; btn_add_custom = None
    quick_row = QHBoxLayout()
    btn_all = QPushButton(_bk_tab_tr(self, "export_table_columns_all"), dlg)
    btn_none = QPushButton(_bk_tab_tr(self, "export_table_columns_none_button"), dlg)
    btn_remember = QPushButton(_bk_tab_tr(self, "export_table_columns_remember"), dlg)
    for btn in (btn_all, btn_none, btn_remember): quick_row.addWidget(btn)
    quick_row.addStretch(1); columns_layout.addLayout(quick_row)
    if columns_box is not None: layout.addWidget(columns_box)
    else: layout.addLayout(columns_layout)

    zone_box = QGroupBox(_bk_tab_tr(self, "export_zones_group"), dlg) if QGroupBox is not None else None
    zone_layout = QHBoxLayout(zone_box) if zone_box is not None else QHBoxLayout()
    cb_zones = QCheckBox(_bk_tab_tr(self, "export_table_use_zones"), dlg)
    try: remembered_zones = self.settings.value("export/table_use_zones", bool(getattr(self, "_bk_export_use_zones", False)), type=bool)
    except Exception: remembered_zones = bool(getattr(self, "_bk_export_use_zones", False))
    cb_zones.setChecked(bool(remembered_zones))
    btn_zones = QPushButton(_bk_tab_tr(self, "export_table_define_zones"), dlg)
    zone_layout.addWidget(cb_zones, 1); zone_layout.addWidget(btn_zones)
    if zone_box is not None: layout.addWidget(zone_box)
    else: layout.addLayout(zone_layout)

    hint = QLabel(_bk_tab_tr(self, "export_table_simple_hint"), dlg)
    hint.setWordWrap(True); layout.addWidget(hint)

    def current_checked_keys(): return [key for key, cb in checkboxes.items() if cb.isChecked()]
    def set_all():
        for cb in checkboxes.values(): cb.setChecked(True)
    def set_none():
        for cb in checkboxes.values(): cb.setChecked(False)
    def add_custom_column():
        if custom_edit is None: return
        title = _bk_tab_clean(custom_edit.text())
        if not title: return
        existing_titles = {_bk_tabular_column_title(self, key).casefold() for key in _BK_TABULAR_KEYS}
        if title.casefold() in existing_titles:
            QMessageBox.warning(dlg, _bk_tab_tr(self, "warn_title"), _bk_tab_tr(self, "export_table_custom_column_exists")); return
        base = _bk_custom_column_slug(title)
        key = "custom_" + base
        used = set(_BK_TABULAR_KEYS)
        i = 2
        while key in used:
            key = f"custom_{base}_{i}"; i += 1
        items = _bk_load_custom_columns(self) + [{"key": key, "title": title}]
        _bk_save_custom_columns(self, items)
        _bk_ensure_custom_columns(self)
        cb = QCheckBox(title, dlg); cb.setChecked(True); checkboxes[key] = cb
        if QGridLayout is not None:
            idx = len(checkboxes) - 1; grid.addWidget(cb, idx % 8, idx // 8)
        else:
            columns_layout.insertWidget(max(0, columns_layout.count()-2), cb)
        custom_edit.clear(); sync_enabled(); dlg.adjustSize()
    if btn_add_custom is not None:
        btn_add_custom.clicked.connect(add_custom_column)
        try: custom_edit.returnPressed.connect(add_custom_column)
        except Exception: pass
    result = {"mode": "table", "columns": selected_keys, "remembered": False, "use_zones": cb_zones.isChecked(), "cancelled": False}
    def remember_selection():
        keys = current_checked_keys()
        if not keys:
            QMessageBox.warning(dlg, _bk_tab_tr(self, "warn_title"), _bk_tab_tr(self, "export_table_columns_none")); return
        result["remembered"] = True; result["columns"] = _bk_save_column_keys(self, keys)
        try: self.settings.setValue("export/table_use_zones", bool(cb_zones.isChecked())); self.settings.sync()
        except Exception: pass
    def define_zones():
        _bk_ensure_custom_columns(self)
        try: task = self._current_task()
        except Exception: task = None
        zones = _bk_open_export_zones_dialog(self, task)
        if zones is not None: cb_zones.setChecked(bool(zones))
    btn_all.clicked.connect(set_all); btn_none.clicked.connect(set_none); btn_remember.clicked.connect(remember_selection); btn_zones.clicked.connect(define_zones)
    def sync_enabled():
        detailed = (rb_table is not None and rb_table.isChecked())
        show_columns = detailed
        if columns_box is not None: columns_box.setVisible(show_columns); columns_box.setEnabled(show_columns)
        if zone_box is not None: zone_box.setVisible(show_columns); zone_box.setEnabled(show_columns)
        cb_zones.setEnabled(show_columns); btn_zones.setEnabled(show_columns)
        hint.setVisible(rb_simple is not None and rb_simple.isChecked())
        try: dlg.adjustSize()
        except Exception: pass
    for rb in (rb_original, rb_lines, rb_layout_tables, rb_simple, rb_table):
        if rb is not None:
            try: rb.toggled.connect(sync_enabled)
            except Exception: pass
    sync_enabled()
    buttons = QDialogButtonBox(QDialogButtonBox.Ok | QDialogButtonBox.Cancel, dlg)
    try:
        buttons.button(QDialogButtonBox.Ok).setText(_bk_tab_tr(self, "btn_ok")); buttons.button(QDialogButtonBox.Cancel).setText(_bk_tab_tr(self, "btn_cancel"))
    except Exception: pass
    def cancel_dialog(): result["cancelled"] = True; dlg.done(QDialog.Rejected)
    def accept_checked():
        result["cancelled"] = False
        if rb_layout_tables is not None and rb_layout_tables.isChecked(): mode = "layout_tables"; keys = []
        elif rb_simple is not None and rb_simple.isChecked(): mode = "table_simple"; keys = []
        elif include_text_modes and rb_lines is not None and rb_lines.isChecked(): mode = "lines"; keys = []
        elif include_text_modes and rb_original is not None and rb_original.isChecked(): mode = "original"; keys = []
        else:
            mode = "table"; keys = current_checked_keys()
            if not keys:
                QMessageBox.warning(dlg, _bk_tab_tr(self, "warn_title"), _bk_tab_tr(self, "export_table_columns_none")); return
        result["mode"] = mode; result["columns"] = _bk_normalize_column_keys(keys); result["use_zones"] = bool(cb_zones.isChecked()) if mode == "table" else False
        try: self._bk_export_use_zones = result["use_zones"]
        except Exception: pass
        dlg.accept()
    buttons.accepted.connect(accept_checked); buttons.rejected.connect(cancel_dialog); layout.addWidget(buttons)
    dlg.setMinimumSize(420, 260)
    try:
        dlg.resize(600 if txt_only or (rb_simple is not None and rb_simple.isChecked()) else 760, 320 if txt_only else 540)
    except Exception: pass
    try: exec_result = dlg.exec()
    except Exception:
        result["cancelled"] = True; return None
    if exec_result != QDialog.Accepted:
        result["cancelled"] = True; return None
    return result


def _bk_export_single_interactive_final(self, item: TaskItem, fmt: str):
    fmt_l = str(fmt or "").lower().lstrip(".")
    if fmt_l in _BK_SIMPLE_TABLE_SPREADSHEET_FMTS:
        result = _bk_column_choice_dialog(self, fmt_l, include_text_modes=False)
        if result is None or result.get("cancelled"): return None
        if result.get("mode") == "table_simple":
            self._bk_export_table_mode = "simple"; self._bk_export_use_zones = False
        else:
            self._bk_export_table_mode = "table"; self._bk_export_current_column_keys = result.get("columns") or []; self._bk_export_use_zones = bool(result.get("use_zones", False))
            if result.get("remembered"): self._bk_export_selected_column_keys = self._bk_export_current_column_keys
        if callable(_BK_TABULAR_PREV_EXPORT_SINGLE): return _BK_TABULAR_PREV_EXPORT_SINGLE(self, item, fmt)
        return None
    if fmt_l in _BK_TEXT_LAYOUT_FMTS:
        result = _bk_column_choice_dialog(self, fmt_l, include_text_modes=True)
        if result is None or result.get("cancelled"): return None
        self._bk_export_text_layout_mode = str(result.get("mode") or "original")
        self._bk_export_current_column_keys = result.get("columns") or []
        self._bk_export_use_zones = bool(result.get("use_zones", False)) if self._bk_export_text_layout_mode == "table" else False
        if result.get("remembered"): self._bk_export_selected_column_keys = self._bk_export_current_column_keys
        if callable(_BK_TABULAR_PREV_EXPORT_SINGLE): return _BK_TABULAR_PREV_EXPORT_SINGLE(self, item, fmt)
        return None
    if callable(_BK_SIMPLE_PREV_EXPORT_SINGLE): return _BK_SIMPLE_PREV_EXPORT_SINGLE(self, item, fmt)
    return None


def _bk_export_batch_final(self, items, fmt: str):
    fmt_l = str(fmt or "").lower().lstrip(".")
    if fmt_l in _BK_SIMPLE_TABLE_SPREADSHEET_FMTS:
        result = _bk_column_choice_dialog(self, fmt_l, include_text_modes=False)
        if result is None or result.get("cancelled"): return None
        if result.get("mode") == "table_simple": self._bk_export_table_mode = "simple"; self._bk_export_use_zones = False
        else:
            self._bk_export_table_mode = "table"; self._bk_export_current_column_keys = result.get("columns") or []; self._bk_export_use_zones = bool(result.get("use_zones", False))
            if result.get("remembered"): self._bk_export_selected_column_keys = self._bk_export_current_column_keys
        if callable(_BK_TABULAR_PREV_EXPORT_BATCH): return _BK_TABULAR_PREV_EXPORT_BATCH(self, items, fmt)
        return None
    if fmt_l in _BK_TEXT_LAYOUT_FMTS:
        result = _bk_column_choice_dialog(self, fmt_l, include_text_modes=True)
        if result is None or result.get("cancelled"): return None
        self._bk_export_text_layout_mode = str(result.get("mode") or "original")
        self._bk_export_current_column_keys = result.get("columns") or []
        self._bk_export_use_zones = bool(result.get("use_zones", False)) if self._bk_export_text_layout_mode == "table" else False
        if result.get("remembered"): self._bk_export_selected_column_keys = self._bk_export_current_column_keys
        if callable(_BK_TABULAR_PREV_EXPORT_BATCH): return _BK_TABULAR_PREV_EXPORT_BATCH(self, items, fmt)
        return None
    if callable(_BK_SIMPLE_PREV_EXPORT_BATCH): return _BK_SIMPLE_PREV_EXPORT_BATCH(self, items, fmt)
    return None

try:
    _bk_column_choice_dialog = _bk_column_choice_dialog_final
    MainWindow._export_single_interactive = _bk_export_single_interactive_final
    MainWindow._export_batch = _bk_export_batch_final
    register_globals('bk', globals(), sorted(set(__all__)))
except Exception:
    pass

# ---------------------------------------------------------------------------
# Finaler Export-Feinschliff: kompakte Dialoge, CSV ohne Varianten,
# einfache Tabelle mit lesbaren Zellbreiten, Custom-Datentypen löschen/resetten.
# ---------------------------------------------------------------------------
try:
    _BK_FINAL_PREV_EXPORT_SINGLE_2 = MainWindow._export_single_interactive
except Exception:
    _BK_FINAL_PREV_EXPORT_SINGLE_2 = None
try:
    _BK_FINAL_PREV_EXPORT_BATCH_2 = MainWindow._export_batch
except Exception:
    _BK_FINAL_PREV_EXPORT_BATCH_2 = None


def _bk_write_plain_overlay_csv_final(path, item):
    try:
        _text, _kr, _pil, record_views = item.results
    except Exception:
        record_views = []
    matrix = _bk_simple_spatial_layout_from_item(item).get("matrix") if item and getattr(item, "results", None) else []
    if not matrix:
        matrix = [[_bk_tab_clean(getattr(rv, "text", ""))] for rv in (record_views or []) if _bk_tab_clean(getattr(rv, "text", ""))]
    with open(path, "w", encoding="utf-8-sig", newline="") as handle:
        writer = csv.writer(handle, delimiter=",", quotechar='"', quoting=csv.QUOTE_MINIMAL, lineterminator="\r\n")
        for row in matrix or []:
            writer.writerow([_bk_tab_clean(cell) for cell in (row or [])])


def _bk_simple_layout_cell_is_numeric(value):
    v = _bk_tab_clean(value)
    if not v or not re.search(r"\d", v):
        return False
    return bool(re.fullmatch(r"[\d IVXLCDMivxlcdm.,;:/\-]+", v))


def _bk_simple_layout_row_is_section(row):
    """Abschnitts-/Kopfzeilen wie 'Seite - 52 -' oder 'unter F. weitersuchen!'
    erkennen: eine einzelne Box, deren Text wie eine Überschrift aussieht."""
    if not row or len(row) != 1:
        return False
    text = _bk_tab_clean(row[0].get("text", ""))
    if not text:
        return False
    try:
        if _bk_simple_is_page_heading(text) or _bk_simple_is_subheading(text):
            return True
    except Exception:
        pass
    if text.endswith("!"):
        return True
    if re.fullmatch(r"[-–—\s]*Seite\s*[-–—]?\s*\d+\s*[-–—]?\s*", text, re.I):
        return True
    return False


def _bk_simple_spatial_layout_from_records(records):
    """Einfache Tabelle: Overlay-Boxen in ein sauberes, menschlich wirkendes
    Calc/Excel-Gitter.

    Prinzip statt Naechster-Cluster-Zuordnung: Die linken Kanten aller Boxen
    werden sortiert und an natuerlichen LUECKEN getrennt - so entstehen genau
    die Spalten, die ein Mensch beim Blick auf die Seite ziehen wuerde. Kaum
    besetzte Streuspalten werden in den Nachbarn eingegliedert, Abschnitts-
    zeilen (Seitenkopf, Zwischenueberschriften) werden als verbundene, fett
    zentrierte Zeilen markiert und Zahlen-/Datumsspalten rechtsbuendig gesetzt.
    """
    records = [dict(rec) for rec in (records or []) if _bk_tab_clean(rec.get("text", ""))]
    if not records:
        return {"matrix": [], "widths": [12.0], "heights": [18.0], "spans": set(), "aligns": []}
    row_groups = _bk_simple_group_rows(records)
    if not row_groups:
        return {"matrix": [], "widths": [12.0], "heights": [18.0], "spans": set(), "aligns": []}
    heights_px = sorted(max(1.0, float(rec.get("h", 1.0) or 1.0)) for rec in records)
    med_h = heights_px[len(heights_px) // 2] if heights_px else 12.0

    # Abschnittszeilen aus der Spaltenermittlung heraushalten, damit ein
    # zentrierter Seitenkopf keine eigene Geisterspalte erzeugt.
    section_flags = [_bk_simple_layout_row_is_section(row) for row in row_groups]
    data_records = [rec for row, is_sec in zip(row_groups, section_flags)
                    if not is_sec for rec in row]
    if not data_records:
        data_records = [rec for row in row_groups for rec in row]

    # 1) Spalten an natuerlichen Luecken der linken Kanten trennen.
    xs = sorted(float(rec.get("x0", 0.0) or 0.0) for rec in data_records)
    char_w = max(4.0, med_h * 0.55)          # grobe Zeichenbreite der Vorlage
    gap_threshold = max(3.0 * char_w, 26.0)  # Luecke, die eine neue Spalte beginnt
    boundaries = [xs[0]]
    for prev, cur in zip(xs, xs[1:]):
        if cur - prev > gap_threshold:
            boundaries.append(cur)
    columns = [{"x": b, "count": 0} for b in boundaries]

    def _column_index(x0):
        idx = 0
        for i, col in enumerate(columns):
            if x0 >= col["x"] - gap_threshold * 0.5:
                idx = i
        return idx

    for rec in data_records:
        columns[_column_index(float(rec.get("x0", 0.0) or 0.0))]["count"] += 1

    # 2) Streuspalten (kaum besetzt) mit dem Nachbarn verschmelzen.
    total_rows = max(1, sum(1 for f in section_flags if not f))
    min_fill = max(2, int(round(total_rows * 0.06)))
    keep_cols = []
    for i, col in enumerate(columns):
        if col["count"] >= min_fill or len(columns) == 1:
            keep_cols.append(col)
        elif keep_cols:
            keep_cols[-1]["count"] += col["count"]
        else:
            col_next = columns[i + 1] if i + 1 < len(columns) else None
            if col_next is not None:
                col_next["count"] += col["count"]
                col_next["x"] = col["x"]
            else:
                keep_cols.append(col)
    columns = keep_cols or columns

    def _final_column_index(x0):
        best, best_i = None, 0
        for i, col in enumerate(columns):
            if x0 >= col["x"] - gap_threshold * 0.5:
                best, best_i = col, i
        return best_i if best is not None else 0

    # 3) Matrix aufbauen; Abschnittszeilen als Span-Zeilen (Zelle 0, Rest leer).
    matrix, row_heights, spans = [], [], set()
    for row, is_sec in zip(row_groups, section_flags):
        if is_sec:
            text = _bk_tab_clean(row[0].get("text", ""))
            matrix.append([text] + [""] * (len(columns) - 1))
            spans.add(len(matrix) - 1)
            row_heights.append(max(17.0, min(30.0, med_h * 1.05 + 5.0)))
            continue
        values = [""] * len(columns)
        for rec in sorted(row or [], key=lambda r: (float(r.get("x0", 0.0) or 0.0), float(r.get("y0", 0.0) or 0.0))):
            text = _bk_simple_record_text(rec)
            if not text:
                continue
            idx = _final_column_index(float(rec.get("x0", 0.0) or 0.0))
            values[idx] = _bk_tab_join_text_fragments([values[idx], text]) if values[idx] else text
        if any(_bk_tab_clean(v) for v in values):
            matrix.append(values)
            try:
                max_h = max(max(1.0, float(rec.get("h", med_h) or med_h)) for rec in row)
            except Exception:
                max_h = med_h
            row_heights.append(max(15.0, min(28.0, max_h * 0.9 + 3.0)))

    if not matrix:
        return {"matrix": [], "widths": [12.0], "heights": [18.0], "spans": set(), "aligns": []}

    # 4) Komplett leere Spalten entfernen (Span-Zeilen zaehlen nicht mit).
    col_count = max(len(r) for r in matrix)
    keep = [c for c in range(col_count)
            if any(_bk_tab_clean(row[c] if c < len(row) else "")
                   for r_i, row in enumerate(matrix) if r_i not in spans or c == 0)]
    if not keep:
        keep = list(range(col_count))
    if 0 not in keep:
        keep = [0] + keep
    matrix = [[row[c] if c < len(row) else "" for c in keep] for row in matrix]
    kept_anchors = [float(columns[c].get("x", c) or c) for c in keep]
    matrix, kept_anchors = _bk_layout_normalize_sparse_text_columns(matrix, kept_anchors)
    matrix = _bk_layout_repair_matrix(matrix)

    # 5) Ausrichtung je Spalte: ueberwiegend Zahlen/Daten -> rechtsbuendig.
    aligns = []
    final_col_count = max([len(row) for row in matrix] or [1])
    for c in range(final_col_count):
        vals = [_bk_tab_clean(row[c]) for r_i, row in enumerate(matrix)
                if r_i not in spans and c < len(row) and _bk_tab_clean(row[c])]
        numeric = sum(1 for v in vals if _bk_simple_layout_cell_is_numeric(v))
        aligns.append("right" if vals and numeric / len(vals) >= 0.6 else "left")

    # Breiten zuerst aus der tatsächlichen Bildgeometrie ableiten. Dadurch
    # bleiben die beiden breiten Namensfelder und die schmalen Zahlenspalten
    # proportional wie in der Vorlage; reine Textlängen machten bisher alle
    # Spalten ähnlich breit und verursachten extreme Umbrüche.
    min_record_x = min(float(rec.get("x0", 0.0) or 0.0) for rec in data_records)
    max_record_x = max(float(rec.get("x1", rec.get("x0", 0.0)) or 0.0) for rec in data_records)
    col_widths = []
    for idx in range(final_col_count):
        if len(kept_anchors) == final_col_count:
            left = min_record_x if idx == 0 else (kept_anchors[idx - 1] + kept_anchors[idx]) / 2.0
            right = max_record_x if idx == final_col_count - 1 else (kept_anchors[idx] + kept_anchors[idx + 1]) / 2.0
            geometry_width = max(8.0, min(64.0, (right - left) / max(3.5, char_w) * 1.05))
        else:
            geometry_width = 12.0
        vals = [str(row[idx] if idx < len(row) else "")
                for r_i, row in enumerate(matrix) if r_i not in spans]
        non_empty = [v for v in vals if _bk_tab_clean(v)]
        longest = max([len(v) for v in non_empty] or [5])
        content_floor = max(8.0, min(42.0, longest * 0.55 + 3.0))
        col_widths.append(max(geometry_width, content_floor))
    return {"matrix": matrix, "widths": col_widths,
            "heights": row_heights or [18.0] * len(matrix),
            "spans": spans, "aligns": aligns}


def _bk_write_simple_layout_xlsx(path, layout, window=None):
    matrix = (layout or {}).get("matrix") or []
    widths = (layout or {}).get("widths") or _bk_simple_col_widths(matrix)
    heights = (layout or {}).get("heights") or [18.0] * max(1, len(matrix))
    spans = set((layout or {}).get("spans") or set())
    aligns = list((layout or {}).get("aligns") or [])
    bold_rows = set((layout or {}).get("bold_rows") or set())
    col_count = max(1, max([len(row) for row in matrix] or [1]))
    # Stil-Indizes: 1=links, 2=rechts, 3=Abschnittszeile (fett, zentriert),
    # 4=Kopfzeile (fett, links) - fuer Tabellenkoepfe des KI-Exports.
    row_xml = []
    merges = []
    for r_idx, row in enumerate(matrix or [[]], start=1):
        cells = []
        is_span = (r_idx - 1) in spans
        is_bold = (r_idx - 1) in bold_rows
        for c_idx in range(1, col_count + 1):
            value = row[c_idx - 1] if c_idx - 1 < len(row) else ""
            ref = f"{_bk_xlsx_col_name(c_idx)}{r_idx}"
            if is_span:
                style = 3
            elif is_bold:
                style = 4
            elif c_idx - 1 < len(aligns) and aligns[c_idx - 1] == "right":
                style = 2
            else:
                style = 1
            cells.append(f'<c r="{ref}" s="{style}" t="inlineStr"><is><t xml:space="preserve">{_bk_xml(value)}</t></is></c>')
        if is_span and col_count > 1:
            merges.append(f'<mergeCell ref="A{r_idx}:{_bk_xlsx_col_name(col_count)}{r_idx}"/>')
        height = heights[r_idx - 1] if r_idx - 1 < len(heights) else 18.0
        row_xml.append(f'<row r="{r_idx}" ht="{float(height):.2f}" customHeight="1">{"".join(cells)}</row>')
    cols = ''.join(f'<col min="{i}" max="{i}" width="{max(8.0, min(64.0, float(widths[i-1] if i-1 < len(widths) else 14.0))):.2f}" customWidth="1"/>' for i in range(1, col_count + 1))
    merge_xml = f'<mergeCells count="{len(merges)}">{"".join(merges)}</mergeCells>' if merges else ''
    sheet = ''.join([
        '<?xml version="1.0" encoding="UTF-8"?>',
        '<worksheet xmlns="http://schemas.openxmlformats.org/spreadsheetml/2006/main" xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships">',
        '<sheetViews><sheetView workbookViewId="0"/></sheetViews>',
        '<sheetFormatPr defaultRowHeight="18"/>',
        '<cols>', cols, '</cols>',
        '<sheetData>', ''.join(row_xml), '</sheetData>',
        merge_xml,
        '<pageMargins left="0.25" right="0.25" top="0.25" bottom="0.25" header="0" footer="0"/>',
        bk_xlsx_page_setup_xml(True, window),
        '</worksheet>',
    ])
    content_types = '<?xml version="1.0" encoding="UTF-8"?><Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types"><Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/><Default Extension="xml" ContentType="application/xml"/><Override PartName="/xl/workbook.xml" ContentType="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet.main+xml"/><Override PartName="/xl/worksheets/sheet1.xml" ContentType="application/vnd.openxmlformats-officedocument.spreadsheetml.worksheet+xml"/><Override PartName="/xl/styles.xml" ContentType="application/vnd.openxmlformats-officedocument.spreadsheetml.styles+xml"/><Override PartName="/docProps/core.xml" ContentType="application/vnd.openxmlformats-package.core-properties+xml"/><Override PartName="/docProps/app.xml" ContentType="application/vnd.openxmlformats-officedocument.extended-properties+xml"/></Types>'
    root_rels = '<?xml version="1.0" encoding="UTF-8"?><Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships"><Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument" Target="xl/workbook.xml"/></Relationships>'
    workbook = '<?xml version="1.0" encoding="UTF-8"?><workbook xmlns="http://schemas.openxmlformats.org/spreadsheetml/2006/main" xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships"><sheets><sheet name="Transkription" sheetId="1" r:id="rId1"/></sheets></workbook>'
    workbook_rels = '<?xml version="1.0" encoding="UTF-8"?><Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships"><Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/worksheet" Target="worksheets/sheet1.xml"/><Relationship Id="rId2" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/styles" Target="styles.xml"/></Relationships>'
    # Font 0 = normal, Font 1 = fett (Abschnittszeilen). Kein wrapText, damit
    # Zellen wie handgesetzt breit bleiben statt sofort mehrzeilig umzubrechen.
    styles = ('<?xml version="1.0" encoding="UTF-8"?><styleSheet xmlns="http://schemas.openxmlformats.org/spreadsheetml/2006/main">'
              '<fonts count="2"><font><sz val="10"/><name val="Calibri"/></font><font><b/><sz val="10"/><name val="Calibri"/></font></fonts>'
              '<fills count="2"><fill><patternFill patternType="none"/></fill><fill><patternFill patternType="gray125"/></fill></fills>'
              '<borders count="2"><border><left/><right/><top/><bottom/><diagonal/></border>'
              '<border><left style="thin"><color rgb="FFB7B7B7"/></left><right style="thin"><color rgb="FFB7B7B7"/></right><top style="thin"><color rgb="FFB7B7B7"/></top><bottom style="thin"><color rgb="FFB7B7B7"/></bottom><diagonal/></border></borders>'
              '<cellStyleXfs count="1"><xf numFmtId="0" fontId="0" fillId="0" borderId="0"/></cellStyleXfs>'
              '<cellXfs count="5">'
              '<xf numFmtId="0" fontId="0" fillId="0" borderId="0" xfId="0"/>'
              '<xf numFmtId="0" fontId="0" fillId="0" borderId="1" xfId="0" applyBorder="1" applyAlignment="1"><alignment horizontal="left" vertical="center" wrapText="0"/></xf>'
              '<xf numFmtId="0" fontId="0" fillId="0" borderId="1" xfId="0" applyBorder="1" applyAlignment="1"><alignment horizontal="right" vertical="center" wrapText="0"/></xf>'
              '<xf numFmtId="0" fontId="1" fillId="0" borderId="1" xfId="0" applyBorder="1" applyFont="1" applyAlignment="1"><alignment horizontal="center" vertical="center" wrapText="0"/></xf>'
              '<xf numFmtId="0" fontId="1" fillId="0" borderId="1" xfId="0" applyBorder="1" applyFont="1" applyAlignment="1"><alignment horizontal="left" vertical="center" wrapText="0"/></xf>'
              '</cellXfs>'
              '<cellStyles count="1"><cellStyle name="Normal" xfId="0" builtinId="0"/></cellStyles></styleSheet>')
    core = '<?xml version="1.0" encoding="UTF-8"?><cp:coreProperties xmlns:cp="http://schemas.openxmlformats.org/package/2006/metadata/core-properties" xmlns:dc="http://purl.org/dc/elements/1.1/"><dc:creator>Bottled Kraken</dc:creator></cp:coreProperties>'
    app = '<?xml version="1.0" encoding="UTF-8"?><Properties xmlns="http://schemas.openxmlformats.org/officeDocument/2006/extended-properties"><Application>Bottled Kraken</Application></Properties>'
    with zipfile.ZipFile(path, "w", zipfile.ZIP_DEFLATED) as archive:
        archive.writestr("[Content_Types].xml", content_types); archive.writestr("_rels/.rels", root_rels)
        archive.writestr("xl/workbook.xml", workbook); archive.writestr("xl/_rels/workbook.xml.rels", workbook_rels)
        archive.writestr("xl/worksheets/sheet1.xml", sheet); archive.writestr("xl/styles.xml", styles)
        archive.writestr("docProps/core.xml", core); archive.writestr("docProps/app.xml", app)


def _bk_write_simple_layout_ods(path, layout, window=None):
    matrix = (layout or {}).get("matrix") or []
    widths = (layout or {}).get("widths") or _bk_simple_col_widths(matrix)
    heights = (layout or {}).get("heights") or [18.0] * max(1, len(matrix))
    spans = set((layout or {}).get("spans") or set())
    aligns = list((layout or {}).get("aligns") or [])
    bold_rows = set((layout or {}).get("bold_rows") or set())
    col_count = max(1, max([len(row) for row in matrix] or [1]))
    column_styles = []
    row_styles = []
    columns = []
    # Spaltenbreiten auf die tatsaechliche Druckausrichtung normieren. Breite
    # historische Tabellen erhalten in Calc A4-Querformat; die frueher stets
    # verwendeten 19 cm entsprachen Hochformat und quetschten sieben Spalten
    # trotz vorhandener Seitenbreite unnoetig zusammen.
    desired = []
    for idx in range(1, col_count + 1):
        desired.append(max(1.2, min(14.0, float(widths[idx-1] if idx-1 < len(widths) else 14.0) / 3.2)))
    total = sum(desired)
    landscape = bool(col_count >= 5)
    max_page_cm = 27.5 if landscape else 19.0
    if total > max_page_cm:
        factor = max_page_cm / total
        desired = [max(1.0, w * factor) for w in desired]
    font_pt = 8.5 if col_count >= 7 else (9.0 if col_count >= 5 else 10.0)
    for idx, width in enumerate(desired, start=1):
        column_styles.append('<style:style style:name="co%d" style:family="table-column"><style:table-column-properties style:column-width="%.3fcm"/></style:style>' % (idx, width))
        columns.append('<table:table-column table:style-name="co%d"/>' % idx)
    for idx, height in enumerate(heights or [18.0], start=1):
        cm = max(0.42, min(1.1, float(height) / 30.0))
        # use-optimal-row-height: umgebrochene Zellen duerfen wachsen statt zu clippen.
        row_styles.append('<style:style style:name="ro%d" style:family="table-row"><style:table-row-properties style:min-row-height="%.3fcm" style:use-optimal-row-height="true" fo:break-before="auto"/></style:style>' % (idx, cm))
    table_rows = []
    for r_idx, row in enumerate(matrix or [[]], start=1):
        cells = []
        if (r_idx - 1) in spans and col_count > 1:
            # Abschnittszeile: eine verbundene, fett zentrierte Zelle ueber alle Spalten.
            text = row[0] if row else ""
            cells.append('<table:table-cell table:style-name="ceSection" office:value-type="string" table:number-columns-spanned="%d"><text:p>%s</text:p></table:table-cell>' % (col_count, _bk_ods_text(text)))
            cells.append('<table:covered-table-cell table:number-columns-repeated="%d"/>' % (col_count - 1))
        else:
            for c_idx in range(col_count):
                text = row[c_idx] if c_idx < len(row) else ""
                if (r_idx - 1) in bold_rows:
                    style = "ceHead"
                elif c_idx < len(aligns) and aligns[c_idx] == "right":
                    style = "ceRight"
                else:
                    style = "ceBody"
                cells.append('<table:table-cell table:style-name="%s" office:value-type="string"><text:p>%s</text:p></table:table-cell>' % (style, _bk_ods_text(text)))
        table_rows.append('<table:table-row table:style-name="ro%d">%s</table:table-row>' % (min(r_idx, len(row_styles) or 1), ''.join(cells)))
    mimetype = "application/vnd.oasis.opendocument.spreadsheet"
    content = ''.join(['<?xml version="1.0" encoding="UTF-8"?>', '<office:document-content xmlns:office="urn:oasis:names:tc:opendocument:xmlns:office:1.0" xmlns:style="urn:oasis:names:tc:opendocument:xmlns:style:1.0" xmlns:text="urn:oasis:names:tc:opendocument:xmlns:text:1.0" xmlns:table="urn:oasis:names:tc:opendocument:xmlns:table:1.0" xmlns:fo="urn:oasis:names:tc:opendocument:xmlns:xsl-fo-compatible:1.0" xmlns:svg="urn:oasis:names:tc:opendocument:xmlns:svg-compatible:1.0" office:version="1.2">', '<office:font-face-decls><style:font-face style:name="Arial" svg:font-family="Arial"/></office:font-face-decls>', '<office:automatic-styles>', ''.join(column_styles), ''.join(row_styles),
        '<style:style style:name="ta1" style:family="table" style:master-page-name="mp1"><style:table-properties table:display="true" style:writing-mode="lr-tb"/></style:style>',
        '<style:style style:name="ceBody" style:family="table-cell"><style:table-cell-properties fo:border="0.05pt solid #B7B7B7" fo:padding="0.04cm" style:vertical-align="middle" fo:wrap-option="wrap"/><style:text-properties fo:font-size="%.1fpt" style:font-name="Arial"/><style:paragraph-properties fo:margin-top="0cm" fo:margin-bottom="0cm"/></style:style>' % font_pt,
        '<style:style style:name="ceRight" style:family="table-cell"><style:table-cell-properties fo:border="0.05pt solid #B7B7B7" fo:padding="0.04cm" style:vertical-align="middle" fo:wrap-option="wrap"/><style:text-properties fo:font-size="%.1fpt" style:font-name="Arial"/><style:paragraph-properties fo:text-align="end" fo:margin-top="0cm" fo:margin-bottom="0cm"/></style:style>' % font_pt,
        '<style:style style:name="ceSection" style:family="table-cell"><style:table-cell-properties fo:border="0.05pt solid #B7B7B7" fo:padding="0.05cm" style:vertical-align="middle" fo:wrap-option="wrap"/><style:text-properties fo:font-size="%.1fpt" fo:font-weight="bold" style:font-name="Arial"/><style:paragraph-properties fo:text-align="center" fo:margin-top="0cm" fo:margin-bottom="0cm"/></style:style>' % max(font_pt, 9.0),
        '<style:style style:name="ceHead" style:family="table-cell"><style:table-cell-properties fo:border="0.05pt solid #808080" fo:padding="0.05cm" style:vertical-align="middle" fo:wrap-option="wrap"/><style:text-properties fo:font-size="%.1fpt" fo:font-weight="bold" style:font-name="Arial"/><style:paragraph-properties fo:margin-top="0cm" fo:margin-bottom="0cm"/></style:style>' % font_pt,
        '</office:automatic-styles>', '<office:body><office:spreadsheet><table:table table:name="Transkription" table:style-name="ta1">', ''.join(columns), ''.join(table_rows), '</table:table></office:spreadsheet></office:body></office:document-content>'])
    page_w, page_h = ((29.7, 21.0) if landscape else (21.0, 29.7))
    orientation = "landscape" if landscape else "portrait"
    styles = ('<?xml version="1.0" encoding="UTF-8"?><office:document-styles '
              'xmlns:office="urn:oasis:names:tc:opendocument:xmlns:office:1.0" '
              'xmlns:style="urn:oasis:names:tc:opendocument:xmlns:style:1.0" '
              'xmlns:text="urn:oasis:names:tc:opendocument:xmlns:text:1.0" '
              'xmlns:fo="urn:oasis:names:tc:opendocument:xmlns:xsl-fo-compatible:1.0" '
              'xmlns:svg="urn:oasis:names:tc:opendocument:xmlns:svg-compatible:1.0" office:version="1.2">'
              '<office:font-face-decls><style:font-face style:name="Arial" svg:font-family="Arial"/></office:font-face-decls>'
              '<office:styles><style:default-style style:family="paragraph"><style:text-properties fo:font-size="%.1fpt" style:font-name="Arial"/></style:default-style></office:styles>'
              '<office:automatic-styles><style:page-layout style:name="pm1"><style:page-layout-properties '
              'fo:page-width="%.1fcm" fo:page-height="%.1fcm" style:print-orientation="%s" '
              'fo:margin-top="0.7cm" fo:margin-bottom="0.7cm" fo:margin-left="0.7cm" fo:margin-right="0.7cm"/>'
              '</style:page-layout></office:automatic-styles><office:master-styles><style:master-page '
              'style:name="mp1" style:page-layout-name="pm1"/></office:master-styles></office:document-styles>') % (font_pt, page_w, page_h, orientation)
    meta = '<?xml version="1.0" encoding="UTF-8"?><office:document-meta xmlns:office="urn:oasis:names:tc:opendocument:xmlns:office:1.0" xmlns:meta="urn:oasis:names:tc:opendocument:xmlns:meta:1.0" office:version="1.2"><office:meta><meta:generator>Bottled Kraken</meta:generator></office:meta></office:document-meta>'
    settings = '<?xml version="1.0" encoding="UTF-8"?><office:document-settings xmlns:office="urn:oasis:names:tc:opendocument:xmlns:office:1.0" office:version="1.2"><office:settings/></office:document-settings>'
    manifest = '<?xml version="1.0" encoding="UTF-8"?><manifest:manifest xmlns:manifest="urn:oasis:names:tc:opendocument:xmlns:manifest:1.0" manifest:version="1.2"><manifest:file-entry manifest:full-path="/" manifest:media-type="%s"/><manifest:file-entry manifest:full-path="content.xml" manifest:media-type="text/xml"/><manifest:file-entry manifest:full-path="styles.xml" manifest:media-type="text/xml"/><manifest:file-entry manifest:full-path="meta.xml" manifest:media-type="text/xml"/><manifest:file-entry manifest:full-path="settings.xml" manifest:media-type="text/xml"/></manifest:manifest>' % mimetype
    with zipfile.ZipFile(path, "w") as archive:
        info = zipfile.ZipInfo("mimetype"); info.date_time = (2020, 1, 1, 0, 0, 0); info.compress_type = zipfile.ZIP_STORED
        archive.writestr(info, mimetype)
        for name, data in (("content.xml", content), ("styles.xml", styles), ("meta.xml", meta), ("settings.xml", settings), ("META-INF/manifest.xml", manifest)):
            zi = zipfile.ZipInfo(name); zi.date_time = (2020, 1, 1, 0, 0, 0); zi.compress_type = zipfile.ZIP_DEFLATED
            archive.writestr(zi, data.encode("utf-8"))


def _bk_write_table_odt(path, rows, column_keys=None, window=None):
    keys = _bk_normalize_column_keys(column_keys)
    if not keys:
        keys = ["original_line"] if "original_line" in _BK_TABULAR_COLUMN_BY_KEY else list(_BK_TABULAR_KEYS[:1])
    matrix = _bk_table_matrix_from_rows(rows, keys, window)
    widths = []
    for key in keys:
        meta_width = float(_BK_TABULAR_COLUMN_BY_KEY.get(key, (None, None, None, 16.0, 3.0))[4])
        widths.append(max(2.2, min(8.0, meta_width)))
    column_styles = []
    columns = []
    for idx, width in enumerate(widths, start=1):
        column_styles.append('<style:style style:name="co%d" style:family="table-column"><style:table-column-properties style:column-width="%.3fcm"/></style:style>' % (idx, width))
        columns.append('<table:table-column table:style-name="co%d"/>' % idx)
    table_rows = []
    for r_idx, row in enumerate(matrix, start=1):
        style = "ceHeader" if r_idx == 1 else "ceBody"
        cells = []
        for c_idx in range(len(widths)):
            text = row[c_idx] if c_idx < len(row) else ""
            cells.append('<table:table-cell table:style-name="%s" office:value-type="string"><text:p>%s</text:p></table:table-cell>' % (style, _bk_ods_text(text)))
        table_rows.append('<table:table-row>%s</table:table-row>' % ''.join(cells))
    content = ''.join(['<?xml version="1.0" encoding="UTF-8"?>', '<office:document-content xmlns:office="urn:oasis:names:tc:opendocument:xmlns:office:1.0" xmlns:style="urn:oasis:names:tc:opendocument:xmlns:style:1.0" xmlns:text="urn:oasis:names:tc:opendocument:xmlns:text:1.0" xmlns:table="urn:oasis:names:tc:opendocument:xmlns:table:1.0" xmlns:fo="urn:oasis:names:tc:opendocument:xmlns:xsl-fo-compatible:1.0" xmlns:svg="urn:oasis:names:tc:opendocument:xmlns:svg-compatible:1.0" office:version="1.2">', '<office:font-face-decls><style:font-face style:name="Arial" svg:font-family="Arial"/></office:font-face-decls>', '<office:automatic-styles>', ''.join(column_styles), '<style:style style:name="ceHeader" style:family="table-cell"><style:table-cell-properties fo:border="0.05pt solid #808080" fo:background-color="#E9EEF6" fo:padding="0.05cm"/><style:text-properties fo:font-size="8pt" fo:font-weight="bold" style:font-name="Arial"/><style:paragraph-properties fo:margin-top="0cm" fo:margin-bottom="0cm"/></style:style>', '<style:style style:name="ceBody" style:family="table-cell"><style:table-cell-properties fo:border="0.05pt solid #B7B7B7" fo:padding="0.05cm"/><style:text-properties fo:font-size="8pt" style:font-name="Arial"/><style:paragraph-properties fo:margin-top="0cm" fo:margin-bottom="0cm"/></style:style>', '</office:automatic-styles>', '<office:body><office:text><table:table table:name="Transkription">', ''.join(columns), ''.join(table_rows), '</table:table></office:text></office:body></office:document-content>'])
    _ods_page_w, _ods_page_h = bk_page_size_cm(True, window)
    styles = '<?xml version="1.0" encoding="UTF-8"?><office:document-styles xmlns:office="urn:oasis:names:tc:opendocument:xmlns:office:1.0" xmlns:style="urn:oasis:names:tc:opendocument:xmlns:style:1.0" xmlns:text="urn:oasis:names:tc:opendocument:xmlns:text:1.0" xmlns:fo="urn:oasis:names:tc:opendocument:xmlns:xsl-fo-compatible:1.0" xmlns:svg="urn:oasis:names:tc:opendocument:xmlns:svg-compatible:1.0" office:version="1.2"><office:font-face-decls><style:font-face style:name="Arial" svg:font-family="Arial"/></office:font-face-decls><office:styles><style:default-style style:family="paragraph"><style:text-properties fo:font-size="8pt" style:font-name="Arial"/></style:default-style></office:styles><office:automatic-styles><style:page-layout style:name="pm1"><style:page-layout-properties fo:page-width="%.1fcm" fo:page-height="%.1fcm" style:print-orientation="%s" fo:margin-top="1.2cm" fo:margin-bottom="1.2cm" fo:margin-left="1.2cm" fo:margin-right="1.2cm"/></style:page-layout></office:automatic-styles><office:master-styles><style:master-page style:name="Standard" style:page-layout-name="pm1"/></office:master-styles></office:document-styles>' % (_ods_page_w, _ods_page_h, bk_odf_orientation_name(True, window))
    meta = '<?xml version="1.0" encoding="UTF-8"?><office:document-meta xmlns:office="urn:oasis:names:tc:opendocument:xmlns:office:1.0" xmlns:meta="urn:oasis:names:tc:opendocument:xmlns:meta:1.0" office:version="1.2"><office:meta><meta:generator>Bottled Kraken</meta:generator></office:meta></office:document-meta>'
    settings = '<?xml version="1.0" encoding="UTF-8"?><office:document-settings xmlns:office="urn:oasis:names:tc:opendocument:xmlns:office:1.0" office:version="1.2"><office:settings/></office:document-settings>'
    mimetype = "application/vnd.oasis.opendocument.text"
    manifest = '<?xml version="1.0" encoding="UTF-8"?><manifest:manifest xmlns:manifest="urn:oasis:names:tc:opendocument:xmlns:manifest:1.0" manifest:version="1.2"><manifest:file-entry manifest:full-path="/" manifest:media-type="%s"/><manifest:file-entry manifest:full-path="content.xml" manifest:media-type="text/xml"/><manifest:file-entry manifest:full-path="styles.xml" manifest:media-type="text/xml"/><manifest:file-entry manifest:full-path="meta.xml" manifest:media-type="text/xml"/><manifest:file-entry manifest:full-path="settings.xml" manifest:media-type="text/xml"/></manifest:manifest>' % mimetype
    with zipfile.ZipFile(path, "w") as archive:
        info = zipfile.ZipInfo("mimetype"); info.date_time = (2020, 1, 1, 0, 0, 0); info.compress_type = zipfile.ZIP_STORED
        archive.writestr(info, mimetype)
        for name, data in (("content.xml", content), ("styles.xml", styles), ("meta.xml", meta), ("settings.xml", settings), ("META-INF/manifest.xml", manifest)):
            zi = zipfile.ZipInfo(name); zi.date_time = (2020, 1, 1, 0, 0, 0); zi.compress_type = zipfile.ZIP_DEFLATED
            archive.writestr(zi, data.encode("utf-8"))



# ---------------------------------------------------------------------------
# Layout mit Tabellen: Fließtext mit Leerzeichen positionieren, erkannte
# Tabellenbereiche aber als echte Office-Tabellen ausgeben. Keine Textfelder.

_BK_LAYOUT_TABLE_TEXT_FMTS = {"docx", "word", "odt"}

def _bk_layout_text_records(record_views):
    records = _records_from_views(record_views)
    return [rec for rec in records if _bk_tab_clean(rec.get("text", ""))]


def _bk_layout_bbox_for_rows(rows):
    boxes = []
    for row in rows or []:
        for rec in row.get("items", []) if isinstance(row, dict) else row:
            if rec.get("bbox"):
                boxes.append((float(rec.get("x0", 0.0)), float(rec.get("y0", 0.0)), float(rec.get("x1", 0.0)), float(rec.get("y1", 0.0))))
    if not boxes:
        return None
    return (min(bb[0] for bb in boxes), min(bb[1] for bb in boxes), max(bb[2] for bb in boxes), max(bb[3] for bb in boxes))


def _bk_layout_median_char_px(records, fallback=6.0):
    values = []
    for rec in records or []:
        txt = _bk_tab_clean(rec.get("text", ""))
        if txt and rec.get("bbox"):
            w = max(1.0, float(rec.get("w", 0.0) or 0.0))
            values.append(w / max(1, len(txt)))
    if not values:
        return float(fallback)
    values.sort()
    return max(2.5, min(14.0, float(values[len(values)//2])))


def _bk_layout_group_rows(records):
    if not records:
        return []
    tolerance = max(2.0, _median_height(records) * 0.55)
    rows = _group_rows(records, tolerance)
    out = []
    for row in rows:
        items = [rec for rec in row.get("items", []) if _bk_tab_clean(rec.get("text", ""))]
        if not items:
            continue
        items.sort(key=lambda rec: (float(rec.get("x0", 0.0) or 0.0), int(rec.get("index", 0) or 0)))
        bb = _bk_layout_bbox_for_rows([{"items": items}])
        if bb:
            cy = (bb[1] + bb[3]) / 2.0
        else:
            cy = float(row.get("cy", 0.0) or 0.0)
        out.append({"items": items, "bbox": bb, "cy": cy})
    out.sort(key=lambda row: (float(row.get("bbox", (0, 0, 0, 0))[1] if row.get("bbox") else row.get("cy", 0.0)), float(row.get("bbox", (0, 0, 0, 0))[0] if row.get("bbox") else 0.0)))
    return out


def _bk_layout_row_is_table_like(row):
    items = [rec for rec in (row or {}).get("items", []) if _bk_tab_clean(rec.get("text", ""))]
    count = len(items)
    if count >= 4:
        return True
    if count < 3:
        return False
    numeric = 0
    shortish = 0
    for rec in items:
        txt = _bk_tab_clean(rec.get("text", ""))
        if re.search(r"\d|%|/|,\d", txt):
            numeric += 1
        if len(txt) <= 18:
            shortish += 1
    return numeric >= 1 or shortish >= max(2, count - 1)


def _bk_layout_column_anchors(table_rows):
    xs = []
    widths = []
    for row in table_rows or []:
        for rec in row.get("items", []):
            if not _bk_tab_clean(rec.get("text", "")):
                continue
            xs.append(float(rec.get("x0", 0.0) or 0.0))
            if rec.get("bbox"):
                widths.append(max(1.0, float(rec.get("w", 1.0) or 1.0)))
    if not xs:
        return []
    widths.sort()
    med_w = widths[len(widths)//2] if widths else 40.0
    tol = max(12.0, min(28.0, med_w * 0.35))
    anchors = []
    counts = []
    for x in sorted(xs):
        if not anchors or abs(x - anchors[-1]) > tol:
            anchors.append(x)
            counts.append(1)
        else:
            n = counts[-1]
            anchors[-1] = (anchors[-1] * n + x) / (n + 1)
            counts[-1] = n + 1
    return anchors


def _bk_layout_column_is_numeric(values):
    values = [_bk_tab_clean(value) for value in (values or []) if _bk_tab_clean(value)]
    if not values:
        return False
    numeric = sum(1 for value in values if _bk_simple_layout_cell_is_numeric(value))
    return numeric / max(1, len(values)) >= 0.55


def _bk_layout_normalize_sparse_text_columns(matrix, anchors=None):
    """Entfernt Geisterspalten, die durch eingerückte Abschnittsüberschriften
    entstehen.

    Historische Doppeltabellen enthalten häufig normale Datenzeilen am linken
    Rand, während Abschnittstitel etwas eingerückt beginnen. Eine reine
    x-Anker-Erkennung deutet diese Einrückung als zusätzliche Spalte. Dasselbe
    passiert bei kurzen Fortsetzungen wie ``von`` auf der rechten Tabellen-
    hälfte. Solche Spalten sind dünn besetzt, textuell und überschneiden sich
    praktisch nie mit der benachbarten eigentlichen Namensspalte. Sie werden
    deshalb in diese Nachbarspalte zurückgeführt. Zahlen-/Datenspalten werden
    ausdrücklich nicht zusammengelegt.
    """
    rows = [list(row or []) for row in (matrix or [])]
    if not rows:
        return [], list(anchors or [])
    width = max(len(row) for row in rows)
    rows = [row + [""] * (width - len(row)) for row in rows]
    xs = list(anchors or [])
    if len(xs) != width:
        xs = [float(i) for i in range(width)]

    # Vollständig leere Spalten zuerst entfernen.
    keep = [c for c in range(width)
            if any(_bk_tab_clean(row[c]) for row in rows)]
    if not keep:
        return rows, xs
    rows = [[row[c] for c in keep] for row in rows]
    xs = [xs[c] for c in keep]

    def stats(col):
        vals = [_bk_tab_clean(row[col]) for row in rows if _bk_tab_clean(row[col])]
        return {
            "count": len(vals),
            "numeric": _bk_layout_column_is_numeric(vals),
        }

    # Mehrfach laufen, weil nach dem Entfernen einer Geisterspalte eine zweite
    # kurze Fortsetzung direkt an ihre Zielspalte rücken kann.
    while len(xs) > 2:
        row_count = max(1, len(rows))
        table_span = max(1.0, float(xs[-1]) - float(xs[0]))
        merge_choice = None
        for c in range(len(xs)):
            current = stats(c)
            if current["count"] == 0 or current["numeric"]:
                continue
            # Abschnittsspalten sind im Verhältnis zur Zahl der Tabellenzeilen
            # klar dünn besetzt. Bei sehr kleinen Tabellen mindestens eine
            # einzelne Zeile zulassen.
            sparse_limit = max(1, int(round(row_count * 0.30)))
            if current["count"] > sparse_limit:
                continue
            for neighbor in (c - 1, c + 1):
                if neighbor < 0 or neighbor >= len(xs):
                    continue
                target = stats(neighbor)
                if target["numeric"] or target["count"] < max(3, current["count"] * 2):
                    continue
                overlap = sum(1 for row in rows
                              if _bk_tab_clean(row[c]) and _bk_tab_clean(row[neighbor]))
                if overlap > max(1, int(round(current["count"] * 0.15))):
                    continue
                gap = abs(float(xs[c]) - float(xs[neighbor]))
                if gap > max(45.0, table_span * 0.18):
                    continue
                score = (gap, -target["count"])
                if merge_choice is None or score < merge_choice[0]:
                    merge_choice = (score, c, neighbor)
        if merge_choice is None:
            break
        _score, source, target = merge_choice
        source_left = float(xs[source]) < float(xs[target])
        for row in rows:
            a = _bk_tab_clean(row[source])
            b = _bk_tab_clean(row[target])
            if a and b:
                row[target] = _bk_tab_join_text_fragments([a, b] if source_left else [b, a])
            elif a:
                row[target] = a
            del row[source]
        del xs[source]
    return rows, xs


def _bk_layout_digit_groups(value):
    """Liefert die Zahlengruppen einer historischen Tabellenzelle.

    Gedruckte Werte wie ``199 575`` oder ``16 029 349`` werden von OCR
    gelegentlich als eine gemeinsame Box gelesen. Die Gruppen werden nur fuer
    die strukturelle Reparatur ausgewertet; der eigentliche Text bleibt
    ansonsten unveraendert.
    """
    return re.findall(r"\d+", _bk_tab_clean(value))


def _bk_layout_typical_digit_group_count(matrix, column):
    counts = []
    for row in matrix or []:
        if column >= len(row):
            continue
        value = _bk_tab_clean(row[column])
        if not value or not _bk_simple_layout_cell_is_numeric(value):
            continue
        groups = _bk_layout_digit_groups(value)
        if groups and len(groups) <= 4:
            counts.append(len(groups))
    if not counts:
        return 0
    frequency = {}
    for count in counts:
        frequency[count] = frequency.get(count, 0) + 1
    return max(sorted(frequency), key=lambda count: (frequency[count], -count))


def _bk_layout_split_merged_numeric_cells(matrix):
    """Teilt versehentlich zusammengezogene Nachbarzahlen wieder auf.

    Bei historischen Statistiktafeln liest Kraken beispielsweise
    ``199 575 16 029 349`` gelegentlich als eine Box, obwohl die Werte in den
    Spalten ``qkm`` und ``Einw.`` stehen. Eine Teilung erfolgt ausschliesslich,
    wenn rechts unmittelbar eine ueberwiegend numerische, aber in dieser Zeile
    leere Spalte folgt und das uebrige Tabellenmaterial eine stabile
    Gruppenzahl erkennen laesst.
    """
    rows = [list(row) for row in (matrix or [])]
    if not rows:
        return rows
    width = max(len(row) for row in rows)
    for row in rows:
        row.extend([""] * (width - len(row)))
    numeric_columns = []
    for column in range(width):
        values = [_bk_tab_clean(row[column]) for row in rows if _bk_tab_clean(row[column])]
        numeric_columns.append(_bk_layout_column_is_numeric(values))
    typical = [_bk_layout_typical_digit_group_count(rows, column) for column in range(width)]

    for row in rows:
        for column in range(width - 1):
            if not numeric_columns[column] or not numeric_columns[column + 1]:
                continue
            value = _bk_tab_clean(row[column])
            if not value or _bk_tab_clean(row[column + 1]):
                continue
            groups = _bk_layout_digit_groups(value)
            if len(groups) < 4:
                continue
            right_count = typical[column + 1]
            left_count = typical[column]
            split_at = 0
            # Die rechte Spalte ist bei Einwohnerzahlen meist stabiler als die
            # Flaechenspalte (zwei oder drei Dreiergruppen). Deshalb von rechts
            # teilen, sobald die verbleibende linke Zahl plausibel ist.
            if 1 <= right_count <= 4 and 1 <= len(groups) - right_count <= 4:
                split_at = len(groups) - right_count
            elif 1 <= left_count <= 4 and 1 <= len(groups) - left_count <= 4:
                split_at = left_count
            if not split_at:
                continue
            left_groups = groups[:split_at]
            right_groups = groups[split_at:]
            if not left_groups or not right_groups:
                continue
            row[column] = " ".join(left_groups)
            row[column + 1] = " ".join(right_groups)
    return rows


def _bk_layout_merge_heading_continuations(matrix):
    """Fuegt kurze, in die Folgezeile gerutschte Ueberschriftsreste an.

    Typisches Beispiel der Vorlage: ``B. Unter der Provinzialregierung`` und
    das allein in der naechsten Zeile erkannte Wort ``von``. Datenzellen werden
    nicht verschoben; die Reparatur gilt nur fuer Textspalten ohne Zahlen auf
    derselben Tabellenseite.
    """
    rows = [list(row) for row in (matrix or [])]
    if len(rows) < 2:
        return rows
    width = max(len(row) for row in rows)
    for row in rows:
        row.extend([""] * (width - len(row)))
    numeric_columns = []
    for column in range(width):
        values = [_bk_tab_clean(row[column]) for row in rows if _bk_tab_clean(row[column])]
        numeric_columns.append(_bk_layout_column_is_numeric(values))
    text_columns = [column for column in range(width) if not numeric_columns[column]]
    connectors = re.compile(r"(?:\b(?:dem|der|des|den|von|unter|de|du|of|the)\b)[ .,:;\-–—]*$", re.IGNORECASE)
    connector_word = re.compile(r"^(?:von|de|du|of|the)[ .,:;\-–—]*$", re.IGNORECASE)
    for column in text_columns:
        next_text_column = min([c for c in text_columns if c > column] or [width])
        side_numeric = [c for c in range(column + 1, next_text_column) if numeric_columns[c]]
        if not side_numeric:
            side_numeric = [c for c in range(column + 1, width) if numeric_columns[c]][:3]
        for index in range(len(rows) - 1):
            current = _bk_tab_clean(rows[index][column])
            continuation = _bk_tab_clean(rows[index + 1][column])
            if not current or not continuation or re.search(r"\d", continuation):
                continue
            if len(continuation) > 24:
                continue
            if not connectors.search(current) and not connector_word.fullmatch(continuation):
                continue
            current_has_values = any(_bk_tab_clean(rows[index][c]) for c in side_numeric)
            next_has_values = any(_bk_tab_clean(rows[index + 1][c]) for c in side_numeric)
            if current_has_values or next_has_values:
                continue
            rows[index][column] = _bk_tab_join_text_fragments([current, continuation])
            rows[index + 1][column] = ""
    return rows


def _bk_layout_repair_matrix(matrix):
    rows = _bk_layout_split_merged_numeric_cells(matrix)
    rows = _bk_layout_merge_heading_continuations(rows)
    return rows


def _bk_layout_table_header_rows(matrix):
    """Markiert nur echte textuelle Kopfzeilen.

    Die fruehere Annahme, die erste erkannte Tabellenzeile sei immer ein
    Tabellenkopf, machte bei historischen Vorlagen die erste Datenzeile fett.
    """
    if not matrix:
        return set()
    first = [_bk_tab_clean(value) for value in matrix[0] if _bk_tab_clean(value)]
    if len(first) < 2:
        return set()
    joined = " ".join(first)
    numeric_cells = sum(1 for value in first if _bk_simple_layout_cell_is_numeric(value))
    if not re.search(r"\d", joined) and numeric_cells == 0:
        return {0}
    return set()


def _bk_layout_table_matrix(table_rows, anchors, return_anchors=False):
    if not anchors:
        return ([], []) if return_anchors else []
    matrix = []
    for row in table_rows or []:
        values = [""] * len(anchors)
        for rec in sorted(row.get("items", []), key=lambda item: float(item.get("x0", 0.0) or 0.0)):
            txt = _bk_tab_clean(rec.get("text", ""))
            if not txt:
                continue
            x = float(rec.get("x0", 0.0) or 0.0)
            col = min(range(len(anchors)), key=lambda idx: abs(float(anchors[idx]) - x))
            values[col] = _bk_tab_join_text_fragments([values[col], txt]) if values[col] else txt
        if any(_bk_tab_clean(v) for v in values):
            matrix.append(values)
    if not matrix:
        return ([], []) if return_anchors else []
    matrix, normalized_anchors = _bk_layout_normalize_sparse_text_columns(matrix, anchors)
    matrix = _bk_layout_repair_matrix(matrix)
    if return_anchors:
        return matrix, normalized_anchors
    return matrix


def _bk_layout_repair_block_heading_fragments(blocks):
    """Zieht einen am Tabellenanfang gelandeten Ueberschriftsrest zurueck.

    Mehrzeilige Gegenueberschriften wie ``A. Unmittelbar unter dem /
    Vicekoenig.`` werden von OCR mitunter so getrennt, dass die letzte Zeile in
    der ersten Datenzeile steht. Nur ein kurzer Text ohne Ziffern wird bewegt,
    und nur wenn die unmittelbar vorhergehende Textzeile mit einem klaren
    Bindewort endet.
    """
    repaired = list(blocks or [])
    connector = re.compile(r"(?:\b(?:dem|der|des|den|von|unter|de|du|of|the)\b)[ .,:;\-–—]*$", re.IGNORECASE)
    for index in range(len(repaired) - 1):
        text_block = repaired[index]
        table_block = repaired[index + 1]
        if text_block.get("type") != "text" or table_block.get("type") != "table":
            continue
        text_rows = text_block.get("rows") or []
        matrix = table_block.get("matrix") or []
        if not text_rows or not matrix:
            continue
        items = [item for item in (text_rows[-1].get("items") or []) if _bk_tab_clean(item.get("text", ""))]
        if not items:
            continue
        items.sort(key=lambda item: float(item.get("x0", 0.0) or 0.0))
        tail_item = items[-1]
        tail = _bk_tab_clean(tail_item.get("text", ""))
        if not connector.search(tail):
            continue
        first = list(matrix[0])
        width = len(first)
        values_by_column = [[_bk_tab_clean(row[c]) for row in matrix if c < len(row) and _bk_tab_clean(row[c])]
                            for c in range(width)]
        numeric_columns = [_bk_layout_column_is_numeric(values) for values in values_by_column]
        text_columns = [c for c in range(width) if not numeric_columns[c]]
        if len(text_columns) < 2:
            continue
        # Bei einer Doppeltabelle ist die zweite Textspalte die rechte
        # Bezeichnungsspalte. Eine etwaige Kopfspalte ganz links bleibt unberuehrt.
        column = text_columns[1]
        fragment = _bk_tab_clean(first[column])
        if not fragment or len(fragment) > 36 or re.search(r"\d", fragment):
            continue
        right_numeric = [c for c in range(column + 1, width) if numeric_columns[c]]
        if any(_bk_tab_clean(first[c]) for c in right_numeric):
            continue
        tail_item["text"] = _bk_tab_join_text_fragments([tail, fragment])
        first[column] = ""
        matrix[0] = first
        table_block["matrix"] = matrix
    return repaired


def _bk_layout_blocks_with_tables(record_views, image_size=None):
    records = _bk_layout_text_records(record_views)
    if not records:
        return {"blocks": [], "records": [], "bounds": (0.0, 0.0, 1.0, 1.0), "page_size": (1.0, 1.0), "median_height": 12.0, "char_px": 6.0}
    page_w, page_h = _page_size(image_size, records)
    min_x, min_y, max_x, max_y = _content_bounds(records, page_w, page_h)
    rows = _bk_layout_group_rows(records)
    blocks = []
    idx = 0
    median_height = _median_height(records)
    while idx < len(rows):
        row = rows[idx]
        if not _bk_layout_row_is_table_like(row):
            blocks.append({"type": "text", "rows": [row], "bbox": row.get("bbox")})
            idx += 1
            continue
        end = idx
        while end < len(rows):
            if _bk_layout_row_is_table_like(rows[end]):
                if end > idx:
                    prev_bb = rows[end - 1].get("bbox")
                    cur_bb = rows[end].get("bbox")
                    if prev_bb and cur_bb and (float(cur_bb[1]) - float(prev_bb[3])) > median_height * 1.65:
                        break
                end += 1
                continue
            # Ueberbruecken: Eine EINZELNE nicht-tabellenartige Zeile mitten im
            # Register (z. B. von Kraken als eine lange Box gelesen) soll die
            # Tabelle nicht zerreissen. Sie wird in den Cluster aufgenommen,
            # wenn direkt danach wieder tabellenartige Zeilen folgen und die
            # vertikalen Abstaende klein sind.
            if (end > idx and end + 1 < len(rows)
                    and _bk_layout_row_is_table_like(rows[end + 1])):
                prev_bb = rows[end - 1].get("bbox")
                cur_bb = rows[end].get("bbox")
                next_bb = rows[end + 1].get("bbox")
                close_above = (prev_bb and cur_bb
                               and (float(cur_bb[1]) - float(prev_bb[3])) <= median_height * 1.2)
                close_below = (cur_bb and next_bb
                               and (float(next_bb[1]) - float(cur_bb[3])) <= median_height * 1.2)
                if close_above and close_below:
                    end += 1
                    continue
            break
        cluster = rows[idx:end]
        max_cols = max((len(r.get("items", [])) for r in cluster), default=0)
        if len(cluster) >= 3 and max_cols >= 3:
            anchors = _bk_layout_column_anchors(cluster)
            matrix, anchors = _bk_layout_table_matrix(cluster, anchors, return_anchors=True)
            if matrix and max((len(r) for r in matrix), default=0) >= 3:
                blocks.append({"type": "table", "rows": cluster, "matrix": matrix, "anchors": anchors,
                               "header_rows": _bk_layout_table_header_rows(matrix),
                               "bbox": _bk_layout_bbox_for_rows(cluster)})
            else:
                for r in cluster:
                    blocks.append({"type": "text", "rows": [r], "bbox": r.get("bbox")})
        else:
            for r in cluster:
                blocks.append({"type": "text", "rows": [r], "bbox": r.get("bbox")})
        idx = end
    blocks = _bk_layout_repair_block_heading_fragments(blocks)
    return {"blocks": blocks, "records": records, "bounds": (min_x, min_y, max_x, max_y), "page_size": (page_w, page_h), "median_height": median_height, "char_px": _bk_layout_median_char_px(records)}


def _bk_layout_row_text(row):
    return " ".join(_bk_tab_clean(rec.get("text", "")) for rec in (row or {}).get("items", []) if _bk_tab_clean(rec.get("text", ""))).strip()


def _bk_layout_is_centered(row, content_bounds):
    bb = (row or {}).get("bbox") or content_bounds
    try:
        min_x, _min_y, max_x, _max_y = content_bounds
        content_w = max(1.0, float(max_x) - float(min_x))
        row_w = max(1.0, float(bb[2]) - float(bb[0]))
        row_c = (float(bb[0]) + float(bb[2])) / 2.0
        content_c = (float(min_x) + float(max_x)) / 2.0
        return abs(row_c - content_c) <= content_w * 0.18 or (row_w <= content_w * 0.62 and float(bb[0]) > float(min_x) + content_w * 0.12)
    except Exception:
        return False


def _bk_layout_row_style(row, content_bounds=None, median_height=12.0):
    """Stable formatting for Layout mit Tabellen.

    The first implementation inferred size/bold/italic/underline independently for
    every OCR fragment from box height and pixel density. Historical scans have
    irregular OCR boxes, so this created random-looking Office formatting.  This
    version keeps body text uniform and only promotes clear layout headings.
    """
    text = _bk_layout_row_text(row)
    bb = (row or {}).get("bbox")
    style = {"bold": False, "italic": False, "underline": False, "size_ratio": 1.0}
    if not text:
        return style
    try:
        row_h = max(1.0, float(bb[3]) - float(bb[1])) if bb else float(median_height or 12.0)
    except Exception:
        row_h = float(median_height or 12.0)
    med = max(1.0, float(median_height or 12.0))
    centered = _bk_layout_is_centered(row, content_bounds) if content_bounds else False
    compact = len(text) <= 90
    roman_page = re.match(r"^[\s—–-]*[IVXLCDMivxlcdm]+[\s—–-]*$", text) is not None
    section_heading = re.match(r"^[a-z]\)\s+von\b", text, re.IGNORECASE) is not None
    table_title = re.search(r"Rabatt\s*[- ]?Tabelle|Freiexemplaren|Tabelle", text, re.IGNORECASE) is not None
    small_notice_heading = re.match(r"^[A-Za-zÄÖÜäöüßÀ-ÿ]+ische\s+kirchliche\s+Festtage", text) is not None

    if roman_page:
        style["size_ratio"] = 1.0
        return style
    if table_title:
        style["bold"] = True
        style["size_ratio"] = 1.28
        return style
    if section_heading:
        style["bold"] = True
        style["size_ratio"] = 1.22
        return style
    if small_notice_heading:
        style["bold"] = True
        style["size_ratio"] = 1.04
        return style
    if centered and compact and row_h >= med * 1.32:
        style["bold"] = True
        style["size_ratio"] = 1.14
        return style
    if centered and compact and row_h >= med * 1.18:
        style["bold"] = True
        style["size_ratio"] = 1.08
        return style
    # Body text stays one uniform style. Avoid automatic italic/underline detection;
    # OCR/image heuristics are too noisy for scans like this.
    return style



def _bk_layout_row_segments(row, min_x, char_px):
    segments = []
    cursor = 0
    for rec in sorted(row.get("items", []), key=lambda item: float(item.get("x0", 0.0) or 0.0)):
        txt = _bk_tab_clean(rec.get("text", ""))
        if not txt:
            continue
        target = max(0, int(round((float(rec.get("x0", min_x) or min_x) - float(min_x)) / max(1.0, float(char_px)))))
        if target > cursor:
            segments.append({"text": " " * (target - cursor), "record": None})
            cursor = target
        elif target < cursor:
            # Bei leicht überlappenden OCR-Boxen wenigstens ein Trennzeichen einfügen.
            segments.append({"text": " ", "record": None})
            cursor += 1
        segments.append({"text": txt, "record": rec})
        cursor += len(txt)
    return segments


def _bk_docx_set_fixed_table(table):
    try:
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        tbl_pr = table._tbl.tblPr
        layout = tbl_pr.first_child_found_in("w:tblLayout")
        if layout is None:
            layout = OxmlElement("w:tblLayout")
            tbl_pr.append(layout)
        layout.set(qn("w:type"), "fixed")
    except Exception:
        pass


def _bk_docx_set_table_indent(table, twips):
    try:
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        tbl_pr = table._tbl.tblPr
        ind = tbl_pr.first_child_found_in("w:tblInd")
        if ind is None:
            ind = OxmlElement("w:tblInd")
            tbl_pr.append(ind)
        ind.set(qn("w:w"), str(max(0, int(twips))))
        ind.set(qn("w:type"), "dxa")
    except Exception:
        pass


def _bk_docx_set_cell_width(cell, twips):
    try:
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        tc_pr = cell._tc.get_or_add_tcPr()
        tc_w = tc_pr.first_child_found_in("w:tcW")
        if tc_w is None:
            tc_w = OxmlElement("w:tcW")
            tc_pr.append(tc_w)
        tc_w.set(qn("w:w"), str(max(100, int(twips))))
        tc_w.set(qn("w:type"), "dxa")
    except Exception:
        pass


def _bk_docx_set_cell_borders(cell, top=False, bottom=False):
    """Vertikale Tabellenlinien plus optionaler oberer/unterer Abschluss.

    Historische Satzspiegel besitzen meist keine horizontale Linie zwischen
    jeder Datenzeile. Die direkte OOXML-Konfiguration ueberschreibt deshalb
    den vollstaendigen Word-Rasterstil.
    """
    try:
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        tc_pr = cell._tc.get_or_add_tcPr()
        borders = tc_pr.first_child_found_in("w:tcBorders")
        if borders is None:
            borders = OxmlElement("w:tcBorders")
            tc_pr.append(borders)
        specs = {
            "left": ("single", "4", "777777"),
            "right": ("single", "4", "777777"),
            "top": (("single", "8", "555555") if top else ("nil", "0", "auto")),
            "bottom": (("single", "8", "555555") if bottom else ("nil", "0", "auto")),
            "insideH": ("nil", "0", "auto"),
        }
        for edge, (val, size, color) in specs.items():
            tag = "w:" + edge
            element = borders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                borders.append(element)
            element.set(qn("w:val"), val)
            element.set(qn("w:sz"), size)
            element.set(qn("w:space"), "0")
            element.set(qn("w:color"), color)
    except Exception:
        pass


def _bk_layout_table_col_widths_px(block, content_bounds):
    matrix = block.get("matrix") or []
    col_count = max([len(row) for row in matrix] or [1])
    anchors = list(block.get("anchors") or [])
    bb = block.get("bbox") or content_bounds
    if len(anchors) == col_count and col_count > 1:
        widths = []
        for idx, anchor in enumerate(anchors):
            left = float(bb[0]) if idx == 0 else (float(anchors[idx-1]) + float(anchor)) / 2.0
            right = float(bb[2]) if idx == col_count - 1 else (float(anchor) + float(anchors[idx+1])) / 2.0
            widths.append(max(8.0, right - left))
        return widths
    table_width = max(1.0, float(bb[2]) - float(bb[0]))
    return [table_width / max(1, col_count)] * col_count


def _bk_write_layout_tables_docx(path, item, export_image, record_views):
    try:
        from docx import Document
        from docx.shared import Inches, Pt
        from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
    except Exception as exc:
        raise RuntimeError(_bk_registry_lookup("err_no_docx_package_short") or "python-docx") from exc
    image_size = getattr(export_image, "size", None) or (0, 0)
    layout = _bk_layout_blocks_with_tables(record_views, image_size)
    blocks = layout.get("blocks") or []
    records = layout.get("records") or []
    doc = Document()
    if not blocks:
        doc.save(path)
        return
    min_x, min_y, max_x, max_y = layout.get("bounds") or (0.0, 0.0, 1.0, 1.0)
    content_w = max(1.0, float(max_x) - float(min_x))
    content_h = max(1.0, float(max_y) - float(min_y))
    page_w_px, page_h_px = layout.get("page_size") or image_size or (1.0, 1.0)
    portrait_auto = float(page_h_px or 1.0) >= float(page_w_px or 1.0)
    portrait = bk_resolve_portrait(portrait_auto, None)
    page_w_in, page_h_in = (8.27, 11.69) if portrait else (11.69, 8.27)
    margin_in = 0.35
    section = doc.sections[0]
    section.page_width = Inches(page_w_in); section.page_height = Inches(page_h_in)
    section.left_margin = section.right_margin = Inches(margin_in)
    section.top_margin = section.bottom_margin = Inches(margin_in)
    usable_w_pt = (page_w_in - margin_in * 2.0) * 72.0
    usable_h_pt = (page_h_in - margin_in * 2.0) * 72.0
    scale_x = usable_w_pt / content_w
    scale_y = usable_h_pt / content_h
    char_px = max(2.5, float(layout.get("char_px") or _bk_layout_median_char_px(records)))
    target_cols = max(60.0, content_w / char_px)
    base_font_pt = max(5.0, min(9.5, usable_w_pt / (target_cols * 0.60)))
    median_h = float(layout.get("median_height") or _median_height(records) or 12.0)
    try:
        normal = doc.styles["Normal"]
        normal.font.name = "Courier New"
        normal.font.size = Pt(base_font_pt)
        normal.paragraph_format.space_before = Pt(0)
        normal.paragraph_format.space_after = Pt(0)
    except Exception:
        pass
    last_bottom = min_y
    for block in blocks:
        bb = block.get("bbox") or (min_x, last_bottom, max_x, last_bottom + median_h)
        gap = max(0.0, float(bb[1]) - float(last_bottom))
        if block.get("type") == "table":
            if gap > median_h * 0.35:
                sp = doc.add_paragraph()
                sp.paragraph_format.space_before = Pt(max(0.0, min(72.0, gap * scale_y)))
                sp.paragraph_format.space_after = Pt(0)
            matrix = block.get("matrix") or []
            if matrix:
                col_count = max(1, max(len(row) for row in matrix))
                table = doc.add_table(rows=len(matrix), cols=col_count)
                table.style = "Table Grid"
                table.alignment = WD_TABLE_ALIGNMENT.LEFT
                table.autofit = False
                _bk_docx_set_fixed_table(table)
                _bk_docx_set_table_indent(table, (float(bb[0]) - float(min_x)) * scale_x * 20.0)
                widths_px = _bk_layout_table_col_widths_px(block, (min_x, min_y, max_x, max_y))
                header_rows = set(block.get("header_rows") or set())
                for r_idx, row in enumerate(matrix):
                    cells = table.rows[r_idx].cells
                    for c_idx in range(col_count):
                        value = row[c_idx] if c_idx < len(row) else ""
                        cell = cells[c_idx]
                        cell.text = ""
                        p = cell.paragraphs[0]
                        p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(0)
                        run = p.add_run(str(value or ""))
                        run.font.name = "Arial"
                        run.font.size = Pt(max(5.0, min(8.5, base_font_pt * 0.95)))
                        run.bold = r_idx in header_rows
                        _bk_docx_set_cell_borders(
                            cell,
                            top=(r_idx == 0 or r_idx in header_rows),
                            bottom=(r_idx == len(matrix) - 1 or r_idx in header_rows),
                        )
                        if c_idx < len(widths_px):
                            _bk_docx_set_cell_width(cell, widths_px[c_idx] * scale_x * 20.0)
                        try:
                            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                        except Exception:
                            pass
            last_bottom = float(bb[3])
            continue
        # Normale Textzeilen: echte Absätze, keine Textfelder; horizontale Lage über Leerzeichen.
        for row in block.get("rows") or []:
            row_bb = row.get("bbox") or bb
            gap = max(0.0, float(row_bb[1]) - float(last_bottom))
            p = doc.add_paragraph()
            pf = p.paragraph_format
            pf.space_after = Pt(0)
            pf.space_before = Pt(max(0.0, min(36.0, gap * scale_y))) if gap > median_h * 0.35 else Pt(0)
            row_h = max(1.0, float(row_bb[3]) - float(row_bb[1])) if row_bb else median_h
            st = _bk_layout_row_style(row, (min_x, min_y, max_x, max_y), median_h)
            row_font = max(5.0, min(13.0, base_font_pt * float(st.get("size_ratio", 1.0))))
            pf.line_spacing = Pt(max(row_font * 1.08, min(18.0, row_h * scale_y * 0.92)))
            for seg in _bk_layout_row_segments(row, min_x, char_px):
                seg_text = seg.get("text", "")
                if not seg_text:
                    continue
                run = p.add_run(seg_text)
                run.font.name = "Courier New"
                run.font.size = Pt(row_font)
                run.bold = bool(st.get("bold"))
                run.italic = bool(st.get("italic"))
                run.underline = bool(st.get("underline"))
            last_bottom = max(float(last_bottom), float(row_bb[3]) if row_bb else float(last_bottom) + median_h)
    doc.save(path)


def _bk_odt_spaces_text(text):
    parts = []
    i = 0
    s = str(text or "")
    while i < len(s):
        if s[i] == " ":
            j = i
            while j < len(s) and s[j] == " ":
                j += 1
            count = j - i
            if count == 1:
                parts.append('<text:s/>')
            else:
                parts.append('<text:s text:c="%d"/>' % count)
            i = j
        else:
            j = i
            while j < len(s) and s[j] != " ":
                j += 1
            parts.append(_bk_xml(s[i:j]))
            i = j
    return ''.join(parts)


def _bk_odt_style_name(style, size_pt):
    flags = ''.join(ch for ch, enabled in (("b", style.get("bold")), ("i", style.get("italic")), ("u", style.get("underline"))) if enabled) or "r"
    return "T_%s_%d" % (flags, int(round(float(size_pt) * 10)))


def _bk_odt_text_style_xml(name, style, size_pt):
    props = ['fo:font-size="%.1fpt"' % float(size_pt), 'style:font-name="Courier New"']
    if style.get("bold"):
        props.append('fo:font-weight="bold"')
    if style.get("italic"):
        props.append('fo:font-style="italic"')
    if style.get("underline"):
        props.append('style:text-underline-style="solid" style:text-underline-width="auto" style:text-underline-color="font-color"')
    return '<style:style style:name="%s" style:family="text"><style:text-properties %s/></style:style>' % (name, ' '.join(props))


def _bk_write_layout_tables_odt(path, item, export_image, record_views):
    image_size = getattr(export_image, "size", None) or (0, 0)
    layout = _bk_layout_blocks_with_tables(record_views, image_size)
    blocks = layout.get("blocks") or []
    records = layout.get("records") or []
    min_x, min_y, max_x, max_y = layout.get("bounds") or (0.0, 0.0, 1.0, 1.0)
    content_w = max(1.0, float(max_x) - float(min_x))
    content_h = max(1.0, float(max_y) - float(min_y))
    page_w_px, page_h_px = layout.get("page_size") or image_size or (1.0, 1.0)
    portrait_auto = float(page_h_px or 1.0) >= float(page_w_px or 1.0)
    portrait = bk_resolve_portrait(portrait_auto, None)
    page_w_cm, page_h_cm = (21.0, 29.7) if portrait else (29.7, 21.0)
    margin_cm = 0.7
    usable_w_cm = page_w_cm - margin_cm * 2.0
    usable_h_cm = page_h_cm - margin_cm * 2.0
    scale_x = usable_w_cm / content_w
    scale_y = usable_h_cm / content_h
    char_px = max(2.5, float(layout.get("char_px") or _bk_layout_median_char_px(records)))
    target_cols = max(60.0, content_w / char_px)
    usable_w_pt = usable_w_cm / 2.54 * 72.0
    base_font_pt = max(5.0, min(9.5, usable_w_pt / (target_cols * 0.60)))
    median_h = float(layout.get("median_height") or _median_height(records) or 12.0)
    automatic_styles = []
    body_parts = []
    text_styles = {}
    table_idx = 1
    last_bottom = min_y
    def ensure_text_style(style, size_pt):
        name = _bk_odt_style_name(style, size_pt)
        if name not in text_styles:
            text_styles[name] = _bk_odt_text_style_xml(name, style, size_pt)
        return name
    for block in blocks:
        bb = block.get("bbox") or (min_x, last_bottom, max_x, last_bottom + median_h)
        gap = max(0.0, float(bb[1]) - float(last_bottom))
        margin_top = max(0.0, min(2.5, gap * scale_y)) if gap > median_h * 0.35 else 0.0
        if block.get("type") == "table":
            matrix = block.get("matrix") or []
            if matrix:
                col_count = max(1, max(len(row) for row in matrix))
                widths_px = _bk_layout_table_col_widths_px(block, (min_x, min_y, max_x, max_y))
                col_styles = []
                table_cols = []
                for c in range(col_count):
                    width_cm = max(0.35, min(6.0, (widths_px[c] if c < len(widths_px) else (float(bb[2])-float(bb[0]))/col_count) * scale_x))
                    cname = "BKLTco%d_%d" % (table_idx, c + 1)
                    col_styles.append('<style:style style:name="%s" style:family="table-column"><style:table-column-properties style:column-width="%.3fcm"/></style:style>' % (cname, width_cm))
                    table_cols.append('<table:table-column table:style-name="%s"/>' % cname)
                tname = "BKLTtbl%d" % table_idx
                automatic_styles.extend(col_styles)
                automatic_styles.append('<style:style style:name="%s" style:family="table"><style:table-properties table:align="left" fo:margin-left="%.3fcm" fo:margin-top="%.3fcm"/></style:style>' % (tname, max(0.0, (float(bb[0])-float(min_x))*scale_x), margin_top))
                rows_xml = []
                header_rows = set(block.get("header_rows") or set())
                for r_idx, row in enumerate(matrix):
                    cells = []
                    for c_idx in range(col_count):
                        value = row[c_idx] if c_idx < len(row) else ""
                        if r_idx in header_rows:
                            style = "BKLTceHead"
                        elif len(matrix) == 1:
                            style = "BKLTceOnly"
                        elif r_idx == 0:
                            style = "BKLTceTop"
                        elif r_idx == len(matrix) - 1:
                            style = "BKLTceBottom"
                        else:
                            style = "BKLTceBody"
                        cells.append('<table:table-cell table:style-name="%s" office:value-type="string"><text:p text:style-name="BKLTCellP">%s</text:p></table:table-cell>' % (style, _bk_odt_spaces_text(value)))
                    rows_xml.append('<table:table-row>%s</table:table-row>' % ''.join(cells))
                body_parts.append('<table:table table:name="LayoutTabelle%d" table:style-name="%s">%s%s</table:table>' % (table_idx, tname, ''.join(table_cols), ''.join(rows_xml)))
                table_idx += 1
            last_bottom = float(bb[3])
            continue
        for row in block.get("rows") or []:
            row_bb = row.get("bbox") or bb
            gap = max(0.0, float(row_bb[1]) - float(last_bottom))
            margin_top = max(0.0, min(1.3, gap * scale_y)) if gap > median_h * 0.35 else 0.0
            row_h = max(1.0, float(row_bb[3]) - float(row_bb[1])) if row_bb else median_h
            st = _bk_layout_row_style(row, (min_x, min_y, max_x, max_y), median_h)
            row_font = max(5.0, min(13.0, base_font_pt * float(st.get("size_ratio", 1.0))))
            pname = "BKLTp_%d_%d" % (len(automatic_styles) + len(body_parts) + 1, int(round(margin_top * 100)))
            automatic_styles.append('<style:style style:name="%s" style:family="paragraph"><style:paragraph-properties fo:margin-top="%.3fcm" fo:margin-bottom="0cm" fo:line-height="110%%"/></style:style>' % (pname, margin_top))
            spans = []
            sname = ensure_text_style(st, row_font)
            for seg in _bk_layout_row_segments(row, min_x, char_px):
                seg_text = seg.get("text", "")
                if not seg_text:
                    continue
                spans.append('<text:span text:style-name="%s">%s</text:span>' % (sname, _bk_odt_spaces_text(seg_text)))
            body_parts.append('<text:p text:style-name="%s">%s</text:p>' % (pname, ''.join(spans)))
            last_bottom = max(float(last_bottom), float(row_bb[3]) if row_bb else float(last_bottom) + median_h)
    automatic_styles.extend(text_styles.values())
    automatic_styles.extend([
        '<style:style style:name="BKLTCellP" style:family="paragraph"><style:paragraph-properties fo:margin-top="0cm" fo:margin-bottom="0cm" fo:line-height="100%%"/><style:text-properties fo:font-size="7.2pt" style:font-name="Arial"/></style:style>',
        '<style:style style:name="BKLTceHead" style:family="table-cell"><style:table-cell-properties fo:border-left="0.05pt solid #777777" fo:border-right="0.05pt solid #777777" fo:border-top="0.5pt solid #555555" fo:border-bottom="0.5pt solid #555555" fo:padding="0.03cm"/><style:text-properties fo:font-weight="bold" fo:font-size="7.2pt" style:font-name="Arial"/></style:style>',
        '<style:style style:name="BKLTceTop" style:family="table-cell"><style:table-cell-properties fo:border-left="0.05pt solid #888888" fo:border-right="0.05pt solid #888888" fo:border-top="0.5pt solid #555555" fo:border-bottom="none" fo:padding="0.03cm"/><style:text-properties fo:font-size="7.2pt" style:font-name="Arial"/></style:style>',
        '<style:style style:name="BKLTceBody" style:family="table-cell"><style:table-cell-properties fo:border-left="0.05pt solid #888888" fo:border-right="0.05pt solid #888888" fo:border-top="none" fo:border-bottom="none" fo:padding="0.03cm"/><style:text-properties fo:font-size="7.2pt" style:font-name="Arial"/></style:style>',
        '<style:style style:name="BKLTceBottom" style:family="table-cell"><style:table-cell-properties fo:border-left="0.05pt solid #888888" fo:border-right="0.05pt solid #888888" fo:border-top="none" fo:border-bottom="0.5pt solid #555555" fo:padding="0.03cm"/><style:text-properties fo:font-size="7.2pt" style:font-name="Arial"/></style:style>',
        '<style:style style:name="BKLTceOnly" style:family="table-cell"><style:table-cell-properties fo:border-left="0.05pt solid #888888" fo:border-right="0.05pt solid #888888" fo:border-top="0.5pt solid #555555" fo:border-bottom="0.5pt solid #555555" fo:padding="0.03cm"/><style:text-properties fo:font-size="7.2pt" style:font-name="Arial"/></style:style>',
    ])
    content = ''.join([
        '<?xml version="1.0" encoding="UTF-8"?>',
        '<office:document-content xmlns:office="urn:oasis:names:tc:opendocument:xmlns:office:1.0" xmlns:style="urn:oasis:names:tc:opendocument:xmlns:style:1.0" xmlns:text="urn:oasis:names:tc:opendocument:xmlns:text:1.0" xmlns:table="urn:oasis:names:tc:opendocument:xmlns:table:1.0" xmlns:fo="urn:oasis:names:tc:opendocument:xmlns:xsl-fo-compatible:1.0" xmlns:svg="urn:oasis:names:tc:opendocument:xmlns:svg-compatible:1.0" office:version="1.2">',
        '<office:font-face-decls><style:font-face style:name="Courier New" svg:font-family="Courier New"/><style:font-face style:name="Arial" svg:font-family="Arial"/></office:font-face-decls>',
        '<office:automatic-styles>', ''.join(automatic_styles), '</office:automatic-styles>',
        '<office:body><office:text>', ''.join(body_parts), '</office:text></office:body></office:document-content>',
    ])
    styles = ''.join([
        '<?xml version="1.0" encoding="UTF-8"?>',
        '<office:document-styles xmlns:office="urn:oasis:names:tc:opendocument:xmlns:office:1.0" xmlns:style="urn:oasis:names:tc:opendocument:xmlns:style:1.0" xmlns:text="urn:oasis:names:tc:opendocument:xmlns:text:1.0" xmlns:fo="urn:oasis:names:tc:opendocument:xmlns:xsl-fo-compatible:1.0" xmlns:svg="urn:oasis:names:tc:opendocument:xmlns:svg-compatible:1.0" office:version="1.2">',
        '<office:font-face-decls><style:font-face style:name="Courier New" svg:font-family="Courier New"/><style:font-face style:name="Arial" svg:font-family="Arial"/></office:font-face-decls>',
        '<office:styles><style:default-style style:family="paragraph"><style:text-properties fo:font-size="%.1fpt" style:font-name="Courier New"/></style:default-style></office:styles>' % base_font_pt,
        '<office:automatic-styles><style:page-layout style:name="pm1"><style:page-layout-properties fo:page-width="%.2fcm" fo:page-height="%.2fcm" fo:margin-top="%.2fcm" fo:margin-bottom="%.2fcm" fo:margin-left="%.2fcm" fo:margin-right="%.2fcm"/></style:page-layout></office:automatic-styles>' % (page_w_cm, page_h_cm, margin_cm, margin_cm, margin_cm, margin_cm),
        '<office:master-styles><style:master-page style:name="Standard" style:page-layout-name="pm1"/></office:master-styles></office:document-styles>',
    ])
    meta = '<?xml version="1.0" encoding="UTF-8"?><office:document-meta xmlns:office="urn:oasis:names:tc:opendocument:xmlns:office:1.0" xmlns:meta="urn:oasis:names:tc:opendocument:xmlns:meta:1.0" office:version="1.2"><office:meta><meta:generator>Bottled Kraken</meta:generator></office:meta></office:document-meta>'
    settings = '<?xml version="1.0" encoding="UTF-8"?><office:document-settings xmlns:office="urn:oasis:names:tc:opendocument:xmlns:office:1.0" office:version="1.2"><office:settings/></office:document-settings>'
    manifest = '<?xml version="1.0" encoding="UTF-8"?><manifest:manifest xmlns:manifest="urn:oasis:names:tc:opendocument:xmlns:manifest:1.0" manifest:version="1.2"><manifest:file-entry manifest:full-path="/" manifest:media-type="application/vnd.oasis.opendocument.text"/><manifest:file-entry manifest:full-path="content.xml" manifest:media-type="text/xml"/><manifest:file-entry manifest:full-path="styles.xml" manifest:media-type="text/xml"/><manifest:file-entry manifest:full-path="meta.xml" manifest:media-type="text/xml"/><manifest:file-entry manifest:full-path="settings.xml" manifest:media-type="text/xml"/></manifest:manifest>'
    with zipfile.ZipFile(path, "w") as archive:
        info = zipfile.ZipInfo("mimetype"); info.date_time = (2020, 1, 1, 0, 0, 0); info.compress_type = zipfile.ZIP_STORED
        archive.writestr(info, "application/vnd.oasis.opendocument.text")
        for name, data in (("content.xml", content), ("styles.xml", styles), ("meta.xml", meta), ("settings.xml", settings), ("META-INF/manifest.xml", manifest)):
            zi = zipfile.ZipInfo(name); zi.date_time = (2020, 1, 1, 0, 0, 0); zi.compress_type = zipfile.ZIP_DEFLATED
            archive.writestr(zi, data.encode("utf-8"))


def _bk_column_choice_dialog_final2(self, fmt=None, include_text_modes=False):
    _bk_ensure_custom_columns(self)
    fmt_l = str(fmt or "").lower().lstrip(".")
    txt_only = fmt_l in {"txt", "text", "txt_plain"}
    csv_pure = fmt_l == "csv"
    if csv_pure:
        return {"mode": "csv", "columns": [], "remembered": False, "use_zones": False, "cancelled": False}
    supports_simple = fmt_l in {"xlsx", "excel", "ods", "calc", "docx", "word", "odt"} and not txt_only
    supports_layout_tables = include_text_modes and fmt_l in _BK_LAYOUT_TABLE_TEXT_FMTS and not txt_only
    supports_extended = fmt_l in {"xlsx", "excel", "ods", "calc", "docx", "word", "odt", "json"} and not txt_only
    dlg = QDialog(self)
    dlg.setWindowTitle(_bk_tab_tr(self, "export_text_layout_title" if include_text_modes else "export_table_columns_title"))
    dlg.setModal(True)
    layout = QVBoxLayout(dlg); layout.setContentsMargins(14, 12, 14, 12); layout.setSpacing(8)
    intro_key = "export_text_layout_intro_extended" if include_text_modes else "export_table_columns_intro"
    label = QLabel(_bk_tab_tr(self, intro_key), dlg); label.setWordWrap(True); layout.addWidget(label)

    rb_original = rb_lines = rb_layout_tables = rb_simple = rb_table = None
    mode_box = QGroupBox(_bk_tab_tr(self, "export_layout_mode_group"), dlg) if QGroupBox is not None else None
    mode_layout = QVBoxLayout(mode_box) if mode_box is not None else QVBoxLayout(); mode_layout.setContentsMargins(10, 8, 10, 8)
    if include_text_modes:
        rb_original = QRadioButton(_bk_tab_tr(self, "export_text_layout_original"), dlg)
        rb_lines = QRadioButton(_bk_tab_tr(self, "export_text_layout_lines"), dlg)
        mode = str(getattr(self, "_bk_export_text_layout_mode", "original") or "original").lower()
        rb_original.setChecked(mode not in {"lines", "table", "table_simple", "layout_tables"} or txt_only)
        rb_lines.setChecked(mode == "lines")
        mode_layout.addWidget(rb_original); mode_layout.addWidget(rb_lines)
    if supports_layout_tables:
        rb_layout_tables = QRadioButton(_bk_tab_tr(self, "export_text_layout_with_tables"), dlg); mode_layout.addWidget(rb_layout_tables)
    if supports_simple:
        rb_simple = QRadioButton(_bk_tab_tr(self, "export_text_layout_table_simple"), dlg); mode_layout.addWidget(rb_simple)
    if supports_extended:
        rb_table = QRadioButton(_bk_tab_tr(self, "export_text_layout_table"), dlg); mode_layout.addWidget(rb_table)
    if not include_text_modes and rb_simple is not None:
        rb_simple.setChecked(True)
    elif not include_text_modes and rb_table is not None:
        rb_table.setChecked(True)
    elif include_text_modes and not txt_only:
        mode = str(getattr(self, "_bk_export_text_layout_mode", "original") or "original").lower()
        if rb_layout_tables is not None: rb_layout_tables.setChecked(mode == "layout_tables")
        if rb_simple is not None: rb_simple.setChecked(mode == "table_simple")
        if rb_table is not None: rb_table.setChecked(mode == "table")
    if mode_box is not None: layout.addWidget(mode_box)
    else: layout.addLayout(mode_layout)

    _bk_add_export_orientation_group(self, dlg, layout)

    selected_keys = _bk_load_saved_column_keys_for_dialog(self)
    checkboxes = {}
    columns_box = QGroupBox(_bk_tab_tr(self, "export_table_columns_label"), dlg) if QGroupBox is not None else None
    columns_layout = QVBoxLayout(columns_box) if columns_box is not None else QVBoxLayout()
    grid = QGridLayout() if QGridLayout is not None else None
    if grid is not None:
        grid.setHorizontalSpacing(18); grid.setVerticalSpacing(4); columns_layout.addLayout(grid)
    def add_cb_to_grid(key, title, checked=False):
        cb = QCheckBox(title, dlg); cb.setChecked(bool(checked)); checkboxes[key] = cb
        if grid is not None:
            idx = len(checkboxes) - 1; grid.addWidget(cb, idx % 8, idx // 8)
        else:
            columns_layout.addWidget(cb)
        return cb
    for key in list(_BK_TABULAR_KEYS):
        add_cb_to_grid(key, _bk_tabular_column_title(self, key), key in selected_keys)

    custom_edit = None
    if _BK_QLineEdit is not None:
        custom_row = QHBoxLayout()
        custom_label = QLabel(_bk_tab_tr(self, "export_table_custom_column_label"), dlg)
        custom_edit = _BK_QLineEdit(dlg); custom_edit.setPlaceholderText(_bk_tab_tr(self, "export_table_custom_column_placeholder"))
        btn_add_custom = QPushButton(_bk_tab_tr(self, "export_table_custom_column_add"), dlg)
        btn_delete_custom = QPushButton(_bk_tab_tr(self, "export_table_custom_column_delete"), dlg)
        btn_reset_custom = QPushButton(_bk_tab_tr(self, "export_table_custom_column_reset"), dlg)
        custom_row.addWidget(custom_label); custom_row.addWidget(custom_edit, 1); custom_row.addWidget(btn_add_custom); custom_row.addWidget(btn_delete_custom); custom_row.addWidget(btn_reset_custom)
        columns_layout.addLayout(custom_row)
    else:
        btn_add_custom = btn_delete_custom = btn_reset_custom = None

    quick_row = QHBoxLayout()
    btn_all = QPushButton(_bk_tab_tr(self, "export_table_columns_all"), dlg)
    btn_none = QPushButton(_bk_tab_tr(self, "export_table_columns_none_button"), dlg)
    btn_remember = QPushButton(_bk_tab_tr(self, "export_table_columns_remember"), dlg)
    for btn in (btn_all, btn_none, btn_remember): quick_row.addWidget(btn)
    quick_row.addStretch(1); columns_layout.addLayout(quick_row)
    if columns_box is not None: layout.addWidget(columns_box)
    else: layout.addLayout(columns_layout)

    zone_box = QGroupBox(_bk_tab_tr(self, "export_zones_group"), dlg) if QGroupBox is not None else None
    zone_layout = QHBoxLayout(zone_box) if zone_box is not None else QHBoxLayout()
    cb_zones = QCheckBox(_bk_tab_tr(self, "export_table_use_zones"), dlg)
    try: remembered_zones = self.settings.value("export/table_use_zones", bool(getattr(self, "_bk_export_use_zones", False)), type=bool)
    except Exception: remembered_zones = bool(getattr(self, "_bk_export_use_zones", False))
    cb_zones.setChecked(bool(remembered_zones))
    btn_zones = QPushButton(_bk_tab_tr(self, "export_table_define_zones"), dlg)
    zone_layout.addWidget(cb_zones, 1); zone_layout.addWidget(btn_zones)
    if zone_box is not None: layout.addWidget(zone_box)
    else: layout.addLayout(zone_layout)

    hint = QLabel(_bk_tab_tr(self, "export_table_simple_hint"), dlg)
    hint.setWordWrap(True); layout.addWidget(hint)

    def current_checked_keys(): return [key for key, cb in checkboxes.items() if cb.isChecked() and key in _BK_TABULAR_COLUMN_BY_KEY]
    def set_all():
        for cb in checkboxes.values(): cb.setChecked(True)
    def set_none():
        for cb in checkboxes.values(): cb.setChecked(False)
    def add_custom_column():
        if custom_edit is None: return
        title = _bk_tab_clean(custom_edit.text())
        if not title: return
        existing_titles = {_bk_tabular_column_title(self, key).casefold() for key in _BK_TABULAR_KEYS}
        if title.casefold() in existing_titles:
            QMessageBox.warning(dlg, _bk_tab_tr(self, "warn_title"), _bk_tab_tr(self, "export_table_custom_column_exists")); return
        base = _bk_custom_column_slug(title); key = "custom_" + base; used = set(_BK_TABULAR_KEYS); i = 2
        while key in used:
            key = f"custom_{base}_{i}"; i += 1
        items = _bk_load_custom_columns(self) + [{"key": key, "title": title}]
        _bk_save_custom_columns(self, items); _bk_ensure_custom_columns(self)
        add_cb_to_grid(key, title, True); custom_edit.clear(); sync_enabled(); dlg.adjustSize()
    def delete_custom_columns():
        selected_custom = {key for key, cb in checkboxes.items() if key.startswith("custom_") and cb.isChecked()}
        if not selected_custom: return
        items = [entry for entry in _bk_load_custom_columns(self) if entry.get("key") not in selected_custom]
        _bk_save_custom_columns(self, items); _bk_ensure_custom_columns(self)
        for key in list(selected_custom):
            cb = checkboxes.pop(key, None)
            if cb is not None:
                try:
                    if grid is not None: grid.removeWidget(cb)
                    cb.setParent(None); cb.deleteLater()
                except Exception: pass
        sync_enabled(); dlg.adjustSize()
    def reset_custom_columns():
        _bk_save_custom_columns(self, []); _bk_ensure_custom_columns(self)
        for key in list(checkboxes.keys()):
            if key.startswith("custom_"):
                cb = checkboxes.pop(key, None)
                if cb is not None:
                    try:
                        if grid is not None: grid.removeWidget(cb)
                        cb.setParent(None); cb.deleteLater()
                    except Exception: pass
        sync_enabled(); dlg.adjustSize()
    if btn_add_custom is not None:
        btn_add_custom.clicked.connect(add_custom_column); btn_delete_custom.clicked.connect(delete_custom_columns); btn_reset_custom.clicked.connect(reset_custom_columns)
        try: custom_edit.returnPressed.connect(add_custom_column)
        except Exception: pass

    result = {"mode": "table", "columns": selected_keys, "remembered": False, "use_zones": cb_zones.isChecked(), "cancelled": False}
    def remember_selection():
        keys = current_checked_keys()
        if not keys:
            QMessageBox.warning(dlg, _bk_tab_tr(self, "warn_title"), _bk_tab_tr(self, "export_table_columns_none")); return
        result["remembered"] = True; result["columns"] = _bk_save_column_keys(self, keys)
        try: self.settings.setValue("export/table_use_zones", bool(cb_zones.isChecked())); self.settings.sync()
        except Exception: pass
    def define_zones():
        _bk_ensure_custom_columns(self)
        try: task = self._current_task()
        except Exception: task = None
        zones = _bk_open_export_zones_dialog(self, task)
        if zones is not None: cb_zones.setChecked(bool(zones))
    btn_all.clicked.connect(set_all); btn_none.clicked.connect(set_none); btn_remember.clicked.connect(remember_selection); btn_zones.clicked.connect(define_zones)
    def sync_enabled():
        detailed = bool(rb_table is not None and rb_table.isChecked())
        if columns_box is not None: columns_box.setVisible(detailed); columns_box.setEnabled(detailed)
        if zone_box is not None: zone_box.setVisible(detailed); zone_box.setEnabled(detailed)
        cb_zones.setEnabled(detailed); btn_zones.setEnabled(detailed)
        if rb_layout_tables is not None and rb_layout_tables.isChecked():
            hint.setText(_bk_tab_tr(self, "export_layout_with_tables_hint"))
            hint.setVisible(True)
        else:
            hint.setText(_bk_tab_tr(self, "export_table_simple_hint"))
            hint.setVisible(bool(rb_simple is not None and rb_simple.isChecked()))
        try:
            _bk_resize_export_dialog_for_mode(dlg, detailed=detailed, txt_only=txt_only)
        except Exception:
            pass
    for rb in (rb_original, rb_lines, rb_layout_tables, rb_simple, rb_table):
        if rb is not None:
            try: rb.toggled.connect(sync_enabled)
            except Exception: pass
    sync_enabled()
    buttons = QDialogButtonBox(QDialogButtonBox.Ok | QDialogButtonBox.Cancel, dlg)
    try:
        buttons.button(QDialogButtonBox.Ok).setText(_bk_tab_tr(self, "btn_ok")); buttons.button(QDialogButtonBox.Cancel).setText(_bk_tab_tr(self, "btn_cancel"))
    except Exception: pass
    def cancel_dialog(): result["cancelled"] = True; dlg.done(QDialog.Rejected)
    def accept_checked():
        result["cancelled"] = False
        if rb_layout_tables is not None and rb_layout_tables.isChecked(): mode = "layout_tables"; keys = []
        elif rb_simple is not None and rb_simple.isChecked(): mode = "table_simple"; keys = []
        elif include_text_modes and rb_lines is not None and rb_lines.isChecked(): mode = "lines"; keys = []
        elif include_text_modes and rb_original is not None and rb_original.isChecked(): mode = "original"; keys = []
        else:
            mode = "table"; keys = current_checked_keys()
            if not keys:
                QMessageBox.warning(dlg, _bk_tab_tr(self, "warn_title"), _bk_tab_tr(self, "export_table_columns_none")); return
        result["mode"] = mode; result["columns"] = _bk_normalize_column_keys(keys); result["use_zones"] = bool(cb_zones.isChecked()) if mode == "table" else False
        try: self._bk_export_use_zones = result["use_zones"]
        except Exception: pass
        dlg.accept()
    buttons.accepted.connect(accept_checked); buttons.rejected.connect(cancel_dialog); layout.addWidget(buttons)
    try:
        sync_enabled()
        _bk_keep_dialog_inside_screen(dlg)
    except Exception:
        pass
    try: exec_result = dlg.exec()
    except Exception: result["cancelled"] = True; return None
    if exec_result != QDialog.Accepted:
        result["cancelled"] = True; return None
    return result


def _bk_render_file_final3(self, path: str, fmt: str, item: TaskItem):
    fmt_l = str(fmt or "").lower().lstrip(".")
    if item and getattr(item, "results", None):
        if fmt_l == "csv":
            return _bk_write_plain_overlay_csv_final(path, item)
        if fmt_l in {"xlsx", "excel", "ods", "calc"} and str(getattr(self, "_bk_export_table_mode", "simple") or "simple").lower() == "simple":
            # "Tabelle (einfach)" ist strikt raeumlich. Der fruehere Versuch,
            # beliebige Seiten automatisch als Personenregister zu deuten,
            # erzeugte bei historischen Sach- und Statistiktabellen erfundene
            # Spalten wie Familienname/Zusatz/Alter. Semantische Benutzer-
            # spalten bleiben ausschliesslich im Modus "Tabelle (erweitert)".
            layout = _bk_simple_spatial_layout_from_item(item)
            if fmt_l in {"xlsx", "excel"}: return _bk_write_simple_layout_xlsx(path, layout, self)
            if fmt_l in {"ods", "calc"}: return _bk_write_simple_layout_ods(path, layout, self)
        text_layout_mode = str(getattr(self, "_bk_export_text_layout_mode", "original") or "original").lower()
        if fmt_l in {"docx", "word", "odt"} and text_layout_mode == "layout_tables":
            _text, _kr, pil_image, record_views = item.results
            try:
                export_image = _load_image_color(item.path)
            except Exception:
                export_image = pil_image
            if fmt_l in {"docx", "word"}: return _bk_write_layout_tables_docx(path, item, export_image, record_views)
            if fmt_l == "odt": return _bk_write_layout_tables_odt(path, item, export_image, record_views)
        if fmt_l in {"docx", "word", "odt"} and text_layout_mode == "table_simple":
            # Auch bei Writer/Word muss "Tabelle (einfach)" die vorhandenen
            # Overlay-Positionen abbilden und darf keine genealogische
            # Registerstruktur hineininterpretieren.
            layout = _bk_simple_spatial_layout_from_item(item); matrix = layout.get("matrix") or []
            if fmt_l in {"docx", "word"}: return _bk_write_simple_matrix_docx(path, matrix, self, layout=layout)
            if fmt_l == "odt": return _bk_write_simple_matrix_odt(path, matrix, self, layout=layout)
    return RENDER_NOT_HANDLED


def _bk_export_single_interactive_final3(self, item: TaskItem, fmt: str):
    fmt_l = str(fmt or "").lower().lstrip(".")
    if fmt_l == "csv":
        self._bk_export_table_mode = "csv"; self._bk_export_use_zones = False
        if callable(_BK_TABULAR_PREV_EXPORT_SINGLE): return _BK_TABULAR_PREV_EXPORT_SINGLE(self, item, fmt)
        if callable(_BK_FINAL_PREV_EXPORT_SINGLE_2): return _BK_FINAL_PREV_EXPORT_SINGLE_2(self, item, fmt)
        return None
    if fmt_l in {"xlsx", "excel", "ods", "calc"}:
        result = _bk_column_choice_dialog_final2(self, fmt_l, include_text_modes=False)
        if result is None or result.get("cancelled"): return None
        if result.get("mode") == "table_simple":
            self._bk_export_table_mode = "simple"; self._bk_export_use_zones = False
        else:
            self._bk_export_table_mode = "table"; self._bk_export_current_column_keys = result.get("columns") or []; self._bk_export_use_zones = bool(result.get("use_zones", False))
            if result.get("remembered"): self._bk_export_selected_column_keys = self._bk_export_current_column_keys
        if callable(_BK_TABULAR_PREV_EXPORT_SINGLE): return _BK_TABULAR_PREV_EXPORT_SINGLE(self, item, fmt)
        return None
    if fmt_l in _BK_TEXT_LAYOUT_FMTS:
        result = _bk_column_choice_dialog_final2(self, fmt_l, include_text_modes=True)
        if result is None or result.get("cancelled"): return None
        self._bk_export_text_layout_mode = str(result.get("mode") or "original")
        self._bk_export_current_column_keys = result.get("columns") or []
        self._bk_export_use_zones = bool(result.get("use_zones", False)) if self._bk_export_text_layout_mode == "table" else False
        if result.get("remembered"): self._bk_export_selected_column_keys = self._bk_export_current_column_keys
        if callable(_BK_TABULAR_PREV_EXPORT_SINGLE): return _BK_TABULAR_PREV_EXPORT_SINGLE(self, item, fmt)
        return None
    if callable(_BK_FINAL_PREV_EXPORT_SINGLE_2): return _BK_FINAL_PREV_EXPORT_SINGLE_2(self, item, fmt)
    return None


def _bk_export_batch_final3(self, items, fmt: str):
    fmt_l = str(fmt or "").lower().lstrip(".")
    if fmt_l == "csv":
        self._bk_export_table_mode = "csv"; self._bk_export_use_zones = False
        if callable(_BK_TABULAR_PREV_EXPORT_BATCH): return _BK_TABULAR_PREV_EXPORT_BATCH(self, items, fmt)
        if callable(_BK_FINAL_PREV_EXPORT_BATCH_2): return _BK_FINAL_PREV_EXPORT_BATCH_2(self, items, fmt)
        return None
    if fmt_l in {"xlsx", "excel", "ods", "calc"}:
        result = _bk_column_choice_dialog_final2(self, fmt_l, include_text_modes=False)
        if result is None or result.get("cancelled"): return None
        if result.get("mode") == "table_simple": self._bk_export_table_mode = "simple"; self._bk_export_use_zones = False
        else:
            self._bk_export_table_mode = "table"; self._bk_export_current_column_keys = result.get("columns") or []; self._bk_export_use_zones = bool(result.get("use_zones", False))
            if result.get("remembered"): self._bk_export_selected_column_keys = self._bk_export_current_column_keys
        if callable(_BK_TABULAR_PREV_EXPORT_BATCH): return _BK_TABULAR_PREV_EXPORT_BATCH(self, items, fmt)
        return None
    if fmt_l in _BK_TEXT_LAYOUT_FMTS:
        result = _bk_column_choice_dialog_final2(self, fmt_l, include_text_modes=True)
        if result is None or result.get("cancelled"): return None
        self._bk_export_text_layout_mode = str(result.get("mode") or "original")
        self._bk_export_current_column_keys = result.get("columns") or []
        self._bk_export_use_zones = bool(result.get("use_zones", False)) if self._bk_export_text_layout_mode == "table" else False
        if result.get("remembered"): self._bk_export_selected_column_keys = self._bk_export_current_column_keys
        if callable(_BK_TABULAR_PREV_EXPORT_BATCH): return _BK_TABULAR_PREV_EXPORT_BATCH(self, items, fmt)
        return None
    if callable(_BK_FINAL_PREV_EXPORT_BATCH_2): return _BK_FINAL_PREV_EXPORT_BATCH_2(self, items, fmt)
    return None

try:
    __all__ = sorted(set(list(__all__) + [
        '_bk_semantic_blocks_from_item', '_bk_semantic_blocks_write',
        '_bk_simple_table_records', '_bk_simple_group_rows',
        # Gemeinsame Layoutanalyse und Writer fuer den normalen und den
        # lokalen-LM-Export. Ohne Registrierung faellt lm_structured_export
        # still auf die alte ungeordnete Zeilenliste zurueck.
        '_bk_layout_blocks_with_tables', '_bk_layout_row_text',
        '_bk_simple_spatial_layout_from_item', '_bk_simple_col_widths',
        '_bk_write_simple_layout_xlsx',
        '_bk_add_export_orientation_group', '_bk_orientation_from_radios',
        '_bk_registry_lookup', '_bk_tr_registry', '_bk_dialog_content_min_size',
        '_bk_ai_zone_orientation_hint',
        '_bk_write_simple_layout_ods', '_bk_write_simple_matrix_odt',
    ]))
    _bk_column_choice_dialog = _bk_column_choice_dialog_final2
    register_render_handler(_bk_render_file_final3)
    MainWindow._export_single_interactive = _bk_export_single_interactive_final3
    MainWindow._export_batch = _bk_export_batch_final3
    register_globals('bk', globals(), sorted(set(__all__)))
except Exception:
    # Diese Registrierung entscheidet, ob die Exportfunktionen der letzten
    # Generation ueberhaupt aktiv sind. Ein stilles pass liess das Programm
    # frueher scheinbar normal starten - nur eben mit alten/fehlenden
    # Exportpfaden und ohne jede Fehlermeldung. Der Fehler wird deshalb
    # jetzt sichtbar gemacht (App startet weiter, aber diagnostizierbar).
    try:
        import sys as _sys
        import traceback as _traceback
        print("[bottled_kraken] FEHLER: Finale Export-Registrierung fehlgeschlagen - "
              "Exportfunktionen sind moeglicherweise unvollstaendig!", file=_sys.stderr)
        _traceback.print_exc(file=_sys.stderr)
    except Exception:
        pass
